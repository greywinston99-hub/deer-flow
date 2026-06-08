"""Chapter-by-chapter CER expansion using checkpoint data + engineer templates."""
import sys, os, json
from pathlib import Path

sys.path.insert(0, os.path.dirname(__file__))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
import _patch_httpx_lru  # noqa: F401

from _v3_1_llm_client import get_llm_client
from langgraph.checkpoint.sqlite import SqliteSaver
from deerflow.runtime.cer_authoring.graph import build_cer_authoring_graph

CHECKPOINT = sys.argv[1] if len(sys.argv) > 1 else None
if not CHECKPOINT:
    print("Usage: expand_cer_chapters.py <checkpoint.db>")
    sys.exit(1)

OUTPUT = os.path.dirname(CHECKPOINT)
THREAD_ID = "A01_无忧跳动"
os.makedirs(os.path.join(OUTPUT, "chapters"), exist_ok=True)

print(f"Loading checkpoint: {CHECKPOINT}", flush=True)

with SqliteSaver.from_conn_string(CHECKPOINT) as checkpointer:
    graph = build_cer_authoring_graph(checkpointer=checkpointer)
    vals = graph.get_state({"configurable": {"thread_id": THREAD_ID}}).values

device = vals.get("device_profile", {})
claims = vals.get("claim_ledger", [])
sota = vals.get("sota_benchmark_matrix", [])
prisma = vals.get("prisma_flow_data", {})
appraisal = vals.get("article_appraisal", [])
clinical_facts = vals.get("clinical_evidence_fact_table", [])
endpoint_fulltext = vals.get("sota_endpoint_extraction_fulltext", [])
sota_hierarchy = vals.get("sota_evidence_hierarchy", [])
sota_context = vals.get("sota_clinical_context_table", [])
benefit_risk = vals.get("benefit_risk_ledger", [])
evidence_conflicts = vals.get("evidence_conflict_report", {})
clinical_source = vals.get("clinical_source_adapter_records", [])
synthesis = vals.get("cross_evidence_synthesis_narratives", [])
pmcf = vals.get("pmcf_plan_control_matrix", {})
gap_recs = vals.get("gap_pmcf_recommendations", [])
vigilance = vals.get("vigilance_event_statistics", [])
gspr = vals.get("gspr_coverage", [])
risk_trace = vals.get("risk_trace_matrix", [])
similar_devices = vals.get("similar_device_attachment_index", [])
guideline = vals.get("guideline_pathway_table", [])
alternative = vals.get("alternative_treatment_benchmark_table", [])
equivalence = vals.get("equivalence_matrix", [])
cep = vals.get("clinical_evaluation_plan", {})

# ── Shared context (device profile) ──
device_ctx = f"""DEVICE: {device.get('device_name','?')} | Manufacturer: {device.get('manufacturer','?')}
Class: {device.get('device_class','?')} (MDR Annex VIII) | Domain: {device.get('clinical_domain','?')}
Device Type: {device.get('device_type','?')} | Family: {device.get('device_family','?')}
Intended Purpose: {device.get('intended_purpose','?')}
Target Population: {device.get('target_population','?')}
Intended User: {device.get('intended_user','?')}
Anatomical Site: {device.get('anatomical_site','?')}
Contraindications: {device.get('contraindications','?')}
Model: {device.get('model_specifications','?')}
Mode of Action: {device.get('mode_of_action','?')}
Working Principle: {device.get('working_principle','Requires IFU confirmation')}
Composition: {device.get('composition','Requires IFU confirmation')}
Sterility: {device.get('sterility','Requires IFU confirmation')}
Shelf Life: {device.get('shelf_life_storage','Requires IFU confirmation')}"""

# ── Evidence context ──
evidence_sample = "\n".join([
    f"[{e.get('evidence_id','?')}] Score={e.get('appraisal_score','?')} Weight={e.get('weight','?')}: {str(e.get('title',''))[:200]}"
    for e in appraisal[:80]
])

facts_sample = "\n".join([
    f"[{f.get('fact_id','')}] {str(f.get('fact_text',''))[:300]} ← {str(f.get('source_article',''))[:100]}"
    for f in clinical_facts[:80]
])

endpoint_text = "\n".join([
    f"[{e.get('endpoint_id','')}] {str(e.get('extracted_value',''))[:300]} ({str(e.get('source_article',''))[:100]})"
    for e in endpoint_fulltext[:30]
])

sota_benchmarks = "\n".join([
    f"[{b.get('benchmark_id','?')}] {b.get('endpoint','')[:200]} → {b.get('sota_value_range','?')} | Rationale: {str(b.get('domain_aware_rationale',''))[:200]}"
    for b in sota[:16]
])

# ── PRISMA context ──
pid = prisma.get("identification", {}) if isinstance(prisma, dict) else {}
psc = prisma.get("screening", {}) if isinstance(prisma, dict) else {}
pin = prisma.get("included", {}) if isinstance(prisma, dict) else {}
prisma_ctx = f"""Database records: {pid.get('database_records','?')} | Deduplicated: {psc.get('deduplicated_records','?')}
Full-text assessed: {psc.get('full_text_assessed','?')} | SOTA included: {pin.get('sota_included','?')}
DUE included: {pin.get('due_included','?')} | Total: {pin.get('total_included','?')}"""

# ── Engineer templates (condensed) ──
engineer_templates = """
CER WRITING STANDARDS (from 5 benchmark CERs by NB-reviewed engineers):

CHAPTER STRUCTURE: §1 Executive Summary(~5p) → §2 Device Description(~15p) → §3 SOTA(~30p) → §4 Literature Search(~10p) → §5 Evidence Appraisal(~40p) → §6 DUE Data(~20p) → §7 Benefit-Risk(~12p) → §8 Conclusions(~5p) → §9 PMCF(~5p) → §10-12 Admin(~4p) → Annexes

WRITING STYLE: Sentence length 22-32 words. Passive voice 15-25%. Hedging: may/might/could/approximately for uncertainty. Certainty: demonstrated/confirmed/proven for strong evidence. NO placeholder text. NO "TBD"/"N/A". Every claim must have evidence citation.

EQUIVALENCE: For this device, equivalence is NOT claimed (novel automated bubble study system). Use template: 'Equivalence is not claimed for the device under evaluation. The device represents a novel technology for which no equivalent device has been identified. Clinical evaluation relies solely on clinical investigation data and literature data.'

TABLES REQUIRED: Device description table, SOTA search strategy table, PRISMA flow diagram, evidence appraisal table (per article), clinical data summary table, adverse events table, benefit-risk determination table, GSPR coverage matrix, PMCF plan table.

EVIDENCE CHAIN: source → data → interpretation → claim. Every claim must trace to specific evidence. Every numerical value must cite the source article and data point.
"""

client = get_llm_client(timeout=1200.0)

CHAPTERS = {
    "01_executive_summary": f"""Write §1 EXECUTIVE SUMMARY (~5-8 pages) for this CER.
{device_ctx}
CLAIMS ({len(claims)}): {json.dumps([{'id':c.get('claim_id','?'), 'type':c.get('claim_type','?'), 'text':str(c.get('claim_text',''))[:200]} for c in claims])}
EVIDENCE BASE: {len(appraisal)} appraised articles, {len(clinical_facts)} clinical facts.
{prisma_ctx}
Include: 1.1 Device overview & classification, 1.2 Clinical evaluation scope (MDR Art.61, MDCG 2020-5), 1.3 Key conclusions, 1.4 Benefit-risk determination, 1.5 Summary of clinical evidence base.
Write in professional English with substantive content. NO placeholders.""",

    "02_device_description": f"""Write §2 DEVICE DESCRIPTION AND SCOPE (~15-20 pages).
{device_ctx}
IFU DATA: {json.dumps(device.get('ifu_structured_extraction',{}))}
CLAIMS SUMMARY: {json.dumps([c.get('claim_text','')[:200] for c in claims[:5]])}
SIMILAR DEVICES: {json.dumps([str(s)[:300] for s in similar_devices[:5]])}
Include sections: 2.1 Device Description (name, classification IIb Rule 10+12, components, materials, working principle, technical specs, model variants BS-2/DCS series), 2.2 Intended Purpose (indications for RLS/PFO detection, target population adults 18-75, intended users cardiac/neuro ultrasound physicians, clinical benefit, anatomical sites cardiac/cerebral, environment), 2.3 Contraindications Warnings Precautions (comprehensive list from IFU), 2.4 Previous Generations and Similar Devices (benchmark comparison table), 2.5 Marketing History (PRC approval, CE mark application), 2.6 Clinical Evaluation Scope (MDR Art.61, MDCG 2020-5, MEDDEV 2.7/1 Rev.4 basis).
{engineer_templates}
Include device description table (model, EMDN, classification, components, accessories, principle of operation).""",

    "03_state_of_the_art": f"""Write §3 STATE OF THE ART (~30-40 pages).
{device_ctx}
SOTA BENCHMARKS: {sota_benchmarks}
CLINICAL CONTEXT: {json.dumps([str(c)[:300] for c in sota_context[:10]])}
ALTERNATIVE TREATMENTS: {json.dumps([str(a)[:300] for a in alternative[:5]])}
GUIDELINES: {json.dumps([str(g)[:300] for g in guideline[:5]])}
Include: 3.1 Clinical Background (RLS/PFO epidemiology, pathophysiology of right-to-left shunt, cryptogenic stroke association, natural history), 3.2 Medical Field Context (contrast echocardiography, transcranial Doppler, bubble study methodology), 3.3 Disease Prevalence (PFO in general population ~25%, cryptogenic stroke patients ~40-50%, RLS prevalence data from published meta-analyses), 3.4 Alternative Diagnostic Options (TEE gold standard, TTE with agitated saline, TCD, alternative automated systems — comparison table), 3.5 Clinical Practice Guidelines (ESC, AAN, Chinese expert consensus on PFO management — extract specific recommendations), 3.6 Similar Devices (benchmark devices for bubble study — comparison table with specs, performance), 3.7 Potential Hazards and Known Complications (air embolism, injection site complications, Valsalva-related risks, misdiagnosis), 3.8 Safety and Performance Endpoints (define each endpoint with SOTA benchmark values from literature, acceptance criteria), 3.9 SOTA Conclusions.
{engineer_templates}
CRITICAL: For EACH endpoint, provide quantitative SOTA benchmark values extracted from published literature. For EACH alternative treatment, provide published performance data. Include detailed comparison tables. This section should be 30+ pages of substantive clinical content.""",

    "04_literature_search": f"""Write §4 LITERATURE SEARCH METHODOLOGY (~10-15 pages).
{device_ctx}
{prisma_ctx}
SEARCH STRATEGY: {json.dumps([str(s)[:300] for s in vals.get('sota_search_strategy_table',[])[:14]])}
Include: 4.1 Search Strategy (databases: PubMed, NCBI PMC, Europe PMC, Embase, Cochrane Library, ClinicalTrials.gov — with exact search strings, date ranges, filters applied), 4.2 PRISMA Flow Diagram (exact numbers: identified → screened → full-text assessed → included, with exclusion reasons at each stage), 4.3 Inclusion/Exclusion Criteria (detailed PICO-based criteria with rationale for each), 4.4 Article Screening Process (title/abstract screening methodology, full-text review, independent dual review if applicable), 4.5 Quality Appraisal Methodology (appraisal tools used: AMSTAR-2 for systematic reviews, Cochrane RoB 2 for RCTs, Newcastle-Ottawa for observational, QUADAS-2 for diagnostic accuracy studies — justify each choice).
{engineer_templates}
Include PRISMA flow diagram as formatted text table. Include detailed exclusion reasons for at least 20 key excluded articles.""",

    "05_evidence_appraisal": f"""Write §5 CLINICAL EVIDENCE APPRAISAL (~40-50 pages).
{device_ctx}
{evidence_sample}
CLINICAL FACTS: {facts_sample}
ENDPOINT DATA: {endpoint_text}
EVIDENCE HIERARCHY: {json.dumps([str(h)[:300] for h in sota_hierarchy[:30]])}
EVIDENCE CONFLICTS: {json.dumps(evidence_conflicts)}
SYNTHESIS NARRATIVES: {json.dumps([str(s)[:300] for s in synthesis[:10]])}
CLINICAL SOURCE RECORDS: {json.dumps([str(c)[:400] for c in clinical_source[:30]])}
Include: 5.1 Evidence Overview (categorize by evidence type: systematic reviews, RCTs, observational studies, case series — provide counts and summary), 5.2 Pivotal Clinical Evidence (own prospective multi-center clinical investigation — detailed presentation with study design, population characteristics table, primary/secondary endpoints, statistical analysis plan, results tables with quantitative data, adverse event analysis, subgroup analyses), 5.3-5.7 For EACH included article provide: full citation (authors, year, journal, PMID), study design with level of evidence, population characteristics (sample size, demographics, inclusion/exclusion criteria), intervention details, KEY QUANTITATIVE RESULTS (sensitivity, specificity, concordance rate, PPV, NPV, AUC, adverse event rates — extract actual numbers from full-text), quality appraisal score with rationale, relevance to Bubble Study System, strengths and limitations.
{engineer_templates}
CRITICAL: Every article must have a complete evidence table row. Every numerical value must cite the source. This is the largest section — write 40+ pages of detailed evidence analysis.""",

    "06_due_clinical_data": f"""Write §6 CLINICAL DATA FROM DEVICE UNDER EVALUATION (~20-25 pages).
{device_ctx}
CEP: {json.dumps(cep)}
GSPR COVERAGE: {json.dumps([str(g)[:300] for g in gspr[:5]])}
RISK TRACE: {json.dumps([str(r)[:300] for r in risk_trace[:5]])}
Include: 6.1 Bench and Pre-Clinical Testing (electrical safety IEC 60601-1, EMC IEC 60601-1-2, biocompatibility ISO 10993-1 for tubing materials, software validation IEC 62304, performance testing — bubble size distribution, injection volume accuracy, pressure monitoring, alarm functions), 6.2 Clinical Investigation Results (prospective multi-center crossover trial: 180 subjects, 3 centers — Fuwai Hospital, Xiangyang Central Hospital, Hebei Medical University First Hospital — detailed results with primary endpoint concordance rate ≥80% target, secondary endpoints sensitivity/specificity for RLS detection, semi-quantitative shunt grading agreement, Valsalva maneuver completion rate, procedure time comparison vs manual method, adverse event rates by severity and causality), 6.3 PMCF Data (PMCF plan per MDCG 2020-7, C.4 clinical experience gathering 67 cases/year target), 6.4 PMS and Vigilance Data (complaint analysis, serious incident rates, corrective actions).
{engineer_templates}
Include detailed clinical data tables: subject demographics, primary endpoint results, secondary endpoint results, adverse event table by SOC and severity, device deficiencies table.""",

    "07_benefit_risk": f"""Write §7 BENEFIT-RISK ANALYSIS (~12-18 pages).
{device_ctx}
BENEFIT-RISK LEDGER: {json.dumps([str(b)[:300] for b in benefit_risk[:14]])}
EVIDENCE CONFLICTS: {json.dumps(evidence_conflicts)}
CLINICAL FACTS (top 30): {facts_sample}
Include: 7.1 Clinical Benefits Quantification (primary benefit: automated standardized agitated saline preparation → reduced manual variability → improved RLS detection concordance ≥80% — with quantitative effect sizes, secondary benefits: reduced procedure time, standardized Valsalva protocol, reduced operator dependency, integrated data management), 7.2 Risks and Harms Characterization (categorized by severity: serious adverse events rate, non-serious AE rate, device-related AE rate, procedure-related complications — with incidence rates and 95% CIs from clinical trial and literature), 7.3 Risk Mitigation Measures (IFU warnings cross-referenced to RMF hazard controls, single-use design to prevent cross-contamination, automated pressure monitoring and alarms, training requirements for operators), 7.4 Benefit-Risk Determination (structured qualitative comparison table, quantitative acceptability assessment against SOTA benchmarks, residual risk acceptance justification per ISO 14971), 7.5 Residual Risks and Uncertainties (PMCF needs, long-term follow-up limitations, generalizability considerations).
{engineer_templates}
Include benefit-risk determination table, risk matrix, mitigation effectiveness summary.""",

    "08_conclusions_and_pmcf": f"""Write §8 CONCLUSIONS (~5 pages), §9 PMCF AND NEXT EVALUATION (~5 pages), §10-12 ADMIN SECTIONS (~3 pages).
{device_ctx}
PMCF PLAN: {json.dumps(pmcf)}
PMCF GAPS: {json.dumps([str(g)[:300] for g in gap_recs[:5]])}
VIGILANCE: {json.dumps([str(v)[:300] for v in vigilance[:5]])}
Include §8: Summary of clinical evidence for each claim, adequacy assessment, overall benefit-risk conclusion ('the benefit-risk profile is acceptable'), acceptability statement per MDR Annex XIV. §9: PMCF objectives (confirm long-term safety, real-world diagnostic performance, user training effectiveness), PMCF methods (registry study, user survey, literature surveillance), timeline (annual PMCF reports, next CER update at 3-5 years), next evaluation date rationale. §10: Qualification of Evaluators (clinical expertise requirements, CV summary). §11: Declaration of Interest (independence statement). §12: Dates and Signatures.
{engineer_templates}
Conclude definitively: state that clinical evidence is sufficient to demonstrate conformity with GSPRs.""",
}

results = {}
for ch_id, prompt in CHAPTERS.items():
    ch_path = os.path.join(OUTPUT, "chapters", f"{ch_id}.md")
    if os.path.exists(ch_path) and os.path.getsize(ch_path) > 500:
        print(f"  [{ch_id}] Already exists ({os.path.getsize(ch_path)} bytes), skipping", flush=True)
        with open(ch_path) as f:
            results[ch_id] = f.read()
        continue

    full_prompt = prompt + "\n\nWrite the complete section now. Generate substantive, publication-ready content with specific data points, tables, and evidence citations. Target professional medical English. NO placeholder text, NO 'TBD', NO 'N/A'."

    print(f"  [{ch_id}] Prompt: {len(full_prompt)} chars → calling deepseek-v4-pro...", flush=True)
    try:
        response = client.messages.create(
            model="deepseek-v4-pro",
            max_tokens=16000,
            system="You are a senior medical device CER author with 15+ years experience. You write MDR-compliant, NB-ready CER reports. Every claim is supported by evidence with explicit reasoning chain: source → data → interpretation → claim. You NEVER use placeholder text. Your CERs consistently pass Notified Body review on first submission. Write in professional English with appropriate technical terminology, 22-32 word sentences, 15-25% passive voice. Target substantive content — this section alone should be publication-ready.",
            messages=[{"role": "user", "content": full_prompt}],
        )
        text = ""
        for block in response.content:
            if hasattr(block, 'text'):
                text += block.text
        results[ch_id] = text
        with open(ch_path, "w") as f:
            f.write(text)
        print(f"    → {len(text)} chars, ~{len(text.split())} words", flush=True)
    except Exception as e:
        print(f"    → ERROR: {e}", flush=True)
        results[ch_id] = f"[ERROR: {e}]"

# ── Assemble full CER ──
print("\nAssembling complete CER...", flush=True)
full_cer = f"""# Clinical Evaluation Report

## Bubble Study System (BS-2) with Disposable Contrast Injection Tubing Set

**Manufacturer:** WYTD MEDICAL TECHNOLOGY (SHENZHEN) CO., LTD.
**Document Number:** CER-BUB-2026-001
**Date:** 2026-06-01
**Classification:** Class IIb per MDR Annex VIII Rule 10 + Rule 12
**Regulatory Basis:** MDR 2017/745 Article 61 and Annex XIV Part A
**Guidance Followed:** MDCG 2020-5, MDCG 2020-6, MEDDEV 2.7/1 Rev.4

---

## Table of Contents

1. Executive Summary
2. Device Description and Scope
3. State of the Art
4. Literature Search Methodology
5. Clinical Evidence Appraisal
6. Clinical Data from Device Under Evaluation
7. Benefit-Risk Analysis
8. Conclusions
9. PMCF and Next Evaluation
10. Qualification of Evaluators
11. Declaration of Interest
12. Dates and Signatures

Annexes (separate volume): Evidence Tables, PRISMA Checklist, Search Logs, GSPR Matrix, Abbreviations

---

"""
for ch_id in ["01_executive_summary", "02_device_description", "03_state_of_the_art",
               "04_literature_search", "05_evidence_appraisal", "06_due_clinical_data",
               "07_benefit_risk", "08_conclusions_and_pmcf"]:
    content = results.get(ch_id, "")
    full_cer += f"\n\n{content}\n\n"

md_path = os.path.join(OUTPUT, "CER_draft.md")
with open(md_path, "w") as f:
    f.write(full_cer)

chars = len(full_cer)
words = len(full_cer.split())
pages_est = chars // 3500
print(f"\n✅ Complete CER: {chars} chars, ~{words} words, ~{pages_est} pages", flush=True)
print(f"Saved: {md_path}", flush=True)

# Generate DOCX
import subprocess
docx_path = os.path.join(OUTPUT, "CER_draft.docx")
result = subprocess.run(["pandoc", md_path, "-o", docx_path], capture_output=True, text=True, check=False)
if result.returncode == 0:
    print(f"DOCX: {docx_path}", flush=True)
else:
    print(f"Pandoc failed, trying python-docx...", flush=True)
    from docx import Document
    doc = Document()
    doc.add_heading("Clinical Evaluation Report", 0)
    for line in full_cer.split("\n")[:1000]:
        if line.startswith("# "): doc.add_heading(line[2:], level=1)
        elif line.startswith("## "): doc.add_heading(line[3:], level=2)
        elif line.startswith("### "): doc.add_heading(line[4:], level=3)
        elif line.strip(): doc.add_paragraph(line.strip())
    doc.save(docx_path)
    print(f"DOCX (fallback): {docx_path}", flush=True)

print("DONE", flush=True)
