"""Generate CER draft from checkpoint data — V2 enhanced version with full evidence context."""
import sys, os, json
from pathlib import Path

sys.path.insert(0, os.path.dirname(__file__))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
import _patch_httpx_lru  # noqa: F401

from _v3_1_llm_client import get_llm_client
from langgraph.checkpoint.sqlite import SqliteSaver
from deerflow.runtime.cer_authoring.graph import build_cer_authoring_graph

db_path = sys.argv[1] if len(sys.argv) > 1 else None
if not db_path:
    print("Usage: generate_cer_draft_v2.py <checkpoint.db>")
    sys.exit(1)

output_dir = os.path.dirname(db_path)
checkpoint_path = db_path

print(f"Loading checkpoint: {checkpoint_path}", flush=True)
print(f"Output dir: {output_dir}", flush=True)

with SqliteSaver.from_conn_string(checkpoint_path) as checkpointer:
    graph = build_cer_authoring_graph(checkpointer=checkpointer)
    state = graph.get_state({"configurable": {"thread_id": "A01_无忧跳动"}})
    vals = state.values

# ── Extract all key data ──
device = vals.get("device_profile", {})
claims = vals.get("claim_ledger", [])
sota = vals.get("sota_benchmark_matrix", [])
sota_hierarchy = vals.get("sota_evidence_hierarchy", [])
sota_context = vals.get("sota_clinical_context_table", [])
endpoint_fulltext = vals.get("sota_endpoint_extraction_fulltext", [])
prisma = vals.get("prisma_flow_data", {})
appraisal = vals.get("article_appraisal", [])
evidence_registry = vals.get("evidence_registry", [])
clinical_facts = vals.get("clinical_evidence_fact_table", [])
endpoint_extraction = vals.get("endpoint_extraction", [])
cep_pico = vals.get("cep_pico_matrix", [])
gap_recommendations = vals.get("gap_pmcf_recommendations", [])
writing_brief = vals.get("writing_brief", {})
device_identity = vals.get("device_identity_lock", {})
risk_trace = vals.get("risk_trace_matrix", [])
rmf_hazard = vals.get("rmf_hazard_trace", {})
sota_deduction = vals.get("sota_deduction_chain", [])
evidence_conflicts = vals.get("evidence_conflict_report", {})
equivalence = vals.get("equivalence_matrix", [])
guideline_pathway = vals.get("guideline_pathway_table", [])
similar_devices = vals.get("similar_device_attachment_index", [])
intended_purpose_claims = vals.get("intended_purpose_claim_table", [])
fact_anchored = vals.get("fact_anchored_claims", [])
benefit_risk = vals.get("benefit_risk_ledger", [])
vigilance = vals.get("vigilance_event_statistics", [])
sota_conclusion = vals.get("sota_conclusion_strength_guard", [])
clinical_source = vals.get("clinical_source_adapter_records", [])
synthesis = vals.get("cross_evidence_synthesis_narratives", [])
cer_sections = vals.get("cer_chapter_drafts", {})
pmcf_gap = vals.get("pmcf_plan_control_matrix", {})

print(f"Device: {device.get('device_name', '?')} | Claims: {len(claims)} | Evidence: {len(appraisal)} | Facts: {len(clinical_facts)}", flush=True)

# ── Build comprehensive prompt ──
device_text = f"""DEVICE INFORMATION
Name: {device.get('device_name', 'N/A')}
Manufacturer: {device.get('manufacturer', 'N/A')}
Class: {device.get('device_class', 'N/A')} (MDR Annex VIII)
Device Type: {device.get('device_type', 'N/A')}
Device Family: {device.get('device_family', 'N/A')}
Clinical Domain: {device.get('clinical_domain', 'N/A')}
Intended Purpose: {device.get('intended_purpose', 'N/A')}
Model/Specifications: {device.get('model_specifications', 'N/A')}
Target Population: {device.get('target_population', 'N/A')}
Intended User: {device.get('intended_user', 'N/A')}
Anatomical Site: {device.get('anatomical_site', 'N/A')}
Mode of Action: {device.get('mode_of_action', 'N/A')}
Contraindications: {device.get('contraindications', 'N/A')}
Sterility: {device.get('sterility', 'N/A')}
Shelf Life / Storage: {device.get('shelf_life_storage', 'N/A')}
Performance Summary: {device.get('performance_summary', 'N/A')}
Working Principle: {device.get('working_principle', 'N/A')}
Composition: {device.get('composition', 'N/A')}"""

# Claims text
claims_text_parts = []
for c in claims[:20]:
    cid = c.get('claim_id', c.get('id', '?'))
    ctype = c.get('claim_type', c.get('type', '?'))
    ctext = c.get('claim_text', c.get('text', str(c)[:300]))
    claims_text_parts.append(f"[{cid}] ({ctype}): {ctext}")
claims_text = "\n".join(claims_text_parts)

# SOTA benchmarks
sota_text_parts = []
for b in sota[:20]:
    bid = b.get('benchmark_id', b.get('id', '?'))
    endpoint = b.get('endpoint', str(b)[:200])
    claim = b.get('corresponding_claim_id', '?')
    value = b.get('sota_value_range', b.get('acceptance_criterion', ''))
    rationale = b.get('domain_aware_rationale', '')
    sota_text_parts.append(f"[{bid}] Endpoint: {endpoint}\n  → Claim: {claim} | SOTA Value: {value}\n  Rationale: {rationale}")
sota_text = "\n\n".join(sota_text_parts)

# Endpoint extraction from full-text
endpoint_fulltext_parts = []
for e in endpoint_fulltext[:20]:
    eid = e.get('endpoint_id', '')
    value = e.get('extracted_value', e.get('value', str(e)[:200]))
    source = e.get('source_article', '')
    endpoint_fulltext_parts.append(f"[{eid}] {value} (Source: {source})")
endpoint_fulltext_text = "\n".join(endpoint_fulltext_parts)

# PRISMA flow
pid = prisma.get("identification", {}) if isinstance(prisma, dict) else {}
psc = prisma.get("screening", {}) if isinstance(prisma, dict) else {}
pin = prisma.get("included", {}) if isinstance(prisma, dict) else {}
prisma_text = f"""PRISMA FLOW:
- Database records identified: {pid.get('database_records', '?')}
- After deduplication: {psc.get('deduplicated_records', '?')}
- Full-text assessed: {psc.get('full_text_assessed', '?')}
- SOTA included: {pin.get('sota_included', '?')}
- DUE included: {pin.get('due_included', '?')}
- Total included in CER: {pin.get('total_included', '?')}"""

# Evidence appraisal — top 50
evidence_parts = []
for e in appraisal[:50]:
    eid = e.get('evidence_id', e.get('id', '?'))
    score = e.get('appraisal_score', e.get('score', '?'))
    weight = e.get('weight', '?')
    title = e.get('title', e.get('article_title', str(e)[:150]))
    evidence_parts.append(f"[{eid}] Score={score} Weight={weight}: {title}")
evidence_text = "\n".join(evidence_parts)

# Clinical facts
facts_parts = []
for f in clinical_facts[:50]:
    fid = f.get('fact_id', f.get('id', ''))
    fact_text = f.get('fact_text', f.get('text', str(f)[:200]))
    source = f.get('source_article', f.get('source', ''))
    facts_parts.append(f"[{fid}] {fact_text} ← {source}")
facts_text = "\n".join(facts_parts)

# Clinical source records — top 30
cs_parts = []
for r in clinical_source[:30]:
    cs_parts.append(str(r)[:300])
cs_text = "\n---\n".join(cs_parts)

# Benefit-risk
br_text = "\n".join([str(b)[:200] for b in benefit_risk[:10]])

# Evidence conflicts
conflict_text = ""
if isinstance(evidence_conflicts, dict):
    conflicts = evidence_conflicts.get("conflicts", [])
    conflict_text = f"Evidence Conflicts: {len(conflicts)} found.\n" + "\n".join([str(c)[:200] for c in conflicts[:10]])

# Writing brief
wb_text = ""
if isinstance(writing_brief, dict):
    wb_text = writing_brief.get('writing_brief', str(writing_brief)[:500])

# SOTA evidence hierarchy — top 30
hierarchy_text = "\n".join([str(h)[:200] for h in sota_hierarchy[:30]])

# Deduction chain
deduction_text = "\n".join([str(d)[:300] for d in sota_deduction[:10]])

# Build the full prompt
prompt = f"""You are a senior medical device CER author. Write a COMPLETE, SUBSTANTIVE Clinical Evaluation Report (CER) in professional English following MDR 2017/745 and MDCG 2020-5 guidance.

The CER must be thorough, detailed, and publication-ready — targeting 150+ pages of content. Every claim must be supported by evidence. Every endpoint must have quantitative or qualitative data. Every reasoning step must be explicit.

{device_text}

─── CLAIMS ({len(claims)}) ───
{claims_text}

─── SOTA BENCHMARKS ({len(sota)}) ───
{sota_text}

─── SOTA ENDPOINT FULL-TEXT EXTRACTIONS ({len(endpoint_fulltext)}) ───
{endpoint_fulltext_text}

─── SOTA EVIDENCE HIERARCHY ({len(sota_hierarchy)} entries, showing top 30) ───
{hierarchy_text}

─── SOTA DEDUCTION CHAIN ───
{deduction_text}

─── {prisma_text} ───

─── EVIDENCE APPRAISAL ({len(appraisal)} records, showing top 50) ───
{evidence_text}

─── CLINICAL FACTS ({len(clinical_facts)} extracted from full-text) ───
{facts_text}

─── BENEFIT-RISK LEDGER ({len(benefit_risk)} items) ───
{br_text}

─── {conflict_text} ───

─── WRITING BRIEF ───
{wb_text}

───

Write the COMPLETE CER now. Follow this structure EXACTLY:

§1 EXECUTIVE SUMMARY (~5-8 pages)
- Device overview and classification
- Clinical evaluation scope
- Key conclusions and benefit-risk determination
- Summary of clinical evidence base

§2 DEVICE DESCRIPTION AND SCOPE (~10-15 pages)
- §2.1 Device Description (name, classification, components, materials, working principle, technical specifications)
- §2.2 Intended Purpose (clinical indications, target population, intended users, clinical benefit, anatomical site, environment)
- §2.3 Contraindications, Warnings, and Precautions
- §2.4 Previous Generations and Similar Devices
- §2.5 Marketing History and Regulatory Status
- §2.6 Clinical Evaluation Scope and Regulatory Basis

§3 STATE OF THE ART (~25-35 pages)
- §3.1 Clinical Background and Disease Overview (epidemiology, pathophysiology, natural history)
- §3.2 Related Medical Field and Clinical Context
- §3.3 Medical Conditions and Disease Prevalence
- §3.4 Alternative Treatment Options (medical therapy, surgical, other interventions — include comparison table)
- §3.5 Clinical Practice Guidelines (major society guidelines with recommendations)
- §3.6 Similar Devices and Benchmark Comparison
- §3.7 Potential Hazards and Known Complications
- §3.8 Safety and Performance Endpoints (define endpoints, SOTA benchmark values, acceptance criteria)
- §3.9 SOTA Conclusions (what is the current standard of care, what gaps exist)

§4 LITERATURE SEARCH METHODOLOGY (~8-12 pages)
- §4.1 Search Strategy (databases, search terms, date ranges, filters)
- §4.2 PRISMA Flow Diagram and Narrative
- §4.3 Inclusion/Exclusion Criteria with detailed rationale for each criterion
- §4.4 Article Screening Process (title/abstract → full-text → final inclusion)
- §4.5 Quality Appraisal Methodology (appraisal tools used, scoring system)

§5 CLINICAL EVIDENCE APPRAISAL (~30-50 pages)
- For EACH included article, provide:
  * Full citation (authors, year, journal, DOI/PMID)
  * Study design and level of evidence
  * Population characteristics (sample size, demographics, inclusion/exclusion)
  * Intervention/exposure details
  * Key results with quantitative data (sensitivity, specificity, concordance rate, adverse event rates, etc.)
  * Quality appraisal score and rationale
  * Relevance to the device under evaluation
  * Strengths and limitations
- Organize articles by evidence theme (diagnostic performance, safety, clinical benefit, etc.)
- Include detailed evidence tables with extracted data points

§6 CLINICAL DATA FROM DEVICE UNDER EVALUATION (~15-25 pages)
- §6.1 Pre-Clinical Data (bench testing, biocompatibility, electrical safety)
- §6.2 Clinical Investigation (trial design, methods, results, statistical analysis)
  * Detailed presentation of the prospective clinical trial data
  * Concordance rates, sensitivity, specificity
  * Adverse event analysis
  * Subgroup analyses if applicable
- §6.3 Post-Market Clinical Follow-up (PMCF) Data
- §6.4 PMS and Vigilance Data

§7 BENEFIT-RISK ANALYSIS (~10-15 pages)
- §7.1 Clinical Benefits (quantified, with evidence strength assessment)
- §7.2 Risks and Harms (categorized by severity, frequency, and certainty)
- §7.3 Risk Mitigation Measures (IFU warnings, design controls, training requirements)
- §7.4 Benefit-Risk Determination (structured qualitative/quantitative comparison)
- §7.5 Residual Risks and Uncertainties

§8 CONCLUSIONS (~3-5 pages)
- Summary of clinical evidence
- Adequacy of clinical evidence for each claim
- Overall benefit-risk conclusion
- Acceptability of the benefit-risk profile

§9 PMCF AND NEXT EVALUATION (~3-5 pages)
- PMCF objectives and methods
- Next CER evaluation date and rationale

§10 QUALIFICATION OF EVALUATORS (~1-2 pages)
§11 DECLARATION OF INTEREST (~1 page)
§12 DATES AND SIGNATURES (~1 page)

ANNEXES
- Annex A: Full Evidence Tables
- Annex B: PRISMA Checklist
- Annex C: Literature Search Logs
- Annex D: GSPR Coverage Matrix
- Annex E: Abbreviations and Glossary

CRITICAL RULES:
1. EVERY claim must be explicitly linked to supporting evidence with reasoning chain: source → data → interpretation → claim
2. EVERY numerical endpoint must include the actual value, not placeholders
3. For EACH included article, provide: full citation, study design, sample size, key quantitative results, appraisal score, and relevance analysis
4. The PRISMA flow must show exact numbers at each stage with exclusion reasons
5. Benefit-risk analysis must be quantitative where possible, qualitative where necessary
6. Do NOT write "to be determined", "TBD", "N/A", or any placeholder
7. Use professional medical English with appropriate technical terminology
8. Write substantive paragraphs, not outlines or bullet points (except for tables)
9. Tables must have descriptive captions and column headers
10. Target output: 150+ pages of substantive content

BEGIN WRITING THE FULL CER NOW:"""

print(f"Prompt size: {len(prompt)} chars (~{len(prompt)//4} tokens)", flush=True)
print("=" * 60)
print("Calling LLM via local router (kimi-k2.6, max_tokens=64000)...", flush=True)

client = get_llm_client(timeout=1200.0)

try:
    response = client.messages.create(
        model="deepseek-v4-pro",
        max_tokens=32000,
        system="You are a senior medical device CER (Clinical Evaluation Report) author with 15+ years of experience. You write MDR-compliant, NB-ready CER reports. Your writing is thorough, evidence-based, and publication-quality. You NEVER use placeholder text — every statement is supported by data. Your CERs consistently pass Notified Body review on first submission.",
        messages=[{"role": "user", "content": prompt}],
    )

    # Handle multiple content blocks (thinking + text from DeepSeek)
    cer_content = ""
    for block in response.content:
        if hasattr(block, 'text'):
            cer_content += block.text
    if not cer_content:
        raise ValueError(f"No text content in response. Blocks: {[type(b).__name__ for b in response.content]}")
    wc = len(cer_content.split())
    print(f"Generated: {len(cer_content)} chars, ~{wc} words", flush=True)

    # Save
    md_path = os.path.join(output_dir, "CER_draft.md")
    with open(md_path, "w") as f:
        f.write(cer_content)
    print(f"Saved: {md_path}", flush=True)

    # Generate DOCX with pandoc
    import subprocess
    docx_path = os.path.join(output_dir, "CER_draft.docx")
    result = subprocess.run(
        ["pandoc", md_path, "-o", docx_path, "--from=markdown", "--to=docx",
         "--reference-doc=/Users/winstonwei/Documents/Playground/deer-flow/backend/scripts/_v4_table_templates.py" if False else ""],
        capture_output=True, text=True, check=False
    )
    # Simple pandoc call without reference-doc since that's a Python file
    result = subprocess.run(
        ["pandoc", md_path, "-o", docx_path],
        capture_output=True, text=True, check=False
    )
    if result.returncode == 0:
        print(f"Saved DOCX: {docx_path}", flush=True)
    else:
        print(f"Pandoc failed: {result.stderr[:200]}", flush=True)
        # Try python-docx fallback
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Clinical Evaluation Report", 0)
            for line in cer_content.split("\n")[:500]:
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line.strip())
            doc.save(docx_path)
            print(f"Saved DOCX (python-docx fallback): {docx_path}", flush=True)
        except Exception as e2:
            print(f"DOCX fallback also failed: {e2}", flush=True)

    print("DONE", flush=True)

except Exception as e:
    print(f"ERROR: {type(e).__name__}: {e}", flush=True)
    import traceback
    traceback.print_exc()
    sys.exit(1)
