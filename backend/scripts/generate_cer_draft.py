"""Generate CER draft directly from checkpoint data via LLM."""
import sys, os, json

sys.path.insert(0, os.path.dirname(__file__))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
import _patch_httpx_lru  # noqa: F401

from _v3_1_llm_client import get_llm_client
from langgraph.checkpoint.sqlite import SqliteSaver
from deerflow.runtime.cer_authoring.graph import build_cer_authoring_graph

db_path = sys.argv[1] if len(sys.argv) > 1 else "/Users/winstonwei/CER-RAG/升级 CCD-3 个项目文件/2026.5.28- 试运行项目/A04_永新/02_CER_OUTPUT/.checkpoints.db"
output_dir = os.path.dirname(db_path)

print(f"Loading checkpoint: {db_path}", flush=True)

with SqliteSaver.from_conn_string(db_path) as checkpointer:
    graph = build_cer_authoring_graph(checkpointer=checkpointer)
    vals = graph.get_state({"configurable": {"thread_id": "YONGXIN_MSOFT_001"}}).values

device = vals["device_profile"]
claims = vals["claim_ledger"]
sota = vals["sota_benchmark_matrix"]
prisma = vals["prisma_flow_data"]
appraisal = vals.get("article_appraisal", [])

# Build prompt
device_text = f"""DEVICE: {device.get('device_name')} by {device.get('manufacturer')}
Class: {device.get('device_class')} (MDR Rule 11) | Model: {device.get('model_specifications')}
Clinical Domain: {device.get('clinical_domain')}
Intended Purpose: {device.get('intended_purpose')}
Target Population: {device.get('target_population')}
Intended User: {device.get('intended_user')}
Contraindications: {device.get('contraindications', 'None')}"""

claims_text = "\n".join(
    f"- [{c.get('claim_id','?')}] ({c.get('claim_type','?')}): {c.get('claim_text','')[:300]}"
    for c in claims
)

sota_text = "\n".join(
    f"- [{b.get('benchmark_id','?')}] {b.get('endpoint','')[:200]} → Claim {b.get('corresponding_claim_id','?')}"
    for b in sota
)

pid = prisma.get("identification", {})
psc = prisma.get("screening", {})
pin = prisma.get("included", {})
prisma_text = f"PRISMA: {pid.get('database_records','?')} identified → {psc.get('deduplicated_records','?')} dedup → {psc.get('full_text_assessed','?')} assessed → {pin.get('sota_included','?')} included"

evidence_sample = "\n".join(
    f"- [{e.get('evidence_id','?')}] {e.get('title','')[:120]}"
    for e in appraisal[:15]
)

prompt = f"""Write a Clinical Evaluation Report (CER) for:

{device_text}

CLAIMS ({len(claims)}):
{claims_text}

SOTA BENCHMARKS ({len(sota)}):
{sota_text}

{prisma_text}

EVIDENCE ({len(appraisal)} records, title/abstract level):
{evidence_sample}

Write a complete CER with sections: §0 Executive Summary, §1 Device Description, §2 SOTA, §3 Literature Search (PRISMA), §4 Evidence Appraisal, §5 Benefit-Risk, §6 Conclusions, §7 PMCF.

Write in professional English. Generate substantive content for each section. Do NOT use placeholder text."""

print(f"Prompt size: {len(prompt)} chars", flush=True)
print("Calling LLM...", flush=True)

client = get_llm_client(timeout=600.0)

try:
    response = client.messages.create(
        model="kimi-k2.6",
        max_tokens=12000,
        system="You are a medical device CER authoring specialist. Write professional, MDR-compliant CER reports.",
        messages=[{"role": "user", "content": prompt}],
    )

    cer_content = response.content[0].text
    wc = len(cer_content.split())
    print(f"Generated: {len(cer_content)} chars, ~{wc} words", flush=True)

    # Save
    md_path = os.path.join(output_dir, "CER_draft.md")
    with open(md_path, "w") as f:
        f.write(cer_content)
    print(f"Saved: {md_path}", flush=True)

    # Generate DOCX
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Clinical Evaluation Report", 0)
        doc.add_paragraph(f"Device: {device.get('device_name')}")
        doc.add_paragraph(f"Manufacturer: {device.get('manufacturer')}")
        doc.add_paragraph(f"Class: {device.get('device_class')} (MDR)")
        doc.add_paragraph("")

        for section in cer_content.split("\n#"):
            if not section.strip():
                continue
            lines = section.strip().split("\n")
            heading = lines[0].strip("# ").strip()
            if heading:
                doc.add_heading(heading, level=1)
            for line in lines[1:]:
                if line.strip():
                    doc.add_paragraph(line.strip())

        docx_path = os.path.join(output_dir, "CER_draft.docx")
        doc.save(docx_path)
        print(f"Saved DOCX: {docx_path}", flush=True)
    except Exception as e:
        print(f"DOCX error: {e}", flush=True)
        import subprocess
        subprocess.run(["pandoc", md_path, "-o", os.path.join(output_dir, "CER_draft.docx")], check=False)

    print("DONE", flush=True)

except Exception as e:
    print(f"ERROR: {type(e).__name__}: {e}", flush=True)
    sys.exit(1)
