"""MCP tool client helpers for CER authoring.

The authoring graph keeps deterministic Python transformations for reliability,
but every external knowledge/evidence capability is reached through the same
tool-call contract used by the DeerFlow MCP layer.  KIMI CODE servers are local
stdio MCP servers; the public-evidence server is imported directly because it is
part of this repo and already exposes the same tool functions.
"""

from __future__ import annotations

import json
import hashlib
import logging
import os
import re
import subprocess
import sys
import threading
import time
import urllib.parse
import urllib.request
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from typing import Any

from deerflow.mcp import cer_public_evidence_server as public_evidence
from deerflow.utils.network import force_direct_api_network

logger = logging.getLogger(__name__)

force_direct_api_network()
_DIRECT_URL_OPENER = urllib.request.build_opener(urllib.request.ProxyHandler({}))

KIMI_CODE_ROOT = Path(os.getenv("CER_AUTHORING_KIMI_CODE_ROOT", "/Users/winstonwei/Documents/KIMI CODE"))
KIMI_MCP_ROOT = KIMI_CODE_ROOT / "mcp-servers"
KIMI_PYTHON = os.getenv("CER_AUTHORING_KIMI_PYTHON", "python3")
MCP_TIMEOUT_SECONDS = int(os.getenv("CER_AUTHORING_MCP_TIMEOUT_SECONDS", "90"))
_MCP_POOL_ENABLED = os.getenv("CER_AUTHORING_ENABLE_MCP_POOL", "0") == "1"
_CIRCUIT_BREAKER_ENABLED = os.getenv("CER_AUTHORING_ENABLE_CIRCUIT_BREAKER", "1") == "1"
NCBI_BASE = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"

_SERVER_FILES = {
    "cer-kb": KIMI_MCP_ROOT / "cer-kb" / "server.py",
    "nb-check": KIMI_MCP_ROOT / "nb-check" / "server.py",
    "doc-proc": KIMI_MCP_ROOT / "doc-proc" / "server.py",
}

# P1-1: MCP version compatibility matrix
# Format: {"server_name": {"min_version": "x.y.z", "max_version": "x.y.z"}}
_MCP_COMPATIBILITY: dict[str, dict[str, str]] = {
    "cer-kb": {"min_version": "1.0.0", "max_version": "99.99.99"},
    "nb-check": {"min_version": "1.0.0", "max_version": "99.99.99"},
    "doc-proc": {"min_version": "1.0.0", "max_version": "99.99.99"},
}


class MCPVersionMismatchError(RuntimeError):
    """Raised when an MCP server version is outside the compatibility matrix."""
    pass


def _parse_version(v: str) -> tuple[int, int, int]:
    """Parse 'x.y.z' into (x, y, z)."""
    parts = v.strip().split(".")
    return (int(parts[0]) if len(parts) > 0 else 0,
            int(parts[1]) if len(parts) > 1 else 0,
            int(parts[2]) if len(parts) > 2 else 0)


def _version_in_range(version: str, min_v: str, max_v: str) -> bool:
    v = _parse_version(version)
    lo = _parse_version(min_v)
    hi = _parse_version(max_v)
    return lo <= v <= hi


def get_mcp_server_version(server: str) -> str | None:
    """Extract __version__ from an MCP server Python file.

    Uses regex parsing (not execution) for safety — avoids side effects
    from running the server module directly.
    """
    server_file = _SERVER_FILES.get(server)
    if not server_file or not server_file.exists():
        return None
    try:
        content = server_file.read_text(encoding="utf-8")
        # Match __version__ = "x.y.z" or __version__ = 'x.y.z'
        match = __import__("re").search(
            r'^\s*__version__\s*=\s*["\']([^"\']+)["\']',
            content,
            __import__("re").MULTILINE,
        )
        return match.group(1) if match else None
    except Exception:
        return None


def check_mcp_server_versions(raise_on_mismatch: bool = False) -> dict[str, dict[str, Any]]:
    """Health-check all MCP servers for version compatibility.

    Returns a dict per server:
      {"version": "x.y.z", "compatible": True/False, "min": "...", "max": "..."}

    NOTE: MCP servers outside this repo may not define __version__.
    Missing version triggers a warning, not a hard error, to avoid
    breaking existing deployments while encouraging version adoption.
    """
    import logging
    logger = logging.getLogger(__name__)
    report: dict[str, dict[str, Any]] = {}
    mismatches: list[str] = []
    for server, constraints in _MCP_COMPATIBILITY.items():
        version = get_mcp_server_version(server)
        if version is None:
            report[server] = {
                "version": None,
                "compatible": True,  # Soft-fail: no version = assume compatible
                "warning": "server_not_found_or_no_version",
            }
            logger.warning(
                "MCP server '%s' has no __version__. Consider adding __version__ to %s for safety.",
                server,
                _SERVER_FILES.get(server),
            )
            continue
        compatible = _version_in_range(version, constraints["min_version"], constraints["max_version"])
        report[server] = {
            "version": version,
            "compatible": compatible,
            "min_version": constraints["min_version"],
            "max_version": constraints["max_version"],
        }
        if not compatible:
            mismatches.append(
                f"{server}: version {version} outside range [{constraints['min_version']}, {constraints['max_version']}]"
            )
    if mismatches and raise_on_mismatch:
        raise MCPVersionMismatchError("MCP version mismatch: " + "; ".join(mismatches))
    return report


# ── PDF Reading Tool ──

def read_pdf(file_path: str, max_pages: int = 50) -> dict[str, Any]:
    """Read PDF file and extract text content.

    Uses pymupdf (fitz) as primary, falls back to pdfplumber.
    Returns {text, pages, metadata}.
    """
    result: dict[str, Any] = {"file_path": file_path, "text": "", "pages": 0, "metadata": {}}
    try:
        import fitz
        doc = fitz.open(file_path)
        result["metadata"] = dict(doc.metadata or {})
        pages_text = []
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            pages_text.append(page.get_text())
        result["text"] = "\n\n".join(pages_text)
        result["pages"] = len(pages_text)
        doc.close()
    except Exception:
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    if i >= max_pages:
                        break
                    pages_text.append(page.extract_text() or "")
                result["text"] = "\n\n".join(pages_text)
                result["pages"] = len(pages_text)
                result["metadata"] = dict(pdf.metadata or {})
        except Exception as e:
            result["error"] = str(e)
    return result


def _call_tool_inner(
    server: str,
    tool: str,
    arguments: dict[str, Any],
    timeout: int | None,
    skip_version_check: bool,
    started: float,
) -> dict[str, Any]:
    """Core MCP tool call logic (pool → subprocess → public)."""

    # ── Phase 4: MCP Process Pool (feature-flagged) ──
    if _MCP_POOL_ENABLED and server in _SERVER_FILES:
        try:
            from deerflow.runtime.cer_authoring.mcp_pool import mcp_pool_call
            return mcp_pool_call(server, tool, arguments, timeout)
        except Exception as exc:
            logger.warning("MCP pool call failed for %s/%s: %s — falling back to subprocess", server, tool, exc)

    if server == "cer-public-evidence":
        return _call_public_tool(tool, arguments, started)
    server_file = _SERVER_FILES.get(server)
    if not server_file or not server_file.exists():
        return _error(server, tool, arguments, "server_unavailable", f"MCP server file not found: {server_file}", started)
    # P1-1: Optional version check on first call per process
    if not skip_version_check and server in _MCP_COMPATIBILITY:
        version = get_mcp_server_version(server)
        constraints = _MCP_COMPATIBILITY[server]
        if version and not _version_in_range(version, constraints["min_version"], constraints["max_version"]):
            return _error(
                server, tool, arguments, "version_mismatch",
                f"MCP server {server} version {version} is outside compatible range "
                f"[{constraints['min_version']}, {constraints['max_version']}]. "
                f"Use skip_version_check=True to bypass or update _MCP_COMPATIBILITY.",
                started,
            )
    request = {
        "jsonrpc": "2.0",
        "id": 1,
        "method": "tools/call",
        "params": {"name": tool, "arguments": arguments},
    }
    try:
        proc = subprocess.Popen(
            [KIMI_PYTHON, str(server_file)],
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        stdout, stderr = proc.communicate(json.dumps(request, ensure_ascii=False) + "\n", timeout=timeout or MCP_TIMEOUT_SECONDS)
    except subprocess.TimeoutExpired:
        try:
            proc.kill()  # type: ignore[name-defined]
        except Exception:
            pass
        return _error(server, tool, arguments, "timeout", f"Tool call exceeded {timeout or MCP_TIMEOUT_SECONDS}s", started)
    except Exception as exc:
        return _error(server, tool, arguments, exc.__class__.__name__, str(exc), started)
    if proc.returncode not in (0, None):
        return _error(server, tool, arguments, "process_error", stderr.strip() or f"return code {proc.returncode}", started)
    response = _parse_jsonrpc_stdout(stdout)
    if not response:
        return _error(server, tool, arguments, "invalid_response", (stderr + "\n" + stdout)[-2000:], started)
    if response.get("error"):
        return _error(server, tool, arguments, "jsonrpc_error", json.dumps(response.get("error"), ensure_ascii=False), started)
    result = response.get("result", {})
    payload = _extract_tool_payload(result)
    if not isinstance(payload, dict):
        payload = {"status": "ok", "value": payload}
    return _with_meta(payload, server, tool, arguments, started, stderr=stderr.strip())


# ── Document Reading with Table Preservation ──

def read_document(file_path: str, max_pages: int = 50) -> dict[str, Any]:
    """Read document (PDF/DOCX) and extract text + tables with structure.

    Uses pdfplumber for table extraction (superior to pymupdf for tables).
    Falls back to pymupdf for text-only extraction.
    """
    ext = Path(file_path).suffix.lower()
    result: dict[str, Any] = {"file_path": file_path, "text": "", "pages": 0, "tables": [], "metadata": {}}
    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                all_text = []
                for i, page in enumerate(pdf.pages):
                    if i >= max_pages:
                        break
                    all_text.append(page.extract_text() or "")
                    page_tables = page.extract_tables()
                    for t in page_tables:
                        if t:
                            result["tables"].append({"page": i + 1, "rows": t})
                result["text"] = "\n\n".join(all_text)
                result["pages"] = len(all_text)
                result["metadata"] = dict(pdf.metadata or {})
        except Exception:
            return read_pdf(file_path, max_pages)
    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(file_path)
            all_text = [p.text for p in doc.paragraphs if p.text.strip()]
            result["text"] = "\n".join(all_text)
            result["pages"] = 1
            for table in doc.tables:
                rows = [[cell.text for cell in row.cells] for row in table.rows]
                result["tables"].append({"rows": rows})
        except Exception:
            result["error"] = "python-docx unavailable"
    else:
        result["text"] = Path(file_path).read_text(encoding="utf-8", errors="replace")
    return result


def call_tool(server: str, tool: str, arguments: dict[str, Any] | None = None, timeout: int | None = None, skip_version_check: bool = False) -> dict[str, Any]:
    """Call an authoring MCP tool and return a structured result.

    Results include a small `_mcp` envelope so downstream gates can prove that a
    tool was actually invoked and can distinguish source failures from empty
    search results.

    If CER_AUTHORING_ENABLE_MCP_POOL=1, subprocess-based servers (nb-check,
    cer-kb, doc-proc) will use a reusable process pool instead of spawning
    a new subprocess for every call.

    If CER_AUTHORING_ENABLE_CIRCUIT_BREAKER=1, calls are protected by a
    per-server circuit breaker that rejects calls after consecutive failures.
    """

    arguments = arguments or {}
    started = time.time()

    def _track(status: str) -> None:
        try:
            from deerflow.runtime.cer_authoring.event_bus.metrics import metrics_collector
            duration_ms = (time.time() - started) * 1000
            metrics_collector.inc(
                "mcp_pool.calls_total",
                labels={"server": server, "status": status},
            )
            metrics_collector.observe(
                "mcp_pool.call_duration_ms",
                duration_ms,
                labels={"server": server, "status": status},
            )
        except Exception:
            pass

    if _CIRCUIT_BREAKER_ENABLED and server not in ("cer-public-evidence",):
        try:
            from deerflow.runtime.cer_authoring.circuit_breaker import (
                get_circuit_breaker,
                CircuitBreakerOpenError,
            )

            breaker = get_circuit_breaker(
                name=f"mcp-{server}",
                failure_threshold=3,
                cooldown_seconds=10.0,
            )
            result = breaker.call(
                _call_tool_inner, server, tool, arguments, timeout, skip_version_check, started
            )
            status = result.get("status", result.get("_mcp", {}).get("status", "ok"))
            _track(status)
            return result
        except CircuitBreakerOpenError as exc:
            logger.warning("Circuit breaker open for %s: %s", server, exc)
            _track("circuit_breaker_open")
            return _error(
                server, tool, arguments,
                "circuit_breaker_open",
                f"MCP server {server} is temporarily unavailable due to recent failures. Retry after {exc.remaining_seconds:.0f}s.",
                started,
            )

    result = _call_tool_inner(server, tool, arguments, timeout, skip_version_check, started)
    status = result.get("status", result.get("_mcp", {}).get("status", "ok"))
    _track(status)
    return result


def call_public(tool: str, arguments: dict[str, Any] | None = None) -> dict[str, Any]:
    return call_tool("cer-public-evidence", tool, arguments or {})


def call_clinical_source_adapter(adapter: str, arguments: dict[str, Any] | None = None) -> dict[str, Any]:
    arguments = arguments or {}
    started = time.time()
    adapters = {
        "pmc_fulltext_search": _adapter_pmc_fulltext_search,
        "europe_pmc_adapter_search": _adapter_europe_pmc_search,
        "clinicaltrials_gov_adapter_search": _adapter_clinicaltrials_gov_search,
    }
    func = adapters.get(adapter)
    if not func:
        return _error("cer-clinical-source-adapters", adapter, arguments, "adapter_unavailable", f"Adapter not found: {adapter}", started)
    try:
        payload = func(**arguments)
        if not isinstance(payload, dict):
            payload = {"status": "ok", "value": payload}
        return _with_meta(payload, "cer-clinical-source-adapters", adapter, arguments, started)
    except Exception as exc:
        return _error("cer-clinical-source-adapters", adapter, arguments, exc.__class__.__name__, str(exc), started)


def mcp_log_entry(result: dict[str, Any], stage: str) -> dict[str, Any]:
    meta = result.get("_mcp", {})
    return {
        "stage": stage,
        "server": meta.get("server"),
        "tool": meta.get("tool"),
        "status": result.get("status") or meta.get("status"),
        "elapsed_ms": meta.get("elapsed_ms"),
        "arguments": meta.get("arguments"),
        "summary": _summarize_result(result),
    }


def _call_public_tool(tool: str, arguments: dict[str, Any], started: float) -> dict[str, Any]:
    func = getattr(public_evidence, tool, None)
    if not callable(func):
        return _error("cer-public-evidence", tool, arguments, "tool_unavailable", f"Tool not found: {tool}", started)
    try:
        payload = func(**arguments)
        if not isinstance(payload, dict):
            payload = {"status": "ok", "value": payload}
        return _with_meta(payload, "cer-public-evidence", tool, arguments, started)
    except Exception as exc:
        return _error("cer-public-evidence", tool, arguments, exc.__class__.__name__, str(exc), started)


def _parse_jsonrpc_stdout(stdout: str) -> dict[str, Any] | None:
    for line in reversed(stdout.splitlines()):
        line = line.strip()
        if not line.startswith("{"):
            continue
        try:
            data = json.loads(line)
        except json.JSONDecodeError:
            continue
        if data.get("jsonrpc") == "2.0":
            return data
    return None


def _extract_tool_payload(result: dict[str, Any]) -> Any:
    content = result.get("content") if isinstance(result, dict) else None
    if isinstance(content, list) and content:
        text = content[0].get("text") if isinstance(content[0], dict) else None
        if isinstance(text, str):
            try:
                return json.loads(text)
            except json.JSONDecodeError:
                return {"status": "ok", "text": text}
    return result


def _with_meta(
    payload: dict[str, Any],
    server: str,
    tool: str,
    arguments: dict[str, Any],
    started: float,
    stderr: str = "",
) -> dict[str, Any]:
    wrapped = dict(payload)
    wrapped["_mcp"] = {
        "server": server,
        "tool": tool,
        "arguments": _safe_arguments(arguments),
        "elapsed_ms": int((time.time() - started) * 1000),
        "status": wrapped.get("status", "ok"),
        "stderr": stderr[-1200:] if stderr else "",
    }
    return wrapped


def _error(server: str, tool: str, arguments: dict[str, Any], error_type: str, message: str, started: float) -> dict[str, Any]:
    return _with_meta(
        {
            "status": "source_unavailable",
            "error_type": error_type,
            "message": message,
        },
        server,
        tool,
        arguments,
        started,
    )


def _safe_arguments(arguments: dict[str, Any]) -> dict[str, Any]:
    safe = {}
    for key, value in arguments.items():
        if isinstance(value, str) and len(value) > 1000:
            safe[key] = value[:1000] + "...[truncated]"
        else:
            safe[key] = value
    return safe


def _summarize_result(result: dict[str, Any]) -> str:
    if "results_count" in result:
        return f"results_count={result.get('results_count')}"
    if "count" in result or "returned_count" in result:
        return f"count={result.get('count')}; returned={result.get('returned_count')}"
    if "template" in result:
        return "template_loaded"
    if "checklist" in result:
        return "checklist_loaded"
    if "pico_framework" in result:
        return "search_strategy_loaded"
    if "predictions" in result:
        return f"predictions={len(result.get('predictions') or [])}"
    if "issues" in result:
        return f"issues={len(result.get('issues') or [])}"
    return str(result.get("status", "ok"))


_EXTERNAL_API_RETRY_ENABLED = os.getenv("CER_AUTHORING_ENABLE_API_RETRY", "1") == "1"
_EXTERNAL_API_MAX_RETRIES = int(os.getenv("CER_AUTHORING_API_MAX_RETRIES", "5"))
_EXTERNAL_API_RETRY_BACKOFF = float(os.getenv("CER_AUTHORING_API_RETRY_BACKOFF", "2.0"))
_CIRCUIT_BREAKER_THRESHOLD = int(os.getenv("CER_AUTHORING_CIRCUIT_BREAKER_THRESHOLD", "3"))
# host → consecutive failure count; cleared on first success for that host
_circuit_state: dict[str, int] = {}


def _host(url: str) -> str:
    return urllib.parse.urlparse(url).hostname or url


def _circuit_trip(url: str) -> bool:
    """Return True if circuit is open for this host (skip request)."""
    h = _host(url)
    count = _circuit_state.get(h, 0)
    if count >= _CIRCUIT_BREAKER_THRESHOLD:
        logger.warning("Circuit OPEN for %s (%d consecutive failures) — skipping", h, count)
        return True
    return False


def _circuit_record(url: str, success: bool) -> None:
    h = _host(url)
    if success:
        if h in _circuit_state:
            _circuit_state.pop(h)
    else:
        _circuit_state[h] = _circuit_state.get(h, 0) + 1

# Rate-limit guard: limit concurrent external API requests to avoid 429s
_api_semaphore = threading.Semaphore(int(os.getenv("CER_AUTHORING_API_MAX_CONCURRENT", "2")))


def _retry_with_backoff(
    func: Callable[..., Any],
    max_retries: int = _EXTERNAL_API_MAX_RETRIES,
    backoff_base: float = _EXTERNAL_API_RETRY_BACKOFF,
    label: str = "",
    *args: Any,
    **kwargs: Any,
) -> Any:
    """Execute func with exponential backoff retry on transient failures.

    Retries on URLError, HTTPError, and TimeoutError.
    """
    import urllib.error

    last_exc: Exception | None = None
    for attempt in range(max_retries + 1):
        try:
            return func(*args, **kwargs)
        except urllib.error.HTTPError as exc:
            # Retry on 5xx or 429 (rate limited)
            if attempt < max_retries and (exc.code >= 500 or exc.code == 429):
                wait = backoff_base ** attempt
                logger.warning(
                    "External API %s attempt %d/%d failed with HTTP %d, retrying in %.1fs...",
                    label,
                    attempt + 1,
                    max_retries + 1,
                    exc.code,
                    wait,
                )
                time.sleep(wait)
                last_exc = exc
                continue
            raise
        except (urllib.error.URLError, TimeoutError, OSError) as exc:
            if attempt < max_retries:
                wait = backoff_base ** attempt
                logger.warning(
                    "External API %s attempt %d/%d failed (%s), retrying in %.1fs...",
                    label,
                    attempt + 1,
                    max_retries + 1,
                    exc.__class__.__name__,
                    wait,
                )
                time.sleep(wait)
                last_exc = exc
                continue
            raise
    # Should never reach here, but just in case
    raise last_exc or RuntimeError("Retry exhausted")


def _adapter_fetch_json(url: str, timeout: int = 20) -> dict[str, Any]:
    if _circuit_trip(url):
        raise urllib.error.URLError(f"Circuit open for {_host(url)} — skipping request")
    req = urllib.request.Request(url, headers={"User-Agent": "DeerFlow-CER-Clinical-Source-Adapters/1.0"})

    def _fetch() -> dict[str, Any]:
        with _api_semaphore:
            with _DIRECT_URL_OPENER.open(req, timeout=timeout) as response:
                return json.loads(response.read().decode("utf-8", errors="replace"))

    try:
        if _EXTERNAL_API_RETRY_ENABLED:
            result = _retry_with_backoff(_fetch, label=url)
        else:
            result = _fetch()
        _circuit_record(url, success=True)
        return result
    except Exception:
        _circuit_record(url, success=False)
        raise


def _adapter_fetch_text(url: str, timeout: int = 20) -> str:
    if _circuit_trip(url):
        raise urllib.error.URLError(f"Circuit open for {_host(url)} — skipping request")
    req = urllib.request.Request(url, headers={"User-Agent": "DeerFlow-CER-Clinical-Source-Adapters/1.0"})

    def _fetch() -> str:
        with _api_semaphore:
            with _DIRECT_URL_OPENER.open(req, timeout=timeout) as response:
                return response.read().decode("utf-8", errors="replace")

    try:
        if _EXTERNAL_API_RETRY_ENABLED:
            result = _retry_with_backoff(_fetch, label=url)
        else:
            result = _fetch()
        _circuit_record(url, success=True)
        return result
    except Exception:
        _circuit_record(url, success=False)
        raise


def _ncbi_url(endpoint: str, params: dict[str, Any]) -> str:
    clean = {key: value for key, value in params.items() if value not in (None, "")}
    return f"{NCBI_BASE}/{endpoint}?{urllib.parse.urlencode(clean, doseq=True)}"


def _adapter_pmc_fulltext_search(query: str, retmax: int = 20, date_from: str = "", date_to: str = "", **_: Any) -> dict[str, Any]:
    search_params: dict[str, Any] = {
        "db": "pmc",
        "term": query,
        "retmode": "json",
        "retmax": max(1, min(int(retmax), 100)),
        "sort": "relevance",
    }
    if date_from or date_to:
        search_params["datetype"] = "pdat"
        search_params["mindate"] = date_from
        search_params["maxdate"] = date_to
    search_url = _ncbi_url("esearch.fcgi", search_params)
    search_data = _adapter_fetch_json(search_url)
    result = search_data.get("esearchresult", {})
    pmcids = [str(value) for value in result.get("idlist", [])]
    if not pmcids:
        return _adapter_response("NCBI PMC", "pmc_fulltext_search", query, search_url, int(result.get("count", 0)), [])
    fetch_url = _ncbi_url("efetch.fcgi", {"db": "pmc", "id": ",".join(pmcids), "retmode": "xml"})
    xml_text = _adapter_fetch_text(fetch_url)
    records = _pmc_records_from_xml(xml_text, query, fetch_url)
    return _adapter_response("NCBI PMC", "pmc_fulltext_search", query, search_url, int(result.get("count", 0)), records, extra={"efetch_url": fetch_url})


def _pmc_records_from_xml(xml_text: str, query: str, source_url: str) -> list[dict[str, Any]]:
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return []
    records = []
    for idx, article in enumerate(root.findall(".//article"), start=1):
        pmcid = (
            _xml_text(article.find(".//article-id[@pub-id-type='pmc']"))
            or _xml_text(article.find(".//article-id[@pub-id-type='pmcid']"))
            or f"PMC-UNKNOWN-{idx:03d}"
        )
        pmid = _xml_text(article.find(".//article-id[@pub-id-type='pmid']"))
        doi = _xml_text(article.find(".//article-id[@pub-id-type='doi']"))
        title = _xml_text(article.find(".//article-title"))
        abstract = " ".join(_xml_text(node) for node in article.findall(".//abstract//p") if _xml_text(node))
        body = " ".join(_xml_text(node) for node in article.findall(".//body//p") if _xml_text(node))
        full_text = "\n\n".join(part for part in (abstract, body) if part)
        records.append(
            _native_clinical_record(
                source_db="NCBI PMC",
                source_type="literature_pmc_fulltext",
                record_id=f"PMC{pmcid}" if pmcid.isdigit() else pmcid,
                query=query,
                title=title,
                abstract=abstract,
                full_text=full_text,
                source_anchor=f"PMC:{pmcid}",
                source_url=source_url,
                pmid=pmid,
                doi=doi,
                raw={"pmcid": pmcid, "pmid": pmid, "doi": doi},
            )
        )
    return records


def _adapter_europe_pmc_search(query: str, page_size: int = 25, **_: Any) -> dict[str, Any]:
    params = {"query": query, "format": "json", "pageSize": max(1, min(int(page_size), 100))}
    url = "https://www.ebi.ac.uk/europepmc/webservices/rest/search?" + urllib.parse.urlencode(params)
    data = _adapter_fetch_json(url)
    raw_records = data.get("resultList", {}).get("result", [])
    records = [_europe_pmc_native_record(row, query, url) for row in raw_records]
    return _adapter_response("Europe PMC", "europe_pmc_adapter_search", query, url, int(data.get("hitCount", len(records))), records)


def _europe_pmc_native_record(row: dict[str, Any], query: str, source_url: str) -> dict[str, Any]:
    full_text_links = []
    for link in row.get("fullTextUrlList", {}).get("fullTextUrl", []) or []:
        if isinstance(link, dict):
            full_text_links.append(link.get("url") or link.get("documentStyle") or "")
    record_id = str(row.get("id") or row.get("pmid") or row.get("doi") or "")
    return _native_clinical_record(
        source_db="Europe PMC",
        source_type="literature_europe_pmc",
        record_id=record_id,
        query=query,
        title=row.get("title", ""),
        abstract=row.get("abstractText", ""),
        source_anchor=f"EUROPE_PMC:{record_id}",
        source_url=source_url,
        pmid=row.get("pmid", ""),
        doi=row.get("doi", ""),
        raw=row,
        extra={
            "journal": row.get("journalTitle", ""),
            "pub_year": row.get("pubYear", ""),
            "full_text_links": [link for link in full_text_links if link],
            "has_full_text": str(row.get("hasFullText", "")).lower() == "y",
            "grants": row.get("grantsList", {}).get("grant", []) if isinstance(row.get("grantsList"), dict) else [],
        },
    )


def _adapter_clinicaltrials_gov_search(
    query: str,
    page_size: int = 25,
    target_keywords: list[str] | str | None = None,
    device_profile: dict[str, Any] | None = None,
    **_: Any,
) -> dict[str, Any]:
    params = {"query.term": query, "format": "json", "pageSize": max(1, min(int(page_size), 100))}
    url = "https://clinicaltrials.gov/api/v2/studies?" + urllib.parse.urlencode(params)
    data = _adapter_fetch_json(url)
    records = [
        _clinicaltrials_native_record(row, query, url, target_keywords=target_keywords, device_profile=device_profile or {})
        for row in data.get("studies", []) or []
    ]
    return _adapter_response("ClinicalTrials.gov", "clinicaltrials_gov_adapter_search", query, url, int(data.get("totalCount", len(records))), records)


def _clinicaltrials_native_record(
    study: dict[str, Any],
    query: str,
    source_url: str,
    *,
    target_keywords: list[str] | str | None,
    device_profile: dict[str, Any],
) -> dict[str, Any]:
    protocol = study.get("protocolSection", {}) if isinstance(study.get("protocolSection"), dict) else {}
    identification = protocol.get("identificationModule", {}) if isinstance(protocol.get("identificationModule"), dict) else {}
    status = protocol.get("statusModule", {}) if isinstance(protocol.get("statusModule"), dict) else {}
    design = protocol.get("designModule", {}) if isinstance(protocol.get("designModule"), dict) else {}
    arms = protocol.get("armsInterventionsModule", {}) if isinstance(protocol.get("armsInterventionsModule"), dict) else {}
    outcomes = protocol.get("outcomesModule", {}) if isinstance(protocol.get("outcomesModule"), dict) else {}
    results = study.get("resultsSection", {}) if isinstance(study.get("resultsSection"), dict) else {}
    adverse = results.get("adverseEventsModule", {}) if isinstance(results.get("adverseEventsModule"), dict) else {}
    nct_id = str(identification.get("nctId") or study.get("nctId") or "")
    brief_title = identification.get("briefTitle", "")
    official_title = identification.get("officialTitle", "")
    interventions = arms.get("interventions", []) or []
    intervention_names = [str(row.get("name") or "") for row in interventions if isinstance(row, dict)]
    intervention_types = [str(row.get("type") or "") for row in interventions if isinstance(row, dict)]
    conditions = protocol.get("conditionsModule", {}).get("conditions", []) if isinstance(protocol.get("conditionsModule"), dict) else []
    result_facts = _clinicaltrials_result_facts(results, adverse)
    relationship = _clinical_trial_relationship(
        " ".join([brief_title, official_title, " ".join(intervention_names), " ".join(conditions)]),
        target_keywords=target_keywords,
        device_profile=device_profile,
    )
    abstract = " | ".join(part for part in (brief_title, official_title, "; ".join(intervention_names), "; ".join(conditions)) if part)
    return _native_clinical_record(
        source_db="ClinicalTrials.gov",
        source_type="clinical_trial_record",
        record_id=nct_id,
        query=query,
        title=brief_title or official_title,
        abstract=abstract,
        source_anchor=nct_id,
        source_url=source_url,
        raw=study,
        extra={
            "nct_id": nct_id,
            "official_title": official_title,
            "study_type": design.get("studyType", ""),
            "phase": "; ".join(design.get("phases", []) or []),
            "enrollment": design.get("enrollmentInfo", {}).get("count") if isinstance(design.get("enrollmentInfo"), dict) else "",
            "completion_date": status.get("completionDateStruct", {}).get("date") if isinstance(status.get("completionDateStruct"), dict) else "",
            "intervention_names": intervention_names,
            "intervention_types": intervention_types,
            "conditions": conditions,
            "primary_outcomes": outcomes.get("primaryOutcomes", []) or [],
            "secondary_outcomes": outcomes.get("secondaryOutcomes", []) or [],
            "result_facts": result_facts,
            "results_status": "RESULTS_AVAILABLE" if result_facts else "NO_RESULTS_AVAILABLE",
            **relationship,
        },
    )


def _clinicaltrials_result_facts(results: dict[str, Any], adverse: dict[str, Any]) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    for module_name in ("outcomeMeasuresModule",):
        module = results.get(module_name, {}) if isinstance(results.get(module_name), dict) else {}
        for idx, measure in enumerate(module.get("outcomeMeasures", []) or [], start=1):
            if not isinstance(measure, dict):
                continue
            facts.append(
                {
                    "fact_id": f"TRIAL-OUTCOME-{idx:03d}",
                    "fact_type": "primary_or_secondary_outcome",
                    "endpoint": measure.get("title") or measure.get("measure"),
                    "value": _first_outcome_value(measure),
                    "unit": measure.get("unitOfMeasure", ""),
                    "extraction_confidence": "high",
                }
            )
    for event_key, fact_type in (("seriousEvents", "serious_adverse_event"), ("otherEvents", "other_adverse_event")):
        for idx, event in enumerate(adverse.get(event_key, []) or [], start=1):
            if isinstance(event, dict):
                facts.append(
                    {
                        "fact_id": f"TRIAL-AE-{fact_type}-{idx:03d}",
                        "fact_type": fact_type,
                        "endpoint": event.get("term") or event.get("organSystem"),
                        "value": event.get("stats", [{}])[0].get("numEvents") if isinstance(event.get("stats"), list) and event.get("stats") else "",
                        "unit": "count",
                        "extraction_confidence": "high",
                    }
                )
    return facts


def _first_outcome_value(measure: dict[str, Any]) -> str:
    for group in measure.get("groups", []) or []:
        if isinstance(group, dict) and group.get("value"):
            return str(group.get("value"))
    for class_row in measure.get("classes", []) or []:
        if not isinstance(class_row, dict):
            continue
        for category in class_row.get("categories", []) or []:
            if not isinstance(category, dict):
                continue
            for measurement in category.get("measurements", []) or []:
                if isinstance(measurement, dict) and measurement.get("value"):
                    return str(measurement.get("value"))
    return ""


def _clinical_trial_relationship(text: str, *, target_keywords: list[str] | str | None, device_profile: dict[str, Any]) -> dict[str, Any]:
    haystack = text.lower()
    keywords = target_keywords
    if isinstance(keywords, str):
        keywords = [part.strip() for part in re.split(r"[,;]", keywords) if part.strip()]
    keyword_hits = [keyword for keyword in (keywords or []) if str(keyword).strip() and str(keyword).lower() in haystack]
    profile_terms = [
        device_profile.get("device_name"),
        device_profile.get("device_type"),
        device_profile.get("device_family"),
        device_profile.get("mode_of_action"),
        device_profile.get("intended_purpose"),
    ]
    profile_hits = [str(term) for term in profile_terms if term and str(term).lower()[:80] in haystack]
    technology_terms = ("catheter", "ablation", "pulsed field", "radiofrequency", "endoscope", "ureteroscope", "sheath", "clip", "pump", "software", "implant")
    technology_hits = [term for term in technology_terms if term in haystack]
    if keyword_hits or profile_hits:
        relationship = "subject"
        confidence = "high" if keyword_hits else "medium"
        rationale = "matched target device/intervention/intended-use signal; sponsor is not used as the deciding signal"
    elif technology_hits:
        relationship = "similar"
        confidence = "medium"
        rationale = "matched device technology/intervention class without subject-device keyword match"
    else:
        relationship = "unrelated"
        confidence = "low"
        rationale = "no device/intervention/technology/intended-use signal matched"
    return {
        "device_relationship": relationship,
        "relationship_rationale": rationale,
        "relationship_confidence": confidence,
        "relationship_signal_terms": keyword_hits + profile_hits + technology_hits,
    }


def _native_clinical_record(
    *,
    source_db: str,
    source_type: str,
    record_id: str,
    query: str,
    title: str,
    abstract: str = "",
    full_text: str = "",
    source_anchor: str = "",
    source_url: str = "",
    pmid: str = "",
    doi: str = "",
    raw: dict[str, Any] | None = None,
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    query_signature = _query_signature(query)
    record = {
        "source_db": source_db,
        "source_database": source_db,
        "source_type": source_type,
        "record_id": str(record_id),
        "stable_record_id": f"{source_db}:{record_id}",
        "retrieval_timestamp": datetime.now().astimezone().isoformat(),
        "query_signature": query_signature,
        "query": query,
        "title": title or "",
        "abstract": abstract or "",
        "full_text": full_text or "",
        "source_anchor": source_anchor or str(record_id),
        "source_url": source_url,
        "pmid": pmid or "",
        "doi": doi or "",
        "raw": raw or {},
    }
    if extra:
        record.update(extra)
    return record


def _adapter_response(
    database: str,
    adapter: str,
    query: str,
    url: str,
    count: int,
    records: list[dict[str, Any]],
    *,
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    payload = {
        "status": "ok",
        "database": database,
        "adapter": adapter,
        "query": query,
        "search_date": datetime.now().astimezone().date().isoformat(),
        "url": url,
        "count": count,
        "returned_count": len(records),
        "records": records,
    }
    if extra:
        payload.update(extra)
    return payload


def _xml_text(node: Any) -> str:
    if node is None:
        return ""
    return " ".join("".join(node.itertext()).split())


def _query_signature(query: str) -> str:
    normalized = re.sub(r"\s+", " ", str(query or "").strip().lower())
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()[:16]
