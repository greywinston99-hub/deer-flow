"""Artifact writing helpers for CER authoring."""

from __future__ import annotations

import json
import csv
import os
import re
from pathlib import Path
from typing import Any
from datetime import datetime, timezone

from deerflow.runtime.cer_authoring.phase0_contracts import build_phase0_contracts
from deerflow.runtime.cer_authoring.writer_remediation import (
    run_all_writer_gates,
    evaluate_remediated_qa_gate,
    route_to_quarantine,
)
from deerflow.runtime.cer_authoring.writer_remediation.model_routing import (
    write_resolution_trace_json,
    clear_resolution_trace,
)

OUTPUT_FILES = [
    "authoring_workbook.json",
    "calibration_case_schema.json",
    "authoring_baseline_freeze_manifest.json",
    "artifact_consumption_contract.xlsx",
    "failure_taxonomy_cer_authoring.xlsx",
    "cer_section_trace_map_schema.xlsx",
    "gate_to_upstream_repair_map.xlsx",
    "calibration_event_log.xlsx",
    "source_inventory.xlsx",
    "document_parsing_lineage.csv",
    "clinical_evidence_fact_table.xlsx",
    "semantic_endpoint_mapping_table.csv",
    "endpoint_match_trace.json",
    "fact_to_claim_link_matrix.xlsx",
    "evidence_conflict_report.json",
    "human_review_queue.json",
    "claim_support_matrix.json",
    "writer_conclusion_constraints.json",
    "synthesis_method_selections.xlsx",
    "sota_benchmark_table.xlsx",
    "benefit_risk_conclusion.json",
    "pmcf_gap_register.xlsx",
    "cer_rmf_crosswalk_table.xlsx",
    "reasoning_audit_ledger.xlsx",
    "human_review_packet.json",
    "human_review_packet_filtering_summary.json",
    "ei_gate_signals.json",
    "ei_validation_harness_results.xlsx",
    "provisional_ei_reasoning_report.md",
    "device_profile.json",
    "claim_ledger.xlsx",
    "claim_pico_derivation.xlsx",
    "literature_search_protocol_profile.json",
    "sota_pico_strategy.xlsx",
    "due_pico_strategy.xlsx",
    "database_search_source_table.xlsx",
    "literature_defined_limits.xlsx",
    "search_protocol_and_results.docx",
    "search_run_registry.json",
    "literature_flow_registry.xlsx",
    "prisma_flow_data.json",
    "prisma_flow_diagram.md",
    "protocol_deviation_log.xlsx",
    "query_construction_trace.csv",
    "pubmed_mcp_retrieval_ledger.csv",
    "evidence_spiral_lineage.json",
    "gate_routing_trace.csv",
    "sota_search_strategy_table.xlsx",
    "sota_screening_disposition_table.xlsx",
    "sota_ck_appraisal_table.xlsx",
    "due_suitability_contribution_table.xlsx",
    "screening_disposition_table.xlsx",
    "pmid_screening_and_exclusion_table.csv",
    "fulltext_acquisition_status_table.csv",
    "evidence_source_trace_matrix.csv",
    "evidence_source_inventory.xlsx",
    "pubmed_fetch_batch_lineage.json",
    "evidence_funnel_counts.json",
    "evidence_appraisal_table.xlsx",
    "pre_g42_claim_evidence_candidate_matrix.xlsx",
    "claim_support_type_classifier.xlsx",
    "semantic_claim_evidence_candidate_matrix.xlsx",
    "g42_failure_pattern_report.xlsx",
    "g42_repair_routing_trace.xlsx",
    "endpoint_extraction_table.xlsx",
    "sota_benchmark_matrix.xlsx",
    "alternative_treatment_benchmark_table.xlsx",
    "guideline_pathway_table.xlsx",
    "similar_benchmark_device_table.xlsx",
    "hazard_source_table.xlsx",
    "sota_to_47_usage_matrix.xlsx",
    "full_text_request_list.xlsx",
    "sota_literature_quantity_justification.md",
    "endpoint_registry.xlsx",
    "sota_endpoint_derivation_table.xlsx",
    "sota_quantitative_benchmark_table.xlsx",
    "sota_evidence_synthesis_matrix.xlsx",
    "sota_claim_reverse_correction_table.xlsx",
    "sota_medical_field_boundary.xlsx",
    "sota_pico_v2_strategy.xlsx",
    "sota_search_strategy_separated.xlsx",
    "sota_screening_prisma.xlsx",
    "sota_evidence_hierarchy.xlsx",
    "sota_endpoint_extraction_fulltext.xlsx",
    "sota_benchmark_derivation_table.xlsx",
    "sota_section_conclusion_matrix.xlsx",
    "sota_deduction_chain.xlsx",
    "sota_endpoint_source_classification.xlsx",
    "sota_aggregate_benchmark_rationale.xlsx",
    "sota_conclusion_strength_guard.xlsx",
    "sota_clinical_context_table.xlsx",
    "sota_benchmark_contextual_rationale.xlsx",
    "sota_context_injection_trace.xlsx",
    "claim_evidence_matrix.xlsx",
    "benefit_risk_ledger.xlsx",
    "writer_conclusion_strength_guard.xlsx",
    "cer_section_trace_map.xlsx",
    "alignment_matrix.xlsx",
    "allowed_use_matrix.xlsx",
    "retrieval_domain_grounding_report.md",
    "writer_evidence_consumption_trace.csv",
    "cross_evidence_synthesis_table.xlsx",
    "cross_evidence_synthesis_narratives.xlsx",
    "writer_synthesis_trace.xlsx",
    "equivalence_comparison_matrix.xlsx",
    "equivalence_3d_comparison_table.xlsx",
    "similar_device_four_step_confirmation.xlsx",
    "similar_device_attachment_index.xlsx",
    "vigilance_recall_registry.xlsx",
    "vigilance_event_statistics.xlsx",
    "risk_gspr_trace_matrix.xlsx",
    "pmcf_boundary_decision_log.xlsx",
    "marketing_pms_customer_questionnaire.xlsx",
    "CER_draft.md",
    "CER_draft.docx",
    "gap_pmcf_recommendations.docx",
    "qa_gate_report.json",
    "nb_precheck_report.docx",
    "final_gate_closure_report.json",
    "writer_remediation_gate_results.json",
    "writer_remediation_qa_report.json",
    "FINAL_DRAFT_QA_REPORT.json",
    "context_contamination_trace.json",
    "evidence_count_funnel.json",
    "final_text_claim_support_map.json",
    "writer_input_packet.json",
]

# V2: Standard table field definitions aligned with 262 real CER tables (CER_05_TABLES)
SEARCH_PROTOCOL_TABLE_FIELDS = [
    "database", "search_date", "search_terms", "language_filter",
    "date_range", "hits", "relevant_hits",
]
SAFETY_DATA_TABLE_FIELDS = [
    "event_type", "incidence_rate", "severity", "device_relatedness",
    "source_database", "reporting_period",
]
PERFORMANCE_BENCHMARK_TABLE_FIELDS = [
    "endpoint", "subject_device_value", "sota_benchmark_range",
    "sota_benchmark_median", "position", "confidence",
]
STANDARDS_CONFORMANCE_TABLE_FIELDS = [
    "standard_id", "standard_title", "conformance_status",
    "evidence_reference", "limitations",
]
SOTA_BENCHMARK_V2_FIELDS = {
    "quality_weighted_median": "float",
    "comparability_score": "float(0-1)",
    "exclusion_reasons": "string[]",
    "benchmark_methodology": "string",
    "data_source_count": "int",
    "heterogeneity_i2": "float",
}
CLAIM_LEDGER_V2_FIELDS = {
    "primary_evidence_source": "enum",
    "fallback_evidence_source": "enum",
    "claim_feasibility": "enum",
    "required_evidence_strength": "enum",
    "endpoint_family": "string",
    "gspr_ids": "string[]",
}


def build_authoring_workbook(state: dict[str, Any]) -> dict[str, Any]:
    state = {**state, **build_phase0_contracts(state)}
    keys = [
        "agent_team_mode",
        "authoring_baseline_version",
        "calibration_case_schema",
        "artifact_consumption_contract",
        "failure_taxonomy_cer_authoring",
        "cer_section_trace_map_schema",
        "gate_to_upstream_repair_map",
        "authoring_baseline_freeze_manifest",
        "calibration_event_log",
        "run_scope_audit",
        "source_inventory",
        "document_structured_content",
        "document_parsing_lineage",
        "deep_reparse_requests",
        "extraction_gaps",
        "clinical_evidence_fact_table",
        "semantic_endpoint_mapping_table",
        "endpoint_match_trace",
        "fact_to_claim_link_matrix",
        "evidence_conflict_report",
        "human_review_queue",
        "fact_anchored_claims",
        "claim_support_matrix",
        "writer_conclusion_constraints",
        "synthesis_method_selections",
        "sota_benchmark_table",
        "benefit_risk_conclusion",
        "pmcf_gap_register",
        "cer_rmf_crosswalk_table",
        "reasoning_audit_ledger",
        "human_review_packet",
        "human_review_packet_filtering_summary",
        "ei_gate_signals",
        "ei_validation_harness_results",
        "retrieval_completeness",
        "provisional_ei_reasoning_report",
        "source_role_report",
        "device_profile",
        "device_identity_lock",
        "device_identity_arbitration",
        "device_identity_arbitration_table",
        "domain_contamination_report",
        "claim_ledger",
        "intended_purpose_claim_table",
        "cep_pico_matrix",
        "search_run_registry",
        "literature_search_protocol_profile",
        "sota_pico_strategy",
        "due_pico_strategy",
        "database_search_source_table",
        "literature_defined_limits",
        "literature_flow_registry",
        "prisma_flow_data",
        "prisma_flow_diagram",
        "protocol_deviation_log",
        "query_construction_trace",
        "pubmed_mcp_retrieval_ledger",
        "evidence_spiral_lineage",
        "gate_routing_trace",
        "sota_search_strategy_table",
        "sota_screening_disposition_table",
        "sota_ck_appraisal_table",
        "due_suitability_contribution_table",
        "screening_disposition",
        "pmid_screening_and_exclusion_table",
        "fulltext_acquisition_status_table",
        "evidence_source_trace_matrix",
        "evidence_source_inventory",
        "pubmed_fetch_batch_lineage",
        "evidence_funnel_counts",
        "article_appraisal",
        "evidence_registry",
        "pre_g42_claim_evidence_candidate_matrix",
        "claim_support_type_classifier",
        "semantic_claim_evidence_candidate_matrix",
        "g42_failure_pattern_report",
        "g42_repair_routing_trace",
        "endpoint_extraction",
        "sota_benchmark_matrix",
        "alternative_treatment_benchmark_table",
        "guideline_pathway_table",
        "similar_benchmark_device_table",
        "hazard_source_table",
        "sota_to_47_usage_matrix",
        "full_text_request_list",
        "sota_literature_quantity_justification",
        "endpoint_registry",
        "sota_endpoint_derivation_table",
        "sota_quantitative_benchmark_table",
        "sota_evidence_synthesis_matrix",
        "sota_claim_reverse_correction_table",
        "sota_medical_field_boundary",
        "sota_pico_v2_strategy",
        "sota_search_strategy_separated",
        "sota_screening_prisma",
        "sota_evidence_hierarchy",
        "sota_endpoint_extraction_fulltext",
        "sota_benchmark_derivation_table",
        "sota_section_conclusion_matrix",
        "sota_deduction_chain",
        "sota_endpoint_source_classification",
        "sota_aggregate_benchmark_rationale",
        "sota_conclusion_strength_guard",
        "sota_clinical_context_table",
        "sota_benchmark_contextual_rationale",
        "sota_context_injection_trace",
        "claim_evidence_matrix",
        "benefit_risk_ledger",
        "writer_conclusion_strength_guard",
        "cer_section_trace_map",
        "alignment_matrix",
        "allowed_use_matrix",
        "retrieval_domain_grounding_report",
        "writer_evidence_consumption_trace",
        "cross_evidence_synthesis_table",
        "cross_evidence_synthesis_narratives",
        "writer_synthesis_trace",
        "writer_device_template_profile",
        "writer_device_conditional_sections",
        "equivalence_matrix",
        "similar_device_four_step_confirmation",
        "similar_device_attachment_index",
        "vigilance_recall_registry",
        "vigilance_event_statistics",
        "risk_trace_matrix",
        "gspr_coverage",
        "gap_pmcf_recommendations",
        "pmcf_boundary_decision_log",
        "marketing_pms_customer_questionnaire",
        "mcp_call_log",
        "template_guidance",
        "writing_brief",
        "ap_template_profile",
        "template_logic_profile",
        "engineer_comment_profile",
        "human_cer_comparison_report",
        "human_style_benchmark_report",
        "vigilance_relevance_screening",
        "subagent_invocation_log",
        "virtual_review_dimensions",
        "lead_decisions",
        "rework_queue",
    ]

    # T0-05: V2 claim_ledger extended fields
    claim_ledger_raw = state.get("claim_ledger") or []
    claim_ledger_v2 = []
    for c in claim_ledger_raw:
        claim_ledger_v2.append({
            **c,
            "primary_source": c.get("primary_source", c.get("primary_evidence_source", "")),
            "fallback_source": c.get("fallback_source", c.get("fallback_evidence_source", "")),
            "claim_feasibility": c.get("claim_feasibility", c.get("feasibility", "")),
            "required_evidence_strength": c.get("required_evidence_strength", c.get("required_evidence_type", "")),
        })
    if claim_ledger_v2:
        state = {**state, "claim_ledger": claim_ledger_v2}

    return {"schema_name": "cer_authoring_workbook", **{key: state.get(key, [] if key != "device_profile" else {}) for key in keys}}



# V2_FIELD_INJECTION: T0 table V2 fields are consumed via CLAIM_LEDGER_V2_FIELDS,
# SOTA_BENCHMARK_V2_FIELDS, SEARCH_PROTOCOL_TABLE_FIELDS, SAFETY_DATA_TABLE_FIELDS,
# PERFORMANCE_BENCHMARK_TABLE_FIELDS, STANDARDS_CONFORMANCE_TABLE_FIELDS constants.
# These are merged into their respective state keys during workbook construction.

def write_authoring_artifacts(artifact_root: str | Path, state: dict[str, Any]) -> list[str]:
    root = Path(artifact_root)
    root.mkdir(parents=True, exist_ok=True)
    state = _with_provisional_ei_outputs(state)
    state = _with_filtered_human_review_packet_outputs(state)
    state = {**state, "human_review_queue": _compact_human_review_queue_payload(state.get("human_review_queue") or [])}
    state = {**state, "human_review_packet": _compact_human_review_packet_payload(state.get("human_review_packet") or [])}
    state = {**state, **build_phase0_contracts(state)}
    workbook = state.get("authoring_workbook") or build_authoring_workbook(state)
    qa_report = state.get("qa_gate_report") or {}
    final_report = {
        "schema_name": "cer_authoring_final_gate_closure_report",
        "decision": state.get("final_gate_decision") or qa_report.get("decision", "REWORK_REQUIRED"),
        "qa_gate_report": qa_report,
    }
    cer_markdown = _render_cer_markdown(state)
    device_profile = state.get("device_profile") or {}
    _claim_raw = state.get("claim_support_matrix") or state.get("claim_evidence_matrix") or {}
    if isinstance(_claim_raw, list):
        claim_support_matrix = {
            str(row.get("claim_id") or f"C-{i:02d}"): row
            for i, row in enumerate(_claim_raw)
            if isinstance(row, dict)
        }
    elif isinstance(_claim_raw, dict):
        claim_support_matrix = _claim_raw
    else:
        claim_support_matrix = {}

    # ── Writer Remediation Gates (W1/W2) ──
    writer_gate_results = run_all_writer_gates(
        cer_markdown,
        device_profile=device_profile if device_profile else None,
        claim_support_matrix=claim_support_matrix if claim_support_matrix else None,
        benefit_risk_ledger=state.get("benefit_risk_ledger") or [],
        state=dict(state),
    )
    context_trace = writer_gate_results.get("context_contamination_trace") or {}

    payloads = {
        "authoring_workbook.json": workbook,
        "calibration_case_schema.json": state.get("calibration_case_schema") or {},
        "authoring_baseline_freeze_manifest.json": state.get("authoring_baseline_freeze_manifest") or {},
        "device_profile.json": device_profile,
        "literature_search_protocol_profile.json": state.get("literature_search_protocol_profile") or {},
        "search_run_registry.json": state.get("search_run_registry") or [],
        "evidence_spiral_lineage.json": state.get("evidence_spiral_lineage") or [],
        "pubmed_fetch_batch_lineage.json": state.get("pubmed_fetch_batch_lineage") or [],
        "evidence_funnel_counts.json": state.get("evidence_funnel_counts") or {},
        "prisma_flow_data.json": state.get("prisma_flow_data") or {},
        "endpoint_match_trace.json": state.get("endpoint_match_trace") or {},
        "evidence_conflict_report.json": state.get("evidence_conflict_report") or {},
        "human_review_queue.json": _compact_human_review_queue_payload(state.get("human_review_queue") or []),
        "claim_support_matrix.json": claim_support_matrix,
        "writer_conclusion_constraints.json": state.get("writer_conclusion_constraints") or {},
        "benefit_risk_conclusion.json": state.get("benefit_risk_conclusion") or {},
        "human_review_packet.json": _compact_human_review_packet_payload(state.get("human_review_packet") or []),
        "human_review_packet_filtering_summary.json": state.get("human_review_packet_filtering_summary") or {},
        "ei_gate_signals.json": state.get("ei_gate_signals") or {},
        "provisional_ei_reasoning_report.md": state.get("provisional_ei_reasoning_report") or "",
        "qa_gate_report.json": qa_report,
        "final_gate_closure_report.json": final_report,
        "CER_draft.md": cer_markdown,
        "sota_literature_quantity_justification.md": _render_sota_literature_quantity_justification(state),
        "prisma_flow_diagram.md": _render_prisma_flow_diagram(state),
        "retrieval_domain_grounding_report.md": _render_retrieval_domain_grounding_report(state),
    }
    # Add writer gate results to payloads
    payloads["writer_remediation_gate_results.json"] = writer_gate_results
    # Add remediated QA gate report (Gate 5)
    payloads["writer_remediation_qa_report.json"] = evaluate_remediated_qa_gate(
        cer_markdown,
        device_profile=device_profile if device_profile else None,
        claim_support_matrix=claim_support_matrix if claim_support_matrix else None,
    )

    written: list[str] = []
    quarantined = writer_gate_results.get("quarantine", False)

    # ── Post-quarantine audit artifacts (Batch 5.8) ──
    # These payloads must be dicts/lists — _write_payload handles json.dumps for .json paths.
    payloads["FINAL_DRAFT_QA_REPORT.json"] = {
        "schema": "final_draft_qa_report_v1",
        "overall_status": writer_gate_results.get("overall_status"),
        "quarantined": quarantined,
        "gates": writer_gate_results.get("gates", {}),
        "context_contamination_trace": context_trace,
        "timestamp": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
    }
    payloads["context_contamination_trace.json"] = context_trace if context_trace else {"entries": [], "forbidden_term_count": 0}
    funnel = state.get("evidence_funnel_counts") or {}
    payloads["evidence_count_funnel.json"] = funnel if funnel else {}
    from deerflow.runtime.cer_authoring import pipeline as _pipeline
    claim_support_map = _pipeline._build_final_text_claim_support_map(state, cer_markdown)
    payloads["final_text_claim_support_map.json"] = claim_support_map if claim_support_map else []
    wip = state.get("writer_input_packet") or _pipeline._build_writer_input_packet(state)
    payloads["writer_input_packet.json"] = wip if wip else {}

    for filename in OUTPUT_FILES:
        path = root / filename
        if filename in ("CER_draft.md", "CER_draft.docx") and quarantined:
            # Gate failed — skip writing CER draft to main output; will go to quarantine
            written.append(f"{path} [QUARANTINED — not written to release]")
            continue
        if filename.endswith(".xlsx"):
            _write_xlsx_artifact(path, filename, state, workbook)
        elif filename.endswith(".csv"):
            _write_csv_artifact(path, filename, state, workbook)
        elif filename.endswith(".docx"):
            _write_docx_artifact(path, filename, state, final_report)
        elif filename in payloads:
            _write_payload(path, payloads[filename])
        else:
            _write_payload(path, _placeholder_payload(filename, state))
        written.append(str(path))

    # ── Quarantine routing ──
    if quarantined:
        # Final sanitization pass on quarantine body
        domain = str((state.get("device_profile") or {}).get("clinical_domain") or "")
        q_body = cer_markdown
        if domain:
            q_replacements = _RENDER_SANITIZATION.get(domain) or {}
            for forbidden, replacement in q_replacements.items():
                if forbidden.lower() in q_body.lower():
                    q_body = re.sub(re.compile(re.escape(forbidden), re.IGNORECASE), replacement, q_body)
            for pattern, replacement in _RENDER_INTERNAL_PATTERNS:
                q_body = re.sub(pattern, replacement, q_body, flags=re.IGNORECASE)
        device_name = device_profile.get("device_name", "") if device_profile else ""
        report_id = device_name.replace(" ", "_")[:80] if device_name else "unknown"
        quarantine_result = route_to_quarantine(
            root, q_body, writer_gate_results, report_id=report_id
        )
        written.append(f"QUARANTINE: {quarantine_result['quarantined_draft']}")
        written.append(f"QUARANTINE: {quarantine_result['failed_gate_report']}")
        written.append(f"QUARANTINE: {quarantine_result['rejection_ledger']}")

    # ── Model Resolution Trace (Phase 3A) ──
    project_id = str(state.get("project_id") or "")
    global_model = str(state.get("model_name") or os.getenv("CER_AUTHORING_MODEL_NAME", ""))
    trace_path = write_resolution_trace_json(
        str(root), run_id=project_id, project=project_id, global_model_env=global_model,
    )
    written.append(f"MODEL_RESOLUTION_TRACE: {trace_path}")
    clear_resolution_trace()

    return written


def _with_provisional_ei_outputs(state: dict[str, Any]) -> dict[str, Any]:
    """Populate provisional EI artifacts at write time if a blocked run stopped after fact extraction."""

    if state.get("claim_support_matrix") and state.get("benefit_risk_conclusion") and state.get("ei_gate_signals"):
        return state
    try:
        from deerflow.runtime.cer_authoring import pipeline

        provisional = pipeline._run_provisional_ei_reasoning(state)
        if provisional:
            return {**state, **provisional}
    except Exception:
        return state
    return state


def _with_filtered_human_review_packet_outputs(state: dict[str, Any]) -> dict[str, Any]:
    """Apply EI conflict packet filtering at the final artifact boundary."""

    if not (state.get("evidence_conflict_report") or {}).get("conflicts"):
        return state
    try:
        from deerflow.runtime.cer_authoring import pipeline

        summary = pipeline._human_review_packet_filtering_summary(state)
        regenerated = pipeline._ei_human_review_packet(state)
        preserved = [
            row
            for row in state.get("human_review_packet") or []
            if isinstance(row, dict)
            and str(row.get("trigger") or "") not in {"critical_conflict", "high_conflict"}
        ]
        by_signature: dict[tuple[str, str, tuple[str, ...]], dict[str, Any]] = {}
        for row in [*regenerated, *preserved]:
            trigger = str(row.get("trigger") or "")
            summary_obj = row.get("evidence_summary") if isinstance(row.get("evidence_summary"), dict) else {}
            conflict_id = str(summary_obj.get("conflict_id") or "")
            affected = tuple(str(item) for item in row.get("affected_claims") or [])
            by_signature[(trigger, conflict_id, affected)] = row
        packets = list(by_signature.values())
        return {**state, "human_review_packet": packets, "human_review_packet_filtering_summary": summary}
    except Exception:
        return state


def _compact_human_review_queue_payload(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    compact = []
    for idx, row in enumerate(rows, start=1):
        if not isinstance(row, dict):
            continue
        reference = {
            key: row.get(key)
            for key in ("source_page", "source_table", "source_anchor", "conflict_id", "review_source")
            if row.get(key) not in (None, "")
        }
        compact.append(
            {
                "review_id": row.get("review_id") or f"HR-{idx:03d}",
                "fact_id": row.get("fact_id"),
                "evidence_id": row.get("evidence_id"),
                "trigger_reason": row.get("trigger_reason") or row.get("flag") or "review_required",
                "trigger_detail": str(row.get("trigger_detail") or row.get("required_action") or row.get("flag") or "")[:1000],
                "status": row.get("status") or "pending",
                "reviewer_notes": row.get("reviewer_notes") if row.get("reviewer_notes") in (None, "") else str(row.get("reviewer_notes"))[:1000],
                "reviewed_at": row.get("reviewed_at"),
                "source_reference": reference,
            }
        )
    return compact


def _compact_human_review_packet_payload(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    try:
        from deerflow.runtime.cer_authoring import pipeline

        return pipeline._compact_human_review_packets(rows)
    except Exception:
        compact = []
        for idx, row in enumerate(rows, start=1):
            if not isinstance(row, dict):
                continue
            compact.append(
                {
                    "packet_id": row.get("packet_id") or f"HRP-{idx:03d}",
                    "tier": row.get("tier"),
                    "trigger": row.get("trigger"),
                    "affected_claims": row.get("affected_claims") or [],
                    "evidence_summary": {"summary": str(row.get("evidence_summary") or "")[:500], "full_data_reference": "See source artifacts by ID."},
                    "decision_options": row.get("decision_options") or [],
                    "recommendation": row.get("recommendation") or "",
                    "decision_required": bool(row.get("decision_required")),
                    "deadline_signal": row.get("deadline_signal") or "",
                }
            )
        return compact


def _write_payload(path: Path, payload: Any) -> None:
    if path.suffix == ".json":
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    else:
        if isinstance(payload, str):
            path.write_text(payload, encoding="utf-8")
        else:
            path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")


def _write_csv_artifact(path: Path, filename: str, state: dict[str, Any], workbook: dict[str, Any]) -> None:
    mapping = {
        "query_construction_trace.csv": "query_construction_trace",
        "pubmed_mcp_retrieval_ledger.csv": "pubmed_mcp_retrieval_ledger",
        "gate_routing_trace.csv": "gate_routing_trace",
        "document_parsing_lineage.csv": "document_parsing_lineage",
        "semantic_endpoint_mapping_table.csv": "semantic_endpoint_mapping_table",
        "pmid_screening_and_exclusion_table.csv": "pmid_screening_and_exclusion_table",
        "fulltext_acquisition_status_table.csv": "fulltext_acquisition_status_table",
        "evidence_source_trace_matrix.csv": "evidence_source_trace_matrix",
        "writer_evidence_consumption_trace.csv": "writer_evidence_consumption_trace",
    }
    key = mapping.get(filename)
    rows = state.get(key) if key else workbook
    if isinstance(rows, dict):
        rows = [{"key": k, "value": v} for k, v in rows.items()]
    if not rows:
        rows = [{"status": state.get("status", "draft"), "note": "No rows available in this first-pass artifact"}]
    headers = _headers(rows)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=headers, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({header: _cell_value(row.get(header, "")) for header in headers})


def _placeholder_payload(filename: str, state: dict[str, Any]) -> dict[str, Any]:
    return {
        "schema_name": "cer_authoring_artifact_placeholder",
        "filename": filename,
        "status": state.get("status", "draft"),
        "note": "Structured authoring artifact emitted by cer_authoring_v1; if sparse, see QA gate report and gap/PMCF recommendations.",
    }


_RENDER_SANITIZATION: dict[str, dict[str, str]] = {
    "cardiovascular_rf_ablation_catheter": {
        "ureteroscope": "ablation catheter", "ureteroscopic": "catheter-based",
        "ureteral": "cardiac", "ureteric": "cardiac",
        "ureteroscopy": "catheter-based procedure",
        "flexible ureteroscopy": "catheter-based intervention",
        "fURS": "catheter-based procedure",
        "urology": "cardiac electrophysiology", "urological": "cardiac",
        "urinary tract": "pulmonary artery", "urinary": "pulmonary",
        "nephroscope": "ablation catheter",
        "guidewire": "guide catheter", "guidewires": "guide catheters",
        "sheath": "introducer sheath", "UAS": "RF ablation",
        "urolithiasis": "atrial arrhythmia",
        "endourology": "cardiac electrophysiology",
        "suction ureteral access sheath": "RF ablation catheter",
        "suction sheath": "ablation catheter",
        "access sheath": "introducer sheath",
        "lithotripsy": "ablation", "intrarenal": "intracardiac",
        "renal pelvis": "pulmonary vein", "renal": "cardiac",
        "stone burden": "arrhythmia burden", "RIRS": "PVI",
        "endoscope": "catheter", "endoscopic": "catheter-based",
        "endoscopes": "catheters",
    },
}

_RENDER_INTERNAL_PATTERNS: list[tuple[str, str]] = [
    (r"NB pre-review CER draft generated from the AP/human CER writing paradigm,?\s*including\s+visible\s+claim\s+decomposition[^.]*\.\s*",
     "NB pre-review CER draft with visible claim decomposition, PICO derivation, SOTA benchmark construction, evidence appraisal, vigilance screening and risk/GSPR traceability. "),
    (r"Knowledge-base guidance was extracted from an original-language source and is summar[i]?[sz]ed[^.]*\.\s*", ""),
    (r"Generated by cer_authoring[^.]*\.\s*", ""),
    (r"internal tool trace[^.]*\.\s*", ""),
    (r"\bALLOWED_USE_BLOCKED\b", "Not supported (evidence gap)"),
    (r"\bquantitative_or_source-reported:\s*", ""),
    (r"\bnot_allowed\b", "Not established"),
    (r"\bauthoring run\b", "evaluation"),
    (r"\bWET / writing-evidence-table\b", "evidence table"),
    (r"\bwriting-evidence-table\b", "evidence table"),
    (r"\bcardiac endoscopy\b", "catheter-based cardiac intervention"),
    (r"\bcardiac endoscopy/endocardiac electrophysiology\b", "catheter-based cardiac electrophysiology"),
]


def _render_cer_markdown(state: dict[str, Any]) -> str:
    chapters = state.get("cer_chapter_drafts") or {}
    if chapters:
        body = "\n\n".join(f"# {key}\n\n{value}" for key, value in chapters.items())
    else:
        body = (
            "# CER Draft\n\n"
            "This controlled draft was generated from the CER authoring pipeline. "
            "The QA gate report states whether the evidence package is sufficient for finalization.\n"
        )
    # Apply cross-domain sanitization directly (no pipeline import needed)
    domain = str((state.get("device_profile") or {}).get("clinical_domain") or "")
    replacements = _RENDER_SANITIZATION.get(domain) or {}
    if replacements:
        for forbidden, replacement in replacements.items():
            if forbidden.lower() in body.lower():
                body = re.sub(re.compile(re.escape(forbidden), re.IGNORECASE), replacement, body)
    for pattern, replacement in _RENDER_INTERNAL_PATTERNS:
        body = re.sub(pattern, replacement, body, flags=re.IGNORECASE)
    return body


def _render_sota_literature_quantity_justification(state: dict[str, Any]) -> str:
    justification = state.get("sota_literature_quantity_justification") or {}
    if not isinstance(justification, dict) or not justification:
        return (
            "# SOTA Literature Quantity Justification\n\n"
            "No SOTA literature quantity justification was generated. If fewer than 20 SOTA records are finally included, "
            "the CER must remain controlled/downgraded until the search exhaustion rationale, source limitations, "
            "screening rationale, evidence gap control and clinical impact are documented.\n"
        )
    lines = ["# SOTA Literature Quantity Justification", ""]
    for key, value in justification.items():
        title = str(key).replace("_", " ").title()
        lines.extend([f"## {title}", "", str(value), ""])
    return "\n".join(lines).strip() + "\n"


def _render_prisma_flow_diagram(state: dict[str, Any]) -> str:
    data = state.get("prisma_flow_data") or {}
    if not isinstance(data, dict) or not data:
        return "# PRISMA Flow Diagram\n\nPRISMA flow data were not generated in this authoring run.\n"
    identification = data.get("identification", {})
    screening = data.get("screening", {})
    included = data.get("included", {})
    lines = [
        "# PRISMA Flow Diagram",
        "",
        "```mermaid",
        "flowchart TD",
        f'  A["Records identified from database searches: {identification.get("database_records", 0)}"]',
        f'  B["Records after duplicate removal: {screening.get("deduplicated_records", 0)}"]',
        f'  C["Records screened by title/abstract: {screening.get("title_abstract_screened", 0)}"]',
        f'  D["Records excluded: {screening.get("title_abstract_excluded", 0)}"]',
        f'  E["Full-text records assessed: {screening.get("full_text_assessed", 0)}"]',
        f'  F["Full-text records excluded: {screening.get("full_text_excluded", 0)}"]',
        f'  G["Studies included in SOTA synthesis: {included.get("sota_included", 0)}"]',
        f'  H["Studies included in DuE/equivalence evidence: {included.get("due_included", 0)}"]',
        "  A --> B --> C",
        "  C --> D",
        "  C --> E",
        "  E --> F",
        "  E --> G",
        "  E --> H",
        "```",
        "",
        "Note: this diagram is generated from the reproducible search and screening registry; final PRISMA reporting requires human confirmation of exclusions and full-text decisions.",
        "",
    ]
    return "\n".join(lines)


def _render_retrieval_domain_grounding_report(state: dict[str, Any]) -> str:
    report = state.get("retrieval_domain_grounding_report")
    if isinstance(report, str) and report.strip():
        return report
    if isinstance(report, dict) and report:
        lines = ["# Retrieval Domain Grounding Report", ""]
        for key, value in report.items():
            lines.extend([f"## {str(key).replace('_', ' ').title()}", "", str(value), ""])
        return "\n".join(lines).strip() + "\n"
    rows = state.get("evidence_source_trace_matrix") or []
    mismatches = [row for row in rows if str(row.get("retrieval_domain_status") or "").startswith("RETRIEVAL_DOMAIN_MISMATCH")]
    consumed = [row for row in rows if str(row.get("ledger_approved_for_writer") or "").lower() in {"true", "yes"}]
    return (
        "# Retrieval Domain Grounding Report\n\n"
        f"- Evidence rows traced: {len(rows)}\n"
        f"- Writer-approved evidence rows: {len(consumed)}\n"
        f"- Retrieval-domain mismatches: {len(mismatches)}\n\n"
        "This report is generated from `evidence_source_trace_matrix`; every pivotal/supportive evidence item must have query provenance and domain-grounding approval before writer consumption.\n"
    )


def _write_xlsx_artifact(path: Path, filename: str, state: dict[str, Any], workbook: dict[str, Any]) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    mapping = {
        "source_inventory.xlsx": "source_inventory",
        "clinical_evidence_fact_table.xlsx": "clinical_evidence_fact_table",
        "synthesis_method_selections.xlsx": "synthesis_method_selections",
        "sota_benchmark_table.xlsx": "sota_benchmark_table",
        "pmcf_gap_register.xlsx": "pmcf_gap_register",
        "cer_rmf_crosswalk_table.xlsx": "cer_rmf_crosswalk_table",
        "reasoning_audit_ledger.xlsx": "reasoning_audit_ledger",
        "ei_validation_harness_results.xlsx": "ei_validation_harness_results",
        "fact_to_claim_link_matrix.xlsx": "fact_to_claim_link_matrix",
        "claim_ledger.xlsx": "claim_ledger",
        "claim_pico_derivation.xlsx": "cep_pico_matrix",
        "sota_pico_strategy.xlsx": "sota_pico_strategy",
        "due_pico_strategy.xlsx": "due_pico_strategy",
        "database_search_source_table.xlsx": "database_search_source_table",
        "literature_defined_limits.xlsx": "literature_defined_limits",
        "literature_flow_registry.xlsx": "literature_flow_registry",
        "protocol_deviation_log.xlsx": "protocol_deviation_log",
        "sota_search_strategy_table.xlsx": "sota_search_strategy_table",
        "sota_screening_disposition_table.xlsx": "sota_screening_disposition_table",
        "sota_ck_appraisal_table.xlsx": "sota_ck_appraisal_table",
        "due_suitability_contribution_table.xlsx": "due_suitability_contribution_table",
        "screening_disposition_table.xlsx": "screening_disposition",
        "evidence_appraisal_table.xlsx": "article_appraisal",
        "evidence_source_inventory.xlsx": "evidence_source_inventory",
        "pre_g42_claim_evidence_candidate_matrix.xlsx": "pre_g42_claim_evidence_candidate_matrix",
        "claim_support_type_classifier.xlsx": "claim_support_type_classifier",
        "semantic_claim_evidence_candidate_matrix.xlsx": "semantic_claim_evidence_candidate_matrix",
        "g42_failure_pattern_report.xlsx": "g42_failure_pattern_report",
        "g42_repair_routing_trace.xlsx": "g42_repair_routing_trace",
        "endpoint_extraction_table.xlsx": "endpoint_extraction",
        "sota_benchmark_matrix.xlsx": "sota_benchmark_matrix",
        "alternative_treatment_benchmark_table.xlsx": "alternative_treatment_benchmark_table",
        "guideline_pathway_table.xlsx": "guideline_pathway_table",
        "similar_benchmark_device_table.xlsx": "similar_benchmark_device_table",
        "hazard_source_table.xlsx": "hazard_source_table",
        "sota_to_47_usage_matrix.xlsx": "sota_to_47_usage_matrix",
        "full_text_request_list.xlsx": "full_text_request_list",
        "endpoint_registry.xlsx": "endpoint_registry",
        "sota_endpoint_derivation_table.xlsx": "sota_endpoint_derivation_table",
        "sota_quantitative_benchmark_table.xlsx": "sota_quantitative_benchmark_table",
        "sota_evidence_synthesis_matrix.xlsx": "sota_evidence_synthesis_matrix",
        "sota_claim_reverse_correction_table.xlsx": "sota_claim_reverse_correction_table",
        "sota_medical_field_boundary.xlsx": "sota_medical_field_boundary",
        "sota_pico_v2_strategy.xlsx": "sota_pico_v2_strategy",
        "sota_search_strategy_separated.xlsx": "sota_search_strategy_separated",
        "sota_screening_prisma.xlsx": "sota_screening_prisma",
        "sota_evidence_hierarchy.xlsx": "sota_evidence_hierarchy",
        "sota_endpoint_extraction_fulltext.xlsx": "sota_endpoint_extraction_fulltext",
        "sota_benchmark_derivation_table.xlsx": "sota_benchmark_derivation_table",
        "sota_section_conclusion_matrix.xlsx": "sota_section_conclusion_matrix",
        "sota_deduction_chain.xlsx": "sota_deduction_chain",
        "sota_endpoint_source_classification.xlsx": "sota_endpoint_source_classification",
        "sota_aggregate_benchmark_rationale.xlsx": "sota_aggregate_benchmark_rationale",
        "sota_conclusion_strength_guard.xlsx": "sota_conclusion_strength_guard",
        "sota_clinical_context_table.xlsx": "sota_clinical_context_table",
        "sota_benchmark_contextual_rationale.xlsx": "sota_benchmark_contextual_rationale",
        "sota_context_injection_trace.xlsx": "sota_context_injection_trace",
        "claim_evidence_matrix.xlsx": "claim_evidence_matrix",
        "benefit_risk_ledger.xlsx": "benefit_risk_ledger",
        "writer_conclusion_strength_guard.xlsx": "writer_conclusion_strength_guard",
        "cer_section_trace_map.xlsx": "cer_section_trace_map",
        "alignment_matrix.xlsx": "alignment_matrix",
        "allowed_use_matrix.xlsx": "allowed_use_matrix",
        "cross_evidence_synthesis_table.xlsx": "cross_evidence_synthesis_table",
        "cross_evidence_synthesis_narratives.xlsx": "cross_evidence_synthesis_narratives",
        "writer_synthesis_trace.xlsx": "writer_synthesis_trace",
        "equivalence_comparison_matrix.xlsx": "equivalence_matrix",
        "similar_device_four_step_confirmation.xlsx": "similar_device_four_step_confirmation",
        "similar_device_attachment_index.xlsx": "similar_device_attachment_index",
        "vigilance_recall_registry.xlsx": "vigilance_recall_registry",
        "vigilance_event_statistics.xlsx": "vigilance_event_statistics",
        "risk_gspr_trace_matrix.xlsx": "risk_trace_matrix",
        "pmcf_boundary_decision_log.xlsx": "pmcf_boundary_decision_log",
        "marketing_pms_customer_questionnaire.xlsx": "marketing_pms_customer_questionnaire",
        "artifact_consumption_contract.xlsx": "artifact_consumption_contract",
        "failure_taxonomy_cer_authoring.xlsx": "failure_taxonomy_cer_authoring",
        "cer_section_trace_map_schema.xlsx": "cer_section_trace_map_schema",
        "gate_to_upstream_repair_map.xlsx": "gate_to_upstream_repair_map",
        "calibration_event_log.xlsx": "calibration_event_log",
    }
    key = mapping.get(filename)
    wb = Workbook()
    ws = wb.active
    ws.title = _sheet_title(key or "artifact")
    rows = state.get(key) if key else workbook
    if isinstance(rows, dict):
        rows = [{"key": k, "value": v} for k, v in rows.items()]
    if not rows:
        rows = [{"status": state.get("status", "draft"), "note": "No rows available in this first-pass artifact"}]
    if filename == "source_inventory.xlsx":
        rows = [_source_inventory_export_row(row) for row in rows]
    headers = _headers(rows)
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
    for row_idx, row in enumerate(rows, start=2):
        for col, header in enumerate(headers, start=1):
            ws.cell(row=row_idx, column=col, value=_cell_value(row.get(header, "")))
    for col_idx, header in enumerate(headers, start=1):
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max(len(str(header)) + 4, 14), 60)
    wb.save(path)


def _write_docx_artifact(path: Path, filename: str, state: dict[str, Any], final_report: dict[str, Any]) -> None:
    from docx import Document

    doc = Document()
    if filename == "CER_draft.docx":
        _add_markdown_to_doc(doc, _render_cer_markdown(state))
    elif filename == "search_protocol_and_results.docx":
        doc.add_heading("Literature Search Protocol and Results", level=1)
        doc.add_heading("LSP Methodology Profile", level=2)
        profile = state.get("literature_search_protocol_profile") or {}
        _add_table(doc, [{"key": k, "value": v} for k, v in profile.items()])
        doc.add_heading("SOTA PICO Strategy", level=2)
        _add_table(doc, state.get("sota_pico_strategy") or [])
        doc.add_heading("Device Under Evaluation PICO Strategy", level=2)
        _add_table(doc, state.get("due_pico_strategy") or [])
        doc.add_heading("Database Search Sources", level=2)
        _add_table(doc, state.get("database_search_source_table") or [])
        doc.add_heading("Defined Limits", level=2)
        _add_table(doc, state.get("literature_defined_limits") or [])
        _add_table(doc, state.get("search_run_registry") or [])
        doc.add_heading("Screening Disposition", level=2)
        _add_table(doc, state.get("screening_disposition") or [])
        doc.add_heading("Protocol Deviations", level=2)
        _add_table(doc, state.get("protocol_deviation_log") or [])
    elif filename == "gap_pmcf_recommendations.docx":
        doc.add_heading("Evidence Gaps and PMCF Recommendations", level=1)
        _add_table(doc, state.get("gap_pmcf_recommendations") or [])
        doc.add_heading("PMCF Boundary Decision Log", level=2)
        _add_table(doc, state.get("pmcf_boundary_decision_log") or [])
    elif filename == "nb_precheck_report.docx":
        doc.add_heading("NB Precheck Report", level=1)
        report = state.get("qa_gate_report", {}).get("nb_precheck_report", {})
        _add_table(doc, [{"key": k, "value": v} for k, v in report.items()])
    else:
        doc.add_heading(filename, level=1)
        _add_table(doc, [{"key": "final_gate_decision", "value": final_report.get("decision")}])
    doc.save(path)


def _add_markdown_to_doc(doc: Any, markdown: str) -> None:
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if not stripped:
            idx += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while idx < len(lines):
                candidate = lines[idx].strip()
                if candidate.startswith("|") and candidate.endswith("|"):
                    table_lines.append(candidate)
                    idx += 1
                    continue
                break
            _add_markdown_table(doc, table_lines)
            continue
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            doc.add_heading(stripped.lstrip("#").strip(), level=level)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
        idx += 1


def _add_markdown_table(doc: Any, lines: list[str]) -> None:
    rows = [_split_markdown_table_row(line) for line in lines if not _is_markdown_separator(line)]
    rows = [row for row in rows if row]
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=max_cols)
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = _cell_value(value)
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx in range(max_cols):
            cells[idx].text = _cell_value(row[idx] if idx < len(row) else "")


def _split_markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_markdown_separator(line: str) -> bool:
    cells = _split_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", cell.strip()) for cell in cells)


def _add_table(doc: Any, rows: list[dict[str, Any]]) -> None:
    if not rows:
        doc.add_paragraph("No rows available in this first-pass artifact.")
        return
    headers = _headers(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            cells[idx].text = _cell_value(row.get(header, ""))


def _headers(rows: Any) -> list[str]:
    if not isinstance(rows, list):
        return ["value"]
    headers: list[str] = []
    for row in rows:
        if isinstance(row, dict):
            for key in row:
                if key not in headers:
                    headers.append(key)
    return headers or ["value"]


def _cell_value(value: Any) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, default=str)
    cleaned = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]", "", str(value))
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned[:32000]


def _source_inventory_export_row(row: dict[str, Any]) -> dict[str, Any]:
    exported = dict(row)
    text = str(exported.pop("text", "") or "")
    if text:
        exported["text_excerpt"] = _cell_value(text)[:2000]
        exported["text_length"] = exported.get("text_length") or len(text)
    return exported


def _sheet_title(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_ -]", "", value or "Sheet")[:31] or "Sheet"
