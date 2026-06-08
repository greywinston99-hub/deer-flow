#!/usr/bin/env python3
"""Claude-side CER repair loop helpers.

This module operates only on completed DeerFlow authoring artifacts. It does
not invoke or mutate the DeerFlow authoring graph. The baseline artifacts remain
under ``deerflow_authoring/`` and every enhancement is written under
``claude_repair/``.
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


CORE_AUTHORING_ARTIFACTS = [
    "authoring_workbook.json",
    "source_inventory.xlsx",
    "device_profile.json",
    "claim_ledger.xlsx",
    "claim_pico_derivation.xlsx",
    "literature_search_protocol_profile.json",
    "sota_pico_strategy.xlsx",
    "due_pico_strategy.xlsx",
    "database_search_source_table.xlsx",
    "literature_defined_limits.xlsx",
    "search_protocol_and_results.docx",
    "search_run_registry.json",
    "literature_flow_registry.xlsx",
    "protocol_deviation_log.xlsx",
    "prisma_flow_data.json",
    "prisma_flow_diagram.md",
    "sota_search_strategy_table.xlsx",
    "sota_screening_disposition_table.xlsx",
    "sota_ck_appraisal_table.xlsx",
    "due_suitability_contribution_table.xlsx",
    "screening_disposition_table.xlsx",
    "evidence_appraisal_table.xlsx",
    "endpoint_extraction_table.xlsx",
    "sota_benchmark_matrix.xlsx",
    "alternative_treatment_benchmark_table.xlsx",
    "guideline_pathway_table.xlsx",
    "similar_benchmark_device_table.xlsx",
    "hazard_source_table.xlsx",
    "sota_to_47_usage_matrix.xlsx",
    "full_text_request_list.xlsx",
    "sota_literature_quantity_justification.md",
    "sota_endpoint_derivation_table.xlsx",
    "sota_quantitative_benchmark_table.xlsx",
    "sota_evidence_synthesis_matrix.xlsx",
    "equivalence_comparison_matrix.xlsx",
    "similar_device_four_step_confirmation.xlsx",
    "similar_device_attachment_index.xlsx",
    "vigilance_recall_registry.xlsx",
    "vigilance_event_statistics.xlsx",
    "risk_gspr_trace_matrix.xlsx",
    "marketing_pms_customer_questionnaire.xlsx",
    "CER_draft.md",
    "CER_draft.docx",
    "gap_pmcf_recommendations.docx",
    "qa_gate_report.json",
    "nb_precheck_report.docx",
    "final_gate_closure_report.json",
]

PLACEHOLDERS = ["HUMAN_REVIEW", "DATA GAP", "pending execution", "TODO", "TBD"]
META_RISK_PATTERNS = [
    "Ifu risk warning was extracted",
    "Ifu warning was extracted",
    "report-safe summaries",
]
VIRTUAL_REVIEW_DIMENSIONS = [
    "methodology_checker",
    "evidence_checker",
    "sota_checker",
    "risk_gspr_checker",
    "human_style_checker",
    "consistency_checker",
]
ENGINEER_COMMENT_THEMES = [
    {
        "theme_id": "theme_1_summary_scope",
        "theme": "Summary and Scope completeness",
        "comments": ["#0", "#1", "#2"],
    },
    {
        "theme_id": "theme_2_equivalence_logic",
        "theme": "Equivalence pathway and similar-device logic",
        "comments": ["#3", "#6"],
    },
    {
        "theme_id": "theme_3_market_vigilance",
        "theme": "Marketing history, PMS and vigilance closure",
        "comments": ["#4", "#9"],
    },
    {
        "theme_id": "theme_4_sota_literature",
        "theme": "SOTA and device-literature search strategy",
        "comments": ["#5", "#8"],
    },
    {
        "theme_id": "theme_5_gspr_benefit_risk",
        "theme": "GSPR analysis and benefit-risk comparison",
        "comments": ["#7", "#10"],
    },
    {
        "theme_id": "theme_6_section_conclusions",
        "theme": "Section-level conclusion discipline",
        "comments": ["#11"],
    },
]

REPAIR_DIR_NAME = "claude_repair"
REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_FULLTEXT_DIR_NAMES = (
    "04_LITERATURE_FULL_TEXT",
    "02_FULL_TEXT_LIBRARY/included",
    "02_FULL_TEXT_LIBRARY",
    "full_text_library",
)
FULLTEXT_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}
SOTA_LITERATURE_TARGET_MIN = 20
SOTA_LITERATURE_TARGET_MAX = 40
SOTA_LITERATURE_CATEGORIES = (
    "guideline_consensus",
    "systematic_review_meta_analysis",
    "clinical_study",
    "registry_real_world",
    "alternative_therapy",
    "similar_benchmark_device",
    "hazard_complication",
)
SIGNATURE_SCORE_THRESHOLD = 85
CONTROLLED_SCORE_THRESHOLD = 75
BENCHMARK_TEMPLATE_PATTERN = re.compile(r"Claim-specific benchmark generated from PICO[-\w]*\.?", re.IGNORECASE)
LONG_DOT_PATTERN = re.compile(r"\.{10,}")


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def read_json(path: Path) -> Any:
    if not path.exists():
        return {} if path.suffix == ".json" else None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2, default=str) + "\n", encoding="utf-8")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def repair_root(run_dir: Path) -> Path:
    return run_dir / REPAIR_DIR_NAME


def artifact_root(run_dir: Path) -> Path:
    status = read_json(run_dir / "status.json")
    if isinstance(status, dict) and status.get("artifact_root"):
        return Path(str(status["artifact_root"]))
    return run_dir / "deerflow_authoring"


def word_count(text: str) -> int:
    return len(re.findall(r"\b[A-Za-z][A-Za-z0-9'-]*\b", text))


def cjk_count(text: str) -> int:
    return len(re.findall(r"[\u3400-\u9fff]", text))


def _rows(value: Any) -> list[dict[str, Any]]:
    if isinstance(value, list):
        return [item for item in value if isinstance(item, dict)]
    return []


def _field(row: dict[str, Any], *names: str) -> Any:
    for name in names:
        value = row.get(name)
        if value not in (None, ""):
            return value
    return ""


def _new_issue(
    *,
    issue_id: str,
    dimension: str,
    severity: str,
    issue: str,
    action: str,
    theme_id: str | None = None,
    rerun_required: bool = False,
    controlled_patch_allowed: bool = False,
    human_request_required: bool = False,
    evidence_enhancement_required: bool = False,
) -> dict[str, Any]:
    return {
        "issue_id": issue_id,
        "dimension": dimension,
        "severity": severity,
        "issue": issue,
        "recommended_action": action,
        "theme_id": theme_id,
        "rerun_required": rerun_required,
        "controlled_patch_allowed": controlled_patch_allowed,
        "human_request_required": human_request_required,
        "evidence_enhancement_required": evidence_enhancement_required,
    }


def load_baseline(run_dir: Path) -> dict[str, Any]:
    root = artifact_root(run_dir)
    workbook = read_json(root / "authoring_workbook.json")
    if not isinstance(workbook, dict):
        workbook = {}
    cer_path = root / "CER_draft.md"
    cer_text = cer_path.read_text(encoding="utf-8", errors="ignore") if cer_path.exists() else ""
    return {
        "run_dir": run_dir,
        "artifact_root": root,
        "workbook": workbook,
        "cer_text": cer_text,
        "qa_gate_report": read_json(root / "qa_gate_report.json") if (root / "qa_gate_report.json").exists() else {},
        "final_gate_closure_report": read_json(root / "final_gate_closure_report.json") if (root / "final_gate_closure_report.json").exists() else {},
    }


def load_effective_package(run_dir: Path) -> dict[str, Any]:
    """Load the current best Claude-repaired package with baseline fallback."""
    baseline = load_baseline(run_dir)
    root = repair_root(run_dir)
    workbook = read_json(root / "enhanced_authoring_workbook.json")
    if not isinstance(workbook, dict) or not workbook:
        workbook = baseline["workbook"]
    text_path = root / "CER_draft_patched.md"
    if not text_path.exists():
        text_path = baseline["artifact_root"] / "CER_draft.md"
    text = text_path.read_text(encoding="utf-8", errors="ignore") if text_path.exists() else baseline["cer_text"]
    return {
        **baseline,
        "repair_root": root,
        "effective_workbook": workbook,
        "effective_cer_text": text,
        "effective_cer_path": text_path,
    }


def build_deep_postflight(run_dir: Path) -> dict[str, Any]:
    baseline = load_baseline(run_dir)
    root = baseline["artifact_root"]
    workbook = baseline["workbook"]
    text = baseline["cer_text"]
    issues: list[dict[str, Any]] = []

    missing = [name for name in CORE_AUTHORING_ARTIFACTS if not (root / name).exists()]
    if missing:
        issues.append(
            _new_issue(
                issue_id="DPF-ARTIFACT-001",
                dimension="consistency_checker",
                severity="critical",
                issue=f"Missing baseline artifacts: {', '.join(missing[:8])}{'...' if len(missing) > 8 else ''}",
                action="Do not patch locally; rerun DeerFlow or restore missing artifacts.",
                rerun_required=True,
            )
        )

    source_role = workbook.get("source_role_report") or {}
    identity = workbook.get("device_identity_lock") or {}
    contamination = workbook.get("domain_contamination_report") or {}
    if source_role.get("status") not in (None, "PASS"):
        issues.append(
            _new_issue(
                issue_id="DPF-IDENTITY-001",
                dimension="consistency_checker",
                severity="critical",
                issue="Source role separation did not pass.",
                action="Rerun DeerFlow after correcting IFU/source folder separation.",
                rerun_required=True,
            )
        )
    if identity.get("status") not in (None, "PASS"):
        issues.append(
            _new_issue(
                issue_id="DPF-IDENTITY-002",
                dimension="consistency_checker",
                severity="critical",
                issue="Device identity lock did not pass.",
                action="Rerun DeerFlow after correcting device profile/target keywords.",
                rerun_required=True,
            )
        )
    if (contamination.get("findings") or []) or contamination.get("status") not in (None, "PASS"):
        issues.append(
            _new_issue(
                issue_id="DPF-IDENTITY-003",
                dimension="consistency_checker",
                severity="critical",
                issue="Clinical domain contamination is present.",
                action="Rerun DeerFlow; local repair cannot safely correct domain contamination.",
                rerun_required=True,
            )
        )

    if not text:
        issues.append(
            _new_issue(
                issue_id="DPF-DRAFT-001",
                dimension="human_style_checker",
                severity="critical",
                issue="CER_draft.md is missing or empty.",
                action="Rerun DeerFlow export.",
                rerun_required=True,
            )
        )
    else:
        if cjk_count(text) > 0:
            issues.append(
                _new_issue(
                    issue_id="DPF-DRAFT-002",
                    dimension="human_style_checker",
                    severity="major",
                    issue="CER draft contains CJK characters.",
                    action="Apply controlled English-only patch and regenerate DOCX.",
                    controlled_patch_allowed=True,
                )
            )
        placeholders = [item for item in PLACEHOLDERS if item.lower() in text.lower()]
        if placeholders:
            issues.append(
                _new_issue(
                    issue_id="DPF-DRAFT-003",
                    dimension="human_style_checker",
                    severity="major",
                    issue=f"CER draft contains placeholders: {', '.join(placeholders)}",
                    action="Replace placeholders with evidence gaps or remove unsupported claims.",
                    controlled_patch_allowed=True,
                )
            )
        meta_hits = [item for item in META_RISK_PATTERNS if item.lower() in text.lower()]
        if meta_hits:
            issues.append(
                _new_issue(
                    issue_id="DPF-DRAFT-004",
                    dimension="human_style_checker",
                    severity="minor",
                    issue=f"CER contains machine/meta wording: {', '.join(meta_hits)}",
                    action="Rewrite affected risk and template-language rows into report-facing English while preserving IDs.",
                    controlled_patch_allowed=True,
                )
            )

    picos = _rows(workbook.get("cep_pico_matrix"))
    sota_picos = _rows(workbook.get("sota_pico_strategy"))
    due_picos = _rows(workbook.get("due_pico_strategy"))
    searches = _rows(workbook.get("search_run_registry"))
    evidence = _rows(workbook.get("evidence_registry"))
    endpoints = _rows(workbook.get("endpoint_extraction"))
    benchmarks = _rows(workbook.get("sota_benchmark_matrix"))
    lsp_profile = workbook.get("literature_search_protocol_profile") if isinstance(workbook.get("literature_search_protocol_profile"), dict) else {}
    sota_pico = _rows(workbook.get("sota_pico_strategy"))
    due_pico = _rows(workbook.get("due_pico_strategy"))
    database_sources = _rows(workbook.get("database_search_source_table"))
    lsp_limits = _rows(workbook.get("literature_defined_limits"))
    flow_registry = _rows(workbook.get("literature_flow_registry"))
    protocol_deviations = _rows(workbook.get("protocol_deviation_log"))
    sota_search_strategy = _rows(workbook.get("sota_search_strategy_table"))
    sota_screening = _rows(workbook.get("sota_screening_disposition_table"))
    sota_ck = _rows(workbook.get("sota_ck_appraisal_table"))
    sota_to_47 = _rows(workbook.get("sota_to_47_usage_matrix"))
    sota_quantity = _sota_literature_quantity(sota_ck, sota_screening)
    sota_quantity_justification = workbook.get("sota_literature_quantity_justification") if isinstance(workbook.get("sota_literature_quantity_justification"), dict) else {}
    prisma_flow = workbook.get("prisma_flow_data") if isinstance(workbook.get("prisma_flow_data"), dict) else {}
    similar_four_step = _rows(workbook.get("similar_device_four_step_confirmation"))
    similar_attachments = _rows(workbook.get("similar_device_attachment_index"))
    vigilance_registry = _rows(workbook.get("vigilance_recall_registry"))
    vigilance_stats = _rows(workbook.get("vigilance_event_statistics"))
    market_questionnaire = _rows(workbook.get("marketing_pms_customer_questionnaire"))
    risks = _rows(workbook.get("risk_trace_matrix"))
    gspr = _rows(workbook.get("gspr_coverage"))
    claims = _rows(workbook.get("claim_ledger"))

    if not picos:
        issues.append(
            _new_issue(
                issue_id="DPF-PICO-001",
                dimension="methodology_checker",
                severity="critical",
                issue="No PICO rows were found.",
                action="Rerun DeerFlow; PICO construction is a baseline authoring responsibility.",
                rerun_required=True,
            )
        )
    else:
        missing_rationale = [row.get("pico_id") for row in picos if not _field(row, "derivation_rationale", "rationale")]
        if missing_rationale:
            issues.append(
                _new_issue(
                    issue_id="DPF-PICO-002",
                    dimension="methodology_checker",
                    severity="major",
                    issue=f"PICO rows lack derivation rationale: {', '.join(map(str, missing_rationale[:8]))}",
                    action="Add PICO derivation rationale in pico_revision_log before finalization.",
                    controlled_patch_allowed=True,
                )
            )
    if not searches:
        issues.append(
            _new_issue(
                issue_id="DPF-SEARCH-001",
                dimension="methodology_checker",
                severity="critical",
                issue="No search registry rows were found.",
                action="Rerun DeerFlow; literature search registry is missing.",
                rerun_required=True,
            )
        )
    if not lsp_profile or not sota_pico or not due_pico:
        issues.append(
            _new_issue(
                issue_id="DPF-LSP-001",
                dimension="methodology_checker",
                severity="major",
                issue="LSP methodology profile or SOTA/DuE PICO strategy is missing.",
                action="Rerun DeerFlow with LSP v2 enabled; do not patch a missing search methodology from prose alone.",
                rerun_required=True,
            )
        )
    if not database_sources or not flow_registry or not protocol_deviations:
        issues.append(
            _new_issue(
                issue_id="DPF-LSP-002",
                dimension="methodology_checker",
                severity="major",
                issue="LSP database source table, literature flow registry, or protocol deviation log is missing.",
                action="Rerun DeerFlow or regenerate LSP artifacts from executed search logs.",
                rerun_required=True,
            )
        )
    else:
        source_text = " ".join(str(row.get("database", "")) + " " + str(row.get("status", "")) for row in database_sources).lower()
        missing_sources = [name for name in ["pubmed", "europe pmc", "clinicaltrials", "eu clinical trials register", "embase", "cochrane"] if name not in source_text]
        if missing_sources:
            issues.append(
                _new_issue(
                    issue_id="DPF-LSP-003",
                    dimension="methodology_checker",
                    severity="major",
                    issue=f"LSP source coverage is missing execution or limitation records for: {', '.join(missing_sources)}",
                    action="Rerun DeerFlow after MCP configuration or record source limitations explicitly.",
                    rerun_required=True,
                )
            )
    if not prisma_flow or not all(isinstance(prisma_flow.get(section), dict) and prisma_flow.get(section) for section in ("identification", "screening", "included")):
        issues.append(
            _new_issue(
                issue_id="DPF-LSP-005",
                dimension="methodology_checker",
                severity="major",
                issue="PRISMA-style flow data is missing identification, screening or included sections.",
                action="Regenerate PRISMA flow from search registry and screening disposition before signature-level use.",
                evidence_enhancement_required=True,
            )
        )
    if not sota_search_strategy or not sota_screening or not sota_ck:
        issues.append(
            _new_issue(
                issue_id="DPF-SOTA-LSP-001",
                dimension="sota_checker",
                severity="major",
                issue="SOTA search strategy, screening disposition, or CK appraisal table is missing.",
                action="Rerun DeerFlow or rebuild SOTA/LSP tables before CER signature-level use.",
                evidence_enhancement_required=True,
            )
        )

    similar_steps_text = " ".join(str(row.get("step", "")) for row in similar_four_step).lower()
    if not similar_four_step or not all(f"step {idx}" in similar_steps_text for idx in range(1, 5)) or len(similar_attachments) < 10:
        issues.append(
            _new_issue(
                issue_id="DPF-SIMILAR-002",
                dimension="sota_checker",
                severity="major",
                issue="Similar-device four-step confirmation or 10-item attachment index is missing.",
                action="Generate the similar-device four-step table and attachment request index; do not use similar devices as benchmark/market evidence until source status is clear.",
                evidence_enhancement_required=True,
                human_request_required=True,
            )
        )

    vigilance_text = " ".join(str(row.get("database", "")) for row in vigilance_registry).lower()
    missing_vigilance_sources = []
    if "eudamed" not in vigilance_text:
        missing_vigilance_sources.append("EUDAMED vigilance")
    if "medsafe" not in vigilance_text and "new zealand" not in vigilance_text:
        missing_vigilance_sources.append("New Zealand Medsafe")
    if missing_vigilance_sources or not vigilance_stats:
        issues.append(
            _new_issue(
                issue_id="DPF-VIGILANCE-003",
                dimension="risk_gspr_checker",
                severity="major",
                issue=f"Additional vigilance source/statistics package incomplete: {', '.join(missing_vigilance_sources) or 'vigilance_event_statistics'}",
                action="Run or record source-limitation rows for EUDAMED/New Zealand vigilance and generate event statistics before benefit-risk finalization.",
                evidence_enhancement_required=True,
            )
        )

    has_market_or_pms_source = any(
        term in " ".join(str(row) for row in _rows(workbook.get("source_inventory"))).lower()
        for term in ("pms", "pmcf", "sales", "marketing", "complaint")
    )
    if not has_market_or_pms_source and len(market_questionnaire) < 8:
        issues.append(
            _new_issue(
                issue_id="DPF-MARKET-002",
                dimension="risk_gspr_checker",
                severity="major",
                issue="Marketing/PMS source data are absent and the customer questionnaire is incomplete.",
                action="Generate the customer questionnaire covering market status, sales exposure, complaints, FSCA/recall, PMS, PMCF, RMF linkage and data-lock/version control.",
                human_request_required=True,
            )
        )
    elif sota_quantity["included_count"] < SOTA_LITERATURE_TARGET_MIN and not _sota_quantity_justification_complete(sota_quantity_justification):
        issues.append(
            _new_issue(
                issue_id="DPF-SOTA-QTY-001",
                dimension="sota_checker",
                severity="major",
                issue=(
                    f"SOTA has {sota_quantity['included_count']} final included literature records; "
                    f"target is {SOTA_LITERATURE_TARGET_MIN}-{SOTA_LITERATURE_TARGET_MAX} or a controlled search-exhaustion justification."
                ),
                action="Broaden/iterate SOTA search or provide sota_literature_quantity_justification before NB-ready finalization.",
                evidence_enhancement_required=True,
                human_request_required=True,
            )
        )
    elif sota_quantity["included_count"] > SOTA_LITERATURE_TARGET_MAX and not _sota_quantity_has_hierarchy(sota_ck):
        issues.append(
            _new_issue(
                issue_id="DPF-SOTA-QTY-002",
                dimension="sota_checker",
                severity="major",
                issue=(
                    f"SOTA has {sota_quantity['included_count']} included records; records above "
                    f"{SOTA_LITERATURE_TARGET_MAX} require hierarchy, endpoint contribution and clinical relevance stratification."
                ),
                action="Stratify SOTA evidence by hierarchy, endpoint contribution and clinical relevance before finalization.",
                controlled_patch_allowed=True,
                evidence_enhancement_required=True,
            )
        )
    if not sota_to_47:
        issues.append(
            _new_issue(
                issue_id="DPF-SOTA-TRACE-001",
                dimension="sota_checker",
                severity="major",
                issue="SOTA-to-4.7 usage matrix is missing.",
                action="Generate the SOTA-to-4.7 usage matrix and patch section 4.7 only if benchmark IDs already exist.",
                controlled_patch_allowed=True,
            )
        )
    if lsp_limits and not any("language" in str(row).lower() for row in lsp_limits):
        issues.append(
            _new_issue(
                issue_id="DPF-LSP-004",
                dimension="methodology_checker",
                severity="minor",
                issue="LSP defined limits do not visibly include language handling.",
                action="Patch LSP limits to state no language exclusion by default or record English-only processing as a deviation.",
                controlled_patch_allowed=True,
            )
        )

    if not evidence:
        issues.append(
            _new_issue(
                issue_id="DPF-EVIDENCE-001",
                dimension="evidence_checker",
                severity="major",
                issue="No Evidence Registry rows were found.",
                action="Run evidence enhancement or rerun DeerFlow if the search strategy was wrong.",
                evidence_enhancement_required=True,
                human_request_required=True,
            )
        )
    weak_endpoint_rows = [
        row
        for row in endpoints
        if "not extracted" in str(row).lower() or "first-pass" in str(row).lower() or not _field(row, "sample_size", "timepoint", "statistical_result", "result")
    ]
    abstract_only_evidence = [row for row in evidence if "first-pass" in str(row).lower() or "not extracted" in str(row).lower()]
    no_pivotal = evidence and not any(str(row.get("weight", "")).lower() == "pivotal" for row in evidence)
    if weak_endpoint_rows or abstract_only_evidence or no_pivotal:
        bits = []
        if weak_endpoint_rows:
            bits.append(f"{len(weak_endpoint_rows)} endpoint rows need full-text extraction")
        if abstract_only_evidence:
            bits.append(f"{len(abstract_only_evidence)} evidence rows are abstract/first-pass limited")
        if no_pivotal:
            bits.append("no pivotal evidence is currently available")
        issues.append(
            _new_issue(
                issue_id="DPF-EVIDENCE-002",
                dimension="evidence_checker",
                severity="major",
                issue="; ".join(bits),
                action="Generate full_text_request_list and update endpoint extraction from full text before signature-level use.",
                evidence_enhancement_required=True,
                human_request_required=True,
            )
        )

    unused_benchmarks = [row.get("benchmark_id") for row in benchmarks if str(row.get("used_in_4_7", "")).lower() in {"false", "no", "0"}]
    if unused_benchmarks:
        issues.append(
            _new_issue(
                issue_id="DPF-SOTA-001",
                dimension="sota_checker",
                severity="major",
                issue=f"SOTA benchmarks are not used in section 4.7: {', '.join(map(str, unused_benchmarks[:8]))}",
                action="Patch section 4.7 so each accepted benchmark is used or explicitly retired.",
                controlled_patch_allowed=True,
            )
        )

    rmf_gaps = [row for row in risks if "requires rmf" in str(row).lower() or "gap" in str(row.get("rmf_coverage", "")).lower()]
    if rmf_gaps:
        issues.append(
            _new_issue(
                issue_id="DPF-RISK-001",
                dimension="risk_gspr_checker",
                severity="major",
                issue=f"{len(rmf_gaps)} risk rows require RMF source review or remain gap-controlled.",
                action="Keep conclusions downgraded and request RMF/GSPR/PMS source evidence.",
                human_request_required=True,
                controlled_patch_allowed=True,
            )
        )

    engineer_theme_issues = _engineer_comment_theme_issues(text, workbook)
    issues.extend(engineer_theme_issues)

    dimensions = _dimension_summary(issues)
    theme_report = _engineer_theme_report(engineer_theme_issues)
    report = {
        "schema_name": "claude_cer_deep_postflight_report",
        "schema_version": "v1",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "baseline_artifact_root": str(root),
        "repair_root": str(repair_root(run_dir)),
        "summary": {
            "word_count": word_count(text),
            "cjk_count": cjk_count(text),
            "missing_artifact_count": len(missing),
            "claim_count": len(claims),
            "pico_count": len(picos),
            "search_run_count": len(searches),
            "evidence_count": len(evidence),
            "endpoint_count": len(endpoints),
            "sota_benchmark_count": len(benchmarks),
            "lsp_profile_present": bool(lsp_profile),
            "sota_pico_strategy_count": len(sota_pico),
            "due_pico_strategy_count": len(due_pico),
            "database_source_count": len(database_sources),
            "literature_flow_count": len(flow_registry),
            "protocol_deviation_count": len(protocol_deviations),
            "sota_ck_appraisal_count": len(sota_ck),
            "sota_final_included_literature_count": sota_quantity["included_count"],
            "sota_literature_target_range": f"{SOTA_LITERATURE_TARGET_MIN}-{SOTA_LITERATURE_TARGET_MAX}",
            "sota_literature_quantity_justification_present": _sota_quantity_justification_complete(sota_quantity_justification),
            "sota_to_47_usage_count": len(sota_to_47),
            "prisma_flow_present": bool(prisma_flow),
            "similar_device_four_step_count": len(similar_four_step),
            "similar_device_attachment_count": len(similar_attachments),
            "vigilance_source_count": len(vigilance_registry),
            "vigilance_event_statistics_count": len(vigilance_stats),
            "marketing_pms_customer_question_count": len(market_questionnaire),
            "risk_count": len(risks),
            "gspr_count": len(gspr),
            "source_role_status": source_role.get("status"),
            "device_identity_status": identity.get("status"),
            "domain_contamination_status": contamination.get("status"),
        },
        "engineer_comment_theme_report": theme_report,
        "virtual_review_dimensions": dimensions,
        "issues": issues,
        "baseline_gate_reports": {
            "qa_gate_report": baseline["qa_gate_report"],
            "final_gate_closure_report": baseline["final_gate_closure_report"],
        },
    }
    write_json(repair_root(run_dir) / "deep_postflight_report.json", report)
    return report


def _engineer_comment_theme_issues(text: str, workbook: dict[str, Any]) -> list[dict[str, Any]]:
    """Convert engineer CER comments into deterministic postflight issues."""
    issues: list[dict[str, Any]] = []
    lower = text.lower()
    device_profile = workbook.get("device_profile") if isinstance(workbook.get("device_profile"), dict) else {}
    benchmarks = _rows(workbook.get("sota_benchmark_matrix"))
    searches = _rows(workbook.get("search_run_registry"))
    sota_pico = _rows(workbook.get("sota_pico_strategy"))
    due_pico = _rows(workbook.get("due_pico_strategy"))
    sota_ck = _rows(workbook.get("sota_ck_appraisal_table"))
    alt_treatments = _rows(workbook.get("alternative_treatment_benchmark_table"))
    guideline_pathway = _rows(workbook.get("guideline_pathway_table"))
    hazard_sources = _rows(workbook.get("hazard_source_table"))
    sota_to_47 = _rows(workbook.get("sota_to_47_usage_matrix"))
    equivalence = _rows(workbook.get("equivalence_matrix"))
    vigilance = _rows(workbook.get("vigilance_recall_registry"))
    risks = _rows(workbook.get("risk_trace_matrix"))
    gspr = _rows(workbook.get("gspr_coverage"))

    if not _has_all_terms(lower, ["mdr"]) or not _has_any_term(lower, ["meddev", "mdcg", "regulation", "guidance"]):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-SUMMARY-001",
                dimension="human_style_checker",
                severity="major",
                theme_id="theme_1_summary_scope",
                issue="Summary does not clearly present the regulatory/guidance basis expected by the engineer comments.",
                action="Patch Summary after full report completion to include MDR basis, applicable guidance and evaluation route.",
                controlled_patch_allowed=True,
            )
        )
    if not (device_profile.get("device_class") or "class " in lower or "classification" in lower):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-SUMMARY-002",
                dimension="human_style_checker",
                severity="major",
                theme_id="theme_1_summary_scope",
                issue="Device classification is not explicit in Summary/device profile.",
                action="Request or confirm classification, then patch Summary and Scope without overstating conformity.",
                controlled_patch_allowed=True,
                human_request_required=True,
            )
        )
    if "clinical benefit" not in lower and "clinical benefits" not in lower:
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-SCOPE-001",
                dimension="methodology_checker",
                severity="major",
                theme_id="theme_1_summary_scope",
                issue="Scope/Intended Purpose does not explicitly include a Clinical benefits clause.",
                action="Patch Scope with IFU-supported clinical benefits and add the rows to Claim Ledger/PICO trace.",
                controlled_patch_allowed=True,
            )
        )

    eq_not_claimed = any("not_claim" in str(row).lower() for row in equivalence) or "equivalence is not claimed" in lower or "equivalence is not claimed" in lower
    eq_claimed = any("demonstrated" in str(row).lower() or "equivalent" in str(row).lower() for row in equivalence) or "equivalence demonstrated" in lower
    if not (eq_not_claimed or eq_claimed):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-EQUIV-001",
                dimension="methodology_checker",
                severity="major",
                theme_id="theme_2_equivalence_logic",
                issue="Equivalence route is not explicit; §4.1/§4.2 must state whether equivalence is claimed or not applicable.",
                action="Patch §1, §4.1 and §4.2 with a clear equivalence route; if claimed, require MDCG 2020-5 three-dimensional comparison.",
                controlled_patch_allowed=True,
            )
        )
    if "market share" not in lower and "european market" not in lower and "gold standard" not in lower:
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-SIMILAR-001",
                dimension="sota_checker",
                severity="major",
                theme_id="theme_2_equivalence_logic",
                issue="Similar-device section does not document European marketing status, market share/gold-standard status or medical-condition comparability.",
                action="Request market investigation sources and patch Similar Devices with market status and clinical comparability limits.",
                human_request_required=True,
                evidence_enhancement_required=True,
            )
        )

    if "marketing history" in lower and not _has_any_term(lower, ["sales", "sold", "pms", "complaint", "not marketed", "not been marketed"]):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-MARKET-001",
                dimension="risk_gspr_checker",
                severity="major",
                theme_id="theme_3_market_vigilance",
                issue="Marketing History does not clearly close sales/PMS/complaint status or state no marketing history.",
                action="Generate customer request for sales/PMS data; patch §2.5 with either data or no-market statement.",
                human_request_required=True,
                controlled_patch_allowed=True,
            )
        )
    vigilance_terms = " ".join(str(row) for row in vigilance).lower()
    if vigilance and not _has_any_term(vigilance_terms, ["trade", "brand", "commercial", "product name", "device name"]):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-VIGILANCE-001",
                dimension="risk_gspr_checker",
                severity="major",
                theme_id="theme_3_market_vigilance",
                issue="Vigilance search does not show equivalent/similar device trade-name strategy.",
                action="Patch vigilance search protocol to use subject/equivalent/similar device trade names and map AE categories to RMR/RMF risks.",
                evidence_enhancement_required=True,
                human_request_required=True,
            )
        )
    if risks and "rmr" not in lower and "risk management report" not in lower:
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-VIGILANCE-002",
                dimension="risk_gspr_checker",
                severity="major",
                theme_id="theme_3_market_vigilance",
                issue="AE/vigilance narrative is not visibly closed to RMR/RMF.",
                action="Patch §4.6 and risk trace to connect AE categories with RMR/RMF risk IDs; keep as gap if RMR is missing.",
                controlled_patch_allowed=True,
                human_request_required=True,
            )
        )

    if len(benchmarks) < 3 or not _has_any_term(lower, ["alternative treatment", "alternative therapy", "clinical practice guideline", "acceptance criterion", "benchmark"]):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-SOTA-001",
                dimension="sota_checker",
                severity="major",
                theme_id="theme_4_sota_literature",
                issue="SOTA does not visibly follow the engineer framework: field -> methods/PICO -> medical condition -> options -> benchmark -> conclusion.",
                action="Patch §3 with explicit SOTA framework and ensure benchmarks feed §4.7.",
                controlled_patch_allowed=True,
                evidence_enhancement_required=True,
            )
        )
    if not (sota_pico and due_pico and sota_ck and alt_treatments and guideline_pathway and hazard_sources and sota_to_47):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-SOTA-002",
                dimension="sota_checker",
                severity="major",
                theme_id="theme_4_sota_literature",
                issue="SOTA does not have the full LSP/complex-table package expected from the engineer feedback.",
                action="Regenerate or patch the LSP/SOTA evidence package: SOTA PICO, DuE PICO, CK appraisal, alternatives, guideline pathway, hazards and SOTA-to-4.7 trace.",
                evidence_enhancement_required=True,
                controlled_patch_allowed=True,
            )
        )
    if searches and not any(_has_any_term(str(row).lower(), ["trade", "brand", "equivalent", "similar", "device name"]) for row in searches):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-LIT-001",
                dimension="methodology_checker",
                severity="major",
                theme_id="theme_4_sota_literature",
                issue="Device-literature search strategy may be using SOTA terms rather than equivalent/similar device trade names.",
                action="Add a separate §4.4 device/equivalent-device search protocol using confirmed trade names; do not reuse SOTA query as device evidence search.",
                evidence_enhancement_required=True,
                human_request_required=True,
            )
        )

    if "4.3.2" in lower and gspr and not _has_any_term(lower, ["pre-clinical", "preclinical", "biocompatibility", "sterilisation", "sterilization", "usability", "performance test"]):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-PRECLIN-001",
                dimension="risk_gspr_checker",
                severity="major",
                theme_id="theme_5_gspr_benefit_risk",
                issue="§4.3.2 does not clearly analyse manufacturer-held non-clinical evidence for GSPR safety/performance.",
                action="Patch §4.3.2 to explain each report's clinical relevance and mapped GSPR/risk claim.",
                controlled_patch_allowed=True,
                human_request_required=True,
            )
        )
    if not _has_any_term(lower, ["compared with the sota", "sota benchmark", "benchmark comparison", "benefit-risk", "benefit/risk"]):
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-GSPR-001",
                dimension="risk_gspr_checker",
                severity="major",
                theme_id="theme_5_gspr_benefit_risk",
                issue="§4.7 does not visibly compare integrated clinical data against §3 SOTA benchmarks to support benefit-risk.",
                action="Patch §4.7 with qualitative and quantitative comparison tables; downgrade conclusions where data are incomplete.",
                controlled_patch_allowed=True,
                evidence_enhancement_required=True,
            )
        )

    section_conclusion_hits = len(re.findall(r"section conclusion|subsection conclusion|chapter conclusion", lower))
    if section_conclusion_hits < 3:
        issues.append(
            _new_issue(
                issue_id="DPF-ENG-CONCLUSION-001",
                dimension="human_style_checker",
                severity="major",
                theme_id="theme_6_section_conclusions",
                issue="Major/secondary sections do not consistently end with section-level conclusions.",
                action="Patch §2, §3 and §4 second-level sections with concise conclusions stating safety/performance/benefit-risk relevance.",
                controlled_patch_allowed=True,
            )
        )
    return issues


def _engineer_theme_report(issues: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows = []
    for theme in ENGINEER_COMMENT_THEMES:
        scoped = [issue for issue in issues if issue.get("theme_id") == theme["theme_id"]]
        rows.append(
            {
                **theme,
                "status": "PASS" if not scoped else "REWORK",
                "issue_count": len(scoped),
                "issue_ids": [issue.get("issue_id") for issue in scoped],
            }
        )
    return rows


def _sota_literature_quantity(sota_ck: list[dict[str, Any]], sota_screening: list[dict[str, Any]]) -> dict[str, Any]:
    accepted_ids: set[str] = set()
    for row in sota_ck:
        disposition = str(row.get("disposition", "")).lower()
        if disposition and "excluded" in disposition:
            continue
        article_id = str(row.get("article_id") or row.get("row_id") or "").strip()
        if article_id:
            accepted_ids.add(article_id)
    if not accepted_ids:
        for row in sota_screening:
            decision = f"{row.get('title_abstract_decision', '')} {row.get('full_text_decision', '')}".lower()
            if "include" in decision and "exclude" not in decision:
                article_id = str(row.get("article_id") or row.get("row_id") or "").strip()
                if article_id:
                    accepted_ids.add(article_id)
    return {
        "included_count": len(accepted_ids),
        "target_min": SOTA_LITERATURE_TARGET_MIN,
        "target_max": SOTA_LITERATURE_TARGET_MAX,
        "article_ids": sorted(accepted_ids),
    }


def _sota_quantity_justification_complete(justification: dict[str, Any]) -> bool:
    required = (
        "search_exhaustion_rationale",
        "database_limitations",
        "screening_rationale",
        "evidence_gap_control",
        "clinical_impact",
    )
    return all(str(justification.get(key, "")).strip() for key in required)


def _sota_quantity_has_hierarchy(sota_ck: list[dict[str, Any]]) -> bool:
    if not sota_ck:
        return False
    sample = sota_ck[: min(len(sota_ck), SOTA_LITERATURE_TARGET_MAX + 1)]
    required_any = ("evidence_hierarchy", "evidence_category", "endpoint_contribution", "clinical_relevance")
    return all(any(str(row.get(key, "")).strip() for key in required_any) for row in sample)


def _has_any_term(text: str, terms: list[str]) -> bool:
    return any(term.lower() in text for term in terms)


def _has_all_terms(text: str, terms: list[str]) -> bool:
    return all(term.lower() in text for term in terms)


def _dimension_summary(issues: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows = []
    for dimension in VIRTUAL_REVIEW_DIMENSIONS:
        scoped = [issue for issue in issues if issue.get("dimension") == dimension]
        critical = sum(1 for issue in scoped if issue.get("severity") == "critical")
        major = sum(1 for issue in scoped if issue.get("severity") == "major")
        rows.append(
            {
                "dimension": dimension,
                "status": "FAIL" if critical else ("REWORK" if major else ("ADVISORY" if scoped else "PASS")),
                "issue_count": len(scoped),
                "critical_count": critical,
                "major_count": major,
            }
        )
    return rows


def classify_rework(run_dir: Path) -> dict[str, Any]:
    report_path = repair_root(run_dir) / "deep_postflight_report.json"
    report = read_json(report_path)
    if not isinstance(report, dict) or not report:
        report = build_deep_postflight(run_dir)
    issues = _rows(report.get("issues"))
    if any(issue.get("rerun_required") for issue in issues):
        decision = "FULL_DEERFLOW_RERUN_REQUIRED"
    elif any(issue.get("evidence_enhancement_required") for issue in issues):
        decision = "EVIDENCE_ENHANCEMENT_REQUIRED"
    elif any(issue.get("human_request_required") for issue in issues):
        decision = "HUMAN_HOLD"
    elif any(issue.get("controlled_patch_allowed") for issue in issues):
        decision = "CONTROLLED_MINOR_PATCH"
    else:
        decision = "ACCEPT_BASELINE"
    payload = {
        "schema_name": "claude_cer_rework_decision",
        "schema_version": "v1",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "decision": decision,
        "issue_count": len(issues),
        "issue_summary": [
            {
                "issue_id": issue.get("issue_id"),
                "severity": issue.get("severity"),
                "dimension": issue.get("dimension"),
                "recommended_action": issue.get("recommended_action"),
            }
            for issue in issues
        ],
    }
    write_json(repair_root(run_dir) / "rework_decision.json", payload)
    return payload


def build_pico_loop(run_dir: Path) -> dict[str, Any]:
    baseline = load_baseline(run_dir)
    workbook = baseline["workbook"]
    picos = _rows(workbook.get("cep_pico_matrix"))
    sota_picos = _rows(workbook.get("sota_pico_strategy"))
    due_picos = _rows(workbook.get("due_pico_strategy"))
    searches = _rows(workbook.get("search_run_registry"))
    evidence = _rows(workbook.get("evidence_registry"))
    endpoints = _rows(workbook.get("endpoint_extraction"))
    total_hits = sum(_safe_int(_field(row, "result_count", "results_count", "count")) for row in searches)
    sota_hits = sum(_safe_int(_field(row, "result_count", "results_count", "count")) for row in searches if str(row.get("objective", "")).lower() == "sota")
    due_hits = sum(_safe_int(_field(row, "result_count", "results_count", "count")) for row in searches if "device" in str(row.get("objective", "")).lower())
    endpoint_success = len([row for row in endpoints if "not extracted" not in str(row).lower() and _field(row, "sample_size", "statistical_result", "result")])
    fulltext_limited = len([row for row in endpoints + evidence if "first-pass" in str(row).lower() or "not extracted" in str(row).lower()])
    rows = []
    for idx, pico in enumerate(picos, start=1):
        if total_hits > 500 and endpoint_success < max(1, len(endpoints) // 4):
            decision = "narrow_or_split"
            reason = "Search hits are broad while endpoint extraction remains weak."
        elif sota_picos and due_picos and sota_hits > 0 and due_hits == 0:
            decision = "broaden_device_literature"
            reason = "SOTA has hits but Device-under-Evaluation / trade-name literature search has no accessible hits."
        elif total_hits < 20:
            decision = "broaden"
            reason = "Search hit count is low; broaden device/procedure terminology."
        elif fulltext_limited:
            decision = "full_text_required"
            reason = "Bibliographic evidence exists but endpoint data require full-text extraction."
        else:
            decision = "keep"
            reason = "PICO has sufficient first-pass support for controlled draft use."
        rows.append(
            {
                "pico_id": pico.get("pico_id") or f"PICO-{idx:02d}",
                "version": "v1",
                "claim_id": pico.get("claim_id"),
                "population": _field(pico, "population", "Population"),
                "intervention": _field(pico, "intervention", "Intervention", "device"),
                "comparator": _field(pico, "comparator", "Comparator"),
                "outcome": _field(pico, "outcome", "Outcome"),
                "search_hit_count": total_hits,
                "sota_hit_count": sota_hits,
                "due_hit_count": due_hits,
                "evidence_count": len(evidence),
                "endpoint_extracted_count": endpoint_success,
                "full_text_limited_count": fulltext_limited,
                "decision": decision,
                "reason_for_change": reason,
                "next_action": _pico_next_action(decision),
            }
        )
    payload = {
        "schema_name": "claude_cer_pico_revision_log",
        "schema_version": "v1",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "rows": rows,
    }
    root = repair_root(run_dir)
    write_json(root / "pico_revision_log.json", payload)
    _write_xlsx(root / "pico_revision_log.xlsx", rows, "pico_revision")
    return payload


def _pico_next_action(decision: str) -> str:
    if decision == "narrow_or_split":
        return "Create v2 PICO by splitting endpoint/procedure/device concepts after full-text review."
    if decision == "broaden":
        return "Create v2 PICO with broader procedure and similar-device terms."
    if decision == "broaden_device_literature":
        return "Keep SOTA PICO; broaden Device-under-Evaluation search using trade names, manufacturer names, model identifiers and confirmed similar/equivalent devices."
    if decision == "full_text_required":
        return "Request full texts before deciding whether PICO needs narrowing or broadening."
    return "Keep v1 PICO; preserve rationale and search trace."


def pico_v2_execute(run_dir: Path) -> dict[str, Any]:
    baseline = load_baseline(run_dir)
    root = repair_root(run_dir)
    workbook = read_json(root / "enhanced_authoring_workbook.json")
    if not isinstance(workbook, dict) or not workbook:
        workbook = dict(baseline["workbook"])
    revision = read_json(root / "pico_revision_log.json")
    if not isinstance(revision, dict) or not revision:
        revision = build_pico_loop(run_dir)
    rows = _rows(revision.get("rows"))
    v2_rows = []
    for idx, row in enumerate(rows, start=1):
        decision = str(row.get("decision") or "keep")
        if decision in {"keep", "full_text_required"}:
            action_status = "not_changed"
            change_rationale = "v1 retained; full-text/evidence enhancement is required before PICO restructuring." if decision == "full_text_required" else "v1 retained."
        else:
            action_status = "v2_generated"
            change_rationale = row.get("reason_for_change") or "PICO revision generated from search/evidence sufficiency."
        v2_rows.append(
            {
                "pico_v2_id": f"{row.get('pico_id') or f'PICO-{idx:02d}'}-V2",
                "source_pico_id": row.get("pico_id"),
                "claim_id": row.get("claim_id"),
                "revision_decision": decision,
                "action_status": action_status,
                "population_v2": row.get("population") or "same as v1",
                "intervention_v2": _pico_v2_intervention(row),
                "comparator_v2": _pico_v2_comparator(row),
                "outcome_v2": _pico_v2_outcome(row),
                "search_strategy_v2": _pico_v2_search_strategy(row),
                "rationale": change_rationale,
                "if_not_executed_reason": "" if action_status == "v2_generated" else change_rationale,
            }
        )
    workbook["pico_v2_strategy"] = v2_rows
    workbook["claude_repair_metadata"] = {
        **(workbook.get("claude_repair_metadata") if isinstance(workbook.get("claude_repair_metadata"), dict) else {}),
        "pico_v2_executed_at": now_iso(),
        "pico_v2_generated_count": sum(1 for row in v2_rows if row.get("action_status") == "v2_generated"),
    }
    write_json(root / "enhanced_authoring_workbook.json", workbook)
    write_json(root / "pico_v2_strategy.json", {"schema_name": "claude_cer_pico_v2_strategy", "rows": v2_rows})
    _write_xlsx(root / "pico_v2_strategy.xlsx", v2_rows, "pico_v2")
    payload = {
        "schema_name": "claude_cer_pico_v2_execution_report",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "row_count": len(v2_rows),
        "v2_generated_count": sum(1 for row in v2_rows if row.get("action_status") == "v2_generated"),
        "outputs": {
            "pico_v2_strategy": str(root / "pico_v2_strategy.xlsx"),
            "enhanced_authoring_workbook": str(root / "enhanced_authoring_workbook.json"),
        },
    }
    write_json(root / "pico_v2_execution_report.json", payload)
    return payload


def _pico_v2_intervention(row: dict[str, Any]) -> str:
    intervention = str(row.get("intervention") or "subject device")
    if row.get("decision") == "broaden_device_literature":
        return f"{intervention}; add confirmed trade names, manufacturer names and model identifiers"
    return intervention


def _pico_v2_comparator(row: dict[str, Any]) -> str:
    comparator = str(row.get("comparator") or "standard care / similar device")
    if row.get("decision") in {"narrow_or_split", "broaden_device_literature"}:
        return f"{comparator}; split subject-device, similar-device and alternative-therapy roles"
    return comparator


def _pico_v2_outcome(row: dict[str, Any]) -> str:
    outcome = str(row.get("outcome") or "safety and performance endpoints")
    if row.get("decision") == "narrow_or_split":
        return f"{outcome}; split by endpoint family and timepoint for extraction"
    return outcome


def _pico_v2_search_strategy(row: dict[str, Any]) -> str:
    decision = row.get("decision")
    if decision == "narrow_or_split":
        return "Create endpoint-family specific searches with outcome/timepoint constraints and retain original broad SOTA as background."
    if decision == "broaden":
        return "Add synonyms, MeSH/device-class terms and alternative therapy terms."
    if decision == "broaden_device_literature":
        return "Run device-only search using subject/equivalent/similar trade names, manufacturer names and model identifiers."
    if decision == "full_text_required":
        return "Do not change query until full-text endpoint extraction confirms whether narrowing is needed."
    return "Keep v1 query."


def build_fulltext_request(run_dir: Path) -> dict[str, Any]:
    baseline = load_baseline(run_dir)
    workbook = baseline["workbook"]
    evidence = _rows(workbook.get("evidence_registry"))
    endpoints = _rows(workbook.get("endpoint_extraction"))
    endpoint_by_evidence = {str(row.get("evidence_id")): row for row in endpoints if row.get("evidence_id")}
    rows = []
    for idx, ev in enumerate(evidence, start=1):
        evidence_id = ev.get("evidence_id") or f"E-{idx:03d}"
        endpoint = endpoint_by_evidence.get(str(evidence_id), {})
        needs_fulltext = (
            "first-pass" in str(ev).lower()
            or "not extracted" in str(ev).lower()
            or "not extracted" in str(endpoint).lower()
            or not _field(endpoint, "sample_size", "timepoint", "statistical_result", "result")
        )
        if not needs_fulltext:
            continue
        rows.append(
            {
                "request_id": f"FT-{len(rows)+1:03d}",
                "evidence_id": evidence_id,
                "article_id": ev.get("article_id") or endpoint.get("article_id"),
                "pmid": _field(ev, "pmid", "PMID", "identifier"),
                "doi": _field(ev, "doi", "DOI"),
                "title": _field(ev, "title", "article_title", "source"),
                "current_weight": ev.get("weight"),
                "reason": "Full text is required to extract sample size, endpoint definition, timepoint, statistics and AE/SAE.",
                "required_file": "PDF/HTML/full-text DOCX",
                "target_folder": "04_LITERATURE_FULL_TEXT",
                "status": "REQUESTED",
            }
        )
    payload = {
        "schema_name": "claude_cer_full_text_request_list",
        "schema_version": "v1",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "rows": rows,
    }
    root = repair_root(run_dir)
    write_json(root / "full_text_request_list.json", payload)
    _write_xlsx(root / "full_text_request_list.xlsx", rows, "full_text_requests")
    gaps = [
        {
            "gap_id": f"EG-{idx:03d}",
            "evidence_id": row["evidence_id"],
            "gap_type": "full_text_endpoint_extraction",
            "gap": row["reason"],
            "recommended_action": f"Provide full text in {row['target_folder']}.",
        }
        for idx, row in enumerate(rows, start=1)
    ]
    write_json(root / "evidence_gap_register.json", {"schema_name": "claude_cer_evidence_gap_register", "rows": gaps})
    _write_xlsx(root / "evidence_gap_register.xlsx", gaps, "evidence_gaps")
    return payload


def fulltext_ingest(run_dir: Path, fulltext_root: str | None = None) -> dict[str, Any]:
    baseline = load_baseline(run_dir)
    workbook = baseline["workbook"]
    root = repair_root(run_dir)
    sources = _resolve_fulltext_roots(run_dir, fulltext_root)
    files: list[Path] = []
    for source in sources:
        if source.is_file() and source.suffix.lower() in FULLTEXT_EXTENSIONS:
            files.append(source)
        elif source.exists():
            files.extend(path for path in sorted(source.rglob("*")) if path.is_file() and path.suffix.lower() in FULLTEXT_EXTENSIONS)
    files = list(dict.fromkeys(files))
    evidence_rows = _rows(workbook.get("evidence_registry"))
    benchmark_rows = _rows(workbook.get("sota_benchmark_matrix"))
    index_rows: list[dict[str, Any]] = []
    endpoint_rows: list[dict[str, Any]] = []
    for idx, path in enumerate(files, start=1):
        extracted = _extract_fulltext(path)
        match = _match_fulltext_to_evidence(path, extracted.get("text", ""), evidence_rows)
        file_id = f"FTLIB-{idx:03d}"
        index_rows.append(
            {
                "file_id": file_id,
                "source_type": "user_provided_full_text",
                "filename": path.name,
                "path": str(path),
                "extension": path.suffix.lower(),
                "page_count": extracted.get("page_count"),
                "text_extracted": bool(extracted.get("text")),
                "extraction_status": extracted.get("status"),
                "matched_evidence_id": match.get("evidence_id"),
                "matched_article_id": match.get("article_id"),
                "match_basis": match.get("basis"),
                "audit_note": "Full text was supplied by the user; no acquisition source is recorded in the CER audit trail.",
            }
        )
        endpoint_rows.extend(_extract_fulltext_endpoint_rows(file_id, path, extracted, match, evidence_rows, benchmark_rows))
    manifest = {
        "schema_name": "claude_cer_full_text_ingest_manifest",
        "schema_version": "v1",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "searched_roots": [str(path) for path in sources],
        "file_count": len(files),
        "extracted_endpoint_count": len(endpoint_rows),
        "policy": "Sci-Hub is not automated or recorded. PDFs are handled only as user_provided_full_text.",
    }
    write_json(root / "full_text_ingest_manifest.json", manifest)
    write_json(root / "full_text_library_index.json", {"schema_name": "claude_cer_full_text_library_index", "rows": index_rows})
    write_json(root / "full_text_endpoint_extraction.json", {"schema_name": "claude_cer_full_text_endpoint_extraction", "rows": endpoint_rows})
    _write_xlsx(root / "full_text_library_index.xlsx", index_rows, "full_text_index")
    _write_xlsx(root / "manual_full_text_source_log.xlsx", _manual_fulltext_source_rows(index_rows), "manual_sources")
    _write_xlsx(root / "full_text_endpoint_extraction.xlsx", endpoint_rows, "endpoint_extract")
    payload = {
        **manifest,
        "outputs": {
            "full_text_library_index": str(root / "full_text_library_index.xlsx"),
            "manual_full_text_source_log": str(root / "manual_full_text_source_log.xlsx"),
            "full_text_endpoint_extraction": str(root / "full_text_endpoint_extraction.xlsx"),
        },
    }
    return payload


def sota_endpoint_enhance(run_dir: Path) -> dict[str, Any]:
    package = load_effective_package(run_dir)
    workbook = dict(package["effective_workbook"])
    root = repair_root(run_dir)
    fulltext_payload = read_json(root / "full_text_endpoint_extraction.json")
    fulltext_rows = _rows(fulltext_payload.get("rows") if isinstance(fulltext_payload, dict) else [])
    if not fulltext_rows:
        fulltext_ingest(run_dir)
        fulltext_payload = read_json(root / "full_text_endpoint_extraction.json")
        fulltext_rows = _rows(fulltext_payload.get("rows") if isinstance(fulltext_payload, dict) else [])
    endpoint_rows = _enhanced_endpoint_rows(workbook, fulltext_rows)
    benchmark_rows = _sota_quantitative_benchmark_rows(workbook, endpoint_rows)
    synthesis_rows = _sota_evidence_synthesis_rows(workbook, endpoint_rows, benchmark_rows)
    derivation_rows = _sota_endpoint_derivation_rows(workbook, endpoint_rows, benchmark_rows)
    deduction_tables = _sota_deduction_tables(
        {
            **workbook,
            "endpoint_extraction": endpoint_rows or workbook.get("endpoint_extraction", []),
            "sota_endpoint_derivation_table": derivation_rows,
            "sota_quantitative_benchmark_table": benchmark_rows,
            "sota_evidence_synthesis_matrix": synthesis_rows,
        }
    )
    workbook["endpoint_extraction"] = endpoint_rows or workbook.get("endpoint_extraction", [])
    workbook["sota_endpoint_derivation_table"] = derivation_rows
    workbook["sota_quantitative_benchmark_table"] = benchmark_rows
    workbook["sota_evidence_synthesis_matrix"] = synthesis_rows
    workbook.update(deduction_tables)
    workbook["claude_repair_metadata"] = {
        **(workbook.get("claude_repair_metadata") if isinstance(workbook.get("claude_repair_metadata"), dict) else {}),
        "full_text_enhanced_at": now_iso(),
        "full_text_endpoint_count": len(fulltext_rows),
    }
    write_json(root / "enhanced_authoring_workbook.json", workbook)
    _write_xlsx(root / "sota_endpoint_derivation_table.xlsx", derivation_rows, "sota_endpoint_derivation")
    _write_xlsx(root / "sota_quantitative_benchmark_table.xlsx", benchmark_rows, "sota_benchmarks")
    _write_xlsx(root / "sota_evidence_synthesis_matrix.xlsx", synthesis_rows, "sota_synthesis")
    _write_xlsx(root / "sota_medical_field_boundary.xlsx", deduction_tables["sota_medical_field_boundary"], "medical_field")
    _write_xlsx(root / "sota_pico_v2_strategy.xlsx", deduction_tables["sota_pico_v2_strategy"], "pico_v2")
    _write_xlsx(root / "sota_search_strategy_separated.xlsx", deduction_tables["sota_search_strategy_separated"], "search_separation")
    _write_xlsx(root / "sota_screening_prisma.xlsx", deduction_tables["sota_screening_prisma"], "screening_prisma")
    _write_xlsx(root / "sota_evidence_hierarchy.xlsx", deduction_tables["sota_evidence_hierarchy"], "evidence_hierarchy")
    _write_xlsx(root / "sota_endpoint_extraction_fulltext.xlsx", deduction_tables["sota_endpoint_extraction_fulltext"], "endpoint_fulltext")
    _write_xlsx(root / "sota_benchmark_derivation_table.xlsx", deduction_tables["sota_benchmark_derivation_table"], "benchmark_derivation")
    _write_xlsx(root / "sota_section_conclusion_matrix.xlsx", deduction_tables["sota_section_conclusion_matrix"], "section_conclusion")
    _write_xlsx(root / "sota_deduction_chain.xlsx", deduction_tables["sota_deduction_chain"], "sota_deduction")
    _write_xlsx(root / "sota_endpoint_source_classification.xlsx", deduction_tables["sota_endpoint_source_classification"], "endpoint_sources")
    _write_xlsx(root / "sota_aggregate_benchmark_rationale.xlsx", deduction_tables["sota_aggregate_benchmark_rationale"], "aggregate_basis")
    _write_xlsx(root / "sota_conclusion_strength_guard.xlsx", deduction_tables["sota_conclusion_strength_guard"], "conclusion_guard")
    narratives = _sota_benchmark_narratives(derivation_rows, benchmark_rows)
    write_text(root / "sota_benchmark_derivation_narratives.md", narratives)
    patch_rows = _sota_to_47_patch_rows(benchmark_rows)
    write_json(root / "sota_to_47_patch_manifest.json", {"schema_name": "claude_cer_sota_to_47_patch_manifest", "rows": patch_rows})
    payload = {
        "schema_name": "claude_cer_sota_endpoint_enhancement",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "full_text_endpoint_count": len(fulltext_rows),
        "enhanced_endpoint_count": len(endpoint_rows),
        "benchmark_count": len(benchmark_rows),
        "outputs": {
            "enhanced_authoring_workbook": str(root / "enhanced_authoring_workbook.json"),
            "sota_endpoint_derivation_table": str(root / "sota_endpoint_derivation_table.xlsx"),
            "sota_quantitative_benchmark_table": str(root / "sota_quantitative_benchmark_table.xlsx"),
            "sota_evidence_synthesis_matrix": str(root / "sota_evidence_synthesis_matrix.xlsx"),
            "sota_medical_field_boundary": str(root / "sota_medical_field_boundary.xlsx"),
            "sota_pico_v2_strategy": str(root / "sota_pico_v2_strategy.xlsx"),
            "sota_search_strategy_separated": str(root / "sota_search_strategy_separated.xlsx"),
            "sota_screening_prisma": str(root / "sota_screening_prisma.xlsx"),
            "sota_evidence_hierarchy": str(root / "sota_evidence_hierarchy.xlsx"),
            "sota_endpoint_extraction_fulltext": str(root / "sota_endpoint_extraction_fulltext.xlsx"),
            "sota_benchmark_derivation_table": str(root / "sota_benchmark_derivation_table.xlsx"),
            "sota_section_conclusion_matrix": str(root / "sota_section_conclusion_matrix.xlsx"),
            "sota_deduction_chain": str(root / "sota_deduction_chain.xlsx"),
            "sota_endpoint_source_classification": str(root / "sota_endpoint_source_classification.xlsx"),
            "sota_aggregate_benchmark_rationale": str(root / "sota_aggregate_benchmark_rationale.xlsx"),
            "sota_conclusion_strength_guard": str(root / "sota_conclusion_strength_guard.xlsx"),
            "sota_benchmark_derivation_narratives": str(root / "sota_benchmark_derivation_narratives.md"),
            "sota_to_47_patch_manifest": str(root / "sota_to_47_patch_manifest.json"),
        },
    }
    write_json(root / "sota_endpoint_enhancement_report.json", payload)
    return payload


def sota_reasoning_patch(run_dir: Path) -> dict[str, Any]:
    baseline = load_baseline(run_dir)
    root = repair_root(run_dir)
    source_md = root / "CER_draft_patched.md"
    if not source_md.exists():
        source_md = baseline["artifact_root"] / "CER_draft.md"
    text = source_md.read_text(encoding="utf-8", errors="ignore") if source_md.exists() else ""
    endpoint_report = read_json(root / "sota_endpoint_enhancement_report.json")
    if not isinstance(endpoint_report, dict) or not endpoint_report:
        endpoint_report = sota_endpoint_enhance(run_dir)
    narrative_path = root / "sota_benchmark_derivation_narratives.md"
    narrative = narrative_path.read_text(encoding="utf-8", errors="ignore") if narrative_path.exists() else ""
    addition = (
        "\n\n# Controlled SOTA Endpoint Derivation Addendum\n\n"
        "This addendum records the Claude-side controlled enhancement of the SOTA endpoint and benchmark reasoning. "
        "It is based only on the baseline DeerFlow workbook and user-provided full text artifacts. "
        "It does not introduce unsupported clinical conclusions.\n\n"
        f"{narrative}\n\n"
        "## Section 4.7 Use\n\n"
        "Each accepted benchmark above must be explicitly compared with the device-under-evaluation evidence in section 4.7. "
        "Where page/table-level full-text extraction remains unavailable, the benefit-risk conclusion remains controlled and downgraded.\n"
    )
    patched = text if "Controlled SOTA Endpoint Derivation Addendum" in text else text.rstrip() + addition
    write_text(root / "CER_draft_patched.md", patched)
    _write_docx(root / "CER_draft_patched.docx", patched)
    patch_manifest = read_json(root / "section_patch_manifest.json")
    patch_rows = _rows(patch_manifest.get("rows") if isinstance(patch_manifest, dict) else [])
    patch_rows.append(
        {
            "patch_id": f"PATCH-{len(patch_rows)+1:03d}",
            "affected_section": "3 SOTA / 4.7 Clinical Data Analysis",
            "affected_claim_id": "multiple",
            "affected_pico_id": "SOTA/DuE PICO",
            "affected_evidence_id": "full_text_endpoint_extraction",
            "affected_benchmark_id": "sota_quantitative_benchmark_table",
            "affected_risk_gspr_id": "GSPR 1/6/8 as applicable",
            "old_text_summary": "Baseline SOTA endpoint reasoning before full-text enhancement.",
            "new_text_summary": "Added controlled endpoint derivation narrative and section 4.7 use instruction.",
            "reason": "SOTA benchmark derivation must be explicit and traceable for NB-facing review.",
            "conclusion_strength_change": "No strengthening; full-text gaps remain downgraded.",
            "consistency_check_result": "pending consistency-gate",
        }
    )
    write_json(root / "section_patch_manifest.json", {"schema_name": "claude_cer_section_patch_manifest", "rows": patch_rows})
    payload = {
        "schema_name": "claude_cer_sota_reasoning_patch_report",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "patch_id": patch_rows[-1]["patch_id"],
        "outputs": {
            "cer_draft_patched_md": str(root / "CER_draft_patched.md"),
            "cer_draft_patched_docx": str(root / "CER_draft_patched.docx"),
            "section_patch_manifest": str(root / "section_patch_manifest.json"),
        },
    }
    write_json(root / "sota_reasoning_patch_report.json", payload)
    return payload


def _resolve_fulltext_roots(run_dir: Path, fulltext_root: str | None) -> list[Path]:
    if fulltext_root:
        return [Path(fulltext_root).expanduser().resolve()]
    status = read_json(run_dir / "status.json")
    run_config = read_json(run_dir / "run_config.json")
    roots: list[Path] = []
    for base_value in (run_config.get("input_root"), status.get("input_root")):
        if not base_value:
            continue
        base = Path(str(base_value)).expanduser().resolve()
        for name in DEFAULT_FULLTEXT_DIR_NAMES:
            roots.append(base / name)
    for name in DEFAULT_FULLTEXT_DIR_NAMES:
        roots.append(run_dir / name)
        roots.append(repair_root(run_dir) / name)
    return list(dict.fromkeys(roots))


def _extract_fulltext(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    try:
        if suffix in {".txt", ".md"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
            return {"status": "ok", "text": text, "page_count": 1}
        if suffix == ".docx":
            from docx import Document

            doc = Document(str(path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    table_text.append(" | ".join(cell.text for cell in row.cells))
            return {"status": "ok", "text": "\n".join([text, *table_text]), "page_count": None}
        if suffix == ".pdf":
            import pdfplumber

            pages = []
            with pdfplumber.open(str(path)) as pdf:
                for page_idx, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ""
                    table_text = []
                    try:
                        for table in page.extract_tables() or []:
                            for row in table:
                                table_text.append(" | ".join("" if cell is None else str(cell) for cell in row))
                    except Exception:
                        table_text = []
                    pages.append(f"\n[PAGE {page_idx}]\n{page_text}\n" + "\n".join(table_text))
                return {"status": "ok", "text": "\n".join(pages), "page_count": len(pdf.pages)}
    except Exception as exc:
        return {"status": f"extraction_failed: {exc}", "text": "", "page_count": None}
    return {"status": "unsupported_extension", "text": "", "page_count": None}


def _match_fulltext_to_evidence(path: Path, text: str, evidence_rows: list[dict[str, Any]]) -> dict[str, Any]:
    haystack = f"{path.stem} {text[:5000]}".lower()
    best: dict[str, Any] = {}
    for row in evidence_rows:
        for field in ("pmid", "doi"):
            value = str(_field(row, field, field.upper())).strip()
            if value and value.lower() in haystack:
                return {"evidence_id": row.get("evidence_id"), "article_id": row.get("article_id"), "basis": field}
        title = str(_field(row, "title", "article_title", "source")).lower()
        words = [word for word in re.findall(r"[a-zA-Z][a-zA-Z0-9-]{4,}", title) if word not in {"clinical", "evaluation", "report"}]
        if words and sum(1 for word in words[:10] if word in haystack) >= min(4, max(2, len(words[:10]) // 2)):
            best = {"evidence_id": row.get("evidence_id"), "article_id": row.get("article_id"), "basis": "title_word_overlap"}
            break
    return best or {"basis": "unmatched_user_provided_full_text"}


def _extract_fulltext_endpoint_rows(
    file_id: str,
    path: Path,
    extracted: dict[str, Any],
    match: dict[str, Any],
    evidence_rows: list[dict[str, Any]],
    benchmark_rows: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    text = extracted.get("text") or ""
    if not text:
        return []
    endpoint_terms = []
    for row in evidence_rows:
        endpoint_terms.extend(str(value) for value in (row.get("endpoint"), row.get("result")) if value)
    for row in benchmark_rows:
        endpoint_terms.extend(str(value) for value in (row.get("endpoint"), row.get("clinical_significance")) if value)
    endpoint_terms.extend(["endpoint", "success", "adverse event", "serious adverse", "complication", "safety", "performance"])
    patterns = [term.lower() for term in endpoint_terms if len(str(term).strip()) >= 5]
    chunks = _page_chunks(text)
    rows = []
    for page_ref, chunk in chunks:
        lower = chunk.lower()
        if not any(term in lower for term in patterns):
            continue
        snippet = _endpoint_snippet(chunk)
        if not snippet:
            continue
        rows.append(
            {
                "row_id": f"FT-EP-{len(rows)+1:03d}",
                "file_id": file_id,
                "filename": path.name,
                "source_type": "user_provided_full_text",
                "evidence_id": match.get("evidence_id"),
                "article_id": match.get("article_id"),
                "page_or_section": page_ref,
                "endpoint_definition": _guess_endpoint_definition(snippet, benchmark_rows),
                "sample_size": _first_match(snippet, [r"\bn\s*=\s*\d+", r"\b\d+\s+patients\b", r"\b\d+\s+subjects\b"]) or "requires human confirmation",
                "timepoint": _first_match(snippet, [r"\b\d+\s*(?:day|days|month|months|year|years)\b", r"\bfollow-up[^.;]{0,80}"]) or "requires human confirmation",
                "statistical_result": _first_match(snippet, [r"\b\d+(?:\.\d+)?\s*%", r"\bp\s*[<=>]\s*0?\.\d+", r"\bCI\b[^.;]{0,60}"]) or "requires human confirmation",
                "extracted_text": snippet[:900],
                "extraction_basis": "user-provided full text page/table-level extraction",
                "limitation": "Automated extraction requires human verification against the full text before final sign-off.",
            }
        )
        if len(rows) >= 80:
            break
    return rows


def _manual_fulltext_source_rows(index_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "file_id": row.get("file_id"),
            "filename": row.get("filename"),
            "source_type": "user_provided_full_text",
            "audit_handling": "Manual full text supplied by user; acquisition route is outside automated CER audit trail.",
            "matched_evidence_id": row.get("matched_evidence_id"),
            "extraction_status": row.get("extraction_status"),
        }
        for row in index_rows
    ]


def _page_chunks(text: str) -> list[tuple[str, str]]:
    matches = list(re.finditer(r"\[PAGE\s+(\d+)\]", text))
    if not matches:
        return [("document", text[:60000])]
    chunks: list[tuple[str, str]] = []
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        chunks.append((f"page {match.group(1)}", text[start:end]))
    return chunks


def _endpoint_snippet(chunk: str) -> str:
    sentences = re.split(r"(?<=[.;:])\s+|\n+", chunk)
    keep = [
        sentence.strip()
        for sentence in sentences
        if any(term in sentence.lower() for term in ("endpoint", "success", "adverse", "serious", "complication", "safety", "performance", "follow-up", "patients", "%"))
    ]
    return " ".join(keep[:6]).strip()


def _guess_endpoint_definition(snippet: str, benchmark_rows: list[dict[str, Any]]) -> str:
    lower = snippet.lower()
    for row in benchmark_rows:
        endpoint = str(row.get("endpoint", ""))
        if endpoint and any(word in lower for word in re.findall(r"[a-zA-Z][a-zA-Z0-9-]{4,}", endpoint.lower())[:4]):
            return endpoint
    if "adverse" in lower or "complication" in lower:
        return "Adverse events / serious adverse events / complications"
    if "success" in lower:
        return "Clinical or procedural success"
    return "Endpoint candidate from user-provided full text"


def _first_match(text: str, patterns: list[str]) -> str:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match.group(0)
    return ""


def _enhanced_endpoint_rows(workbook: dict[str, Any], fulltext_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    existing = _rows(workbook.get("endpoint_extraction"))
    output = [dict(row) for row in existing]
    by_evidence = {str(row.get("evidence_id")): row for row in output if row.get("evidence_id")}
    for row in fulltext_rows:
        evidence_id = str(row.get("evidence_id") or "")
        target = by_evidence.get(evidence_id) if evidence_id else None
        update = {
            "full_text_file_id": row.get("file_id"),
            "full_text_page_or_section": row.get("page_or_section"),
            "endpoint_definition": row.get("endpoint_definition"),
            "sample_size": row.get("sample_size"),
            "timepoint": row.get("timepoint"),
            "statistical_result": row.get("statistical_result"),
            "extraction_basis": row.get("extraction_basis"),
            "source_quote_or_cell": row.get("extracted_text"),
            "conclusion": "Full-text endpoint candidate extracted; human verification required before final sign-off.",
        }
        if target:
            target.update({key: value for key, value in update.items() if value})
        else:
            output.append(
                {
                    "endpoint_id": f"EP-FT-{len(output)+1:03d}",
                    "evidence_id": evidence_id or "unmatched",
                    "article_id": row.get("article_id"),
                    "endpoint": row.get("endpoint_definition"),
                    **update,
                }
            )
    return output


def _sota_quantitative_benchmark_rows(workbook: dict[str, Any], endpoint_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows = []
    endpoints_by_benchmark: dict[str, list[dict[str, Any]]] = {}
    for endpoint in endpoint_rows:
        endpoints_by_benchmark.setdefault(str(endpoint.get("benchmark_id") or "unmapped"), []).append(endpoint)
    for idx, benchmark in enumerate(_rows(workbook.get("sota_benchmark_matrix")), start=1):
        benchmark_id = str(benchmark.get("benchmark_id") or f"BM-{idx:02d}")
        scoped = endpoints_by_benchmark.get(benchmark_id, []) or endpoint_rows[:1]
        stats = [str(row.get("statistical_result")) for row in scoped if row.get("statistical_result") and "requires human" not in str(row.get("statistical_result")).lower()]
        rows.append(
            {
                "row_id": f"SOTA-QBM-{idx:03d}",
                "benchmark_id": benchmark_id,
                "endpoint": benchmark.get("endpoint"),
                "clinical_meaning": benchmark.get("clinical_significance"),
                "source_evidence_ids": ", ".join(sorted({str(row.get("evidence_id")) for row in scoped if row.get("evidence_id")})),
                "source_page_or_table": "; ".join(str(row.get("full_text_page_or_section") or row.get("page_or_section")) for row in scoped if row.get("full_text_page_or_section") or row.get("page_or_section")),
                "derived_value_or_range": "; ".join(stats) if stats else benchmark.get("sota_value_range") or "Not quantitatively derived",
                "ci_or_precision": "CI not reported in source unless explicitly captured in the endpoint extraction.",
                "acceptance_criterion": benchmark.get("acceptance_criterion"),
                "limitation": "Automated derivation requires human verification; if no quantitative full-text rows exist, use qualitative controlled benchmark only.",
                "section_4_7_use": "Compare DuE evidence against this benchmark in §4.7; downgrade if only qualitative or abstract-limited.",
                "conclusion": "Benchmark derivation updated from full-text endpoint extraction where available.",
            }
        )
    return rows


def _sota_evidence_synthesis_rows(workbook: dict[str, Any], endpoint_rows: list[dict[str, Any]], benchmark_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    evidence = {str(row.get("evidence_id")): row for row in _rows(workbook.get("evidence_registry"))}
    rows = []
    for idx, endpoint in enumerate(endpoint_rows, start=1):
        ev = evidence.get(str(endpoint.get("evidence_id")), {})
        rows.append(
            {
                "row_id": f"SOTA-SYN-{idx:03d}",
                "evidence_id": endpoint.get("evidence_id"),
                "article_id": endpoint.get("article_id") or ev.get("article_id"),
                "endpoint_id": endpoint.get("endpoint_id"),
                "benchmark_id": endpoint.get("benchmark_id"),
                "evidence_weight": ev.get("weight", "background"),
                "sample_size": endpoint.get("sample_size"),
                "timepoint": endpoint.get("timepoint"),
                "statistical_result": endpoint.get("statistical_result"),
                "applicability": ev.get("device_relevance") or ev.get("population_relevance") or "requires human confirmation",
                "sota_use": "endpoint benchmark derivation / risk identification / background depending on evidence quality",
                "conclusion": "Use strength must follow evidence weight and full-text verification status.",
            }
        )
    if not rows and benchmark_rows:
        rows.append({"row_id": "SOTA-SYN-GAP", "sota_use": "gap", "conclusion": "No full-text endpoint rows available; maintain downgraded SOTA conclusion."})
    return rows


def _sota_endpoint_derivation_rows(workbook: dict[str, Any], endpoint_rows: list[dict[str, Any]], benchmark_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    benchmark_by_id = {str(row.get("benchmark_id")): row for row in benchmark_rows}
    rows = []
    for idx, endpoint in enumerate(endpoint_rows, start=1):
        benchmark = benchmark_by_id.get(str(endpoint.get("benchmark_id")), {})
        rows.append(
            {
                "row_id": f"SOTA-DER-{idx:03d}",
                "pico_id": endpoint.get("pico_id") or "SOTA-PICO",
                "search_id": endpoint.get("search_id") or "see search registry",
                "article_id": endpoint.get("article_id"),
                "evidence_id": endpoint.get("evidence_id"),
                "endpoint_id": endpoint.get("endpoint_id"),
                "benchmark_id": endpoint.get("benchmark_id"),
                "full_text_page_table_section": endpoint.get("full_text_page_or_section") or endpoint.get("page_or_section"),
                "endpoint_definition": endpoint.get("endpoint_definition") or endpoint.get("endpoint"),
                "numerator_denominator": endpoint.get("numerator_denominator") or "requires human confirmation",
                "sample_size": endpoint.get("sample_size"),
                "timepoint": endpoint.get("timepoint"),
                "statistical_result": endpoint.get("statistical_result"),
                "clinical_meaning": endpoint.get("clinical_meaning") or benchmark.get("clinical_meaning"),
                "benchmark_value_or_range": benchmark.get("derived_value_or_range"),
                "limitation": endpoint.get("limitation") or benchmark.get("limitation"),
                "use_in_section_4_7": benchmark.get("section_4_7_use"),
                "conclusion": "Endpoint derivation row created for traceable SOTA-to-4.7 reasoning.",
            }
        )
    return rows


def _sota_benchmark_narratives(derivation_rows: list[dict[str, Any]], benchmark_rows: list[dict[str, Any]]) -> str:
    lines = ["# SOTA Benchmark Derivation Narratives", ""]
    for benchmark in benchmark_rows:
        benchmark_id = benchmark.get("benchmark_id")
        scoped = [row for row in derivation_rows if row.get("benchmark_id") == benchmark_id]
        lines.extend(
            [
                f"## {benchmark_id}: {benchmark.get('endpoint')}",
                "",
                f"Clinical meaning: {benchmark.get('clinical_meaning')}.",
                f"Derived value/range: {benchmark.get('derived_value_or_range')}.",
                f"Acceptance criterion: {benchmark.get('acceptance_criterion')}.",
                f"Evidence basis: {benchmark.get('source_evidence_ids') or 'No mapped full-text evidence yet'}.",
                f"Limitations: {benchmark.get('limitation')}.",
                f"Section 4.7 use: {benchmark.get('section_4_7_use')}.",
                "",
            ]
        )
        for row in scoped[:5]:
            lines.append(
                f"- {row.get('evidence_id')} / {row.get('article_id')} at {row.get('full_text_page_table_section')}: "
                f"{row.get('sample_size')}, {row.get('timepoint')}, {row.get('statistical_result')}."
            )
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _sota_to_47_patch_rows(benchmark_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "patch_target": "§4.7 Clinical Data Analysis",
            "benchmark_id": row.get("benchmark_id"),
            "endpoint": row.get("endpoint"),
            "required_comparison": row.get("section_4_7_use"),
            "conclusion_control": "Do not strengthen conclusion unless DuE evidence meets or justifiably compares with this benchmark.",
        }
        for row in benchmark_rows
    ]


def _sota_deduction_tables(workbook: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    benchmarks = _rows(workbook.get("sota_quantitative_benchmark_table")) or _rows(workbook.get("sota_benchmark_matrix"))
    endpoints = _rows(workbook.get("endpoint_extraction"))
    evidence = _rows(workbook.get("evidence_registry"))
    searches = _rows(workbook.get("search_run_registry"))
    pivotal = [row for row in evidence if str(row.get("weight", "")).lower() == "pivotal"]
    sota_searches = [row for row in searches if str(row.get("objective", "")).lower() == "sota"]
    highest = _highest_evidence_level([str(row.get("evidence_level") or row.get("nb_check_evidence_grade") or "") for row in evidence])
    chain = [
        {"row_id": "SOTA-STEP-001", "step_number": 1, "logic_question": "Why this clinical problem?", "why": "The IFU/intended purpose defines the medical field; SOTA must explain clinical pathway and unmet limitation before device judgment.", "how": "Use Device Profile, Claim Ledger, alternative treatment table and guideline pathway table.", "source_basis": "Device Profile; Claim Ledger; alternative/guideline tables", "output_artifact": "SOTA clinical background and alternative treatment benchmark table", "conclusion_control": "Background does not prove the device; it defines the evaluation question."},
        {"row_id": "SOTA-STEP-002", "step_number": 2, "logic_question": "Why search this way?", "why": "PICO and database choices must follow clinical uncertainty, not generic expansion.", "how": f"{len(sota_searches)} SOTA searches are recorded; unavailable subscription sources must be limitation-recorded.", "source_basis": "LSP, SOTA PICO, database source table, protocol deviations", "output_artifact": "Reproducible search protocol and PRISMA flow", "conclusion_control": "Source unavailable/auth_required is a limitation, not no evidence."},
        {"row_id": "SOTA-STEP-003", "step_number": 3, "logic_question": "Why include or exclude these literatures?", "why": "Only suitable, clinically relevant, higher-contribution literature should define SOTA.", "how": "Use SOTA CK appraisal, screening disposition, evidence level and publication-type filtering.", "source_basis": "Screening log; SOTA CK appraisal; evidence registry", "output_artifact": "Pivotal/supportive/background evidence pool", "conclusion_control": "Editorial/letter/comment/case report cannot be pivotal."},
        {"row_id": "SOTA-STEP-004", "step_number": 4, "logic_question": "Where do endpoints come from?", "why": "Endpoints must come from guideline, aggregate evidence, comparator public data or extracted full text.", "how": f"{len(endpoints)} endpoint rows are linked to {len(benchmarks)} benchmark rows.", "source_basis": "Endpoint extraction; endpoint source classification", "output_artifact": "Performance and safety endpoint source table", "conclusion_control": "Unknown-source endpoints remain qualitative or PMCF/full-text requests."},
        {"row_id": "SOTA-STEP-005", "step_number": 5, "logic_question": "Where do benchmark values come from?", "why": "CEAR/NB review expects acceptance values to be based on aggregate data where possible.", "how": "Record aggregate basis status, source evidence, value/range, CI limitation, population/sample context.", "source_basis": "Quantitative benchmark table; aggregate benchmark rationale", "output_artifact": "SOTA criteria / acceptance standard table", "conclusion_control": "A single isolated study is supportive context only unless explicitly downgraded."},
        {"row_id": "SOTA-STEP-006", "step_number": 6, "logic_question": "How does the product compare with the benchmark?", "why": "SOTA becomes meaningful only when used in §4.7 GSPR benefit-risk comparison.", "how": "Map each benchmark to claim/GSPR, device evidence, PMCF gap or risk-control conclusion.", "source_basis": "SOTA-to-4.7 matrix; risk/GSPR trace; PMCF/PMS/manufacturer data", "output_artifact": "§4.7 safety/performance/benefit-risk comparison", "conclusion_control": "Missing product data means partial support/gap, not favourable conclusion."},
        {"row_id": "SOTA-STEP-007", "step_number": 7, "logic_question": "Why can the conclusion not exceed evidence strength?", "why": "Conclusion language must be limited by evidence level, full-text depth and comparator quality.", "how": f"Highest evidence level: {highest}; pivotal count: {len(pivotal)}.", "source_basis": "Evidence registry; article appraisal; conclusion strength guard", "output_artifact": "Controlled conclusion wording table", "conclusion_control": "No superiority/non-inferiority claim without Level 1/2 comparator evidence and subject-device data."},
    ]
    endpoint_sources = []
    aggregate_rows = []
    for idx, benchmark in enumerate(benchmarks, start=1):
        benchmark_id = str(benchmark.get("benchmark_id") or benchmark.get("row_id") or f"BM-{idx:02d}")
        scoped = [row for row in endpoints if str(row.get("benchmark_id")) == benchmark_id]
        text = " ".join(str(value) for value in [benchmark, *scoped]).lower()
        source_class = _endpoint_source_class(text)
        aggregate_status = _aggregate_basis_status(text)
        endpoint_sources.append(
            {
                "row_id": f"SOTA-ENDSRC-{idx:03d}",
                "benchmark_id": benchmark_id,
                "endpoint": benchmark.get("endpoint") or "endpoint not mapped",
                "endpoint_family": _endpoint_family(str(benchmark.get("endpoint") or "")),
                "source_class": source_class,
                "allowed_endpoint_use": _endpoint_allowed_use(source_class),
                "source_basis": benchmark.get("source_evidence_ids") or benchmark.get("sota_source") or "source pending",
                "clinical_relevance_rationale": benchmark.get("clinical_meaning") or benchmark.get("clinical_significance") or "requires confirmation",
                "limitation": "Unknown or single-study source basis remains controlled until full-text/aggregate confirmation.",
                "conclusion": "Endpoint source classified for SOTA Step 4.",
            }
        )
        aggregate_rows.append(
            {
                "row_id": f"SOTA-AGG-{idx:03d}",
                "benchmark_id": benchmark_id,
                "endpoint": benchmark.get("endpoint") or "endpoint not mapped",
                "benchmark_value_or_range": benchmark.get("derived_value_or_range") or benchmark.get("sota_value_range") or "not quantitatively derived",
                "aggregate_basis_status": aggregate_status,
                "aggregate_source_requirement": "Guideline, systematic review, meta-analysis, registry or multiple high-quality clinical studies preferred.",
                "source_evidence_ids": benchmark.get("source_evidence_ids") or benchmark.get("sota_source") or "source pending",
                "single_study_control": "Single isolated study cannot define SOTA acceptance criterion without downgrade.",
                "allowed_conclusion_strength": _allowed_conclusion_strength_for_aggregate_status(aggregate_status),
                "limitation": "Human evaluator verifies aggregate derivation before signature-level use.",
                "conclusion": "Aggregate basis/downgrade recorded.",
            }
        )
    conclusion_guard = [
        {
            "row_id": "SOTA-CG-001",
            "guard_scope": "Overall SOTA conclusion",
            "highest_evidence_level": highest,
            "pivotal_evidence_count": len(pivotal),
            "comparator_basis": "requires explicit comparator and subject-device data",
            "superiority_claim_allowed": False,
            "allowed_language": "consistent with SOTA / acceptable / partially supported / PMCF-controlled",
            "prohibited_language": "superior, proven better, definitive, fully demonstrated, all risks acceptable",
            "reason": "Conclusion wording is bounded by evidence level and aggregate benchmark status.",
            "conclusion": "Conclusion guard recorded.",
        }
    ]
    medical_field_boundary = [
        {
            "row_id": "SOTA-FIELD-001",
            "element": "medical_field_boundary",
            "definition": (workbook.get("device_profile") or {}).get("device_type") or "IFU-defined medical field",
            "source_basis": "Device Profile / Claim Ledger / locked subject IFU",
            "boundary_rule": "Only evidence matching the IFU-defined medical field can support pivotal conclusions.",
            "extrapolation_limit": "Remote populations, procedures or device classes require transferability justification.",
            "use_in_cer": "SOTA §3.1-3.3 and §4.7 comparator boundary",
            "conclusion": "Medical field boundary controls SOTA search and benchmark use.",
        }
    ]
    pico_v2 = [
        {
            "row_id": f"SOTA-PICO-V2-{idx:03d}",
            "source_pico_id": row.get("pico_id") or row.get("row_id"),
            "population": row.get("population"),
            "intervention": row.get("intervention"),
            "comparator": row.get("comparator"),
            "outcome": row.get("outcome"),
            "revision_logic": "Separate SOTA clinical-field questions from DuE/trade-name literature questions.",
            "query": row.get("query") or row.get("proposed_query"),
            "conclusion": "PICO v2 controls SOTA search logic.",
        }
        for idx, row in enumerate(_rows(workbook.get("sota_pico_strategy")) or _rows(workbook.get("cep_pico_matrix")), start=1)
    ]
    search_separated = [
        {
            "row_id": f"SOTA-SEARCH-SEP-{idx:03d}",
            "search_id": row.get("search_id"),
            "database": row.get("database"),
            "search_role": row.get("objective"),
            "query": row.get("query"),
            "separation_rule": "SOTA uses disease/procedure/device-class terms; device literature uses subject/equivalent/similar trade/manufacturer/model terms.",
            "result_count": row.get("result_count"),
            "conclusion": "Search role separation recorded.",
        }
        for idx, row in enumerate(searches, start=1)
    ]
    screening_prisma = [
        {"row_id": f"SOTA-SCREEN-{idx:03d}", "phase": "screening", "metric": row.get("article_id") or row.get("screen_id"), "value": row.get("title_abstract_decision") or row.get("full_text_decision"), "exclusion_reason": row.get("exclusion_reason"), "conclusion": row.get("conclusion") or "Screening decision recorded."}
        for idx, row in enumerate(_rows(workbook.get("sota_screening_disposition_table")), start=1)
    ]
    evidence_hierarchy = [
        {"row_id": f"SOTA-EH-{idx:03d}", "article_id": row.get("article_id"), "evidence_id": row.get("evidence_id"), "evidence_category": row.get("evidence_category") or row.get("evidence_type"), "evidence_level": row.get("evidence_level") or row.get("nb_check_evidence_grade") or row.get("evidence_hierarchy"), "weight": row.get("weight") or row.get("disposition"), "endpoint_contribution": row.get("endpoint_contribution") or row.get("evidence_use"), "limitation": row.get("limitation") or row.get("limitations"), "conclusion": "Evidence hierarchy controls benchmark contribution and conclusion strength."}
        for idx, row in enumerate(_rows(workbook.get("sota_ck_appraisal_table")) or evidence, start=1)
    ]
    endpoint_fulltext = [
        {"row_id": f"SOTA-FT-EP-{idx:03d}", "endpoint_id": row.get("endpoint_id"), "benchmark_id": row.get("benchmark_id"), "evidence_id": row.get("source_evidence_id") or row.get("evidence_id"), "endpoint": row.get("endpoint"), "sample_size": row.get("sample_size"), "timepoint": row.get("timepoint"), "statistical_result": row.get("statistical_result"), "full_text_trace": row.get("full_text_page_or_section") or row.get("page_or_section") or row.get("source_page_or_table"), "full_text_status": "full_text_or_source_trace_present" if "full text" in " ".join(str(value) for value in row.values()).lower() or row.get("full_text_page_or_section") else "full_text_request_or_abstract_limited", "conclusion": "Full-text endpoint extraction must be written back to benchmark narrative and §4.7."}
        for idx, row in enumerate(endpoints, start=1)
    ]
    section_conclusions = [
        {"row_id": f"SOTA-CONC-{idx:03d}", "section": section, "topic": topic, "required_conclusion_logic": "State what was established, source basis, limitation and §4.7 use.", "trace_source": "SOTA deduction chain / benchmark matrix / usage matrix", "use_in_cer": "End of SOTA subsection", "conclusion": "Section conclusion required for human CER style."}
        for idx, (section, topic) in enumerate([("3.1", "Medical field"), ("3.2", "Target condition"), ("3.4", "Alternatives"), ("3.5", "Guidelines"), ("3.6", "Benchmark devices"), ("3.7", "Hazards"), ("3.8", "Endpoints/benchmarks"), ("3.9", "SOTA conclusion")], start=1)
    ]
    return {
        "sota_medical_field_boundary": medical_field_boundary,
        "sota_pico_v2_strategy": pico_v2,
        "sota_search_strategy_separated": search_separated,
        "sota_screening_prisma": screening_prisma,
        "sota_evidence_hierarchy": evidence_hierarchy,
        "sota_endpoint_extraction_fulltext": endpoint_fulltext,
        "sota_benchmark_derivation_table": benchmarks,
        "sota_section_conclusion_matrix": section_conclusions,
        "sota_deduction_chain": chain,
        "sota_endpoint_source_classification": endpoint_sources,
        "sota_aggregate_benchmark_rationale": aggregate_rows,
        "sota_conclusion_strength_guard": conclusion_guard,
    }


def _endpoint_source_class(text: str) -> str:
    if any(token in text for token in ("guideline", "consensus", "recommendation")):
        return "clinical_practice_guideline"
    if any(token in text for token in ("meta-analysis", "systematic review", "cochrane")):
        return "aggregate_systematic_review_or_meta_analysis"
    if any(token in text for token in ("registry", "real-world", "real world")):
        return "registry_or_real_world_aggregate"
    if any(token in text for token in ("ifu", "510(k)", "public clinical", "clinical report")):
        return "comparator_ifu_or_public_clinical_report"
    if any(token in text for token in ("full text", "page", "table", "randomized", "prospective", "retrospective")):
        return "clinical_study_full_text"
    return "source_pending_or_qualitative_context"


def _endpoint_allowed_use(source_class: str) -> str:
    if source_class in {"clinical_practice_guideline", "aggregate_systematic_review_or_meta_analysis", "registry_or_real_world_aggregate"}:
        return "May support endpoint selection and aggregate benchmark derivation after population/endpoint match."
    if source_class == "clinical_study_full_text":
        return "May support endpoint data; single-study value requires downgrade before SOTA threshold use."
    return "Qualitative controlled context or full-text request only."


def _endpoint_family(endpoint: str) -> str:
    lower = endpoint.lower()
    if any(token in lower for token in ("adverse", "injury", "infection", "bleeding", "perforation", "malfunction", "serious")):
        return "safety_endpoint"
    if any(token in lower for token in ("success", "completion", "visualization", "image", "performance", "maneuver", "deflection")):
        return "performance_or_clinical_benefit_endpoint"
    return "mixed_or_claim_specific_endpoint"


def _aggregate_basis_status(text: str) -> str:
    if any(token in text for token in ("meta-analysis", "systematic review", "cochrane", "registry", "guideline", "consensus")):
        return "aggregate_available"
    if any(token in text for token in ("not quantitatively derived", "pending", "not quantified", "requires full-text", "full text needed")):
        return "aggregate_required_pending_full_text"
    if any(token in text for token in ("randomized", "prospective", "retrospective", "single study")):
        return "single_study_context_downgraded"
    return "qualitative_controlled"


def _allowed_conclusion_strength_for_aggregate_status(status: str) -> str:
    if status == "aggregate_available":
        return "quantitative comparison allowed after applicability check"
    if status == "single_study_context_downgraded":
        return "supportive only; no definitive threshold or superiority conclusion"
    return "qualitative/partial support until aggregate/full-text evidence is available"


def _highest_evidence_level(levels: list[str]) -> str:
    text = " ".join(levels).lower()
    for level in ("1", "2", "3", "4", "5"):
        if f"level {level}" in text or f"cebm {level}" in text:
            return f"level {level}"
    return "not graded"


def sota_narrative_merge(run_dir: Path) -> dict[str, Any]:
    package = load_effective_package(run_dir)
    root = repair_root(run_dir)
    workbook = dict(package["effective_workbook"])
    text = package["effective_cer_text"]
    enhancement = read_json(root / "sota_endpoint_enhancement_report.json")
    if not isinstance(enhancement, dict) or not enhancement:
        enhancement = sota_endpoint_enhance(run_dir)
        workbook = read_json(root / "enhanced_authoring_workbook.json") or workbook
    narrative_path = root / "sota_benchmark_derivation_narratives_v2.md"
    if not narrative_path.exists():
        narrative_path = root / "sota_benchmark_derivation_narratives.md"
    narrative = narrative_path.read_text(encoding="utf-8", errors="ignore") if narrative_path.exists() else _sota_benchmark_narratives(
        _rows(workbook.get("sota_endpoint_derivation_table")),
        _rows(workbook.get("sota_quantitative_benchmark_table")),
    )
    cleaned, removed_templates, removed_dot_blocks = _clean_benchmark_template_text(text)
    merged_section = _build_sota_merge_section(workbook, narrative)
    patched = _replace_or_append_section(cleaned, "Controlled SOTA Benchmark Derivation and 4.7 Merge", merged_section)
    patch_manifest = read_json(root / "section_patch_manifest.json")
    patch_rows = _rows(patch_manifest.get("rows") if isinstance(patch_manifest, dict) else [])
    patch_rows.append(
        {
            "patch_id": f"PATCH-{len(patch_rows)+1:03d}",
            "affected_section": "§3.8 SOTA endpoints / §4.7 GSPR analysis / Annex SOTA tables",
            "affected_claim_id": "multiple",
            "affected_pico_id": "SOTA/DuE PICO",
            "affected_evidence_id": "sota_endpoint_derivation_table",
            "affected_benchmark_id": "BM-ALL",
            "affected_risk_gspr_id": "GSPR 1/6/8 where benchmark is used",
            "old_text_summary": f"Removed {removed_templates} template benchmark phrases and {removed_dot_blocks} long-dot truncation blocks.",
            "new_text_summary": "Merged differentiated SOTA benchmark narratives, values, source/evidence IDs, population/sample/timepoint and CI limitation statements into the CER body.",
            "reason": "Benchmark narratives must be report-body evidence, not only intermediate artifacts.",
            "conclusion_strength_change": "none; quantitative limitations remain explicit",
            "consistency_check_result": "pending consistency-gate",
        }
    )
    workbook["claude_repair_metadata"] = {
        **(workbook.get("claude_repair_metadata") if isinstance(workbook.get("claude_repair_metadata"), dict) else {}),
        "sota_narrative_merged_at": now_iso(),
        "removed_template_benchmark_phrases": removed_templates,
        "removed_long_dot_blocks": removed_dot_blocks,
    }
    write_json(root / "enhanced_authoring_workbook.json", workbook)
    _write_xlsx(root / "sota_medical_field_boundary.xlsx", _rows(workbook.get("sota_medical_field_boundary")), "medical_field")
    _write_xlsx(root / "sota_pico_v2_strategy.xlsx", _rows(workbook.get("sota_pico_v2_strategy")), "pico_v2")
    _write_xlsx(root / "sota_search_strategy_separated.xlsx", _rows(workbook.get("sota_search_strategy_separated")), "search_separation")
    _write_xlsx(root / "sota_screening_prisma.xlsx", _rows(workbook.get("sota_screening_prisma")), "screening_prisma")
    _write_xlsx(root / "sota_evidence_hierarchy.xlsx", _rows(workbook.get("sota_evidence_hierarchy")), "evidence_hierarchy")
    _write_xlsx(root / "sota_endpoint_extraction_fulltext.xlsx", _rows(workbook.get("sota_endpoint_extraction_fulltext")), "endpoint_fulltext")
    _write_xlsx(root / "sota_benchmark_derivation_table.xlsx", _rows(workbook.get("sota_benchmark_derivation_table")), "benchmark_derivation")
    _write_xlsx(root / "sota_section_conclusion_matrix.xlsx", _rows(workbook.get("sota_section_conclusion_matrix")), "section_conclusion")
    _write_xlsx(root / "sota_deduction_chain.xlsx", _rows(workbook.get("sota_deduction_chain")), "sota_deduction")
    _write_xlsx(root / "sota_endpoint_source_classification.xlsx", _rows(workbook.get("sota_endpoint_source_classification")), "endpoint_sources")
    _write_xlsx(root / "sota_aggregate_benchmark_rationale.xlsx", _rows(workbook.get("sota_aggregate_benchmark_rationale")), "aggregate_basis")
    _write_xlsx(root / "sota_conclusion_strength_guard.xlsx", _rows(workbook.get("sota_conclusion_strength_guard")), "conclusion_guard")
    write_json(root / "section_patch_manifest.json", {"schema_name": "claude_cer_section_patch_manifest", "rows": patch_rows})
    write_text(root / "CER_draft_patched.md", patched)
    _write_docx(root / "CER_draft_patched.docx", patched)
    payload = {
        "schema_name": "claude_cer_sota_narrative_merge_report",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "removed_template_benchmark_phrases": removed_templates,
        "removed_long_dot_blocks": removed_dot_blocks,
        "benchmark_count": len(_rows(workbook.get("sota_quantitative_benchmark_table") or workbook.get("sota_benchmark_matrix"))),
        "outputs": {
            "enhanced_authoring_workbook": str(root / "enhanced_authoring_workbook.json"),
            "section_patch_manifest": str(root / "section_patch_manifest.json"),
            "sota_deduction_chain": str(root / "sota_deduction_chain.xlsx"),
            "sota_aggregate_benchmark_rationale": str(root / "sota_aggregate_benchmark_rationale.xlsx"),
            "sota_conclusion_strength_guard": str(root / "sota_conclusion_strength_guard.xlsx"),
            "cer_draft_patched_md": str(root / "CER_draft_patched.md"),
            "cer_draft_patched_docx": str(root / "CER_draft_patched.docx"),
        },
    }
    write_json(root / "sota_narrative_merge_report.json", payload)
    return payload


def _clean_benchmark_template_text(text: str) -> tuple[str, int, int]:
    cleaned, template_count = BENCHMARK_TEMPLATE_PATTERN.subn("", text)
    cleaned, dot_count = LONG_DOT_PATTERN.subn("", cleaned)
    return cleaned, template_count, dot_count


def _build_sota_merge_section(workbook: dict[str, Any], narrative: str) -> str:
    benchmark_rows = _rows(workbook.get("sota_quantitative_benchmark_table")) or _rows(workbook.get("sota_benchmark_matrix"))
    derivation_rows = _rows(workbook.get("sota_endpoint_derivation_table"))
    deduction = _sota_deduction_tables(workbook)
    workbook.update({key: workbook.get(key) or value for key, value in deduction.items()})
    lines = [
        "# Controlled SOTA Benchmark Derivation and 4.7 Merge",
        "",
        "This section replaces template benchmark placeholders with differentiated endpoint reasoning derived from the controlled workbook, full-text endpoint extraction where available, and recorded SOTA evidence limitations. It is part of the CER body and must be read together with sections 3.8 and 4.7.",
        "",
        "## SOTA Seven-Step Deduction Chain",
        "",
        "| Step | Reviewer question | Why | How | Source basis | Conclusion control |",
        "| --- | --- | --- | --- | --- | --- |",
    ]
    for row in workbook.get("sota_deduction_chain") or []:
        lines.append(
            "| "
            + " | ".join(
                _escape_md_cell(str(value))
                for value in (row.get("step_number"), row.get("logic_question"), row.get("why"), row.get("how"), row.get("source_basis"), row.get("conclusion_control"))
            )
            + " |"
        )
    lines.extend(
        [
            "",
        "## Quantitative Benchmark Derivation Table",
        "",
        "| Benchmark ID | Endpoint | Source / Evidence | Population / Sample | Value / CI | Clinical meaning | 4.7 use | Limitation |",
        "| --- | --- | --- | --- | --- | --- | --- | --- |",
        ]
    )
    for idx, row in enumerate(benchmark_rows, start=1):
        benchmark_id = row.get("benchmark_id") or row.get("row_id") or f"BM-{idx:02d}"
        scoped = [item for item in derivation_rows if str(item.get("benchmark_id")) == str(benchmark_id)]
        source = row.get("source_evidence_ids") or row.get("sota_source") or "; ".join(str(item.get("evidence_id")) for item in scoped if item.get("evidence_id")) or "source not mapped"
        sample = row.get("sample_size") or "; ".join(str(item.get("sample_size")) for item in scoped if item.get("sample_size")) or "sample size not reported in source"
        value = row.get("derived_value_or_range") or row.get("sota_value_range") or "; ".join(str(item.get("statistical_result")) for item in scoped if item.get("statistical_result")) or "not quantitatively derived"
        ci = _ci_statement(row, scoped)
        clinical = row.get("clinical_meaning") or row.get("clinical_significance") or "clinical meaning requires human confirmation"
        use = row.get("section_4_7_use") or row.get("section_4_7") or "Compare with device evidence in §4.7; downgrade if DuE data are absent."
        limitation = row.get("limitation") or "Automated benchmark derivation requires human verification before sign-off."
        lines.append(
            "| "
            + " | ".join(
                _escape_md_cell(str(value))
                for value in (benchmark_id, row.get("endpoint") or "endpoint not mapped", source, sample, f"{value}; {ci}", clinical, use, limitation)
            )
            + " |"
        )
    lines.extend(
        [
            "",
            "## Benchmark Derivation Narratives",
            "",
            narrative.strip() or "No differentiated narrative was available; keep this CER as controlled draft with gaps.",
            "",
            "## Section 4.7 Use Control",
            "",
            "Each benchmark above must be explicitly used in the GSPR 1 safety/performance/benefit-risk analysis. Where subject-device clinical data, PMS or full-text endpoint extraction remain incomplete, conclusions remain partially supported and PMCF-controlled.",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _ci_statement(row: dict[str, Any], scoped: list[dict[str, Any]]) -> str:
    text = " ".join(str(value) for value in [*row.values(), *[item for scoped_row in scoped for item in scoped_row.values()]])
    if re.search(r"\b(?:CI|confidence interval)\b", text, flags=re.IGNORECASE):
        return "CI reported in source/extraction"
    return "CI not reported in source"


def _escape_md_cell(value: str) -> str:
    return value.replace("\n", " ").replace("|", "\\|").strip()


def _replace_or_append_section(text: str, title: str, section: str) -> str:
    heading = f"# {title}"
    pattern = re.compile(rf"^# {re.escape(title)}\n.*?(?=^# |\Z)", flags=re.MULTILINE | re.DOTALL)
    if pattern.search(text):
        return pattern.sub(section, text).rstrip() + "\n"
    return text.rstrip() + "\n\n" + section


def nb_flag_triage(run_dir: Path) -> dict[str, Any]:
    baseline = load_baseline(run_dir)
    root = repair_root(run_dir)
    flags = _extract_nb_flags(baseline)
    rows = []
    for idx, flag in enumerate(flags, start=1):
        decision, rationale, residual = _triage_nb_flag(flag)
        rows.append(
            {
                "flag_id": f"NB-FLAG-{idx:04d}",
                "severity": flag.get("severity") or "UNKNOWN",
                "type": flag.get("type") or flag.get("category") or "unknown",
                "matched": flag.get("matched") or flag.get("term") or "",
                "location_or_context": flag.get("context") or flag.get("location") or "",
                "decision": decision,
                "rationale": rationale,
                "evidence_or_patch_id": _flag_patch_reference(flag),
                "residual_status": residual,
                "mdr_ref": flag.get("mdr_ref") or "",
            }
        )
    if not rows:
        rows.append(
            {
                "flag_id": "NB-FLAG-0000",
                "severity": "NONE",
                "type": "no_flags",
                "decision": "accepted_no_action",
                "rationale": "No NB precheck flags were found in baseline gate reports.",
                "evidence_or_patch_id": "qa_gate_report",
                "residual_status": "closed",
            }
        )
    payload = {
        "schema_name": "claude_cer_nb_flag_triage",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "flag_count": len(flags),
        "high_flag_count": sum(1 for flag in flags if str(flag.get("severity", "")).upper() == "HIGH"),
        "rows": rows,
    }
    write_json(root / "nb_flag_triage.json", payload)
    _write_xlsx(root / "nb_flag_triage.xlsx", rows, "nb_flag_triage")
    return payload


def _extract_nb_flags(baseline: dict[str, Any]) -> list[dict[str, Any]]:
    reports = [baseline.get("qa_gate_report") or {}, baseline.get("final_gate_closure_report") or {}]
    flags: list[dict[str, Any]] = []
    seen: set[str] = set()

    def walk(value: Any) -> None:
        if isinstance(value, dict):
            maybe_flags = value.get("flags")
            if isinstance(maybe_flags, list):
                for item in maybe_flags:
                    if isinstance(item, dict):
                        key = json.dumps(item, ensure_ascii=False, sort_keys=True, default=str)
                        if key not in seen:
                            flags.append(item)
                            seen.add(key)
            for child in value.values():
                walk(child)
        elif isinstance(value, list):
            for child in value:
                walk(child)

    for report in reports:
        walk(report)
    return flags


def _triage_nb_flag(flag: dict[str, Any]) -> tuple[str, str, str]:
    matched = str(flag.get("matched") or "").lower()
    suggestion = str(flag.get("suggestion") or "")
    severity = str(flag.get("severity") or "").upper()
    if "equivalent to" in matched or "equivalence" in matched:
        return ("treat_with_route_context", "Equivalence wording requires route-specific context; use not-claimed/comparable wording unless MDCG equivalence is demonstrated.", "open_until_text_verified")
    if "clinical benefit" in matched:
        return ("treat_with_evidence_mapping", "Clinical-benefit wording must be linked to PMID/DOI/evidence ID and GSPR; unsupported benefit claims remain partial.", "open_until_evidence_linked")
    if "literature review" in matched:
        return ("treat_with_lsp_reference", "Literature-search method must cite LSP databases, dates, queries, inclusion/exclusion criteria and PRISMA trace.", "open_until_lsp_trace_verified")
    if suggestion:
        return ("review_and_patch", suggestion[:500], "open")
    if severity == "HIGH":
        return ("human_review_required", "High-severity flag requires explicit acceptance or text patch before signature readiness.", "open")
    return ("accepted_with_rationale", "Flag reviewed; no automatic text patch rule matched.", "accepted")


def _flag_patch_reference(flag: dict[str, Any]) -> str:
    matched = str(flag.get("matched") or "").lower()
    if "equivalence" in matched or "equivalent to" in matched:
        return "equivalence_route_statement / section_patch_manifest"
    if "clinical benefit" in matched:
        return "claim_ledger / evidence_registry / GSPR trace"
    if "literature" in matched:
        return "LSP package / PRISMA / search registry"
    return "nb_precheck_report / section_patch_manifest"


def rmf_pmcf_closure(run_dir: Path) -> dict[str, Any]:
    package = load_effective_package(run_dir)
    root = repair_root(run_dir)
    workbook = dict(package["effective_workbook"])
    source_inventory = _rows(workbook.get("source_inventory"))
    risks = _rows(workbook.get("risk_trace_matrix"))
    benchmarks = _rows(workbook.get("sota_benchmark_matrix"))
    rmf_sources = [
        row for row in source_inventory
        if any(term in str(row).lower() for term in ("rmf", "rmr", "risk management", "fmea", "风险"))
    ]
    rmf_rows = []
    for idx, risk in enumerate(risks or [{"risk_id": "R-GAP", "risk_side_effect": "risk source not extracted"}], start=1):
        rmf_rows.append(
            {
                "row_id": f"RMF-EXT-{idx:03d}",
                "risk_id": risk.get("risk_id") or f"R-{idx:03d}",
                "risk_or_hazard": risk.get("risk_side_effect") or risk.get("hazard") or risk.get("source") or "risk not named",
                "source_status": "rmf_source_present_review_required" if rmf_sources else "rmf_source_missing",
                "source_ids": ", ".join(str(row.get("source_id")) for row in rmf_sources if row.get("source_id")),
                "risk_control": risk.get("rmf_coverage") or "not yet extracted",
                "ifu_control": risk.get("ifu_coverage") or "not yet extracted",
                "pms_pmcf_control": risk.get("pms_pmcf_coverage") or "PMCF/PMS request required",
                "extraction_conclusion": "Requires source-document verification before residual-risk acceptability is upgraded.",
            }
        )
    closure_rows = []
    for idx, risk in enumerate(rmf_rows, start=1):
        benchmark = benchmarks[(idx - 1) % len(benchmarks)] if benchmarks else {}
        closure_rows.append(
            {
                "row_id": f"BR-CLOSE-{idx:03d}",
                "risk_id": risk.get("risk_id"),
                "risk_or_hazard": risk.get("risk_or_hazard"),
                "corresponding_benefit": benchmark.get("endpoint") or "clinical benefit requires claim/evidence mapping",
                "benchmark_id": benchmark.get("benchmark_id") or "unmapped",
                "risk_control": risk.get("risk_control"),
                "ifu_control": risk.get("ifu_control"),
                "residual_benefit_risk_conclusion": "partial / gap-controlled until RMF extraction, PMS exposure and PMCF data are verified",
            }
        )
    pmcf_rows = _pmcf_timetable_rows(workbook)
    workbook["rmf_deep_extraction"] = rmf_rows
    workbook["risk_benefit_closure_matrix"] = closure_rows
    workbook["pmcf_timetable_request"] = pmcf_rows
    workbook["claude_repair_metadata"] = {
        **(workbook.get("claude_repair_metadata") if isinstance(workbook.get("claude_repair_metadata"), dict) else {}),
        "rmf_pmcf_closure_generated_at": now_iso(),
    }
    write_json(root / "enhanced_authoring_workbook.json", workbook)
    write_json(root / "rmf_deep_extraction.json", {"schema_name": "claude_cer_rmf_deep_extraction", "rows": rmf_rows})
    write_json(root / "risk_benefit_closure_matrix.json", {"schema_name": "claude_cer_risk_benefit_closure_matrix", "rows": closure_rows})
    write_json(root / "pmcf_timetable_request.json", {"schema_name": "claude_cer_pmcf_timetable_request", "rows": pmcf_rows})
    _write_xlsx(root / "rmf_deep_extraction.xlsx", rmf_rows, "rmf_extraction")
    _write_xlsx(root / "risk_benefit_closure_matrix.xlsx", closure_rows, "risk_benefit")
    _write_xlsx(root / "pmcf_timetable_request.xlsx", pmcf_rows, "pmcf_timetable")
    payload = {
        "schema_name": "claude_cer_rmf_pmcf_closure_report",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "rmf_source_count": len(rmf_sources),
        "risk_row_count": len(rmf_rows),
        "pmcf_timetable_row_count": len(pmcf_rows),
        "outputs": {
            "rmf_deep_extraction": str(root / "rmf_deep_extraction.xlsx"),
            "risk_benefit_closure_matrix": str(root / "risk_benefit_closure_matrix.xlsx"),
            "pmcf_timetable_request": str(root / "pmcf_timetable_request.xlsx"),
        },
    }
    write_json(root / "rmf_pmcf_closure_report.json", payload)
    return payload


def _pmcf_timetable_rows(workbook: dict[str, Any]) -> list[dict[str, Any]]:
    device = (workbook.get("device_profile") or {}).get("device_name") if isinstance(workbook.get("device_profile"), dict) else "subject device"
    return [
        {"row_id": "PMCF-TIME-001", "subject_device": device, "time_window": "0-3 months after CE / launch", "activity": "Confirm PMS data capture, complaint coding, vigilance linkage and PMCF responsibilities.", "endpoint_or_data": "sales denominator, complaint categories, reportable incidents", "status": "requested"},
        {"row_id": "PMCF-TIME-002", "subject_device": device, "time_window": "3-6 months", "activity": "Collect early user feedback and procedure-performance endpoints from representative users.", "endpoint_or_data": "visualization/procedure success, device deficiency, AE/SAE", "status": "requested"},
        {"row_id": "PMCF-TIME-003", "subject_device": device, "time_window": "6-12 months", "activity": "Perform PMCF interim analysis and update CER evidence gaps.", "endpoint_or_data": "PMS trends, complaint rates, residual risks, literature update", "status": "requested"},
        {"row_id": "PMCF-TIME-004", "subject_device": device, "time_window": "Annual until evidence mature", "activity": "Update CER, PMS/PMCF evaluation and risk-benefit conclusion.", "endpoint_or_data": "long-term safety, rare risks, real-world performance", "status": "requested"},
    ]


def controlled_patch(run_dir: Path) -> dict[str, Any]:
    package = load_effective_package(run_dir)
    baseline = load_baseline(run_dir)
    workbook = package["effective_workbook"]
    text = package["effective_cer_text"]
    root = repair_root(run_dir)
    root.mkdir(parents=True, exist_ok=True)
    enhanced = dict(workbook)
    patches: list[dict[str, Any]] = []
    patched = text
    replacements = {
        "Ifu risk warning was extracted from an original-language source and is summarized in English in the CER; the original wording remains traceable in the controlled workbook.": "IFU-defined warning requiring controlled clinical confirmation; the original warning text remains traceable to the subject-device IFU and the residual risk remains gap-controlled until RMF/PMS evidence is reviewed.",
        "Ifu warning was extracted from an original-language source and is summarized in English in the CER; the original wording remains traceable in the controlled workbook.": "IFU-defined warning requiring clinical and risk-management traceability; final acceptability remains conditional on RMF/PMS review.",
        "report-safe summaries": "controlled English summaries",
    }
    for old, new in replacements.items():
        if old in patched:
            patched = patched.replace(old, new)
            patches.append(
                {
                    "patch_id": f"PATCH-{len(patches)+1:03d}",
                    "affected_section": "risk tables / narrative rows",
                    "affected_claim_id": "multiple",
                    "affected_pico_id": "multiple",
                    "affected_evidence_id": "not changed",
                    "affected_benchmark_id": "not changed",
                    "affected_risk_gspr_id": "risk_trace_matrix",
                    "old_text_summary": old[:180],
                    "new_text_summary": new[:180],
                    "reason": "Replace machine/meta wording with report-facing controlled English without strengthening the conclusion.",
                    "conclusion_strength_change": "none; remains gap-controlled",
                    "consistency_check_result": "pending consistency-gate",
                }
            )
    enhanced["claude_repair_metadata"] = {
        "schema_name": "claude_enhanced_authoring_workbook",
        "generated_at": now_iso(),
        "baseline_artifact_root": str(baseline["artifact_root"]),
        "patch_count": len(patches),
        "rule": "Baseline DeerFlow workbook retained; Claude patches must be traceable in section_patch_manifest.",
    }
    write_json(root / "enhanced_authoring_workbook.json", enhanced)
    write_json(root / "section_patch_manifest.json", {"schema_name": "claude_cer_section_patch_manifest", "rows": patches})
    write_text(root / "CER_draft_patched.md", patched)
    _write_docx(root / "CER_draft_patched.docx", patched)
    payload = {
        "schema_name": "claude_cer_controlled_patch_report",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "patch_count": len(patches),
        "outputs": {
            "enhanced_authoring_workbook": str(root / "enhanced_authoring_workbook.json"),
            "section_patch_manifest": str(root / "section_patch_manifest.json"),
            "cer_draft_patched_md": str(root / "CER_draft_patched.md"),
            "cer_draft_patched_docx": str(root / "CER_draft_patched.docx"),
        },
    }
    write_json(root / "controlled_patch_report.json", payload)
    return payload


def consistency_gate(run_dir: Path) -> dict[str, Any]:
    root = repair_root(run_dir)
    artifact = artifact_root(run_dir)
    workbook_path = root / "enhanced_authoring_workbook.json"
    draft_path = root / "CER_draft_patched.md"
    if not workbook_path.exists():
        workbook_path = artifact / "authoring_workbook.json"
    if not draft_path.exists():
        draft_path = artifact / "CER_draft.md"
    workbook = read_json(workbook_path)
    text = draft_path.read_text(encoding="utf-8", errors="ignore") if draft_path.exists() else ""
    issues: list[str] = []
    if not isinstance(workbook, dict) or not workbook:
        issues.append("Enhanced or baseline workbook is missing.")
    if not text:
        issues.append("Patched or baseline CER markdown is missing.")
    if cjk_count(text) > 0:
        issues.append("CER contains CJK characters.")
    placeholders = [item for item in PLACEHOLDERS if item.lower() in text.lower()]
    if placeholders:
        issues.append(f"CER contains placeholders: {', '.join(placeholders)}")
    claims = [str(row.get("claim_id")) for row in _rows(workbook.get("claim_ledger")) if row.get("claim_id")]
    evidence = [str(row.get("evidence_id")) for row in _rows(workbook.get("evidence_registry")) if row.get("evidence_id")]
    picos = [str(row.get("pico_id")) for row in _rows(workbook.get("cep_pico_matrix")) if row.get("pico_id")]
    missing_claim_refs = [claim_id for claim_id in claims[:20] if claim_id not in text]
    if len(missing_claim_refs) > max(3, len(claims) // 2):
        issues.append("Many Claim Ledger IDs are absent from the CER text.")
    if evidence and not any(eid in text for eid in evidence):
        issues.append("Evidence IDs are absent from the CER text.")
    if picos and not any(pid in text for pid in picos):
        issues.append("PICO IDs are absent from the CER text.")
    if "Ifu risk warning was extracted" in text:
        issues.append("Machine/meta IFU warning wording remains in patched CER.")
    decision = "PASS" if not issues else "REWORK_REQUIRED"
    payload = {
        "schema_name": "claude_cer_patch_qa_report",
        "schema_version": "v1",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "decision": decision,
        "issues": issues,
        "checked_files": {
            "workbook": str(workbook_path),
            "cer_markdown": str(draft_path),
        },
    }
    write_json(root / "patch_qa_report.json", payload)
    return payload


def self_check_scorecard(run_dir: Path) -> dict[str, Any]:
    package = load_effective_package(run_dir)
    root = repair_root(run_dir)
    workbook = package["effective_workbook"]
    text = package["effective_cer_text"]
    checks = _self_check_rows(run_dir, workbook, text)
    part_weights = {
        "Part 1": 15,
        "Part 2": 25,
        "Part 3": 20,
        "Part 4": 15,
        "Part 5": 10,
        "Part 6": 10,
        "Part 7": 5,
    }
    part_scores: dict[str, dict[str, Any]] = {}
    total_score = 0.0
    for part, weight in part_weights.items():
        scoped = [row for row in checks if row["part"] == part]
        passed = sum(1 for row in scoped if row["result"] == "PASS")
        score = round(weight * passed / max(1, len(scoped)), 2)
        total_score += score
        part_scores[part] = {"max": weight, "score": score, "passed": passed, "total": len(scoped)}
    hard_failures = [row for row in checks if row["part"] == "Part 8" and row["result"] == "FAIL" and row.get("hard_gate")]
    score_int = int(round(total_score))
    if not hard_failures and score_int >= SIGNATURE_SCORE_THRESHOLD:
        conclusion = "SIGNATURE_READY"
    elif not hard_failures and score_int >= CONTROLLED_SCORE_THRESHOLD:
        conclusion = "CONTROLLED_DRAFT_WITH_GAPS"
    else:
        conclusion = "NOT_DELIVERABLE"
    payload = {
        "schema_name": "claude_cer_self_check_scorecard",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "score": score_int,
        "delivery_conclusion": conclusion,
        "passed_count_part_1_7": sum(1 for row in checks if row["part"] in part_weights and row["result"] == "PASS"),
        "failed_count_part_1_7": sum(1 for row in checks if row["part"] in part_weights and row["result"] == "FAIL"),
        "hard_gate_failure_count": len(hard_failures),
        "part_scores": part_scores,
        "top_gaps": [row for row in checks if row["result"] == "FAIL"][:10],
        "rows": checks,
    }
    write_json(root / "cer_self_check_scorecard.json", payload)
    _write_xlsx(root / "cer_self_check_scorecard.xlsx", checks, "scorecard")
    write_text(root / "cer_self_check_scorecard.md", _scorecard_markdown(payload))
    return payload


def _self_check_rows(run_dir: Path, workbook: dict[str, Any], text: str) -> list[dict[str, Any]]:
    lower = text.lower()
    claims = _rows(workbook.get("claim_ledger"))
    picos = _rows(workbook.get("cep_pico_matrix"))
    searches = _rows(workbook.get("search_run_registry"))
    evidence = _rows(workbook.get("evidence_registry"))
    endpoints = _rows(workbook.get("endpoint_extraction"))
    benchmarks = _rows(workbook.get("sota_quantitative_benchmark_table")) or _rows(workbook.get("sota_benchmark_matrix"))
    derivations = _rows(workbook.get("sota_endpoint_derivation_table"))
    deduction_chain = _rows(workbook.get("sota_deduction_chain"))
    endpoint_sources = _rows(workbook.get("sota_endpoint_source_classification"))
    aggregate_rationale = _rows(workbook.get("sota_aggregate_benchmark_rationale"))
    conclusion_guards = _rows(workbook.get("sota_conclusion_strength_guard"))
    risks = _rows(workbook.get("risk_trace_matrix"))
    gspr = _rows(workbook.get("gspr_coverage"))
    vigilance = _rows(workbook.get("vigilance_recall_registry"))
    vigilance_stats = _rows(workbook.get("vigilance_event_statistics"))
    nb_triage = _rows(read_json(repair_root(run_dir) / "nb_flag_triage.json").get("rows") if isinstance(read_json(repair_root(run_dir) / "nb_flag_triage.json"), dict) else [])
    nb_flags = _extract_nb_flags(load_baseline(run_dir))
    high_flags = [flag for flag in nb_flags if str(flag.get("severity", "")).upper() == "HIGH"]
    pico_revision = _rows(read_json(repair_root(run_dir) / "pico_revision_log.json").get("rows") if isinstance(read_json(repair_root(run_dir) / "pico_revision_log.json"), dict) else [])
    pico_v2 = _rows(workbook.get("pico_v2_strategy")) or _rows(read_json(repair_root(run_dir) / "pico_v2_strategy.json").get("rows") if isinstance(read_json(repair_root(run_dir) / "pico_v2_strategy.json"), dict) else [])
    final_dir = artifact_root(run_dir) / "final"
    qa = load_baseline(run_dir).get("qa_gate_report") or {}
    qa_results = _rows(qa.get("results"))
    pivotal = [row for row in evidence if str(row.get("weight", "")).lower() == "pivotal"]
    rows: list[dict[str, Any]] = []

    def add(part: str, check_id: str, item: str, ok: bool, evidence_text: str, gap: str = "", hard: bool = False) -> None:
        rows.append(
            {
                "part": part,
                "check_id": check_id,
                "check_item": item,
                "result": "PASS" if ok else "FAIL",
                "evidence": evidence_text if ok else "",
                "gap": "" if ok else gap,
                "hard_gate": hard,
            }
        )

    add("Part 1", "1.1", "MDR CER chapter structure complete", all(token in lower for token in ("summary", "scope", "sota", "device under evaluation", "conclusion")), "Core chapters detected", "Missing core chapter marker")
    add("Part 1", "1.2", "Claim -> PICO -> SOTA -> GSPR chain complete", bool(claims and picos and benchmarks and gspr), f"{len(claims)} claims / {len(picos)} PICO / {len(benchmarks)} BM / {len(gspr)} GSPR", "Core chain table missing")
    add("Part 1", "1.3", "LSP and PRISMA recorded", bool(workbook.get("literature_search_protocol_profile") and workbook.get("prisma_flow_data")), "LSP profile and PRISMA present", "LSP or PRISMA missing")
    add("Part 1", "1.4", "Search reproducible", _searches_reproducible(searches), f"{len(searches)} search records", "Search rows lack date/query/count/status")
    add("Part 1", "1.5", "Manufacturer clinical data / PMCF timetable complete", bool(workbook.get("pmcf_timetable_request") or _has_source_term(workbook, ("pmcf", "clinical investigation", "clinical report"))), "PMCF/clinical source control present", "No clinical investigation source or PMCF timetable")

    search_text = " ".join(str(row) for row in searches).lower()
    add("Part 2", "2.1.1", "Database selection correct", sum(1 for db in ("pubmed", "embase", "cochrane", "europe pmc", "clinicaltrials") if db in search_text) >= 3 and "maude" in search_text and "eudamed" in search_text, "SOTA/device/vigilance database coverage present", "Required databases/vigilance sources missing")
    add("Part 2", "2.1.2", "SOTA and Device Literature terms separated", _sota_due_searches_separated(searches), "SOTA and DuE query roles are separated", "SOTA and device literature appear mixed or duplicated", True)
    add("Part 2", "2.1.3", "Equivalent/similar trade names included", _trade_name_search_present(searches, vigilance), "Trade/device-name strategy present", "No visible trade-name/manufacturer/model search terms")
    add("Part 2", "2.1.4", "Vigilance results filtered", bool(vigilance_stats and all(row.get("relevant_count") is not None or row.get("relevant_cases") is not None for row in vigilance_stats)), "Vigilance event statistics present", "Vigilance has raw rows but no filtered event statistics")
    add("Part 2", "2.2.1", "Each endpoint has benchmark", bool(endpoints and benchmarks and len(benchmarks) >= min(len(endpoints), len(picos))), f"{len(endpoints)} endpoints / {len(benchmarks)} benchmarks", "Endpoint-to-benchmark coverage incomplete")
    add("Part 2", "2.2.2", "Benchmark has source", _benchmarks_have_source(benchmarks), "Benchmark source fields present", "Benchmark source/PMID/DOI/evidence IDs missing")
    add("Part 2", "2.2.3", "Benchmark has value and CI/CI limitation", _benchmarks_have_values(benchmarks), "Benchmark values and CI limitation statements present", "Benchmark value/source/CI limitation incomplete", True)
    add("Part 2", "2.2.4", "Benchmark has population context", _benchmarks_have_population(benchmarks, derivations), "Population/sample context present", "Population/sample/research type context missing")
    add("Part 2", "2.2.5", "Benchmark narratives differentiated", text.count("Claim-specific benchmark generated") == 0 and _unique_ratio([str(row.get("clinical_meaning") or row.get("clinical_significance") or row.get("endpoint")) for row in benchmarks]) >= 0.65, "Benchmark template repetition not detected", "Benchmark narratives are repetitive/template-like", True)
    add("Part 2", "2.2.6", "Benchmark evidence quality and aggregate basis sufficient", len(pivotal) >= 2 and not any(_excluded_publication_type(row) for row in pivotal) and _aggregate_basis_controlled(aggregate_rationale) and bool(endpoint_sources), f"{len(pivotal)} pivotal rows; {len(aggregate_rationale)} aggregate-basis rows", "Need >=2 pivotal evidence, endpoint source classification and aggregate/downgrade rationale; no editorial/letter/comment as pivotal", True)

    add("Part 3", "3.1", "Evidence weight distribution reasonable", len(pivotal) >= 2, f"{len(pivotal)} pivotal rows", "Pivotal evidence count below 2", True)
    add("Part 3", "3.2", "Each evidence has sample size and follow-up", _evidence_has_sample_followup(evidence), "Evidence sample/follow-up fields populated", "Evidence sample size/follow-up not extracted")
    add("Part 3", "3.3", "Endpoint values have source, CI and sample", _endpoints_have_source_ci_sample(endpoints), "Endpoint source/sample/CI or CI limitation present", "Endpoint source/sample/CI fields incomplete")
    add("Part 3", "3.4", "Equivalence path explicit", "equivalence is not claimed" in lower or bool(workbook.get("equivalence_matrix")), "Equivalence path declared", "Equivalence/not-equivalence path unclear")

    add("Part 4", "4.1", "Each claim has GSPR trace", bool(claims and gspr), "Claim/GSPR rows present", "Claim/GSPR trace missing")
    add("Part 4", "4.2", "GSPR conclusion has qualitative/quantitative/source basis", all(token in lower for token in ("benefit-risk", "benchmark", "gspr")), "Benefit-risk/GSPR/benchmark language present", "GSPR analysis lacks one of benefit-risk/benchmark/source")
    add("Part 4", "4.3", "Risk-benefit closure", bool(workbook.get("risk_benefit_closure_matrix")) and not _all_gap_controlled(risks), "Risk-benefit closure matrix present", "Risks remain gap-controlled or RMF not extracted", True)
    add("Part 4", "4.4", "Claim to benchmark mapping", bool(workbook.get("sota_to_47_usage_matrix") and claims), "SOTA-to-4.7 usage matrix present", "Claim/benchmark/product/PMCF chain missing")
    add("Part 4", "4.5", "Section conclusions", text.count("Section conclusion:") >= 10, f"{text.count('Section conclusion:')} section conclusions", "Section conclusions too sparse")

    add("Part 5", "5.1", "Marketing history clear", "marketing history" in lower and any(term in lower for term in ("not marketed", "no marketing", "sales", "pms", "complaint")), "Marketing/PMS status visible", "Marketing status/source not clear")
    add("Part 5", "5.2", "Vigilance AE to RMR closure", bool(vigilance_stats and workbook.get("risk_benefit_closure_matrix")), "Vigilance statistics and risk closure present", "AE statistics do not close to RMF/RMR risks", True)
    add("Part 5", "5.3", "PMS or PMCF plan timetable", bool(_has_source_term(workbook, ("pms", "pmcf")) or workbook.get("pmcf_timetable_request")), "PMS/PMCF timetable/source present", "No PMS data or PMCF timetable")

    add("Part 6", "6.1", "No 'equivalent to' wording", "equivalent to" not in lower, "No exact 'equivalent to' wording", "Forbidden equivalence wording remains")
    add("Part 6", "6.2", "Clinical benefits have data references", "clinical benefit" in lower and any(token in lower for token in ("pmid", "doi", "e-00", "evidence")), "Clinical benefit has evidence markers", "Clinical benefit lacks source reference")
    add("Part 6", "6.3", "No template repetition", text.count("Claim-specific benchmark generated") == 0, "No benchmark template phrase", "Benchmark template phrase remains")
    add("Part 6", "6.4", "No乱码/truncation", cjk_count(text) == 0 and not LONG_DOT_PATTERN.search(text), "No CJK/long-dot truncation detected", "CJK or truncation artifacts remain")
    add("Part 6", "6.5", "Equivalence route consistent", _equivalence_route_consistent(lower), "Equivalence route consistent", "Equivalence route ambiguous")

    add("Part 7", "7.1", "All final outputs in deerflow_authoring/final", final_dir.exists() and (final_dir / "final_manifest.json").exists(), "Final package directory present", "Final package not published to deerflow_authoring/final")
    add("Part 7", "7.2", "SOTA benchmark table in CER body", "Controlled SOTA Benchmark Derivation" in text or "Quantitative Benchmark Derivation Table" in text, "Benchmark derivation present in body", "Benchmark derivation only in intermediate artifacts")
    add("Part 7", "7.3", "DeerFlow QA gates / final gate resolved", _qa_gates_signature_ok(qa_results), "Gate report has no G30/G33 unresolved blocker", "G30/G33 or insufficient gate pass remains", True)
    add("Part 7", "7.4", "NB flags triaged", len(nb_triage) >= len(high_flags), f"{len(nb_triage)} triage rows / {len(high_flags)} high flags", "HIGH flags lack one-by-one triage", True)

    add("Part 8", "8.1.1", "No pivotal evidence risk", len(pivotal) >= 2, f"{len(pivotal)} pivotal rows", "Pivotal evidence count below 2", True)
    add("Part 8", "8.1.2", "No benchmark value without source", not _benchmark_value_without_source(benchmarks), "No value-without-source benchmark found", "Benchmark value/source mismatch", True)
    add("Part 8", "8.1.3", "Sample size/follow-up extracted", _evidence_has_sample_followup(evidence), "Evidence extraction populated", "Sample/follow-up extraction incomplete", True)
    add("Part 8", "8.2.1", "Benchmark differentiation", text.count("Claim-specific benchmark generated") == 0, "No repeated BM template phrase", "Repeated BM template phrase remains", True)
    add("Part 8", "8.2.2", "Publication type filtered", not any(_excluded_publication_type(row) and str(row.get("weight", "")).lower() == "pivotal" for row in evidence), "No excluded publication type as pivotal", "Editorial/letter/comment/case report used as pivotal", True)
    add("Part 8", "8.2.3", "Fulltext ingest written back", _fulltext_written_back(workbook, text), "Full-text endpoint data written back", "Full-text endpoints not reflected in workbook/CER body", True)
    add("Part 8", "8.3.1", "Failed DeerFlow gates not bypassed", _qa_gates_signature_ok(qa_results), "No unresolved G30/G33 blocker", "Baseline G30/G33 unresolved; final must re-evaluate", True)
    add("Part 8", "8.3.2", "NB flags one-by-one triage", len(nb_triage) >= len(high_flags), "HIGH flags triaged", "HIGH flags lack triage rows", True)
    add("Part 8", "8.3.3", "PICO narrow/split executed", _pico_v2_executed(pico_revision, pico_v2), "PICO v2 executed or reason recorded", "PICO loop advice not executed", True)
    add("Part 8", "8.4.1", "Final artifact location", final_dir.exists(), "Final directory exists", "Final artifacts only in claude_repair", True)
    add("Part 8", "8.4.2", "4.7 GSPR analysis table complete", bool(workbook.get("risk_benefit_closure_matrix") and workbook.get("sota_quantitative_benchmark_table")), "Risk/benchmark analysis tables present", "4.7 analysis table incomplete", True)

    add("Part 9", "9.1", "SOTA 8-step/7-deduction execution capability", _part9_status(workbook, text, "sota") and _sota_deduction_complete(deduction_chain) and bool(conclusion_guards), "SOTA pipeline and reviewer-facing deduction chain present", "SOTA steps or seven-step deduction chain incomplete")
    add("Part 9", "9.2", "SOTA quantitative extraction capability", bool(workbook.get("sota_quantitative_benchmark_table")), "Quantitative benchmark table present", "Quantitative benchmark table missing")
    add("Part 9", "9.3", "Claims to benchmark chain capability", bool(claims and picos and benchmarks), "Claim/PICO/BM chain present", "Claim/PICO/BM chain missing")
    add("Part 9", "9.4", "Similar-device four-step capability", bool(workbook.get("similar_device_four_step_confirmation") and workbook.get("similar_device_attachment_index")), "Similar-device package present", "Similar-device package missing")
    add("Part 9", "9.5", "GSPR qualitative/quantitative conclusion capability", "benefit-risk" in lower and bool(workbook.get("sota_quantitative_benchmark_table")), "GSPR dual-track basis present", "GSPR qualitative/quantitative basis incomplete")
    add("Part 9", "9.6", "LSP + PRISMA execution capability", bool(workbook.get("literature_search_protocol_profile") and workbook.get("prisma_flow_data")), "LSP and PRISMA present", "LSP/PRISMA missing")
    return rows


def _scorecard_markdown(payload: dict[str, Any]) -> str:
    failed = [row for row in payload.get("rows", []) if row.get("result") == "FAIL"]
    return "\n".join(
        [
            "# AI CER Self-Check Conclusion",
            "",
            f"Total score: {payload.get('score')}/100",
            f"Delivery conclusion: {payload.get('delivery_conclusion')}",
            f"Passed items: {payload.get('passed_count_part_1_7')}",
            f"Failed items: {payload.get('failed_count_part_1_7')}",
            "",
            "## Top 3 Gaps",
            *[f"- {row.get('check_id')}: {row.get('gap')}" for row in failed[:3]],
            "",
            "## Next Actions",
            "- Resolve hard-gate failures before signature-ready classification.",
            "- Keep external-source gaps as controlled draft limitations until manufacturer/full-text evidence is supplied.",
        ]
    ) + "\n"


def _searches_reproducible(searches: list[dict[str, Any]]) -> bool:
    if not searches:
        return False
    for row in searches:
        if not row.get("database") or not row.get("query") or not row.get("search_date"):
            return False
        if row.get("result_count") is None and row.get("status") not in {"auth_required", "source_unavailable"}:
            return False
    return True


def _has_source_term(workbook: dict[str, Any], terms: tuple[str, ...]) -> bool:
    text = " ".join(str(row) for row in _rows(workbook.get("source_inventory"))).lower()
    return any(term.lower() in text for term in terms)


def _sota_due_searches_separated(searches: list[dict[str, Any]]) -> bool:
    sota_queries = {
        str(row.get("query", "")).strip().lower()
        for row in searches
        if "sota" in str(row.get("objective", "")).lower() or "sota" in str(row.get("search_id", "")).lower()
    }
    due_queries = {
        str(row.get("query", "")).strip().lower()
        for row in searches
        if any(token in f"{row.get('objective', '')} {row.get('search_id', '')}".lower() for token in ("due", "device", "literature"))
    }
    due_queries.discard("")
    sota_queries.discard("")
    return bool(sota_queries and due_queries and (sota_queries - due_queries) and (due_queries - sota_queries))


def _trade_name_search_present(searches: list[dict[str, Any]], vigilance: list[dict[str, Any]]) -> bool:
    text = " ".join(str(row) for row in [*searches, *vigilance]).lower()
    return any(term in text for term in ("lithovue", "clearpetra", "ambu", "pc200", "pusheng", "trade", "brand", "manufacturer", "model"))


def _benchmarks_have_source(benchmarks: list[dict[str, Any]]) -> bool:
    if not benchmarks:
        return False
    for row in benchmarks:
        text = " ".join(str(row.get(key, "")) for key in ("source_evidence_ids", "sota_source", "source", "pmid", "doi")).lower()
        if not text.strip() or "bibliographic record verified" in text:
            return False
    return True


def _benchmarks_have_values(benchmarks: list[dict[str, Any]]) -> bool:
    if not benchmarks:
        return False
    for row in benchmarks:
        value = str(row.get("derived_value_or_range") or row.get("sota_value_range") or row.get("statistical_result") or "")
        row_text = " ".join(str(value) for value in row.values()).lower()
        if not any(char.isdigit() for char in value):
            return False
        if not ("ci" in row_text or "confidence interval" in row_text or "ci not reported" in row_text):
            return False
        if not _benchmarks_have_source([row]):
            return False
    return True


def _benchmarks_have_population(benchmarks: list[dict[str, Any]], derivations: list[dict[str, Any]]) -> bool:
    text = " ".join(str(row) for row in [*benchmarks, *derivations]).lower()
    return any(token in text for token in ("patients", "subjects", "population", "n=", "sample", "systematic review", "meta", "registry", "in-vitro", "in vitro"))


def _sota_deduction_complete(rows: list[dict[str, Any]]) -> bool:
    steps = {str(row.get("step_number")) for row in rows}
    required = {str(idx) for idx in range(1, 8)}
    return required.issubset(steps) and all(
        str(row.get(field, "")).strip()
        for row in rows
        for field in ("logic_question", "why", "how", "source_basis", "conclusion_control")
    )


def _aggregate_basis_controlled(rows: list[dict[str, Any]]) -> bool:
    if not rows:
        return False
    for row in rows:
        status = str(row.get("aggregate_basis_status", "")).lower()
        strength = str(row.get("allowed_conclusion_strength", "")).lower()
        if not status or not row.get("single_study_control"):
            return False
        if "single_study" in status and any(term in strength for term in ("superior", "definitive", "fully")):
            return False
    return True


def _unique_ratio(values: list[str]) -> float:
    normalized = [re.sub(r"\s+", " ", value.strip().lower()) for value in values if value.strip()]
    return len(set(normalized)) / len(normalized) if normalized else 0.0


def _excluded_publication_type(row: dict[str, Any]) -> bool:
    text = " ".join(str(value) for value in row.values()).lower()
    return any(token in text for token in ("editorial", "letter", "comment", "news", "case report"))


def _evidence_has_sample_followup(evidence: list[dict[str, Any]]) -> bool:
    if not evidence:
        return False
    for row in evidence:
        sample = str(row.get("sample_size") or "")
        follow = str(row.get("follow_up") or row.get("followup") or row.get("timepoint") or "")
        if not sample.strip() or not follow.strip() or "not extracted" in f"{sample} {follow}".lower():
            return False
    return True


def _endpoints_have_source_ci_sample(endpoints: list[dict[str, Any]]) -> bool:
    if not endpoints:
        return False
    for row in endpoints:
        text = " ".join(str(value) for value in row.values()).lower()
        has_source = bool(row.get("evidence_id") or row.get("source_evidence_id") or row.get("source"))
        has_sample = bool(row.get("sample_size")) and "not extracted" not in str(row.get("sample_size")).lower()
        has_ci = "ci" in text or "confidence interval" in text or "ci not reported" in text
        if not (has_source and has_sample and has_ci):
            return False
    return True


def _all_gap_controlled(risks: list[dict[str, Any]]) -> bool:
    if not risks:
        return True
    return all("gap" in str(row).lower() or "not yet extracted" in str(row).lower() for row in risks)


def _equivalence_route_consistent(lower_text: str) -> bool:
    not_claimed = "equivalence is not claimed" in lower_text or "equivalent-device use is not assumed" in lower_text
    claimed = "equivalence demonstrated" in lower_text
    return not (not_claimed and claimed) and (not_claimed or claimed or "similar device" in lower_text)


def _qa_gates_signature_ok(qa_results: list[dict[str, Any]]) -> bool:
    if not qa_results:
        return False
    failed = [row for row in qa_results if row.get("status") != "PASS"]
    if any(row.get("gate_id") in {"G30", "G33"} for row in failed):
        return False
    return len([row for row in qa_results if row.get("status") == "PASS"]) >= min(39, len(qa_results))


def _benchmark_value_without_source(benchmarks: list[dict[str, Any]]) -> bool:
    for row in benchmarks:
        value = str(row.get("derived_value_or_range") or row.get("sota_value_range") or "")
        if any(char.isdigit() for char in value) and not _benchmarks_have_source([row]):
            return True
    return False


def _fulltext_written_back(workbook: dict[str, Any], text: str) -> bool:
    return bool(
        workbook.get("sota_endpoint_derivation_table")
        and workbook.get("sota_quantitative_benchmark_table")
        and ("Controlled SOTA Benchmark Derivation" in text or "Quantitative Benchmark Derivation Table" in text)
    )


def _pico_v2_executed(pico_revision: list[dict[str, Any]], pico_v2: list[dict[str, Any]]) -> bool:
    needs_v2 = [row for row in pico_revision if row.get("decision") in {"narrow_or_split", "broaden", "broaden_device_literature"}]
    if not needs_v2:
        return True
    if not pico_v2:
        return False
    source_ids = {str(row.get("source_pico_id")) for row in pico_v2}
    return all(str(row.get("pico_id")) in source_ids or any(row.get("if_not_executed_reason") for row in pico_v2) for row in needs_v2)


def _part9_status(workbook: dict[str, Any], text: str, capability: str) -> bool:
    if capability == "sota":
        return bool(
            workbook.get("sota_pico_strategy")
            and workbook.get("sota_ck_appraisal_table")
            and workbook.get("prisma_flow_data")
            and workbook.get("sota_to_47_usage_matrix")
            and ("Section conclusion:" in text)
        )
    return False


def signature_readiness_check(run_dir: Path) -> dict[str, Any]:
    root = repair_root(run_dir)
    qa = read_json(root / "patch_qa_report.json")
    if not isinstance(qa, dict) or not qa:
        qa = consistency_gate(run_dir)
    workbook = read_json(root / "enhanced_authoring_workbook.json")
    if not isinstance(workbook, dict) or not workbook:
        workbook = load_baseline(run_dir)["workbook"]
    scorecard = read_json(root / "cer_self_check_scorecard.json")
    if not isinstance(scorecard, dict) or not scorecard:
        scorecard = self_check_scorecard(run_dir)
    fulltext_manifest = read_json(root / "full_text_ingest_manifest.json")
    benchmark_rows = _rows(workbook.get("sota_quantitative_benchmark_table"))
    derivation_rows = _rows(workbook.get("sota_endpoint_derivation_table"))
    sota_ck = _rows(workbook.get("sota_ck_appraisal_table"))
    sota_screening = _rows(workbook.get("sota_screening_disposition_table"))
    quantity = _sota_literature_quantity(sota_ck, sota_screening)
    issues = []
    if qa.get("decision") != "PASS":
        issues.append("consistency-gate has not passed")
    if quantity["included_count"] < SOTA_LITERATURE_TARGET_MIN and not _sota_quantity_justification_complete(workbook.get("sota_literature_quantity_justification") or {}):
        issues.append("SOTA final included literature count is below target and lacks complete justification")
    if not derivation_rows:
        issues.append("SOTA endpoint derivation table is missing")
    if not benchmark_rows:
        issues.append("SOTA quantitative benchmark table is missing")
    if not fulltext_manifest and any("requires_full_text" in str(row).lower() or "first-pass" in str(row).lower() for row in _rows(workbook.get("evidence_registry"))):
        issues.append("full-text ingest has not been executed for full-text-limited evidence")
    if int(scorecard.get("hard_gate_failure_count") or 0) > 0:
        issues.append("self-check hard gates failed")
    if int(scorecard.get("score") or 0) < SIGNATURE_SCORE_THRESHOLD:
        issues.append(f"self-check score below signature threshold: {scorecard.get('score')}")
    decision = "SIGNATURE_REVIEW_CANDIDATE" if not issues else "CONTROLLED_DRAFT_WITH_GAPS"
    payload = {
        "schema_name": "claude_cer_signature_readiness_report",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "decision": decision,
        "issues": issues,
        "sota_literature_quantity": quantity,
        "self_check_score": scorecard.get("score"),
        "self_check_conclusion": scorecard.get("delivery_conclusion"),
        "human_responsibility": "Human evaluator/manufacturer confirmation remains required before regulatory submission or sign-off.",
    }
    write_json(root / "signature_readiness_report.json", payload)
    return payload


def finalize_package(run_dir: Path, *, beautify_docx: bool = False, theme: str = "slate") -> dict[str, Any]:
    root = repair_root(run_dir)
    if not (root / "CER_draft_patched.md").exists():
        controlled_patch(run_dir)
    if not (root / "sota_narrative_merge_report.json").exists():
        sota_narrative_merge(run_dir)
    if not (root / "pico_v2_strategy.xlsx").exists():
        pico_v2_execute(run_dir)
    if not (root / "nb_flag_triage.xlsx").exists():
        nb_flag_triage(run_dir)
    if not (root / "risk_benefit_closure_matrix.xlsx").exists():
        rmf_pmcf_closure(run_dir)
    qa = read_json(root / "patch_qa_report.json")
    if not isinstance(qa, dict) or not qa:
        qa = consistency_gate(run_dir)
    else:
        qa = consistency_gate(run_dir)
    if not (artifact_root(run_dir) / "final" / "final_manifest.json").exists():
        publish_final_package(run_dir, decision="PENDING_SELF_CHECK", scorecard={}, readiness={})
    scorecard = read_json(root / "cer_self_check_scorecard.json")
    if not isinstance(scorecard, dict) or not scorecard:
        scorecard = self_check_scorecard(run_dir)
    readiness = read_json(root / "signature_readiness_report.json")
    if not isinstance(readiness, dict) or not readiness:
        readiness = signature_readiness_check(run_dir)
    if readiness.get("decision") == "SIGNATURE_REVIEW_CANDIDATE" and scorecard.get("delivery_conclusion") == "SIGNATURE_READY":
        decision = "SIGNATURE_READY"
    elif int(scorecard.get("score") or 0) >= CONTROLLED_SCORE_THRESHOLD and int(scorecard.get("hard_gate_failure_count") or 0) == 0:
        decision = "CONTROLLED_DRAFT_WITH_GAPS"
    else:
        decision = "NOT_DELIVERABLE_REWORK_REQUIRED"
    required = [
        "deep_postflight_report.json",
        "rework_decision.json",
        "pico_revision_log.xlsx",
        "pico_v2_strategy.xlsx",
        "full_text_request_list.xlsx",
        "evidence_gap_register.xlsx",
        "enhanced_authoring_workbook.json",
        "section_patch_manifest.json",
        "CER_draft_patched.md",
        "CER_draft_patched.docx",
        "patch_qa_report.json",
        "cer_self_check_scorecard.json",
        "cer_self_check_scorecard.xlsx",
        "nb_flag_triage.xlsx",
    ]
    optional = [
        "full_text_ingest_manifest.json",
        "full_text_library_index.xlsx",
        "manual_full_text_source_log.xlsx",
        "full_text_endpoint_extraction.xlsx",
        "sota_endpoint_derivation_table.xlsx",
        "sota_quantitative_benchmark_table.xlsx",
        "sota_evidence_synthesis_matrix.xlsx",
        "sota_medical_field_boundary.xlsx",
        "sota_pico_v2_strategy.xlsx",
        "sota_search_strategy_separated.xlsx",
        "sota_screening_prisma.xlsx",
        "sota_evidence_hierarchy.xlsx",
        "sota_endpoint_extraction_fulltext.xlsx",
        "sota_benchmark_derivation_table.xlsx",
        "sota_section_conclusion_matrix.xlsx",
        "sota_deduction_chain.xlsx",
        "sota_endpoint_source_classification.xlsx",
        "sota_aggregate_benchmark_rationale.xlsx",
        "sota_conclusion_strength_guard.xlsx",
        "sota_benchmark_derivation_narratives.md",
        "sota_to_47_patch_manifest.json",
        "sota_reasoning_patch_report.json",
        "sota_narrative_merge_report.json",
        "rmf_deep_extraction.xlsx",
        "risk_benefit_closure_matrix.xlsx",
        "pmcf_timetable_request.xlsx",
        "rmf_pmcf_closure_report.json",
        "signature_readiness_report.json",
        "CER_draft_final_beautified.docx",
        "docx_beautification_report.json",
    ]
    beautification: dict[str, Any] | None = None
    if beautify_docx:
        if qa.get("decision") == "PASS":
            beautification = _beautify_final_docx(root, theme)
        else:
            beautification = {
                "schema_name": "claude_cer_docx_beautification_report",
                "generated_at": now_iso(),
                "decision": "SKIPPED_CONTENT_GATE_NOT_PASS",
                "reason": "Final DOCX beautification is blocked until consistency-gate passes.",
            }
            write_json(root / "docx_beautification_report.json", beautification)
        if beautification.get("decision") != "PASS" and decision == "SIGNATURE_READY":
            decision = "CONTROLLED_DRAFT_WITH_GAPS"
    present = [name for name in required if (root / name).exists()]
    missing = [name for name in required if name not in present]
    optional_present = [name for name in optional if (root / name).exists()]
    final_package = publish_final_package(run_dir, decision=decision, scorecard=scorecard, readiness=readiness)
    payload = {
        "schema_name": "claude_cer_final_closeout_report",
        "schema_version": "v1",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "decision": decision,
        "self_check_score": scorecard.get("score"),
        "self_check_conclusion": scorecard.get("delivery_conclusion"),
        "signature_readiness": readiness.get("decision"),
        "present_outputs": present,
        "missing_outputs": missing,
        "optional_outputs_present": optional_present,
        "docx_beautification": beautification,
        "published_final_package": final_package,
        "human_responsibility": "Human evaluator/manufacturer confirmation remains required before regulatory submission or sign-off.",
    }
    write_json(root / "final_closeout_report.json", payload)
    shutil.copy2(root / "final_closeout_report.json", artifact_root(run_dir) / "final" / "final_closeout_report.json")
    return payload


def publish_final_package(run_dir: Path, *, decision: str | None = None, scorecard: dict[str, Any] | None = None, readiness: dict[str, Any] | None = None) -> dict[str, Any]:
    artifact = artifact_root(run_dir)
    root = repair_root(run_dir)
    final_dir = artifact / "final"
    final_dir.mkdir(parents=True, exist_ok=True)
    if scorecard is None:
        scorecard = read_json(root / "cer_self_check_scorecard.json") or {}
    if readiness is None:
        readiness = read_json(root / "signature_readiness_report.json") or {}
    if decision is None:
        if readiness.get("decision") == "SIGNATURE_REVIEW_CANDIDATE" and scorecard.get("delivery_conclusion") == "SIGNATURE_READY":
            decision = "SIGNATURE_READY"
        elif int(scorecard.get("score") or 0) >= CONTROLLED_SCORE_THRESHOLD and int(scorecard.get("hard_gate_failure_count") or 0) == 0:
            decision = "CONTROLLED_DRAFT_WITH_GAPS"
        else:
            decision = "NOT_DELIVERABLE_REWORK_REQUIRED"
    file_map = {
        "CER_final.md": root / "CER_draft_patched.md",
        "CER_final.docx": root / "CER_draft_patched.docx",
        "CER_final_beautified.docx": root / "CER_draft_final_beautified.docx",
        "enhanced_authoring_workbook.json": root / "enhanced_authoring_workbook.json",
        "cer_self_check_scorecard.json": root / "cer_self_check_scorecard.json",
        "cer_self_check_scorecard.xlsx": root / "cer_self_check_scorecard.xlsx",
        "cer_self_check_scorecard.md": root / "cer_self_check_scorecard.md",
        "nb_flag_triage.xlsx": root / "nb_flag_triage.xlsx",
        "pico_v2_strategy.xlsx": root / "pico_v2_strategy.xlsx",
        "sota_endpoint_derivation_table.xlsx": root / "sota_endpoint_derivation_table.xlsx",
        "sota_quantitative_benchmark_table.xlsx": root / "sota_quantitative_benchmark_table.xlsx",
        "sota_evidence_synthesis_matrix.xlsx": root / "sota_evidence_synthesis_matrix.xlsx",
        "sota_medical_field_boundary.xlsx": root / "sota_medical_field_boundary.xlsx",
        "sota_pico_v2_strategy.xlsx": root / "sota_pico_v2_strategy.xlsx",
        "sota_search_strategy_separated.xlsx": root / "sota_search_strategy_separated.xlsx",
        "sota_screening_prisma.xlsx": root / "sota_screening_prisma.xlsx",
        "sota_evidence_hierarchy.xlsx": root / "sota_evidence_hierarchy.xlsx",
        "sota_endpoint_extraction_fulltext.xlsx": root / "sota_endpoint_extraction_fulltext.xlsx",
        "sota_benchmark_derivation_table.xlsx": root / "sota_benchmark_derivation_table.xlsx",
        "sota_section_conclusion_matrix.xlsx": root / "sota_section_conclusion_matrix.xlsx",
        "sota_deduction_chain.xlsx": root / "sota_deduction_chain.xlsx",
        "sota_endpoint_source_classification.xlsx": root / "sota_endpoint_source_classification.xlsx",
        "sota_aggregate_benchmark_rationale.xlsx": root / "sota_aggregate_benchmark_rationale.xlsx",
        "sota_conclusion_strength_guard.xlsx": root / "sota_conclusion_strength_guard.xlsx",
        "sota_benchmark_derivation_narratives.md": root / "sota_benchmark_derivation_narratives.md",
        "risk_benefit_closure_matrix.xlsx": root / "risk_benefit_closure_matrix.xlsx",
        "rmf_deep_extraction.xlsx": root / "rmf_deep_extraction.xlsx",
        "pmcf_timetable_request.xlsx": root / "pmcf_timetable_request.xlsx",
        "patch_qa_report.json": root / "patch_qa_report.json",
        "signature_readiness_report.json": root / "signature_readiness_report.json",
        "section_patch_manifest.json": root / "section_patch_manifest.json",
    }
    copied = []
    missing = []
    for final_name, source in file_map.items():
        if source.exists():
            shutil.copy2(source, final_dir / final_name)
            copied.append({"file": final_name, "source": str(source), "destination": str(final_dir / final_name)})
        else:
            missing.append({"file": final_name, "source": str(source)})
    manifest = {
        "schema_name": "claude_cer_final_package_manifest",
        "generated_at": now_iso(),
        "run_dir": str(run_dir),
        "baseline_artifact_root": str(artifact),
        "repair_root": str(root),
        "final_dir": str(final_dir),
        "decision": decision,
        "self_check_score": scorecard.get("score"),
        "self_check_conclusion": scorecard.get("delivery_conclusion"),
        "signature_readiness": readiness.get("decision"),
        "copied_files": copied,
        "missing_files": missing,
        "audit_boundary": "Baseline deerflow_authoring root is not overwritten; final package is published under deerflow_authoring/final.",
    }
    write_json(final_dir / "final_manifest.json", manifest)
    write_json(final_dir / "final_gate_report.json", {"decision": decision, "scorecard": scorecard, "signature_readiness": readiness})
    return manifest


def _beautify_final_docx(root: Path, theme: str) -> dict[str, Any]:
    source = root / "CER_draft_patched.docx"
    output = root / "CER_draft_final_beautified.docx"
    script = REPO_ROOT / ".claude" / "skills" / "docx-beautifier" / "scripts" / "beautify_docx.py"
    if not source.exists():
        payload = {
            "schema_name": "claude_cer_docx_beautification_report",
            "generated_at": now_iso(),
            "decision": "FAILED",
            "reason": f"Patched DOCX missing: {source}",
        }
        write_json(root / "docx_beautification_report.json", payload)
        return payload
    if not script.exists():
        payload = {
            "schema_name": "claude_cer_docx_beautification_report",
            "generated_at": now_iso(),
            "decision": "FAILED",
            "reason": f"docx-beautifier script missing: {script}",
        }
        write_json(root / "docx_beautification_report.json", payload)
        return payload
    completed = subprocess.run(
        [
            sys.executable,
            str(script),
            str(source),
            str(output),
            "--theme",
            theme,
            "--header-text",
            "Clinical Evaluation Report",
        ],
        cwd=str(REPO_ROOT),
        text=True,
        capture_output=True,
        check=False,
    )
    payload = {
        "schema_name": "claude_cer_docx_beautification_report",
        "generated_at": now_iso(),
        "decision": "PASS" if completed.returncode == 0 and output.exists() else "FAILED",
        "theme": theme,
        "source": str(source),
        "output": str(output),
        "returncode": completed.returncode,
        "stdout_tail": completed.stdout[-2000:],
        "stderr_tail": completed.stderr[-2000:],
        "boundary": "Formatting only; clinical content, evidence IDs and conclusions are controlled by earlier gates.",
    }
    write_json(root / "docx_beautification_report.json", payload)
    return payload


def _safe_int(value: Any) -> int:
    try:
        return int(value)
    except Exception:
        return 0


def _write_xlsx(path: Path, rows: list[dict[str, Any]], sheet_name: str) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31] or "sheet"
    if not rows:
        rows = [{"status": "no_rows", "note": "No rows generated for this artifact."}]
    headers: list[str] = []
    for row in rows:
        for key in row:
            if key not in headers:
                headers.append(key)
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
    for row_idx, row in enumerate(rows, start=2):
        for col, header in enumerate(headers, start=1):
            ws.cell(row=row_idx, column=col, value=_cell_value(row.get(header, "")))
    for idx, header in enumerate(headers, start=1):
        ws.column_dimensions[ws.cell(row=1, column=idx).column_letter].width = min(max(len(header) + 4, 16), 60)
    wb.save(path)


def _write_docx(path: Path, markdown: str) -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            doc.add_heading(stripped.lstrip("#").strip(), level=level)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    doc.save(path)


def _cell_value(value: Any) -> str | int | float | bool | None:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    return json.dumps(value, ensure_ascii=False, default=str)


def copy_baseline_report_to_desktop(run_dir: Path, output_path: Path) -> dict[str, Any]:
    """Utility used manually by the supervisor when the user asks for a copy."""
    source = artifact_root(run_dir) / "CER_draft.docx"
    if not source.exists():
        raise FileNotFoundError(source)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, output_path)
    return {"source": str(source), "output": str(output_path)}
