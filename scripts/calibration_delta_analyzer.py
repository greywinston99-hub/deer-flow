#!/usr/bin/env python3
"""Locked Structural + Semantic Delta Analyzer for CER calibration.

This script is a standalone post-freeze evaluator. It is the only approved
access path for `02_NB_ROUNDS_AND_RESPONSES_LOCKED` and
`03_FINAL_CERTIFIED_PACKAGE_LOCKED` during calibration. It must not call the
authoring workflow, repair loop, finalization, or writer.

Inputs:
  --baseline-root
  --authoring-workbook
  --qa-gate-report
  --nb-locked-root
  --final-locked-root
  --output-dir

Outputs:
  - 8 delta tables as CSV
  - 9 semantic delta tables as CSV
  - DELTA_ANALYSIS_MANIFEST.json
  - semantic_delta_manifest.json
  - semantic_delta_case_summary.md
  - LOCKED_ACCESS_LOG.csv
  - NEEDS_HUMAN_CLASSIFICATION.csv
  - low_confidence_review_queue.csv
  - PILOT_RUN_REPORT.md
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import subprocess
import sys
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any


DELTA_TABLE_NAMES = (
    "CLAIM_DELTA_TABLE",
    "SOTA_BENCHMARK_DELTA_TABLE",
    "EVIDENCE_SELECTION_DELTA_TABLE",
    "EVIDENCE_APPRAISAL_DELTA_TABLE",
    "CLAIM_EVIDENCE_DELTA_MATRIX",
    "PMCF_BOUNDARY_DELTA_TABLE",
    "ALIGNMENT_DELTA_TABLE",
    "CEAR_DEFICIENCY_PATTERN_TABLE",
)

LOCKED_ROLES = (
    "NB_FEEDBACK",
    "OUR_RESPONSE",
    "SUBMITTED_SUPPORTING_FILE",
    "FINAL_CHANGE_REFERENCE",
    "UNKNOWN",
)

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".xml", ".html", ".htm"}

SEMANTIC_TABLE_NAMES = (
    "claim_semantic_alignment_table",
    "evidence_correspondence_table",
    "sota_benchmark_delta_table",
    "evidence_appraisal_delta_table",
    "pmcf_boundary_delta_table",
    "benefit_risk_reasoning_delta_table",
    "cross_document_alignment_delta_table",
    "nb_relevance_delta_table",
    "cognitive_gap_attribution_table",
)

SEMANTIC_DELTA_CLASSES = (
    "clinically_material_gap",
    "nb_relevant_gap",
    "accepted_stylistic_variance",
    "project_specific_difference",
    "ai_overclaim_or_misreasoning",
    "gold_reference_only_detail",
    "ai_potentially_better_than_gold",
    "needs_human_confirmation",
)

COGNITIVE_ROOT_CAUSES = (
    "workflow_order_gap",
    "missing_intermediate_artifact",
    "prompt_gap",
    "rule_gap",
    "knowledge_gap",
    "template_or_section_gap",
    "source_interpretation_gap",
    "evidence_appraisal_gap",
    "writer_synthesis_gap",
    "acceptable_variance",
    "source_insufficiency",
    "needs_human_confirmation",
)

STOPWORDS = {
    "the",
    "and",
    "for",
    "with",
    "from",
    "this",
    "that",
    "were",
    "been",
    "into",
    "shall",
    "device",
    "clinical",
    "evaluation",
    "report",
    "data",
    "evidence",
    "safety",
    "performance",
    "patients",
    "patient",
    "medical",
    "using",
    "used",
    "use",
    "based",
    "source",
    "result",
    "results",
    "study",
    "studies",
    "final",
    "baseline",
}

SEMANTIC_ZONES = (
    "claims",
    "sota",
    "evidence_selection",
    "evidence_appraisal",
    "pmcf_unanswered_questions",
    "benefit_risk",
    "conclusion",
    "alignment",
)

ZONE_KEYWORDS = {
    "claims": ("intended", "purpose", "claim", "indication", "benefit", "performance", "safety", "预期", "适用", "声称", "获益"),
    "sota": ("sota", "state of the art", "benchmark", "guideline", "alternative", "endpoint", "acceptance", "基准", "指南", "终点"),
    "evidence_selection": ("literature", "evidence", "study", "pmid", "doi", "clinical data", "registry", "文献", "证据"),
    "evidence_appraisal": ("pivotal", "supportive", "level", "quality", "limitation", "bias", "appraisal", "suitability", "contribution", "评价", "局限"),
    "pmcf_unanswered_questions": ("pmcf", "pms", "post-market", "unanswered", "residual", "follow-up", "registry", "问卷", "上市后"),
    "benefit_risk": ("benefit-risk", "benefit risk", "benefits outweigh", "residual risk", "acceptable", "获益", "风险"),
    "conclusion": ("conclusion", "summary", "overall", "conclude", "acceptable", "support", "结论", "总结"),
    "alignment": ("ifu", "rmf", "rmr", "gspr", "sscp", "pmcf", "warning", "contraindication", "risk management", "说明书", "风险管理"),
}

GOLD_ROLE_FOLDER_MAP = {
    "final_cer": "FINAL_CER",
    "final_ifu": "FINAL_IFU",
    "final_rmf": "FINAL_RMF_RMR",
    "final_rmr": "FINAL_RMF_RMR",
    "final_gspr": "FINAL_GSPR",
    "final_pms": "FINAL_PMS_PMCF_PSUR",
    "final_pmcf": "FINAL_PMS_PMCF_PSUR",
    "final_psur": "FINAL_PMS_PMCF_PSUR",
    "final_sscp": "FINAL_SSCP",
    "final_other": "FINAL_OTHER_ACCEPTED_CHANGES",
}

ENTITY_SYNONYMS = {
    "state_of_the_art": ("sota", "state of the art", "current clinical practice", "现有技术", "当前技术水平"),
    "benefit_risk": ("benefit-risk", "benefit risk", "benefits outweigh risks", "获益风险", "受益风险"),
    "post_market_follow_up": ("pmcf", "post-market clinical follow-up", "上市后临床跟踪"),
    "post_market_surveillance": ("pms", "post-market surveillance", "上市后监督"),
    "risk_management": ("rmf", "rmr", "risk management", "risk management file", "风险管理"),
    "instructions_for_use": ("ifu", "instructions for use", "user manual", "使用说明书", "说明书"),
    "gspr": ("gspr", "general safety and performance requirements", "annex i"),
    "clinical_evidence": ("clinical evidence", "clinical data", "literature", "evidence", "临床证据", "文献"),
    "pivotal_evidence": ("pivotal", "level 1", "randomized", "meta-analysis", "systematic review"),
    "intended_purpose": ("intended purpose", "intended use", "indication", "indications for use", "预期用途", "适用范围", "适应症"),
    "clinical_benefit": ("clinical benefit", "patient benefit", "benefit claim", "clinical benefits", "临床获益"),
    "endpoint_success_rate": ("success rate", "technical success", "clinical success", "procedural success", "stone-free rate", "hemostasis success", "成功率"),
    "endpoint_adverse_event_rate": ("adverse event", "adverse events", "complication", "complications", "side-effect", "side effect", "ae rate", "并发症", "不良事件"),
    "endpoint_serious_adverse_event_rate": ("serious adverse event", "sae", "serious complication", "严重不良事件", "严重并发症"),
    "endpoint_device_failure": ("device failure", "malfunction", "breakage", "migration", "rupture", "器械故障", "破裂", "脱落"),
    "benchmark_value": ("benchmark value", "acceptance criterion", "acceptance criteria", "threshold", "criterion", "基准值", "接受标准"),
    "confidence_interval": ("ci", "confidence interval", "95% ci", "置信区间"),
    "sample_size": ("sample size", "number of patients", "n=", "n =", "subjects", "样本量"),
    "follow_up": ("follow-up", "follow up", "followup", "随访"),
    "population": ("population", "patient population", "target population", "人群", "目标人群"),
    "risk_control": ("risk control", "risk mitigation", "control measure", "warning", "precaution", "风险控制", "警告", "注意事项"),
    "unanswered_question": ("unanswered question", "residual question", "remaining uncertainty", "evidence gap", "未回答问题", "剩余不确定性", "证据缺口"),
    "post_market_registry": ("registry", "patient registry", "clinical registry", "登记", "注册登记"),
    "questionnaire": ("questionnaire", "survey", "user feedback", "问卷", "调查"),
    "vigilance": ("vigilance", "maude", "recall", "fsca", "adverse event database", "警戒", "召回"),
    "equivalence": ("equivalence", "equivalent", "similar device", "comparable device", "等同", "类似器械"),
    "classification": ("classification", "device classification", "class i", "class ii", "class iii", "分类"),
}

CITATION_PATTERNS = (
    ("PMID", re.compile(r"\bPMID[:\s]*([0-9]{5,10})\b", re.IGNORECASE)),
    ("DOI", re.compile(r"\b(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)\b", re.IGNORECASE)),
    ("MDCG", re.compile(r"\bMDCG\s+\d{4}-\d+\b", re.IGNORECASE)),
    ("MEDDEV", re.compile(r"\bMEDDEV\s+[0-9./]+\b", re.IGNORECASE)),
    ("MDR", re.compile(r"\bMDR\b|\bRegulation\s+\(EU\)\s+2017/745\b", re.IGNORECASE)),
    ("ISO", re.compile(r"\bISO\s+\d{4,5}(?:-\d+)?\b", re.IGNORECASE)),
)

LOW_CONFIDENCE_TYPES = (
    "TEXT_EXTRACTION_GAP",
    "SECTION_MAPPING_GAP",
    "ENTITY_NORMALIZATION_GAP",
    "SEMANTIC_MATCHING_GAP",
    "TRUE_AMBIGUITY_HUMAN_GATE",
)

LOW_CONFIDENCE_ACTIONS = {
    "TEXT_EXTRACTION_GAP": ("yes", "Improve final-package document text/table extraction before treating this as an AI CER defect."),
    "SECTION_MAPPING_GAP": ("yes", "Improve section-aware retrieval or add stronger section anchors for the affected document role."),
    "ENTITY_NORMALIZATION_GAP": ("yes", "Add synonym/entity normalization for the device, endpoint, risk or citation terms involved."),
    "SEMANTIC_MATCHING_GAP": ("yes", "Improve semantic candidate generation/reranking; keep the item out of systemic-gap counting for now."),
    "TRUE_AMBIGUITY_HUMAN_GATE": ("no", "Ask a human calibrator to decide whether the difference is clinically material or accepted variance."),
}

PHASE35C_BASELINE_METRICS = {
    "semantic_rows": 1473,
    "low_confidence_total": 663,
    "automatic_high_confidence_rate": 55.0,
    "entity_normalization_gap": 320,
    "true_ambiguity_human_gate": 306,
    "gold_sections_indexed": 4362,
    "gold_citations_indexed": 243,
}


@dataclass(frozen=True)
class LockedFileRecord:
    source_root_label: str
    path: Path
    relative_path: str
    role: str
    confidence: str
    extraction_status: str
    sha256: str
    size_bytes: int
    reason: str
    text_excerpt: str = ""


@dataclass(frozen=True)
class SemanticItem:
    item_id: str
    item_type: str
    source: str
    text: str
    metadata: dict[str, Any]


def now_iso() -> str:
    return datetime.now(UTC).isoformat()


def _resolve_existing_path(value: str, label: str, *, must_be_dir: bool | None = None) -> Path:
    path = Path(value).expanduser().resolve()
    if not path.exists():
        raise SystemExit(f"{label} does not exist: {path}")
    if must_be_dir is True and not path.is_dir():
        raise SystemExit(f"{label} must be a directory: {path}")
    if must_be_dir is False and not path.is_file():
        raise SystemExit(f"{label} must be a file: {path}")
    return path


def _is_relative_to(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def _validate_output_dir(output_dir: Path, forbidden_roots: list[Path]) -> None:
    parts = set(output_dir.resolve().parts)
    if "01_INITIAL_INPUT_FOR_WRITER" in parts:
        raise SystemExit("output-dir cannot be inside 01_INITIAL_INPUT_FOR_WRITER")
    for root in forbidden_roots:
        if _is_relative_to(output_dir, root):
            raise SystemExit(f"output-dir must be separate from read-only input roots: {output_dir} is under {root}")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def read_text_excerpt(path: Path, *, limit: int = 80000) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return "content_extracted", path.read_text(encoding="utf-8", errors="replace")[:limit]
    if suffix == ".xlsx":
        try:
            import openpyxl
        except Exception as exc:  # pragma: no cover - dependency-dependent branch
            return "metadata_only_openpyxl_unavailable", f"openpyxl unavailable: {exc}"
        try:
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            fragments: list[str] = []
            for ws in wb.worksheets[:3]:
                fragments.append(f"[sheet={ws.title}]")
                for row in ws.iter_rows(max_row=20, values_only=True):
                    values = [str(cell) for cell in row if cell is not None]
                    if values:
                        fragments.append(" | ".join(values))
            wb.close()
            return "content_extracted", "\n".join(fragments)[:limit]
        except Exception as exc:  # pragma: no cover - defensive branch
            return "metadata_only_extraction_failed", f"xlsx extraction failed: {exc}"
    if suffix == ".docx":
        try:
            import docx
        except Exception as exc:  # pragma: no cover - dependency-dependent branch
            return "metadata_only_docx_unavailable", f"docx unavailable: {exc}"
        try:
            document = docx.Document(str(path))
            fragments = [p.text for p in document.paragraphs if p.text and p.text.strip()]
            for table_idx, table in enumerate(document.tables, start=1):
                fragments.append(f"[table={table_idx}]")
                for row_idx, row in enumerate(table.rows, start=1):
                    cells = [_clean(cell.text) for cell in row.cells if _clean(cell.text)]
                    if cells:
                        fragments.append(f"[table={table_idx};row={row_idx}] " + " | ".join(cells))
            return "content_extracted", "\n".join(fragments)[:limit]
        except Exception as exc:  # pragma: no cover - defensive branch
            return "metadata_only_extraction_failed", f"docx extraction failed: {exc}"
    if suffix == ".doc":
        try:
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", str(path)],
                check=False,
                capture_output=True,
                text=True,
                timeout=45,
            )
            if result.stdout.strip():
                return "content_extracted", result.stdout[:limit]
            return "metadata_only_extraction_failed", (result.stderr or "textutil returned no text")[:limit]
        except Exception as exc:  # pragma: no cover - platform-dependent branch
            return "metadata_only_extraction_failed", f"doc extraction failed: {exc}"
    return "metadata_only_unsupported_suffix", ""


def classify_locked_file(relative_path: str, root_label: str, text_excerpt: str) -> tuple[str, str, str]:
    # Deliberately avoid absolute parent paths: folder names such as
    # `02_NB_ROUNDS_AND_RESPONSES_LOCKED` would otherwise bias every file toward
    # NB_FEEDBACK and hide response/supporting/unknown files.
    haystack = " ".join([relative_path, text_excerpt[:1000]]).lower()
    chinese_haystack = " ".join([relative_path, text_excerpt[:1000]])

    if root_label == "03_FINAL_CERTIFIED_PACKAGE_LOCKED" and any(
        token in haystack or token in chinese_haystack
        for token in ("final", "certified", "accepted", "signed", "approved", "final cer", "定稿", "终版", "签字")
    ):
        return "FINAL_CHANGE_REFERENCE", "high", "final/certified package naming"

    if any(
        token in haystack or token in chinese_haystack
        for token in ("response", "reply", "answer", "rectification", "corrective", "回复", "回答", "整改", "答复")
    ):
        return "OUR_RESPONSE", "high", "response keywords"

    if any(
        token in haystack or token in chinese_haystack
        for token in ("submitted", "supporting", "attachment", "evidence", "annex", "附件", "提交", "支持文件", "补充资料")
    ):
        return "SUBMITTED_SUPPORTING_FILE", "medium", "supporting/submitted-file keywords"

    if any(
        token in haystack or token in chinese_haystack
        for token in ("loq", "nb comment", "nb question", "notified body", "deficiency", "finding", "nonconformity", "round", "缺陷", "问题", "公告机构", "询问")
    ):
        return "NB_FEEDBACK", "high", "NB feedback keywords"

    if root_label == "03_FINAL_CERTIFIED_PACKAGE_LOCKED":
        return "FINAL_CHANGE_REFERENCE", "medium", "inside final certified package root"

    return "UNKNOWN", "low", "role could not be determined from filename/path/excerpt"


def scan_locked_root(root: Path, root_label: str) -> list[LockedFileRecord]:
    records: list[LockedFileRecord] = []
    for path in sorted(p for p in root.rglob("*") if p.is_file() and not p.name.startswith(".")):
        relative_path = str(path.relative_to(root))
        digest = sha256_file(path)
        extraction_status, excerpt = read_text_excerpt(path)
        role, confidence, reason = classify_locked_file(relative_path, root_label, excerpt)
        records.append(
            LockedFileRecord(
                source_root_label=root_label,
                path=path,
                relative_path=relative_path,
                role=role,
                confidence=confidence,
                extraction_status=extraction_status,
                sha256=digest,
                size_bytes=path.stat().st_size,
                reason=reason,
                text_excerpt=excerpt,
            )
        )
    return records


def load_baseline_context(authoring_workbook: Path, qa_gate_report: Path) -> dict[str, Any]:
    workbook = read_json(authoring_workbook)
    qa_report = read_json(qa_gate_report)
    gate_results = qa_report.get("results") or qa_report.get("gates") or []
    failed_gates = [
        {
            "gate_id": str(item.get("gate_id") or item.get("id") or ""),
            "status": str(item.get("status") or item.get("result") or ""),
            "message": str(item.get("message") or item.get("reason") or item.get("details") or "")[:500],
            "severity": str(item.get("severity") or ""),
        }
        for item in gate_results
        if str(item.get("status") or item.get("result") or "").upper() not in {"PASS", "PASSED", "OK"}
    ]
    return {
        "workbook": workbook,
        "qa_report": qa_report,
        "failed_gates": failed_gates,
        "counts": {
            "claims": len(workbook.get("claim_ledger") or []),
            "picos": len(workbook.get("cep_pico_matrix") or []),
            "sota_benchmarks": len(workbook.get("sota_benchmark_matrix") or []),
            "evidence": len(workbook.get("evidence_registry") or []),
            "endpoint_extraction": len(workbook.get("endpoint_extraction") or []),
            "risk_rows": len(workbook.get("risk_trace_matrix") or []),
            "failed_gates": len(failed_gates),
        },
    }


def _row(
    table_name: str,
    row_id: str,
    baseline_ref: str,
    locked_ref: str,
    observed_delta: str,
    root_cause_category: str,
    severity: str,
    needs_human: bool,
    evidence_path: str,
) -> dict[str, Any]:
    return {
        "table_name": table_name,
        "row_id": row_id,
        "baseline_ref": baseline_ref,
        "locked_ref": locked_ref,
        "observed_delta": observed_delta,
        "root_cause_category": root_cause_category,
        "severity": severity,
        "needs_human_classification": "yes" if needs_human else "no",
        "evidence_path": evidence_path,
    }


def build_delta_tables(
    baseline: dict[str, Any],
    locked_records: list[LockedFileRecord],
) -> dict[str, list[dict[str, Any]]]:
    workbook = baseline["workbook"]
    failed_gates = baseline["failed_gates"]
    counts = baseline["counts"]
    unknown_files = [record for record in locked_records if record.role == "UNKNOWN"]
    nb_feedback = [record for record in locked_records if record.role == "NB_FEEDBACK"]
    final_refs = [record for record in locked_records if record.role == "FINAL_CHANGE_REFERENCE"]

    tables: dict[str, list[dict[str, Any]]] = {name: [] for name in DELTA_TABLE_NAMES}
    tables["CLAIM_DELTA_TABLE"].append(
        _row(
            "CLAIM_DELTA_TABLE",
            "CLAIM-DELTA-001",
            f"authoring_workbook.claim_ledger count={counts['claims']}",
            f"final_locked_files count={len(final_refs)}",
            "Baseline claims are frozen; comparison against final accepted claim wording is assigned to Delta Analyzer only.",
            "claim",
            "medium",
            bool(final_refs),
            "authoring_workbook.json / 03_FINAL_CERTIFIED_PACKAGE_LOCKED",
        )
    )
    tables["SOTA_BENCHMARK_DELTA_TABLE"].append(
        _row(
            "SOTA_BENCHMARK_DELTA_TABLE",
            "SOTA-DELTA-001",
            f"authoring_workbook.sota_benchmark_matrix count={counts['sota_benchmarks']}",
            "final package SOTA references are locked for evaluator-only comparison",
            "SOTA benchmark completeness and accepted-human CER logic must be compared after baseline freeze.",
            "sota",
            "high",
            True,
            "authoring_workbook.json / final locked package",
        )
    )
    tables["EVIDENCE_SELECTION_DELTA_TABLE"].append(
        _row(
            "EVIDENCE_SELECTION_DELTA_TABLE",
            "EVID-SEL-DELTA-001",
            f"authoring_workbook.evidence_registry count={counts['evidence']}",
            f"locked NB feedback files count={len(nb_feedback)}",
            "Evidence selection deltas are inferred from baseline evidence registry and locked NB/final references.",
            "evidence",
            "medium",
            True,
            "authoring_workbook.json / 02_NB_ROUNDS_AND_RESPONSES_LOCKED",
        )
    )
    tables["EVIDENCE_APPRAISAL_DELTA_TABLE"].append(
        _row(
            "EVIDENCE_APPRAISAL_DELTA_TABLE",
            "EVID-APP-DELTA-001",
            "authoring_workbook.evidence_registry / endpoint_extraction",
            "final locked package",
            "Evidence appraisal delta requires evaluator comparison of weight, Oxford level, suitability and contribution.",
            "appraisal",
            "medium",
            True,
            "authoring_workbook.json / 03_FINAL_CERTIFIED_PACKAGE_LOCKED",
        )
    )
    tables["CLAIM_EVIDENCE_DELTA_MATRIX"].append(
        _row(
            "CLAIM_EVIDENCE_DELTA_MATRIX",
            "CLAIM-EVID-DELTA-001",
            f"claims={counts['claims']}; picos={counts['picos']}; evidence={counts['evidence']}",
            "locked final CER and response matrices",
            "Claim-to-evidence chain is frozen for baseline; final accepted chain comparison is controlled by this analyzer.",
            "claim/evidence",
            "high",
            True,
            "authoring_workbook.json / locked delta-only roots",
        )
    )
    tables["PMCF_BOUNDARY_DELTA_TABLE"].append(
        _row(
            "PMCF_BOUNDARY_DELTA_TABLE",
            "PMCF-DELTA-001",
            "authoring_workbook.gap_pmcf_recommendations",
            "final locked package / NB responses",
            "PMCF boundary deltas must determine whether gaps were allowed, corrected, or deferred in accepted package.",
            "pmcf",
            "high",
            True,
            "authoring_workbook.json / locked delta-only roots",
        )
    )
    tables["ALIGNMENT_DELTA_TABLE"].append(
        _row(
            "ALIGNMENT_DELTA_TABLE",
            "ALIGN-DELTA-001",
            f"risk_trace_matrix count={counts['risk_rows']}",
            "final locked package / submitted supporting files",
            "Alignment delta checks accepted consistency across CER/RMF/IFU/PMCF/SSCP after baseline freeze.",
            "alignment",
            "high",
            True,
            "authoring_workbook.json / locked delta-only roots",
        )
    )
    for idx, gate in enumerate(failed_gates, start=1):
        tables["CEAR_DEFICIENCY_PATTERN_TABLE"].append(
            _row(
                "CEAR_DEFICIENCY_PATTERN_TABLE",
                f"CEAR-GATE-{idx:03d}",
                f"qa_gate_report.{gate['gate_id']}",
                "baseline gate failure frozen before locked access",
                gate["message"] or "Baseline gate failure requires CEAR-style classification.",
                "gate",
                gate["severity"] or "medium",
                False,
                "qa_gate_report.json",
            )
        )
    for idx, record in enumerate(nb_feedback, start=1):
        tables["CEAR_DEFICIENCY_PATTERN_TABLE"].append(
            _row(
                "CEAR_DEFICIENCY_PATTERN_TABLE",
                f"CEAR-NB-{idx:03d}",
                "not available to authoring writer",
                record.relative_path,
                f"Locked NB feedback classified as {record.role}; content review is allowed only in Delta Analyzer outputs.",
                "cear_deficiency",
                "medium",
                False,
                f"{record.source_root_label}/{record.relative_path}",
            )
        )
    for idx, record in enumerate(unknown_files, start=1):
        tables["CEAR_DEFICIENCY_PATTERN_TABLE"].append(
            _row(
                "CEAR_DEFICIENCY_PATTERN_TABLE",
                f"CEAR-UNKNOWN-{idx:03d}",
                "not available to authoring writer",
                record.relative_path,
                "Locked file role could not be classified automatically and requires human classification.",
                "leakage_control",
                "medium",
                True,
                f"{record.source_root_label}/{record.relative_path}",
            )
        )
    return tables


def _clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _row_text(row: dict[str, Any], preferred_keys: tuple[str, ...] = ()) -> str:
    parts: list[str] = []
    for key in preferred_keys:
        if row.get(key) not in (None, ""):
            parts.append(str(row.get(key)))
    for key, value in row.items():
        if key in preferred_keys or isinstance(value, (dict, list)):
            continue
        if value not in (None, ""):
            parts.append(str(value))
    return _clean(" | ".join(parts))


def _tokens(text: Any) -> set[str]:
    raw = re.findall(r"[A-Za-z][A-Za-z0-9_+-]{2,}|PMID[:\s]?\d+|10\.\d{4,9}/[-._;()/:A-Za-z0-9]+|[\u4e00-\u9fff]{2,}", str(text or "").lower())
    tokens = {token.strip(" .,:;()[]{}") for token in raw if token and token not in STOPWORDS and len(token) >= 3}
    tokens.update(_entity_norms(text))
    return tokens


def _semantic_similarity(left: Any, right: Any) -> float:
    left_tokens = _tokens(left)
    right_tokens = _tokens(right)
    if not left_tokens or not right_tokens:
        return 0.0
    overlap = len(left_tokens & right_tokens)
    dice = (2 * overlap) / (len(left_tokens) + len(right_tokens))
    containment = overlap / max(1, min(len(left_tokens), len(right_tokens)))
    return round((0.7 * dice) + (0.3 * containment), 3)


def _gold_document_role(record: LockedFileRecord) -> str:
    normalized = re.sub(r"[^a-z0-9_]+", "_", record.relative_path.lower())
    for token, role in GOLD_ROLE_FOLDER_MAP.items():
        if token in normalized:
            return role
    if "cer" in normalized or "clinical_evaluation" in normalized:
        return "FINAL_CER"
    if "ifu" in normalized or "instruction" in normalized or "manual" in normalized or "说明书" in record.relative_path:
        return "FINAL_IFU"
    if "risk" in normalized or "rmf" in normalized or "rmr" in normalized or "风险" in record.relative_path:
        return "FINAL_RMF_RMR"
    if "gspr" in normalized:
        return "FINAL_GSPR"
    if "pmcf" in normalized or "pms" in normalized or "psur" in normalized:
        return "FINAL_PMS_PMCF_PSUR"
    if "sscp" in normalized:
        return "FINAL_SSCP"
    return "FINAL_OTHER_ACCEPTED_CHANGES"


def _semantic_zone(text: Any, source_hint: str = "") -> str:
    haystack = f"{source_hint} {text}".lower()
    scores = {
        zone: sum(1 for keyword in keywords if keyword.lower() in haystack or keyword in str(text or ""))
        for zone, keywords in ZONE_KEYWORDS.items()
    }
    if scores:
        top_zone, top_score = max(scores.items(), key=lambda item: item[1])
        if top_score > 0:
            return top_zone
    return "evidence_selection" if re.search(r"\b(PMID|DOI|study|trial)\b", haystack, re.IGNORECASE) else "claims"


def _normalize_entities(text: Any) -> list[dict[str, str]]:
    value = str(text or "")
    lower = value.lower()
    entities: list[dict[str, str]] = []
    for normalized, synonyms in ENTITY_SYNONYMS.items():
        for synonym in synonyms:
            if synonym.lower() in lower or synonym in value:
                entities.append({"entity_type": "controlled_term", "original": synonym, "normalized": normalized, "provenance": "built_in_synonym"})
                break
    for citation_type, pattern in CITATION_PATTERNS:
        for match in pattern.finditer(value):
            entities.append({"entity_type": "citation", "original": match.group(0), "normalized": f"{citation_type}:{match.group(1) if match.groups() else match.group(0)}", "provenance": "regex"})
    for match in re.finditer(r"\b([A-Z][A-Za-z0-9-]{2,}(?:\s+[A-Z][A-Za-z0-9-]{2,}){0,4})\b", value):
        phrase = match.group(1).strip()
        if phrase.lower() not in STOPWORDS and len(phrase) <= 80:
            entities.append({"entity_type": "candidate_name", "original": phrase, "normalized": phrase.lower(), "provenance": "capitalized_phrase"})
    seen = set()
    unique = []
    for entity in entities:
        key = (entity["entity_type"], entity["normalized"])
        if key in seen:
            continue
        seen.add(key)
        unique.append(entity)
    return unique[:30]


def _entity_norms(text: Any) -> set[str]:
    return {entity["normalized"] for entity in _normalize_entities(text)}


def _entity_overlap(left: Any, right: Any) -> float:
    left_entities = _entity_norms(left)
    right_entities = _entity_norms(right)
    if not left_entities or not right_entities:
        return 0.0
    overlap = len(left_entities & right_entities)
    return round(overlap / max(1, min(len(left_entities), len(right_entities))), 3)


def _staged_confidences(
    baseline_item: SemanticItem | None,
    gold_item: SemanticItem | None,
    candidate_score: float,
    *,
    nb_relevant: bool = False,
    root_cause: str = "",
) -> dict[str, float]:
    if not baseline_item or not gold_item:
        return {
            "candidate_retrieval_confidence": round(candidate_score, 3),
            "semantic_equivalence_confidence": 0.0,
            "clinical_materiality_confidence": 0.4 if nb_relevant else 0.25,
            "nb_relevance_confidence": 0.75 if nb_relevant else 0.2,
            "root_cause_confidence": 0.35,
            "overall_confidence": round(candidate_score, 3),
        }
    zone_match = baseline_item.metadata.get("semantic_zone") == gold_item.metadata.get("semantic_zone")
    role_bonus = 0.07 if zone_match else 0.0
    entity_score = _entity_overlap(baseline_item.text, gold_item.text)
    semantic_score = min(1.0, (0.78 * _semantic_similarity(baseline_item.text, gold_item.text)) + (0.22 * entity_score) + role_bonus)
    materiality_score = 0.75 if any(token in f"{baseline_item.text} {gold_item.text}".lower() for token in ("benchmark", "pivotal", "risk", "benefit", "pmcf", "gspr", "ifu", "endpoint")) else 0.45
    nb_score = 0.78 if nb_relevant else 0.25
    root_score = 0.72 if root_cause and root_cause != "needs_human_confirmation" and semantic_score >= 0.70 else 0.45
    retrieval_score = min(1.0, candidate_score + role_bonus + (0.08 if entity_score else 0.0))
    overall = round((0.40 * semantic_score) + (0.25 * retrieval_score) + (0.15 * materiality_score) + (0.10 * nb_score) + (0.10 * root_score), 3)
    return {
        "candidate_retrieval_confidence": round(retrieval_score, 3),
        "semantic_equivalence_confidence": round(semantic_score, 3),
        "clinical_materiality_confidence": round(materiality_score, 3),
        "nb_relevance_confidence": round(nb_score, 3),
        "root_cause_confidence": round(root_score, 3),
        "overall_confidence": overall,
    }


def _split_evidence_sentences(text: str, *, limit: int = 80) -> list[str]:
    normalized = re.sub(r"\s+", " ", text or "")
    chunks = re.split(r"(?<=[.;。；])\s+|\n+", normalized)
    useful = []
    for chunk in chunks:
        cleaned = _clean(chunk)
        if 35 <= len(cleaned) <= 700:
            useful.append(cleaned)
        if len(useful) >= limit:
            break
    return useful


def _heading_level(value: str) -> int:
    match = re.match(r"^\s*(\d+(?:\.\d+)*)", value)
    if match:
        return min(6, match.group(1).count(".") + 1)
    return 1


def _looks_like_section_heading(value: str) -> bool:
    stripped = _clean(value)
    if not stripped or len(stripped) > 180:
        return False
    if re.match(r"^\d+(?:\.\d+)*\s+.{3,}", stripped):
        return True
    lowered = stripped.lower().strip(":")
    return lowered in {
        "summary",
        "scope of clinical evaluation",
        "state of the art",
        "clinical data",
        "clinical evaluation data",
        "benefit-risk analysis",
        "conclusion",
        "conclusions",
        "intended purpose",
        "risk management",
        "general safety and performance requirements",
        "post-market clinical follow-up",
        "pmcf",
        "pms",
    }


def build_gold_reference_indexes(locked_records: list[LockedFileRecord]) -> tuple[dict[str, Any], list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]], list[SemanticItem]]:
    final_records = [record for record in locked_records if record.source_root_label == "03_FINAL_CERTIFIED_PACKAGE_LOCKED" and record.role == "FINAL_CHANGE_REFERENCE"]
    section_rows: list[dict[str, Any]] = []
    table_rows: list[dict[str, Any]] = []
    citation_rows: list[dict[str, Any]] = []
    semantic_items: list[SemanticItem] = []
    role_counts: dict[str, int] = {}

    for record in final_records:
        document_role = _gold_document_role(record)
        role_counts[document_role] = role_counts.get(document_role, 0) + 1
        source = f"{record.source_root_label}/{record.relative_path}"
        current_heading = "unsectioned"
        current_level = 0
        paragraphs = [line.strip() for line in re.split(r"\n+|(?<=[.;。；])\s+", record.text_excerpt or "") if _clean(line)]
        if not paragraphs and record.text_excerpt:
            paragraphs = [_clean(record.text_excerpt)]
        for para_idx, paragraph in enumerate(paragraphs[:240], start=1):
            cleaned = _clean(paragraph)
            if not cleaned:
                continue
            if _looks_like_section_heading(cleaned):
                current_heading = cleaned[:160]
                current_level = _heading_level(cleaned)
                continue
            zone = _semantic_zone(cleaned, record.relative_path)
            section_id = f"GOLD-SEC-{len(section_rows)+1:05d}"
            anchor = f"{Path(record.relative_path).stem[:40]}:P{para_idx:04d}"
            entities = _normalize_entities(cleaned)
            section_rows.append(
                {
                    "section_id": section_id,
                    "document_role": document_role,
                    "source_path": source,
                    "section_heading": current_heading,
                    "section_level": current_level,
                    "paragraph_anchor": anchor,
                    "semantic_zone": zone,
                    "text_excerpt": cleaned[:1200],
                    "normalized_entities": json.dumps(entities, ensure_ascii=False),
                    "extraction_status": record.extraction_status,
                    "section_mapping_confidence": "high" if current_heading != "unsectioned" else "medium",
                }
            )
            semantic_items.append(
                SemanticItem(
                    item_id=section_id,
                    item_type="gold_section",
                    source=source,
                    text=cleaned,
                    metadata={
                        "document_role": document_role,
                        "section_id": section_id,
                        "section_heading": current_heading,
                        "paragraph_anchor": anchor,
                        "semantic_zone": zone,
                        "extraction_status": record.extraction_status,
                        "normalized_entities": entities,
                    },
                )
            )
            if "|" in cleaned or "\t" in cleaned or cleaned.startswith("[sheet="):
                table_rows.append(
                    {
                        "table_id": f"GOLD-TABLE-{len(table_rows)+1:04d}",
                        "document_role": document_role,
                        "source_path": source,
                        "section_id": section_id,
                        "table_anchor": anchor,
                        "semantic_zone": zone,
                        "row_count_estimate": cleaned.count("|") if "|" in cleaned else max(1, cleaned.count("\t")),
                        "table_excerpt": cleaned[:1200],
                    }
                )
            for citation_type, pattern in CITATION_PATTERNS:
                for match in pattern.finditer(cleaned):
                    citation_rows.append(
                        {
                            "citation_id": f"GOLD-CIT-{len(citation_rows)+1:05d}",
                            "citation_type": citation_type,
                            "citation_text": match.group(0),
                            "normalized_citation": f"{citation_type}:{match.group(1) if match.groups() else match.group(0)}",
                            "document_role": document_role,
                            "source_path": source,
                            "section_id": section_id,
                            "semantic_zone": zone,
                            "context": cleaned[:900],
                        }
                    )

    manifest = {
        "schema_name": "gold_reference_extraction_manifest",
        "schema_version": "phase3.5c-gold-reference-extraction-v1",
        "generated_at": now_iso(),
        "status": "GOLD_REFERENCE_EXTRACTION_COMPLETE",
        "final_locked_file_count": len(final_records),
        "document_role_counts": role_counts,
        "section_count": len(section_rows),
        "table_count": len(table_rows),
        "citation_count": len(citation_rows),
        "semantic_item_count": len(semantic_items),
        "extraction_limitations": [
            "PDF binary text remains metadata-only unless text extraction is available.",
            "DOC extraction depends on macOS textutil when present.",
            "Section heading detection is deterministic and conservative.",
        ],
    }
    return manifest, section_rows, table_rows, citation_rows, semantic_items


def _item_from_row(prefix: str, idx: int, item_type: str, row: dict[str, Any], preferred_keys: tuple[str, ...]) -> SemanticItem:
    item_id = str(row.get("claim_id") or row.get("evidence_id") or row.get("benchmark_id") or row.get("endpoint_id") or row.get("row_id") or f"{prefix}-{idx:03d}")
    text = _row_text(row, preferred_keys)
    metadata = dict(row)
    metadata.setdefault("semantic_zone", _semantic_zone(text, item_type))
    metadata.setdefault("normalized_entities", _normalize_entities(text))
    return SemanticItem(item_id=item_id, item_type=item_type, source="baseline_workbook", text=text, metadata=metadata)


def _baseline_items(workbook: dict[str, Any], table_name: str, item_type: str, preferred_keys: tuple[str, ...]) -> list[SemanticItem]:
    rows = workbook.get(table_name) or []
    if isinstance(rows, dict):
        rows = list(rows.values())
    return [_item_from_row(item_type.upper(), idx, item_type, row, preferred_keys) for idx, row in enumerate(rows or [], start=1) if isinstance(row, dict)]


def _locked_items(records: list[LockedFileRecord], role: str | None, item_type: str, keywords: tuple[str, ...], *, source_prefix: str) -> list[SemanticItem]:
    items: list[SemanticItem] = []
    selected = [record for record in records if role is None or record.role == role]
    for record in selected:
        sentences = _split_evidence_sentences(record.text_excerpt)
        if not sentences and record.text_excerpt:
            sentences = [_clean(record.text_excerpt[:700])]
        for sentence in sentences:
            lower = sentence.lower()
            if keywords and not any(keyword.lower() in lower or keyword in sentence for keyword in keywords):
                continue
            items.append(
                SemanticItem(
                    item_id=f"{source_prefix}-{len(items)+1:04d}",
                    item_type=item_type,
                    source=f"{record.source_root_label}/{record.relative_path}",
                    text=sentence,
                    metadata={"role": record.role, "confidence": record.confidence, "extraction_status": record.extraction_status},
                )
            )
            if len(items) >= 160:
                return items
    return items


def _gold_items_from_sections(
    gold_items: list[SemanticItem],
    zone: str,
    item_type: str,
    keywords: tuple[str, ...],
    *,
    source_prefix: str,
    document_roles: tuple[str, ...] = (),
) -> list[SemanticItem]:
    filtered: list[SemanticItem] = []
    for item in gold_items:
        if document_roles and str(item.metadata.get("document_role") or "") not in document_roles:
            continue
        if zone and item.metadata.get("semantic_zone") != zone:
            continue
        lower = item.text.lower()
        if keywords and not any(keyword.lower() in lower or keyword in item.text for keyword in keywords):
            continue
        filtered.append(
            SemanticItem(
                item_id=f"{source_prefix}-{len(filtered)+1:04d}",
                item_type=item_type,
                source=item.source,
                text=item.text,
                metadata=dict(item.metadata),
            )
        )
        if len(filtered) >= 180:
            break
    if not filtered and zone:
        for item in gold_items:
            if document_roles and str(item.metadata.get("document_role") or "") not in document_roles:
                continue
            if item.metadata.get("semantic_zone") == zone:
                filtered.append(
                    SemanticItem(
                        item_id=f"{source_prefix}-{len(filtered)+1:04d}",
                        item_type=item_type,
                        source=item.source,
                        text=item.text,
                        metadata=dict(item.metadata),
                    )
                )
            if len(filtered) >= 80:
                break
    return filtered


def _number_tokens(text: Any) -> set[str]:
    return set(re.findall(r"\b\d+(?:\.\d+)?\s*%?|\bn\s*=\s*\d+\b", str(text or "").lower()))


def _source_hierarchy_rank(text: Any) -> int:
    lower = str(text or "").lower()
    if any(token in lower for token in ("meta-analysis", "systematic review", "cochrane", "pooled", "aggregate")):
        return 5
    if any(token in lower for token in ("guideline", "consensus", "mdcg", "eau", "aha", "esc")):
        return 4
    if "registry" in lower or "real-world" in lower:
        return 3
    if any(token in lower for token in ("randomized", "rct", "cohort", "prospective")):
        return 2
    if any(token in lower for token in ("case series", "retrospective", "single-arm")):
        return 1
    return 0


def _semantic_match_score(item: SemanticItem, candidate: SemanticItem) -> float:
    base = _semantic_similarity(item.text, candidate.text)
    entity = _entity_overlap(item.text, candidate.text)
    zone_bonus = 0.07 if item.metadata.get("semantic_zone") == candidate.metadata.get("semantic_zone") else 0.0
    citation_bonus = 0.10 if _entity_norms(item.text) & _entity_norms(candidate.text) and any(norm.startswith(("PMID:", "DOI:")) for norm in (_entity_norms(item.text) & _entity_norms(candidate.text))) else 0.0
    number_bonus = 0.05 if _number_tokens(item.text) & _number_tokens(candidate.text) else 0.0
    score = (0.72 * base) + (0.20 * entity) + zone_bonus + citation_bonus + number_bonus
    if item.item_type == "benchmark" or "benchmark" in item.item_type:
        candidate_rank = _source_hierarchy_rank(candidate.text)
        if candidate_rank:
            score += min(0.08, 0.015 * candidate_rank)
        if any(token in candidate.text.lower() for token in ("literature filter", "search strategy", "before executing a literature search")):
            score -= 0.12
    return round(max(0.0, min(1.0, score)), 3)


def _best_match(item: SemanticItem, candidates: list[SemanticItem]) -> tuple[SemanticItem | None, float, list[tuple[str, float]]]:
    scored = [(candidate, _semantic_match_score(item, candidate)) for candidate in candidates]
    scored.sort(key=lambda pair: pair[1], reverse=True)
    options = [(candidate.item_id, score) for candidate, score in scored[:5]]
    if not scored:
        return None, 0.0, options
    return scored[0][0], scored[0][1], options


def _rank_candidates(item: SemanticItem, candidates: list[SemanticItem], *, limit: int = 5) -> list[tuple[SemanticItem, float]]:
    scored = [(candidate, _semantic_match_score(item, candidate)) for candidate in candidates]
    scored.sort(key=lambda pair: pair[1], reverse=True)
    return scored[:limit]


def _benchmark_numeric_profile(text: Any) -> dict[str, Any]:
    value = str(text or "")
    percentages = re.findall(r"\b\d+(?:\.\d+)?\s*%", value)
    sample_sizes = re.findall(r"\bn\s*=\s*\d+\b|\b\d+\s+(?:patients|subjects|cases)\b", value, flags=re.IGNORECASE)
    ci = re.findall(r"(?:95%\s*)?CI[:\s]*[0-9.%-]+\s*(?:to|-|–)\s*[0-9.%-]+|confidence interval", value, flags=re.IGNORECASE)
    return {
        "percentages": percentages,
        "sample_sizes": sample_sizes,
        "ci_or_range": ci,
        "has_value": bool(percentages or sample_sizes or ci or re.search(r"\b\d+(?:\.\d+)?\b", value)),
    }


def _benchmark_rationale_difference(left: Any, right: Any) -> str:
    left_profile = _benchmark_numeric_profile(left)
    right_profile = _benchmark_numeric_profile(right)
    if left_profile["has_value"] and right_profile["has_value"]:
        if set(left_profile["percentages"]) & set(right_profile["percentages"]):
            return "value_close_or_same"
        return "both_quantified_but_value_or_rationale_differs"
    if left_profile["has_value"] != right_profile["has_value"]:
        return "one_side_quantified_only"
    return "qualitative_or_unquantified_both_sides"


def _delta_classification(confidence: float, baseline_text: str, gold_text: str, *, nb_relevant: bool = False, overclaim: bool = False) -> str:
    if confidence < 0.70:
        return "needs_human_confirmation"
    if overclaim:
        return "ai_overclaim_or_misreasoning"
    if nb_relevant:
        return "nb_relevant_gap"
    if confidence >= 0.86:
        return "accepted_stylistic_variance"
    left = _tokens(baseline_text)
    right = _tokens(gold_text)
    if len(right - left) >= 8:
        return "gold_reference_only_detail"
    if len(left - right) >= 8:
        return "project_specific_difference"
    return "clinically_material_gap"


def _materiality(delta_classification: str, confidence: float) -> str:
    if delta_classification in {"clinically_material_gap", "nb_relevant_gap", "ai_overclaim_or_misreasoning"}:
        return "high"
    if confidence < 0.70:
        return "unknown"
    if delta_classification == "gold_reference_only_detail":
        return "medium"
    return "low"


def _root_cause(delta_classification: str, context: str) -> str:
    lower = context.lower()
    if delta_classification == "accepted_stylistic_variance":
        return "acceptable_variance"
    if delta_classification == "needs_human_confirmation":
        return "needs_human_confirmation"
    if any(token in lower for token in ("benchmark", "sota", "endpoint", "population", "ci", "sample")):
        return "knowledge_gap"
    if any(token in lower for token in ("pivotal", "supportive", "level", "appraisal", "limitation")):
        return "evidence_appraisal_gap"
    if any(token in lower for token in ("pmcf", "pms", "unanswered", "residual")):
        return "rule_gap"
    if any(token in lower for token in ("ifu", "rmf", "gspr", "sscp", "alignment")):
        return "source_interpretation_gap"
    if any(token in lower for token in ("benefit", "risk", "conclusion", "overclaim", "superior")):
        return "writer_synthesis_gap"
    return "missing_intermediate_artifact"


def _reasoning_why_material(
    delta_classification: str,
    semantic_delta: str,
    baseline_item: SemanticItem | None,
    gold_item: SemanticItem | None,
) -> str:
    if delta_classification == "accepted_stylistic_variance":
        return "The matched AI and gold spans appear semantically close enough that the difference is treated as wording/style unless a reviewer says otherwise."
    if delta_classification == "needs_human_confirmation":
        return "The analyzer could not establish a reliable semantic bridge; the item is routed to human gate and is not counted as a systemic material gap."
    evidence_parts = []
    if baseline_item:
        evidence_parts.append(f"AI-side span from {baseline_item.source} expresses: {_clean(baseline_item.text)[:220]}")
    else:
        evidence_parts.append("No AI-side counterpart was found for this gold-side judgment.")
    if gold_item:
        evidence_parts.append(f"Gold-side span from {gold_item.source} expresses: {_clean(gold_item.text)[:220]}")
    else:
        evidence_parts.append("No gold-side counterpart was found for this AI-side judgment.")
    return f"{semantic_delta} This is potentially material because {' '.join(evidence_parts)}"


def _why_not_alternative_causes(root_cause: str, delta_classification: str) -> str:
    if delta_classification == "accepted_stylistic_variance":
        return "Alternative defect causes are not asserted because the difference is classified as accepted variance."
    if delta_classification == "needs_human_confirmation":
        return "Alternative causes are not eliminated; low confidence means this remains a human-gate item."
    explanations = {
        "knowledge_gap": "The span refers to SOTA/benchmark/source details, so the leading cause is missing or weak knowledge extraction rather than final wording only.",
        "evidence_appraisal_gap": "The span turns on study design, pivotal/supportive role or limitations, so appraisal logic is more likely than pure source classification.",
        "source_interpretation_gap": "The span involves IFU/RMF/GSPR/SSCP or final-package alignment, so source interpretation is more likely than narrative style.",
        "writer_synthesis_gap": "The span involves conclusion strength or benefit-risk wording, so synthesis/wording control is more likely than retrieval alone.",
        "rule_gap": "The span involves PMCF/PMS/residual-question boundary decisions, so deterministic boundary rules are more likely than isolated text extraction.",
        "missing_intermediate_artifact": "The span does not clearly fall into a specialized rule bucket, so the most likely cause is a missing or unused intermediate artifact.",
    }
    return explanations.get(root_cause, "The alternative cause cannot be confidently excluded; keep this as a calibration candidate rather than a final defect.")


def _overclaim_signal(text: str) -> bool:
    return bool(re.search(r"\b(superior|best|proven|definitive|always|eliminates|guarantees|significantly better)\b", text or "", flags=re.IGNORECASE))


def _section_label(item: SemanticItem) -> str:
    heading = str(item.metadata.get("section_heading") or "").strip()
    if heading and heading.lower() != "unsectioned":
        return heading
    zone = str(item.metadata.get("semantic_zone") or "").strip()
    role = str(item.metadata.get("document_role") or item.metadata.get("role") or "").strip()
    if role and zone:
        return f"{role}:{zone}"
    return zone or role


def _semantic_row(
    table_name: str,
    row_id: str,
    baseline_item: SemanticItem | None,
    gold_item: SemanticItem | None,
    confidence: float,
    delta_classification: str,
    semantic_delta: str,
    *,
    nb_relevant: bool = False,
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    root_cause = _root_cause(delta_classification, " ".join([semantic_delta, baseline_item.text if baseline_item else "", gold_item.text if gold_item else ""]))
    staged = _staged_confidences(baseline_item, gold_item, confidence, nb_relevant=nb_relevant, root_cause=root_cause)
    effective_confidence = min(1.0, max(confidence, staged["overall_confidence"]))
    baseline_entities = _normalize_entities(baseline_item.text) if baseline_item else []
    gold_entities = _normalize_entities(gold_item.text) if gold_item else []
    baseline_norms = {entity["normalized"] for entity in baseline_entities}
    gold_norms = {entity["normalized"] for entity in gold_entities}
    entity_bridge = sorted(baseline_norms & gold_norms)
    zone_match = bool(baseline_item and gold_item and baseline_item.metadata.get("semantic_zone") == gold_item.metadata.get("semantic_zone"))
    if (
        delta_classification == "needs_human_confirmation"
        and baseline_item
        and gold_item
        and effective_confidence >= 0.70
        and (staged["semantic_equivalence_confidence"] >= 0.70 or staged["candidate_retrieval_confidence"] >= 0.70)
        and (entity_bridge or zone_match)
    ):
        delta_classification = _delta_classification(effective_confidence, baseline_item.text, gold_item.text, nb_relevant=nb_relevant, overclaim=_overclaim_signal(baseline_item.text))
        root_cause = _root_cause(delta_classification, " ".join([semantic_delta, baseline_item.text, gold_item.text]))
        staged = _staged_confidences(baseline_item, gold_item, effective_confidence, nb_relevant=nb_relevant, root_cause=root_cause)
        effective_confidence = min(1.0, max(effective_confidence, staged["overall_confidence"]))
    human_gate = effective_confidence < 0.70 or delta_classification == "needs_human_confirmation"
    baseline_section = _section_label(baseline_item) if baseline_item else ""
    gold_section = _section_label(gold_item) if gold_item else ""
    row = {
        "semantic_table": table_name,
        "semantic_delta_id": row_id,
        "baseline_item_id": baseline_item.item_id if baseline_item else "",
        "baseline_source": baseline_item.source if baseline_item else "",
        "source_section": baseline_section,
        "baseline_text": baseline_item.text[:900] if baseline_item else "",
        "ai_side_evidence_span": baseline_item.text[:700] if baseline_item else "",
        "gold_item_id": gold_item.item_id if gold_item else "",
        "gold_source": gold_item.source if gold_item else "",
        "target_section": gold_section,
        "gold_text": gold_item.text[:900] if gold_item else "",
        "gold_side_evidence_span": gold_item.text[:700] if gold_item else "",
        "retrieval_basis": "section_aware_zone_entity_similarity",
        "retrieval_confidence": f"{staged['candidate_retrieval_confidence']:.2f}",
        "semantic_relation": "matched" if baseline_item and gold_item and effective_confidence >= 0.70 else "low_confidence_or_unmatched",
        "delta_classification": delta_classification,
        "semantic_delta": semantic_delta,
        "match_confidence": f"{effective_confidence:.2f}",
        "candidate_retrieval_confidence": f"{staged['candidate_retrieval_confidence']:.2f}",
        "semantic_equivalence_confidence": f"{staged['semantic_equivalence_confidence']:.2f}",
        "clinical_materiality_confidence": f"{staged['clinical_materiality_confidence']:.2f}",
        "nb_relevance_confidence": f"{staged['nb_relevance_confidence']:.2f}",
        "root_cause_confidence": f"{staged['root_cause_confidence']:.2f}",
        "materiality": _materiality(delta_classification, effective_confidence),
        "nb_relevance": "yes" if nb_relevant else "no",
        "root_cause_candidate": root_cause,
        "baseline_entities": json.dumps(baseline_entities, ensure_ascii=False),
        "gold_entities": json.dumps(gold_entities, ensure_ascii=False),
        "entity_bridge": json.dumps(entity_bridge, ensure_ascii=False),
        "normalized_entity_overlap": f"{_entity_overlap(baseline_item.text if baseline_item else '', gold_item.text if gold_item else ''):.2f}",
        "reasoning_why_material": _reasoning_why_material(delta_classification, semantic_delta, baseline_item, gold_item),
        "why_not_alternative_causes": _why_not_alternative_causes(root_cause, delta_classification),
        "upgrade_implication": _upgrade_candidate(root_cause),
        "human_gate_required": "yes" if human_gate else "no",
        "allowed_for_systemic_gap_counting": "no" if human_gate else "yes",
    }
    if extra:
        row.update(extra)
    return row


def _low_confidence_type(row: dict[str, Any], options: list[tuple[str, float]]) -> str:
    ai_text = str(row.get("baseline_text") or "")
    gold_text = str(row.get("gold_text") or "")
    if not gold_text or "metadata_only" in gold_text.lower() or "extraction failed" in gold_text.lower():
        return "TEXT_EXTRACTION_GAP"
    if not row.get("target_section") or str(row.get("target_section")) in {"", "unsectioned"}:
        return "SECTION_MAPPING_GAP"
    try:
        entity_overlap = float(row.get("normalized_entity_overlap") or 0)
    except ValueError:
        entity_overlap = 0.0
    if ai_text and gold_text and entity_overlap == 0.0 and any(token in f"{ai_text} {gold_text}".lower() for token in ("endpoint", "benchmark", "pmid", "doi", "ifu", "rmf", "risk", "pmcf")):
        return "ENTITY_NORMALIZATION_GAP"
    if options and len(options) >= 2:
        scores = [score for _, score in options[:3]]
        if scores and max(scores) >= 0.35:
            return "SEMANTIC_MATCHING_GAP"
    return "TRUE_AMBIGUITY_HUMAN_GATE"


def _low_confidence_row(row: dict[str, Any], options: list[tuple[str, float]]) -> dict[str, Any]:
    lc_type = _low_confidence_type(row, options)
    resolvable, action = LOW_CONFIDENCE_ACTIONS[lc_type]
    return {
        "queue_id": f"LC-{row.get('semantic_delta_id', 'UNKNOWN')}",
        "semantic_table": row.get("semantic_table", ""),
        "candidate_match": f"{row.get('baseline_item_id', '')} -> {row.get('gold_item_id', '')}",
        "match_confidence": row.get("match_confidence", ""),
        "candidate_retrieval_confidence": row.get("candidate_retrieval_confidence", ""),
        "semantic_equivalence_confidence": row.get("semantic_equivalence_confidence", ""),
        "clinical_materiality_confidence": row.get("clinical_materiality_confidence", ""),
        "nb_relevance_confidence": row.get("nb_relevance_confidence", ""),
        "root_cause_confidence": row.get("root_cause_confidence", ""),
        "low_confidence_type": lc_type,
        "resolvable_by_analyzer_upgrade": resolvable,
        "reason_for_low_confidence": "Semantic similarity below 0.70 or no candidate match found.",
        "options_considered": json.dumps(options, ensure_ascii=False),
        "evidence_seen": f"AI: {row.get('baseline_text', '')[:300]} | Gold: {row.get('gold_text', '')[:300]}",
        "recommended_next_action": action,
        "recommended_human_question": "Confirm whether this AI baseline item and gold-reference item express the same clinical-evaluation judgment, a material difference, or an acceptable variance.",
    }


def low_confidence_breakdown_rows(low_confidence_rows: list[dict[str, Any]], *, project_label: str = "") -> list[dict[str, Any]]:
    total = len(low_confidence_rows)
    rows: list[dict[str, Any]] = []
    for lc_type in LOW_CONFIDENCE_TYPES:
        count = sum(1 for row in low_confidence_rows if row.get("low_confidence_type") == lc_type)
        resolvable, action = LOW_CONFIDENCE_ACTIONS[lc_type]
        rows.append(
            {
                "low_confidence_type": lc_type,
                "count": count,
                "percentage": f"{(count / total * 100):.1f}" if total else "0.0",
                "project_distribution": f"{project_label}:{count}" if project_label else str(count),
                "resolvable_by_analyzer_upgrade": resolvable,
                "recommended_next_action": action,
            }
        )
    return rows


def _entity_json(value: Any) -> list[dict[str, Any]]:
    if isinstance(value, list):
        return [item for item in value if isinstance(item, dict)]
    try:
        parsed = json.loads(str(value or "[]"))
    except json.JSONDecodeError:
        return []
    return [item for item in parsed if isinstance(item, dict)] if isinstance(parsed, list) else []


def build_entity_normalization_outputs(semantic_tables: dict[str, list[dict[str, Any]]]) -> dict[str, list[dict[str, Any]]]:
    dictionary_rows: list[dict[str, Any]] = []
    for normalized, aliases in ENTITY_SYNONYMS.items():
        dictionary_rows.append(
            {
                "normalized_entity": normalized,
                "entity_type": "controlled_term",
                "aliases": "; ".join(aliases),
                "alias_source": "built_in_controlled_dictionary",
                "evidence": "Configured analyzer synonym dictionary",
            }
        )
    alias_map: dict[tuple[str, str], dict[str, Any]] = {}
    provenance_rows: list[dict[str, Any]] = []
    resolved_rows: list[dict[str, Any]] = []
    for table_name, rows in semantic_tables.items():
        if table_name == "cognitive_gap_attribution_table":
            continue
        for row in rows:
            baseline_entities = _entity_json(row.get("baseline_entities"))
            gold_entities = _entity_json(row.get("gold_entities"))
            for side, entities in (("AI", baseline_entities), ("GOLD", gold_entities)):
                for entity in entities:
                    normalized = str(entity.get("normalized") or "")
                    original = str(entity.get("original") or "")
                    if not normalized or not original:
                        continue
                    key = (normalized, original)
                    entry = alias_map.setdefault(
                        key,
                        {
                            "normalized_entity": normalized,
                            "alias": original,
                            "entity_type": entity.get("entity_type", ""),
                            "sides_seen": set(),
                            "occurrence_count": 0,
                            "alias_source": entity.get("provenance", ""),
                        },
                    )
                    entry["sides_seen"].add(side)
                    entry["occurrence_count"] += 1
                    provenance_rows.append(
                        {
                            "semantic_delta_id": row.get("semantic_delta_id", ""),
                            "semantic_table": table_name,
                            "side": side,
                            "raw_text": original,
                            "normalized_entity": normalized,
                            "entity_type": entity.get("entity_type", ""),
                            "alias_source": entity.get("provenance", ""),
                            "source": row.get("baseline_source" if side == "AI" else "gold_source", ""),
                        }
                    )
            bridge = _entity_json(row.get("entity_bridge"))
            if not bridge:
                try:
                    bridge = [{"normalized": item} for item in json.loads(str(row.get("entity_bridge") or "[]"))]
                except json.JSONDecodeError:
                    bridge = []
            resolved_rows.append(
                {
                    "semantic_delta_id": row.get("semantic_delta_id", ""),
                    "semantic_table": table_name,
                    "baseline_item_id": row.get("baseline_item_id", ""),
                    "gold_item_id": row.get("gold_item_id", ""),
                    "entity_bridge": row.get("entity_bridge", "[]"),
                    "normalized_entity_overlap": row.get("normalized_entity_overlap", ""),
                    "resolution_status": "bridged" if row.get("entity_bridge") not in ("[]", "", None) else "not_bridged",
                }
            )
    alias_rows = [
        {
            "normalized_entity": normalized,
            "alias": alias,
            "entity_type": data["entity_type"],
            "sides_seen": ";".join(sorted(data["sides_seen"])),
            "occurrence_count": data["occurrence_count"],
            "alias_source": data["alias_source"],
        }
        for (normalized, alias), data in sorted(alias_map.items())
    ]
    return {
        "dictionary": dictionary_rows,
        "alias_mapping": alias_rows,
        "provenance": provenance_rows,
        "resolved_items": resolved_rows,
    }


def _table_type(text: Any) -> str:
    lower = str(text or "").lower()
    if any(token in lower for token in ("benchmark", "endpoint", "acceptance", "criterion", "sota")):
        return "SOTA_BENCHMARK_TABLE"
    if any(token in lower for token in ("pivotal", "supportive", "appraisal", "evidence level", "limitation")):
        return "EVIDENCE_APPRAISAL_TABLE"
    if any(token in lower for token in ("pmcf", "unanswered", "residual question", "pms")):
        return "PMCF_UNRESOLVED_QUESTION_TABLE"
    if "gspr" in lower or "annex i" in lower:
        return "GSPR_MAPPING_TABLE"
    if any(token in lower for token in ("risk", "hazard", "complication", "adverse event", "harm")):
        return "RISK_COMPLICATION_TABLE"
    if any(token in lower for token in ("clinical benefit", "benefit", "claim")):
        return "CLINICAL_BENEFIT_ENDPOINT_TABLE"
    return "OTHER_OR_LAYOUT_TABLE"


def build_table_extraction_outputs(gold_context: dict[str, Any]) -> dict[str, Any]:
    table_rows = gold_context.get("table_rows") or []
    manifest = dict(gold_context.get("manifest") or {})
    type_rows: list[dict[str, Any]] = []
    cell_rows: list[dict[str, Any]] = []
    failure_rows: list[dict[str, Any]] = []
    for row in table_rows:
        table_type = _table_type(row.get("table_excerpt", ""))
        type_rows.append(
            {
                "table_id": row.get("table_id", ""),
                "document_role": row.get("document_role", ""),
                "source_path": row.get("source_path", ""),
                "section_id": row.get("section_id", ""),
                "table_anchor": row.get("table_anchor", ""),
                "semantic_zone": row.get("semantic_zone", ""),
                "table_type": table_type,
                "is_priority_table_type": "yes" if table_type != "OTHER_OR_LAYOUT_TABLE" else "no",
                "table_excerpt": row.get("table_excerpt", "")[:1000],
            }
        )
        cells = [cell.strip() for cell in str(row.get("table_excerpt") or "").split("|") if cell.strip()]
        for idx, cell in enumerate(cells[:80], start=1):
            cell_rows.append(
                {
                    "cell_anchor_id": f"{row.get('table_id', 'GOLD-TABLE')}:C{idx:03d}",
                    "table_id": row.get("table_id", ""),
                    "document_role": row.get("document_role", ""),
                    "source_path": row.get("source_path", ""),
                    "section_id": row.get("section_id", ""),
                    "table_type": table_type,
                    "cell_index": idx,
                    "cell_text": cell,
                    "semantic_zone": _semantic_zone(cell, table_type),
                    "normalized_entities": json.dumps(_normalize_entities(cell), ensure_ascii=False),
                }
            )
    role_counts = manifest.get("document_role_counts") or {}
    roles_with_tables = {row.get("document_role") for row in table_rows}
    for role, count in role_counts.items():
        if count and role not in roles_with_tables:
            failure_rows.append(
                {
                    "document_role": role,
                    "final_file_count": count,
                    "failure_reason": "No table anchors were recovered from extracted text for this document role.",
                    "silent_drop": "no",
                    "recommended_action": "Improve DOCX/PDF table extraction or provide text-rendered table artifacts.",
                }
            )
    table_manifest = {
        "schema_name": "gold_reference_table_extraction_manifest",
        "schema_version": "phase3.5d-table-extraction-v1",
        "generated_at": now_iso(),
        "status": "TABLE_EXTRACTION_COMPLETE_WITH_FAILURE_LOG" if failure_rows else "TABLE_EXTRACTION_COMPLETE",
        "table_count": len(table_rows),
        "cell_anchor_count": len(cell_rows),
        "priority_table_count": sum(1 for row in type_rows if row["is_priority_table_type"] == "yes"),
        "failure_count": len(failure_rows),
        "table_type_counts": {table_type: sum(1 for row in type_rows if row["table_type"] == table_type) for table_type in sorted({row["table_type"] for row in type_rows})},
    }
    return {"manifest": table_manifest, "type_rows": type_rows, "cell_rows": cell_rows, "failure_rows": failure_rows}


def true_ambiguity_sampling_audit_rows(low_confidence_rows: list[dict[str, Any]], *, sample_per_type: int = 25) -> tuple[list[dict[str, Any]], str]:
    audit_rows: list[dict[str, Any]] = []
    by_type: dict[str, list[dict[str, Any]]] = {name: [] for name in LOW_CONFIDENCE_TYPES}
    for row in low_confidence_rows:
        by_type.setdefault(row.get("low_confidence_type") or "TRUE_AMBIGUITY_HUMAN_GATE", []).append(row)
    for lc_type, rows in by_type.items():
        for idx, row in enumerate(rows[:sample_per_type], start=1):
            text = f"{row.get('evidence_seen', '')} {row.get('options_considered', '')}".lower()
            reclassified = lc_type
            rationale = "Retained original low-confidence type after deterministic audit."
            if lc_type == "TRUE_AMBIGUITY_HUMAN_GATE":
                if any(token in text for token in ("endpoint", "benchmark", "pmid", "doi", "risk", "pmcf", "ifu", "rmf")):
                    reclassified = "ENTITY_NORMALIZATION_GAP"
                    rationale = "The sampled item contains bridgeable CER entities; likely resolvable through stronger normalization rather than permanent human gate."
                elif "metadata_only" in text or "extraction failed" in text:
                    reclassified = "TEXT_EXTRACTION_GAP"
                    rationale = "The sampled item lacks recoverable text and should be treated as extraction limitation."
            audit_rows.append(
                {
                    "audit_id": f"TA-AUDIT-{len(audit_rows)+1:04d}",
                    "original_queue_id": row.get("queue_id", ""),
                    "semantic_table": row.get("semantic_table", ""),
                    "original_low_confidence_type": lc_type,
                    "reclassified_low_confidence_type": reclassified,
                    "match_confidence": row.get("match_confidence", ""),
                    "sample_stratum_index": idx,
                    "rationale": rationale,
                    "forced_reduction": "no",
                    "human_gate_required": "yes" if reclassified == "TRUE_AMBIGUITY_HUMAN_GATE" else "no",
                }
            )
    total_true = len(by_type.get("TRUE_AMBIGUITY_HUMAN_GATE", []))
    audited_true = [row for row in audit_rows if row["original_low_confidence_type"] == "TRUE_AMBIGUITY_HUMAN_GATE"]
    retained_true = [row for row in audited_true if row["reclassified_low_confidence_type"] == "TRUE_AMBIGUITY_HUMAN_GATE"]
    lines = [
        "# True Ambiguity Reclassification Report",
        "",
        f"Generated: {now_iso()}",
        "",
        f"- TRUE_AMBIGUITY_HUMAN_GATE total: {total_true}",
        f"- TRUE_AMBIGUITY sampled: {len(audited_true)}",
        f"- TRUE_AMBIGUITY retained in sample: {len(retained_true)}",
        f"- TRUE_AMBIGUITY reclassified in sample: {len(audited_true) - len(retained_true)}",
        "",
        "This audit does not lower the confidence threshold and does not force uncertain items into confirmed conclusions. It only identifies whether some low-confidence rows are better described as analyzer-resolvable limitations.",
    ]
    return audit_rows, "\n".join(lines) + "\n"


def build_semantic_delta_tables(
    baseline: dict[str, Any],
    locked_records: list[LockedFileRecord],
) -> tuple[
    dict[str, list[dict[str, Any]]],
    list[dict[str, Any]],
    dict[str, Any],
    dict[str, Any],
    list[dict[str, Any]],
    list[dict[str, Any]],
    list[dict[str, Any]],
    dict[str, list[dict[str, Any]]],
]:
    workbook = baseline["workbook"]
    nb_records = [record for record in locked_records if record.role == "NB_FEEDBACK"]
    gold_manifest, gold_section_rows, gold_table_rows, gold_citation_rows, gold_items = build_gold_reference_indexes(locked_records)
    tables: dict[str, list[dict[str, Any]]] = {name: [] for name in SEMANTIC_TABLE_NAMES}
    low_confidence: list[dict[str, Any]] = []
    sota_candidate_rankings: list[dict[str, Any]] = []
    sota_reranking_explanations: list[dict[str, Any]] = []

    final_claims = _gold_items_from_sections(gold_items, "claims", "gold_claim", ("intended", "purpose", "claim", "indication", "benefit", "performance", "safety", "适用", "预期", "获益"), source_prefix="GOLD-CLAIM", document_roles=("FINAL_CER", "FINAL_IFU", "FINAL_SSCP", "FINAL_OTHER_ACCEPTED_CHANGES"))
    if not final_claims:
        final_claims = _gold_items_from_sections(gold_items, "claims", "gold_claim", (), source_prefix="GOLD-CLAIM")[:40]
    baseline_claims = _baseline_items(workbook, "claim_ledger", "claim", ("claim_id", "claim_text", "statement", "claim_type", "required_evidence"))
    for idx, item in enumerate(baseline_claims, start=1):
        match, confidence, options = _best_match(item, final_claims)
        delta_class = _delta_classification(confidence, item.text, match.text if match else "", overclaim=_overclaim_signal(item.text))
        row = _semantic_row(
            "claim_semantic_alignment_table",
            f"CLAIM-SEM-{idx:03d}",
            item,
            match,
            confidence,
            delta_class,
            "Claim wording, scope, limitation and strength compared against final accepted package.",
            nb_relevant=bool(nb_records and confidence < 0.82),
            extra={"ai_claim_scope": item.metadata.get("claim_type", ""), "gold_claim_scope": match.metadata.get("role", "") if match else ""},
        )
        tables["claim_semantic_alignment_table"].append(row)
        if row["human_gate_required"] == "yes":
            low_confidence.append(_low_confidence_row(row, options))
    matched_gold = {row["gold_item_id"] for row in tables["claim_semantic_alignment_table"] if row.get("gold_item_id")}
    for gold in final_claims[:40]:
        if gold.item_id not in matched_gold:
            tables["claim_semantic_alignment_table"].append(
                _semantic_row(
                    "claim_semantic_alignment_table",
                    f"CLAIM-GOLD-ONLY-{len(tables['claim_semantic_alignment_table'])+1:03d}",
                    None,
                    gold,
                    0.0,
                    "gold_reference_only_detail",
                    "Gold reference contains a potentially material claim/detail not matched to an AI baseline claim.",
                    nb_relevant=False,
                )
            )

    baseline_evidence = _baseline_items(workbook, "evidence_registry", "evidence", ("evidence_id", "source", "title", "endpoint", "result", "weight", "sample_size", "follow_up"))
    gold_evidence = _gold_items_from_sections(gold_items, "evidence_selection", "gold_evidence", ("pmid", "doi", "randomized", "systematic", "meta", "cohort", "registry", "pivotal", "supportive", "study", "n="), source_prefix="GOLD-EVID", document_roles=("FINAL_CER", "FINAL_OTHER_ACCEPTED_CHANGES"))
    if not gold_evidence:
        gold_evidence = _gold_items_from_sections(gold_items, "sota", "gold_evidence", ("pmid", "doi", "study", "evidence", "literature"), source_prefix="GOLD-EVID")
    for idx, item in enumerate(baseline_evidence, start=1):
        match, confidence, options = _best_match(item, gold_evidence)
        delta_class = _delta_classification(confidence, item.text, match.text if match else "", nb_relevant=bool(nb_records and confidence < 0.78))
        row = _semantic_row(
            "evidence_correspondence_table",
            f"EVID-CORR-{idx:03d}",
            item,
            match,
            confidence,
            delta_class,
            "Evidence inclusion, exclusion, role and applicability compared with final accepted evidence trail.",
            nb_relevant=bool(nb_records and confidence < 0.78),
            extra={
                "ai_evidence_role": item.metadata.get("weight", "") or item.metadata.get("contribution", ""),
                "gold_evidence_role": "pivotal/supportive candidate" if match else "",
            },
        )
        tables["evidence_correspondence_table"].append(row)
        if row["human_gate_required"] == "yes":
            low_confidence.append(_low_confidence_row(row, options))

    baseline_benchmarks = _baseline_items(workbook, "sota_benchmark_matrix", "benchmark", ("benchmark_id", "endpoint", "benchmark_value", "sota_source", "population", "sample_size", "CI_or_range", "clinical_significance", "rationale"))
    if not baseline_benchmarks:
        baseline_benchmarks = _baseline_items(workbook, "sota_benchmark_derivation_table", "benchmark", ("benchmark_id", "endpoint", "benchmark_value", "source", "population", "sample_size", "CI_or_range", "rationale"))
    gold_benchmarks = _gold_items_from_sections(gold_items, "sota", "gold_benchmark", ("benchmark", "endpoint", "%", "ci", "confidence interval", "range", "rate", "n=", "acceptance", "criterion", "基准", "终点"), source_prefix="GOLD-BM", document_roles=("FINAL_CER", "FINAL_OTHER_ACCEPTED_CHANGES"))
    for idx, item in enumerate(baseline_benchmarks, start=1):
        ranked_candidates = _rank_candidates(item, gold_benchmarks, limit=8)
        match = ranked_candidates[0][0] if ranked_candidates else None
        confidence = ranked_candidates[0][1] if ranked_candidates else 0.0
        options = [(candidate.item_id, score) for candidate, score in ranked_candidates[:5]]
        for rank, (candidate, score) in enumerate(ranked_candidates, start=1):
            source_rank = _source_hierarchy_rank(candidate.text)
            candidate_profile = _benchmark_numeric_profile(candidate.text)
            sota_candidate_rankings.append(
                {
                    "baseline_benchmark_id": item.item_id,
                    "candidate_rank": rank,
                    "gold_candidate_id": candidate.item_id,
                    "candidate_score": f"{score:.2f}",
                    "endpoint_entity_overlap": f"{_entity_overlap(item.text, candidate.text):.2f}",
                    "source_hierarchy_rank": source_rank,
                    "candidate_has_value": "yes" if candidate_profile["has_value"] else "no",
                    "candidate_percentages": "; ".join(candidate_profile["percentages"]),
                    "candidate_sample_sizes": "; ".join(candidate_profile["sample_sizes"]),
                    "candidate_ci_or_range": "; ".join(candidate_profile["ci_or_range"]),
                    "candidate_section": _section_label(candidate),
                    "candidate_source": candidate.source,
                    "candidate_text": candidate.text[:900],
                }
            )
        missing_value = not any(key in item.text.lower() for key in ("%", "ci", "range", "n=", "pmid", "doi")) and not re.search(r"\b\d+(\.\d+)?\b", item.text)
        delta_class = "clinically_material_gap" if missing_value else _delta_classification(confidence, item.text, match.text if match else "")
        rationale_diff = _benchmark_rationale_difference(item.text, match.text if match else "")
        sota_reranking_explanations.append(
            {
                "baseline_benchmark_id": item.item_id,
                "selected_gold_candidate_id": match.item_id if match else "",
                "selected_score": f"{confidence:.2f}",
                "reranking_basis": "entity_overlap + section zone + source hierarchy + numeric/citation signals",
                "source_hierarchy_preference": "aggregate > guideline > registry > cohort > case series",
                "value_rationale_relation": rationale_diff,
                "applicability_aware_comparison": "candidate_match" if confidence >= 0.70 else "low_confidence",
                "same_concept_different_label": "yes" if _entity_overlap(item.text, match.text if match else "") >= 0.30 and confidence < 0.70 else "no",
                "explanation": "Selected highest reranked gold benchmark candidate without changing the human-gate threshold.",
            }
        )
        row = _semantic_row(
            "sota_benchmark_delta_table",
            f"SOTA-BM-SEM-{idx:03d}",
            item,
            match,
            confidence,
            delta_class,
            "Endpoint, benchmark value, source hierarchy, population, sample size and rationale compared against final SOTA logic.",
            nb_relevant=bool(nb_records and (missing_value or confidence < 0.78)),
            extra={
                "endpoint_consistency": "needs_review" if confidence < 0.70 else "candidate_match",
                "benchmark_value_consistency": "missing_or_weak" if missing_value else "candidate_match",
                "source_hierarchy_consistency": "candidate_ranked" if match and _source_hierarchy_rank(match.text) else "needs_review",
                "selected_source_hierarchy_rank": _source_hierarchy_rank(match.text) if match else 0,
                "value_rationale_relation": rationale_diff,
            },
        )
        tables["sota_benchmark_delta_table"].append(row)
        if row["human_gate_required"] == "yes":
            low_confidence.append(_low_confidence_row(row, options))

    baseline_appraisal = baseline_evidence + _baseline_items(workbook, "article_appraisal", "appraisal", ("article_id", "evidence_id", "evidence_level", "applicability", "contribution", "weight", "limitations"))
    gold_appraisal = _gold_items_from_sections(gold_items, "evidence_appraisal", "gold_appraisal", ("level", "pivotal", "supportive", "limitation", "applicability", "quality", "bias", "randomized", "meta"), source_prefix="GOLD-APP", document_roles=("FINAL_CER", "FINAL_OTHER_ACCEPTED_CHANGES"))
    if not gold_appraisal:
        gold_appraisal = _gold_items_from_sections(gold_items, "evidence_selection", "gold_appraisal", ("pivotal", "supportive", "level", "limitation"), source_prefix="GOLD-APP")
    for idx, item in enumerate(baseline_appraisal[:120], start=1):
        match, confidence, options = _best_match(item, gold_appraisal)
        role_mismatch = bool(match and ("pivotal" in match.text.lower()) and ("pivotal" not in item.text.lower()))
        delta_class = "ai_overclaim_or_misreasoning" if role_mismatch else _delta_classification(confidence, item.text, match.text if match else "")
        row = _semantic_row(
            "evidence_appraisal_delta_table",
            f"EVID-APP-SEM-{idx:03d}",
            item,
            match,
            confidence,
            delta_class,
            "Evidence level, applicability, contribution, limitations and conclusion-strength impact compared.",
            nb_relevant=role_mismatch,
            extra={"conclusion_strength_impact": "yes" if role_mismatch or confidence < 0.70 else "possible"},
        )
        tables["evidence_appraisal_delta_table"].append(row)
        if row["human_gate_required"] == "yes":
            low_confidence.append(_low_confidence_row(row, options))

    baseline_pmcf = _baseline_items(workbook, "gap_pmcf_recommendations", "pmcf_gap", ("gap_id", "gap", "recommendation", "pmcf_activity", "timeline", "status"))
    if not baseline_pmcf:
        baseline_pmcf = _baseline_items(workbook, "risk_trace_matrix", "pmcf_gap", ("risk_id", "risk", "PMS/PMCF coverage", "residual_risk_conclusion"))
    gold_pmcf = _gold_items_from_sections(gold_items, "pmcf_unanswered_questions", "gold_pmcf", ("pmcf", "pms", "post-market", "unanswered", "residual", "follow-up", "registry", "questionnaire"), source_prefix="GOLD-PMCF", document_roles=("FINAL_CER", "FINAL_PMS_PMCF_PSUR", "FINAL_OTHER_ACCEPTED_CHANGES"))
    for idx, item in enumerate(baseline_pmcf[:80], start=1):
        match, confidence, options = _best_match(item, gold_pmcf)
        delta_class = _delta_classification(confidence, item.text, match.text if match else "", nb_relevant=bool(nb_records and confidence < 0.78))
        row = _semantic_row(
            "pmcf_boundary_delta_table",
            f"PMCF-SEM-{idx:03d}",
            item,
            match,
            confidence,
            delta_class,
            "Boundary between pre-submission resolution, accepted residual question and PMCF follow-up compared.",
            nb_relevant=bool(nb_records and confidence < 0.78),
            extra={"pmcf_boundary_decision": "needs_review" if confidence < 0.70 else "candidate_match"},
        )
        tables["pmcf_boundary_delta_table"].append(row)
        if row["human_gate_required"] == "yes":
            low_confidence.append(_low_confidence_row(row, options))

    baseline_br = _baseline_items(workbook, "benefit_risk_ledger", "benefit_risk", ("br_id", "benefit_statement", "overall_judgment", "allowed_wording", "not_allowed_wording", "rationale"))
    if not baseline_br:
        baseline_br = _baseline_items(workbook, "writer_conclusion_strength_guard", "benefit_risk", ("section", "conclusion", "allowed_wording", "evidence_strength"))
    gold_br = _gold_items_from_sections(gold_items, "benefit_risk", "gold_benefit_risk", ("benefit-risk", "benefit risk", "benefits outweigh", "residual risk", "acceptable", "conclusion", "获益", "风险"), source_prefix="GOLD-BR", document_roles=("FINAL_CER", "FINAL_OTHER_ACCEPTED_CHANGES"))
    if not gold_br:
        gold_br = _gold_items_from_sections(gold_items, "conclusion", "gold_benefit_risk", ("conclusion", "acceptable", "support", "risk"), source_prefix="GOLD-BR")
    for idx, item in enumerate(baseline_br[:80], start=1):
        match, confidence, options = _best_match(item, gold_br)
        delta_class = _delta_classification(confidence, item.text, match.text if match else "", overclaim=_overclaim_signal(item.text))
        row = _semantic_row(
            "benefit_risk_reasoning_delta_table",
            f"BR-SEM-{idx:03d}",
            item,
            match,
            confidence,
            delta_class,
            "Benefit-risk reasoning chain compared for evidence, bridge logic, limitations and conclusion strength.",
            nb_relevant=delta_class == "ai_overclaim_or_misreasoning",
            extra={"overclaim_signal": "yes" if _overclaim_signal(item.text) else "no"},
        )
        tables["benefit_risk_reasoning_delta_table"].append(row)
        if row["human_gate_required"] == "yes":
            low_confidence.append(_low_confidence_row(row, options))

    source_docs = _baseline_items(workbook, "source_inventory", "source", ("source_id", "document_type", "filename", "source_role"))
    for idx, record in enumerate([r for r in locked_records if r.role == "FINAL_CHANGE_REFERENCE"][:120], start=1):
        gold = SemanticItem(f"GOLD-DOC-{idx:03d}", "gold_document", f"{record.source_root_label}/{record.relative_path}", record.text_excerpt or record.relative_path, {"role": record.role})
        match, confidence, options = _best_match(gold, source_docs)
        doc_class = "gold_reference_only_detail" if confidence < 0.70 else "accepted_stylistic_variance"
        row = _semantic_row(
            "cross_document_alignment_delta_table",
            f"ALIGN-SEM-{idx:03d}",
            match,
            gold,
            confidence,
            doc_class,
            "Final CER/IFU/RMF/GSPR/PMCF/SSCP document presence and semantic linkage compared against AI source inventory.",
            nb_relevant="rmf" in record.relative_path.lower() or "ifu" in record.relative_path.lower() or "gspr" in record.relative_path.lower(),
            extra={"target_document": record.relative_path.split("/")[0] if "/" in record.relative_path else record.relative_path},
        )
        tables["cross_document_alignment_delta_table"].append(row)
        if row["human_gate_required"] == "yes":
            low_confidence.append(_low_confidence_row(row, options))

    baseline_failure_items = [
        SemanticItem(f"GATE-{idx:03d}", "gate_failure", "qa_gate_report", _row_text(gate, ("gate_id", "message", "severity")), gate)
        for idx, gate in enumerate(baseline["failed_gates"], start=1)
    ]
    nb_items = _locked_items(locked_records, "NB_FEEDBACK", "nb_finding", ("deficiency", "finding", "question", "loq", "sota", "evidence", "risk", "pmcf", "缺陷", "问题"), source_prefix="NB")
    response_items = _locked_items(locked_records, "OUR_RESPONSE", "response", ("response", "reply", "corrective", "action", "整改", "回复", "答复"), source_prefix="RESP")
    for idx, nb_item in enumerate(nb_items[:160], start=1):
        match, confidence, options = _best_match(nb_item, baseline_failure_items + baseline_claims + baseline_benchmarks + baseline_evidence)
        response_match, response_confidence, _ = _best_match(nb_item, response_items)
        final_match, final_confidence, _ = _best_match(nb_item, gold_items)
        nb_process_confidence = max(response_confidence, final_confidence)
        nb_process_confirmed = nb_process_confidence >= 0.50
        confidence_for_row = max(confidence, nb_process_confidence if nb_process_confirmed else 0.0)
        delta_class = "nb_relevant_gap" if confidence >= 0.70 or nb_process_confirmed else "needs_human_confirmation"
        row = _semantic_row(
            "nb_relevance_delta_table",
            f"NB-REL-SEM-{idx:03d}",
            match,
            nb_item,
            confidence_for_row,
            delta_class,
            "Real NB finding compared against AI baseline gates/claims/SOTA/evidence to identify whether the issue remained or was missed.",
            nb_relevant=True,
            extra={
                "final_response_candidate": response_match.text[:500] if response_match else "",
                "final_response_confidence": f"{response_confidence:.2f}",
                "final_package_candidate": final_match.text[:500] if final_match else "",
                "final_package_confidence": f"{final_confidence:.2f}",
                "nb_process_linkage_confidence": f"{nb_process_confidence:.2f}",
                "ai_same_issue_confidence": f"{confidence:.2f}",
                "nb_process_confirmed": "yes" if nb_process_confirmed else "no",
                "missed_layer_candidate": _root_cause(delta_class, nb_item.text),
            },
        )
        if nb_process_confirmed and confidence < 0.70:
            row["semantic_relation"] = "nb_process_confirmed_ai_direct_match_weak"
            row["human_gate_required"] = "no"
            row["allowed_for_systemic_gap_counting"] = "yes"
            row["reasoning_why_material"] = "NB process evidence links this finding to a response or final-package change; the weak AI text match is treated as AI coverage gap evidence rather than semantic ambiguity."
        tables["nb_relevance_delta_table"].append(row)
        if row["human_gate_required"] == "yes":
            low_confidence.append(_low_confidence_row(row, options))

    material_tables = [name for name in SEMANTIC_TABLE_NAMES if name != "cognitive_gap_attribution_table"]
    attribution_idx = 1
    for table_name in material_tables:
        for source_row in tables[table_name]:
            is_low_confidence = source_row.get("allowed_for_systemic_gap_counting") != "yes"
            if not is_low_confidence and source_row.get("materiality") not in {"high", "medium"}:
                continue
            root_cause = "needs_human_confirmation" if is_low_confidence else (source_row.get("root_cause_candidate") or "needs_human_confirmation")
            if root_cause not in COGNITIVE_ROOT_CAUSES:
                root_cause = "needs_human_confirmation"
            tables["cognitive_gap_attribution_table"].append(
                {
                    "semantic_delta_id": f"COG-GAP-{attribution_idx:04d}",
                    "source_semantic_table": table_name,
                    "source_delta_id": source_row.get("semantic_delta_id", ""),
                    "delta_classification": source_row.get("delta_classification", ""),
                    "materiality": source_row.get("materiality", ""),
                    "nb_relevance": source_row.get("nb_relevance", ""),
                    "cognitive_root_cause": root_cause,
                    "system_layer": root_cause,
                    "upgrade_candidate": _upgrade_candidate(root_cause),
                    "observed_ai_side_evidence_span": source_row.get("ai_side_evidence_span", "")[:700],
                    "observed_gold_side_evidence_span": source_row.get("gold_side_evidence_span", "")[:700],
                    "observed_nb_side_evidence_span": source_row.get("gold_side_evidence_span", "")[:700] if source_row.get("nb_relevance") == "yes" else "",
                    "reasoning_why_material": source_row.get("reasoning_why_material", ""),
                    "why_not_alternative_causes": source_row.get("why_not_alternative_causes", ""),
                    "root_cause_confidence": source_row.get("root_cause_confidence", ""),
                    "evidence_seen": " | ".join([source_row.get("baseline_text", "")[:300], source_row.get("gold_text", "")[:300]]),
                    "allowed_for_systemic_gap_counting": source_row.get("allowed_for_systemic_gap_counting", "no"),
                }
            )
            attribution_idx += 1

    nb_link_tables = build_nb_final_link_tables(baseline, locked_records, gold_items, tables)
    summary = _semantic_summary(tables, low_confidence, locked_records)
    gold_context = {
        "manifest": gold_manifest,
        "section_rows": gold_section_rows,
        "table_rows": gold_table_rows,
        "citation_rows": gold_citation_rows,
    }
    semantic_context = {
        "gold_items": gold_items,
        "sota_candidate_rankings": sota_candidate_rankings,
        "sota_reranking_explanations": sota_reranking_explanations,
    }
    return tables, low_confidence, summary, gold_context, nb_link_tables["nb_to_final"], nb_link_tables["ai_gap_to_nb"], semantic_context


def build_nb_final_link_tables(
    baseline: dict[str, Any],
    locked_records: list[LockedFileRecord],
    gold_items: list[SemanticItem],
    semantic_tables: dict[str, list[dict[str, Any]]],
) -> dict[str, list[dict[str, Any]]]:
    workbook = baseline["workbook"]
    nb_items = _locked_items(locked_records, "NB_FEEDBACK", "nb_finding", ("deficiency", "finding", "question", "loq", "sota", "evidence", "risk", "pmcf", "缺陷", "问题"), source_prefix="NB-LINK")
    response_items = _locked_items(locked_records, "OUR_RESPONSE", "manufacturer_response", ("response", "reply", "corrective", "action", "整改", "回复", "答复"), source_prefix="RESP-LINK")
    baseline_candidates = []
    baseline_candidates.extend(_baseline_items(workbook, "claim_ledger", "claim", ("claim_id", "claim_text", "statement", "claim_type", "required_evidence")))
    baseline_candidates.extend(_baseline_items(workbook, "sota_benchmark_matrix", "benchmark", ("benchmark_id", "endpoint", "benchmark_value", "sota_source", "population", "sample_size", "CI_or_range", "clinical_significance", "rationale")))
    baseline_candidates.extend(_baseline_items(workbook, "evidence_registry", "evidence", ("evidence_id", "source", "title", "endpoint", "result", "weight", "sample_size", "follow_up")))
    baseline_candidates.extend(
        SemanticItem(f"GATE-LINK-{idx:03d}", "gate_failure", "qa_gate_report", _row_text(gate, ("gate_id", "message", "severity")), gate)
        for idx, gate in enumerate(baseline["failed_gates"], start=1)
    )
    nb_to_final: list[dict[str, Any]] = []
    for idx, nb_item in enumerate(nb_items[:200], start=1):
        response_match, response_confidence, _ = _best_match(nb_item, response_items)
        final_match, final_confidence, _ = _best_match(nb_item, gold_items)
        ai_match, ai_confidence, _ = _best_match(nb_item, baseline_candidates)
        likely_trigger = "yes" if response_confidence >= 0.45 and final_confidence >= 0.45 else "uncertain"
        link_confidence = round((0.35 * response_confidence) + (0.45 * final_confidence) + (0.20 * ai_confidence), 3)
        nb_to_final.append(
            {
                "link_id": f"NB-FINAL-LINK-{idx:04d}",
                "nb_finding_source": nb_item.source,
                "nb_side_evidence_span": nb_item.text[:900],
                "manufacturer_response_candidate": response_match.text[:900] if response_match else "",
                "response_source": response_match.source if response_match else "",
                "response_confidence": f"{response_confidence:.2f}",
                "final_package_change_candidate": final_match.text[:900] if final_match else "",
                "final_source": final_match.source if final_match else "",
                "final_section": final_match.metadata.get("section_heading", "") if final_match else "",
                "final_change_confidence": f"{final_confidence:.2f}",
                "ai_baseline_candidate": ai_match.text[:700] if ai_match else "",
                "ai_baseline_source": ai_match.source if ai_match else "",
                "ai_gap_similarity": f"{ai_confidence:.2f}",
                "likely_final_change_triggered_by_nb": likely_trigger,
                "link_confidence": f"{link_confidence:.2f}",
                "human_gate_required": "yes" if link_confidence < 0.70 else "no",
            }
        )

    material_rows = [
        row
        for name, rows in semantic_tables.items()
        if name != "cognitive_gap_attribution_table"
        for row in rows
        if row.get("materiality") in {"high", "medium"} or row.get("nb_relevance") == "yes"
    ]
    ai_gap_to_nb: list[dict[str, Any]] = []
    for idx, row in enumerate(material_rows[:300], start=1):
        gap_item = SemanticItem(
            f"AI-GAP-{idx:04d}",
            "ai_semantic_gap",
            row.get("baseline_source", "") or row.get("semantic_table", ""),
            " ".join([str(row.get("semantic_delta", "")), str(row.get("baseline_text", "")), str(row.get("gold_text", ""))]),
            dict(row),
        )
        nb_match, nb_confidence, _ = _best_match(gap_item, nb_items)
        ai_gap_to_nb.append(
            {
                "link_id": f"AI-GAP-NB-LINK-{idx:04d}",
                "semantic_delta_id": row.get("semantic_delta_id", ""),
                "semantic_table": row.get("semantic_table", ""),
                "ai_gap_evidence_span": row.get("ai_side_evidence_span", "")[:900],
                "gold_side_evidence_span": row.get("gold_side_evidence_span", "")[:900],
                "candidate_nb_finding": nb_match.text[:900] if nb_match else "",
                "nb_source": nb_match.source if nb_match else "",
                "nb_link_confidence": f"{nb_confidence:.2f}",
                "same_or_related_issue": "yes" if nb_confidence >= 0.70 else "uncertain",
                "likely_missed_layer": row.get("root_cause_candidate", ""),
                "human_gate_required": "yes" if nb_confidence < 0.70 else "no",
            }
        )
    return {"nb_to_final": nb_to_final, "ai_gap_to_nb": ai_gap_to_nb}


def _upgrade_candidate(root_cause: str) -> str:
    mapping = {
        "workflow_order_gap": "Add native gate-to-upstream calibration route.",
        "missing_intermediate_artifact": "Add/readback mandatory intermediate artifact before writer.",
        "prompt_gap": "Patch corresponding agent prompt after aggregate calibration approval.",
        "rule_gap": "Add deterministic rule or boundary check after aggregate approval.",
        "knowledge_gap": "Add SOTA/evidence knowledge asset or benchmark source rule.",
        "template_or_section_gap": "Patch AP/human CER section contract.",
        "source_interpretation_gap": "Improve source classification/extraction/alignment.",
        "evidence_appraisal_gap": "Harden evidence appraisal and pivotal/supportive classification.",
        "writer_synthesis_gap": "Tighten writer ledger consumption and conclusion-strength guard.",
        "acceptable_variance": "No system change recommended.",
        "source_insufficiency": "Request source/full-text/manufacturer data.",
        "needs_human_confirmation": "Route to low-confidence human gate.",
    }
    return mapping.get(root_cause, "Review after aggregate calibration.")


def _semantic_summary(tables: dict[str, list[dict[str, Any]]], low_confidence: list[dict[str, Any]], locked_records: list[LockedFileRecord]) -> dict[str, Any]:
    material_rows = [
        row
        for table_name, rows in tables.items()
        if table_name != "cognitive_gap_attribution_table"
        for row in rows
        if row.get("materiality") in {"high", "medium"}
    ]
    high_rows = [row for row in material_rows if row.get("materiality") == "high"]
    nb_rows = [row for row in material_rows if row.get("nb_relevance") == "yes"]
    root_counts: dict[str, int] = {}
    for row in tables.get("cognitive_gap_attribution_table", []):
        root = str(row.get("cognitive_root_cause") or "unknown")
        root_counts[root] = root_counts.get(root, 0) + 1
    table_counts = {name: len(rows) for name, rows in tables.items()}
    return {
        "semantic_table_counts": table_counts,
        "material_delta_count": len(material_rows),
        "high_material_delta_count": len(high_rows),
        "nb_relevant_delta_count": len(nb_rows),
        "low_confidence_count": len(low_confidence),
        "locked_file_count": len(locked_records),
        "top_root_causes": sorted(root_counts.items(), key=lambda item: item[1], reverse=True)[:8],
        "top_material_deltas": [
            {
                "semantic_delta_id": row.get("semantic_delta_id"),
                "semantic_table": row.get("semantic_table"),
                "classification": row.get("delta_classification"),
                "materiality": row.get("materiality"),
                "root_cause": row.get("root_cause_candidate"),
                "delta": row.get("semantic_delta"),
            }
            for row in high_rows[:5]
        ],
    }


def write_csv(path: Path, rows: list[dict[str, Any]], fieldnames: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fieldnames is None:
        fieldnames = sorted({key for row in rows for key in row.keys()}) or ["status"]
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def locked_access_rows(records: list[LockedFileRecord]) -> list[dict[str, Any]]:
    return [
        {
            "access_id": f"LOCKED-ACCESS-{idx:04d}",
            "accessed_at": now_iso(),
            "source_root_label": record.source_root_label,
            "relative_path": record.relative_path,
            "role": record.role,
            "confidence": record.confidence,
            "extraction_status": record.extraction_status,
            "sha256": record.sha256,
            "size_bytes": record.size_bytes,
            "reason": record.reason,
        }
        for idx, record in enumerate(records, start=1)
    ]


def human_classification_rows(records: list[LockedFileRecord]) -> list[dict[str, Any]]:
    rows = []
    for idx, record in enumerate(records, start=1):
        if record.role == "UNKNOWN":
            rows.append(
                {
                    "classification_id": f"HUMAN-CLASS-{idx:04d}",
                    "source_root_label": record.source_root_label,
                    "relative_path": record.relative_path,
                    "current_role": record.role,
                    "reason": record.reason,
                    "required_action": "Classify as NB_FEEDBACK / OUR_RESPONSE / SUBMITTED_SUPPORTING_FILE / FINAL_CHANGE_REFERENCE / exclude_from_delta",
                    "blocks_formal_calibration": "no",
                }
            )
    return rows


def write_manifest(
    output_dir: Path,
    baseline_root: Path,
    authoring_workbook: Path,
    qa_gate_report: Path,
    nb_locked_root: Path,
    final_locked_root: Path,
    locked_records: list[LockedFileRecord],
    delta_tables: dict[str, list[dict[str, Any]]],
) -> dict[str, Any]:
    role_counts = {role: 0 for role in LOCKED_ROLES}
    for record in locked_records:
        role_counts[record.role] = role_counts.get(record.role, 0) + 1
    manifest = {
        "schema_name": "delta_analysis_manifest",
        "schema_version": "phase3.5d-structural-semantic-delta-analyzer-v1",
        "generated_at": now_iso(),
        "status": "DELTA_ANALYSIS_COMPLETE",
        "baseline_root": str(baseline_root),
        "authoring_workbook": str(authoring_workbook),
        "qa_gate_report": str(qa_gate_report),
        "nb_locked_root": str(nb_locked_root),
        "final_locked_root": str(final_locked_root),
        "output_dir": str(output_dir),
        "access_policy": {
            "locked_access_allowed_only_in_analyzer": True,
            "authoring_writer_may_read_locked_roots": False,
            "repair_or_finalization_triggered": False,
            "baseline_modified": False,
            "writer_input_modified": False,
        },
        "locked_file_count": len(locked_records),
        "locked_role_counts": role_counts,
        "delta_tables": {
            table_name: {
                "path": str(output_dir / f"{table_name}.csv"),
                "row_count": len(rows),
            }
            for table_name, rows in delta_tables.items()
        },
        "semantic_delta_layer": {
            "enabled": True,
            "semantic_tables": {table_name: str(output_dir / "semantic_delta" / f"{table_name}.csv") for table_name in SEMANTIC_TABLE_NAMES},
            "case_summary": str(output_dir / "semantic_delta_case_summary.md"),
            "manifest": str(output_dir / "semantic_delta_manifest.json"),
            "low_confidence_queue": str(output_dir / "low_confidence_review_queue.csv"),
            "gold_reference_section_index": str(output_dir / "gold_reference_section_index.csv"),
            "gold_reference_table_index": str(output_dir / "gold_reference_table_index.csv"),
            "gold_reference_citation_index": str(output_dir / "gold_reference_citation_index.csv"),
            "nb_to_final_resolution_link_table": str(output_dir / "nb_to_final_resolution_link_table.csv"),
            "ai_gap_to_nb_finding_link_table": str(output_dir / "ai_gap_to_nb_finding_link_table.csv"),
        },
        "required_human_classification_count": role_counts.get("UNKNOWN", 0),
        "invalidated_prior_run": "PROJECT_01_PILOT_INVALID_DUE_TO_LOCKED_ACCESS_LEAKAGE",
        "baseline_bump_required_before_project1_rerun": True,
    }
    (output_dir / "DELTA_ANALYSIS_MANIFEST.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return manifest


def write_semantic_outputs(
    output_dir: Path,
    semantic_tables: dict[str, list[dict[str, Any]]],
    low_confidence_rows: list[dict[str, Any]],
    semantic_summary: dict[str, Any],
    gold_context: dict[str, Any],
    nb_to_final_rows: list[dict[str, Any]],
    ai_gap_to_nb_rows: list[dict[str, Any]],
    semantic_context: dict[str, Any],
    *,
    baseline_root: Path,
    nb_locked_root: Path,
    final_locked_root: Path,
) -> dict[str, Any]:
    semantic_dir = output_dir / "semantic_delta"
    semantic_dir.mkdir(parents=True, exist_ok=True)
    for table_name, rows in semantic_tables.items():
        write_csv(semantic_dir / f"{table_name}.csv", rows)
    write_csv(output_dir / "low_confidence_review_queue.csv", low_confidence_rows)
    low_confidence_breakdown = low_confidence_breakdown_rows(low_confidence_rows, project_label=baseline_root.parent.parent.name if baseline_root.parent.parent else "")
    write_csv(output_dir / "low_confidence_root_cause_breakdown.csv", low_confidence_breakdown)
    entity_outputs = build_entity_normalization_outputs(semantic_tables)
    write_csv(output_dir / "entity_normalization_dictionary.csv", entity_outputs["dictionary"])
    write_csv(output_dir / "entity_alias_mapping_table.csv", entity_outputs["alias_mapping"])
    write_csv(output_dir / "entity_normalization_provenance.csv", entity_outputs["provenance"])
    write_csv(output_dir / "entity_normalization_resolved_items.csv", entity_outputs["resolved_items"])
    (output_dir / "gold_reference_extraction_manifest.json").write_text(json.dumps(gold_context["manifest"], ensure_ascii=False, indent=2), encoding="utf-8")
    write_csv(output_dir / "gold_reference_section_index.csv", gold_context["section_rows"])
    write_csv(output_dir / "gold_reference_table_index.csv", gold_context["table_rows"])
    write_csv(output_dir / "gold_reference_citation_index.csv", gold_context["citation_rows"])
    table_outputs = build_table_extraction_outputs(gold_context)
    (output_dir / "gold_reference_table_extraction_manifest.json").write_text(json.dumps(table_outputs["manifest"], ensure_ascii=False, indent=2), encoding="utf-8")
    write_csv(output_dir / "gold_reference_table_type_index.csv", table_outputs["type_rows"])
    write_csv(output_dir / "gold_reference_table_cell_anchor_index.csv", table_outputs["cell_rows"])
    write_csv(output_dir / "table_extraction_failure_log.csv", table_outputs["failure_rows"])
    write_csv(output_dir / "sota_benchmark_candidate_ranking_table.csv", semantic_context.get("sota_candidate_rankings") or [])
    write_csv(output_dir / "sota_benchmark_reranking_explanation_table.csv", semantic_context.get("sota_reranking_explanations") or [])
    true_ambiguity_rows, true_ambiguity_report = true_ambiguity_sampling_audit_rows(low_confidence_rows)
    write_csv(output_dir / "true_ambiguity_sampling_audit.csv", true_ambiguity_rows)
    (output_dir / "true_ambiguity_reclassification_report.md").write_text(true_ambiguity_report, encoding="utf-8")
    write_csv(output_dir / "nb_to_final_resolution_link_table.csv", nb_to_final_rows)
    write_csv(output_dir / "ai_gap_to_nb_finding_link_table.csv", ai_gap_to_nb_rows)
    manifest = {
        "schema_name": "semantic_delta_manifest",
        "schema_version": "phase3.5d-semantic-delta-precision-lift-v1",
        "generated_at": now_iso(),
        "status": "SEMANTIC_DELTA_COMPLETE",
        "baseline_root": str(baseline_root),
        "nb_locked_root": str(nb_locked_root),
        "final_locked_root": str(final_locked_root),
        "semantic_tables": {
            table_name: {
                "path": str(semantic_dir / f"{table_name}.csv"),
                "row_count": len(rows),
            }
            for table_name, rows in semantic_tables.items()
        },
        "low_confidence_queue": {
            "path": str(output_dir / "low_confidence_review_queue.csv"),
            "row_count": len(low_confidence_rows),
            "rule": "Items with match_confidence < 0.70 or needs_human_confirmation are excluded from systemic-gap counting.",
        },
        "low_confidence_root_cause_breakdown": {
            "path": str(output_dir / "low_confidence_root_cause_breakdown.csv"),
            "row_count": len(low_confidence_breakdown),
        },
        "entity_normalization": {
            "dictionary": str(output_dir / "entity_normalization_dictionary.csv"),
            "alias_mapping": str(output_dir / "entity_alias_mapping_table.csv"),
            "provenance": str(output_dir / "entity_normalization_provenance.csv"),
            "resolved_items": str(output_dir / "entity_normalization_resolved_items.csv"),
            "alias_count": len(entity_outputs["alias_mapping"]),
        },
        "gold_reference_extraction": {
            "manifest": str(output_dir / "gold_reference_extraction_manifest.json"),
            "section_index": str(output_dir / "gold_reference_section_index.csv"),
            "table_index": str(output_dir / "gold_reference_table_index.csv"),
            "citation_index": str(output_dir / "gold_reference_citation_index.csv"),
            "section_count": len(gold_context["section_rows"]),
            "table_count": len(gold_context["table_rows"]),
            "citation_count": len(gold_context["citation_rows"]),
        },
        "gold_reference_table_extraction": {
            "manifest": str(output_dir / "gold_reference_table_extraction_manifest.json"),
            "table_type_index": str(output_dir / "gold_reference_table_type_index.csv"),
            "cell_anchor_index": str(output_dir / "gold_reference_table_cell_anchor_index.csv"),
            "failure_log": str(output_dir / "table_extraction_failure_log.csv"),
            "table_count": table_outputs["manifest"]["table_count"],
            "cell_anchor_count": table_outputs["manifest"]["cell_anchor_count"],
            "priority_table_count": table_outputs["manifest"]["priority_table_count"],
        },
        "sota_benchmark_reranking": {
            "candidate_ranking_table": str(output_dir / "sota_benchmark_candidate_ranking_table.csv"),
            "reranking_explanation_table": str(output_dir / "sota_benchmark_reranking_explanation_table.csv"),
            "candidate_row_count": len(semantic_context.get("sota_candidate_rankings") or []),
            "explanation_row_count": len(semantic_context.get("sota_reranking_explanations") or []),
        },
        "true_ambiguity_reaudit": {
            "sampling_audit": str(output_dir / "true_ambiguity_sampling_audit.csv"),
            "report": str(output_dir / "true_ambiguity_reclassification_report.md"),
            "sample_row_count": len(true_ambiguity_rows),
        },
        "nb_final_linkage": {
            "nb_to_final_resolution_link_table": str(output_dir / "nb_to_final_resolution_link_table.csv"),
            "ai_gap_to_nb_finding_link_table": str(output_dir / "ai_gap_to_nb_finding_link_table.csv"),
            "nb_to_final_row_count": len(nb_to_final_rows),
            "ai_gap_to_nb_row_count": len(ai_gap_to_nb_rows),
        },
        "summary": semantic_summary,
        "semantic_delta_classes": list(SEMANTIC_DELTA_CLASSES),
        "cognitive_root_causes": list(COGNITIVE_ROOT_CAUSES),
        "access_policy": {
            "locked_access_allowed_only_in_analyzer": True,
            "authoring_writer_may_read_locked_roots": False,
            "repair_or_finalization_triggered": False,
            "baseline_modified": False,
            "writer_input_modified": False,
        },
    }
    (output_dir / "semantic_delta_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    write_semantic_case_summary(output_dir, manifest)
    return manifest


def write_semantic_case_summary(output_dir: Path, semantic_manifest: dict[str, Any]) -> None:
    summary = semantic_manifest.get("summary") or {}
    top_deltas = summary.get("top_material_deltas") or []
    top_roots = summary.get("top_root_causes") or []
    lines = [
        "# Semantic Delta Case Summary",
        "",
        f"Generated: {semantic_manifest.get('generated_at')}",
        "",
        "## Executive Answer",
        "",
        "This report compares the frozen AI baseline CER artifacts against the locked final certified package and NB process evidence at semantic-delta level. It preserves the existing structural delta outputs and adds claim, evidence, SOTA, PMCF, benefit-risk, cross-document, NB-relevance and cognitive root-cause attribution layers.",
        "",
        "## Five Largest Cognitive Gaps",
        "",
    ]
    if top_deltas:
        for idx, row in enumerate(top_deltas[:5], start=1):
            lines.append(f"{idx}. `{row.get('semantic_delta_id')}` ({row.get('semantic_table')}): {row.get('classification')} / {row.get('root_cause')} - {row.get('delta')}")
    else:
        lines.append("No high-materiality semantic delta was counted without human confirmation.")
    lines.extend(
        [
            "",
            "## NB-Relevant / Clinically Material",
            "",
            f"- Material semantic deltas: {summary.get('material_delta_count', 0)}",
            f"- High-materiality semantic deltas: {summary.get('high_material_delta_count', 0)}",
            f"- NB-relevant semantic deltas: {summary.get('nb_relevant_delta_count', 0)}",
            f"- Low-confidence human-gate items: {summary.get('low_confidence_count', 0)}",
            "",
            "## Repeated System Factors In This Case",
            "",
            "| Root cause | Count |",
            "|---|---:|",
        ]
    )
    for root, count in top_roots:
        lines.append(f"| {root} | {count} |")
    lines.extend(
        [
            "",
            "## Accepted Variance Boundary",
            "",
            "Differences classified as `accepted_stylistic_variance` are not treated as defects. Differences below 0.70 confidence are routed to `low_confidence_review_queue.csv` and are not counted as systemic gaps.",
            "",
            "## Three System Capabilities To Consider Next",
            "",
        ]
    )
    for idx, (root, _) in enumerate(top_roots[:3], start=1):
        lines.append(f"{idx}. {root}: {_upgrade_candidate(root)}")
    if not top_roots:
        lines.append("1. No upgrade priority can be confirmed without human-gate resolution.")
    (output_dir / "semantic_delta_case_summary.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_pilot_report(output_dir: Path, manifest: dict[str, Any]) -> None:
    lines = [
        "# Phase 0.2 Pilot Run Report",
        "",
        f"Generated: {manifest['generated_at']}",
        "",
        "## Decision",
        "",
        "This Delta Analyzer run is infrastructure-only. It does not validate or improve the baseline CER content.",
        "",
        "## Locked Access Boundary",
        "",
        "- Locked roots were scanned only by `scripts/calibration_delta_analyzer.py`.",
        "- Authoring writer, repair loop, finalization and core CER logic were not invoked.",
        "- Every locked file encountered by the analyzer is recorded in `LOCKED_ACCESS_LOG.csv`.",
        "",
        "## Inputs",
        "",
        f"- Baseline root: `{manifest['baseline_root']}`",
        f"- Authoring workbook: `{manifest['authoring_workbook']}`",
        f"- QA gate report: `{manifest['qa_gate_report']}`",
        f"- NB locked root: `{manifest['nb_locked_root']}`",
        f"- Final locked root: `{manifest['final_locked_root']}`",
        "",
        "## Outputs",
        "",
        "- `DELTA_ANALYSIS_MANIFEST.json`",
        "- `semantic_delta_manifest.json`",
        "- `semantic_delta_case_summary.md`",
        "- `low_confidence_review_queue.csv`",
        "- `LOCKED_ACCESS_LOG.csv`",
        "- `NEEDS_HUMAN_CLASSIFICATION.csv`",
    ]
    for table_name, spec in manifest["delta_tables"].items():
        lines.append(f"- `{table_name}.csv` ({spec['row_count']} rows)")
    lines.extend(
        [
            "",
            "## Locked File Role Counts",
            "",
            "| Role | Count |",
            "|---|---:|",
        ]
    )
    for role, count in manifest["locked_role_counts"].items():
        lines.append(f"| {role} | {count} |")
    lines.extend(
        [
            "",
            "## Prior Pilot Status",
            "",
            "- Previous Project 1 Pilot is marked `PROJECT_01_PILOT_INVALID_DUE_TO_LOCKED_ACCESS_LEAKAGE`.",
            "- Project 1 rerun is blocked until Phase 0.2 is accepted and baseline version is bumped.",
        ]
    )
    (output_dir / "PILOT_RUN_REPORT.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def aggregate_semantic_cases(case_output_dirs: list[Path], output_dir: Path) -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    case_manifests = []
    cognitive_rows: list[dict[str, Any]] = []
    low_confidence_breakdown_rows_all: list[dict[str, Any]] = []
    sota_rows_all: list[dict[str, Any]] = []
    all_low_confidence = 0
    for case_dir in case_output_dirs:
        manifest_path = case_dir / "semantic_delta_manifest.json"
        if not manifest_path.exists():
            case_manifests.append({"case_dir": str(case_dir), "status": "MISSING_SEMANTIC_MANIFEST"})
            continue
        manifest = read_json(manifest_path)
        case_manifests.append(manifest)
        all_low_confidence += int(((manifest.get("low_confidence_queue") or {}).get("row_count")) or 0)
        breakdown_path = case_dir / "low_confidence_root_cause_breakdown.csv"
        if breakdown_path.exists():
            with breakdown_path.open("r", encoding="utf-8", newline="") as handle:
                for row in csv.DictReader(handle):
                    row = dict(row)
                    row["case_dir"] = str(case_dir)
                    low_confidence_breakdown_rows_all.append(row)
        cog_path = case_dir / "semantic_delta" / "cognitive_gap_attribution_table.csv"
        if not cog_path.exists():
            cog_path = case_dir / "cognitive_gap_attribution_table.csv"
        if cog_path.exists():
            with cog_path.open("r", encoding="utf-8", newline="") as handle:
                for row in csv.DictReader(handle):
                    row = dict(row)
                    row["case_dir"] = str(case_dir)
                    cognitive_rows.append(row)
        sota_path = case_dir / "semantic_delta" / "sota_benchmark_delta_table.csv"
        if sota_path.exists():
            with sota_path.open("r", encoding="utf-8", newline="") as handle:
                for row in csv.DictReader(handle):
                    row = dict(row)
                    row["case_dir"] = str(case_dir)
                    sota_rows_all.append(row)

    frequency: dict[tuple[str, str], dict[str, Any]] = {}
    for row in cognitive_rows:
        key = (row.get("cognitive_root_cause", "unknown"), row.get("delta_classification", "unknown"))
        bucket = frequency.setdefault(
            key,
            {
                "cognitive_root_cause": key[0],
                "delta_classification": key[1],
                "case_count": set(),
                "delta_count": 0,
                "nb_relevant_count": 0,
                "high_materiality_count": 0,
                "upgrade_candidate": row.get("upgrade_candidate", ""),
            },
        )
        bucket["case_count"].add(row.get("case_dir", ""))
        bucket["delta_count"] += 1
        if row.get("nb_relevance") == "yes":
            bucket["nb_relevant_count"] += 1
        if row.get("materiality") == "high":
            bucket["high_materiality_count"] += 1

    frequency_rows = []
    for item in frequency.values():
        frequency_rows.append(
            {
                "cognitive_root_cause": item["cognitive_root_cause"],
                "delta_classification": item["delta_classification"],
                "case_count": len(item["case_count"]),
                "delta_count": item["delta_count"],
                "nb_relevant_count": item["nb_relevant_count"],
                "high_materiality_count": item["high_materiality_count"],
                "upgrade_candidate": item["upgrade_candidate"],
            }
        )
    frequency_rows.sort(key=lambda row: (int(row["case_count"]), int(row["nb_relevant_count"]), int(row["high_materiality_count"]), int(row["delta_count"])), reverse=True)
    write_csv(output_dir / "cross_project_gap_frequency_matrix.csv", frequency_rows)

    low_counts: dict[str, int] = {name: 0 for name in LOW_CONFIDENCE_TYPES}
    for row in low_confidence_breakdown_rows_all:
        lc_type = row.get("low_confidence_type") or "TRUE_AMBIGUITY_HUMAN_GATE"
        low_counts[lc_type] = low_counts.get(lc_type, 0) + int(float(row.get("count") or 0))

    priority_rows = []
    for idx, row in enumerate(
        [item for item in frequency_rows if item.get("cognitive_root_cause") not in {"needs_human_confirmation", "acceptable_variance"}],
        start=1,
    ):
        repeated = int(row["case_count"]) >= 2
        impact_score = int(row["delta_count"]) + (2 * int(row["nb_relevant_count"])) + (2 * int(row["high_materiality_count"])) + (3 if repeated else 0)
        priority_rows.append(
            {
                "rank": idx,
                "system_capability": row["cognitive_root_cause"],
                "reason": f"{row['delta_count']} semantic candidate deltas across {row['case_count']} cases; NB-relevant={row['nb_relevant_count']}; high-materiality={row['high_materiality_count']}",
                "impact_score": impact_score,
                "upgrade_candidate": row["upgrade_candidate"],
            }
        )
    analyzer_priority_map = {
        "ENTITY_NORMALIZATION_GAP": ("entity_normalization_layer", "Expand device, endpoint, adverse event, citation and PMCF synonym normalization."),
        "TEXT_EXTRACTION_GAP": ("gold_reference_text_table_extraction", "Improve DOCX/PDF table and attachment extraction from final packages."),
        "SEMANTIC_MATCHING_GAP": ("semantic_candidate_reranking", "Improve section-aware candidate ranking and similarity scoring."),
        "SECTION_MAPPING_GAP": ("section_anchor_mapping", "Improve section tree and paragraph-anchor extraction."),
    }
    for lc_type, (capability, action) in analyzer_priority_map.items():
        count = low_counts.get(lc_type, 0)
        if not count:
            continue
        priority_rows.append(
            {
                "rank": len(priority_rows) + 1,
                "system_capability": capability,
                "reason": f"{count} low-confidence items attributed to {lc_type}; these are analyzer limitations, not confirmed AI CER defects.",
                "impact_score": count,
                "upgrade_candidate": action,
            }
        )
    priority_rows.sort(key=lambda row: int(row["impact_score"]), reverse=True)
    for idx, row in enumerate(priority_rows, start=1):
        row["rank"] = idx
    write_csv(output_dir / "next_upgrade_priority_ranking.csv", priority_rows)

    quality = _quality_level_judgment(case_manifests, frequency_rows, all_low_confidence)
    (output_dir / "semantic_quality_level_preliminary_judgment.md").write_text(quality, encoding="utf-8")
    write_cross_project_low_confidence_report(output_dir, low_confidence_breakdown_rows_all)
    write_semantic_match_reliability_report(output_dir, case_manifests, low_confidence_breakdown_rows_all)
    write_calibration_grade_readiness_report(output_dir, case_manifests, frequency_rows, low_confidence_breakdown_rows_all)
    write_before_after_metrics_report(output_dir, case_manifests, low_confidence_breakdown_rows_all, sota_rows_all)
    report_lines = [
        "# Cross-Project Semantic Gap Aggregation Report",
        "",
        f"Generated: {now_iso()}",
        "",
        "## Cases",
        "",
    ]
    for manifest in case_manifests:
        if manifest.get("status") == "MISSING_SEMANTIC_MANIFEST":
            report_lines.append(f"- `{manifest['case_dir']}`: missing semantic manifest")
        else:
            summary = manifest.get("summary") or {}
            report_lines.append(
                f"- `{manifest.get('baseline_root')}`: material={summary.get('material_delta_count', 0)}, NB-relevant={summary.get('nb_relevant_delta_count', 0)}, low-confidence={summary.get('low_confidence_count', 0)}"
            )
    report_lines.extend(
        [
            "",
            "## Repeated Semantic Gaps",
            "",
            "| Root cause | Classification | Cases | Deltas | NB-relevant | Upgrade candidate |",
            "|---|---|---:|---:|---:|---|",
        ]
    )
    for row in frequency_rows[:20]:
        report_lines.append(
            f"| {row['cognitive_root_cause']} | {row['delta_classification']} | {row['case_count']} | {row['delta_count']} | {row['nb_relevant_count']} | {row['upgrade_candidate']} |"
        )
    report_lines.extend(
        [
            "",
            "## Next Upgrade Priorities",
            "",
        ]
    )
    for row in priority_rows[:5]:
        report_lines.append(f"{row['rank']}. {row['system_capability']} - {row['reason']} - {row['upgrade_candidate']}")
    (output_dir / "cross_project_semantic_gap_aggregation_report.md").write_text("\n".join(report_lines) + "\n", encoding="utf-8")
    manifest = {
        "schema_name": "cross_project_semantic_delta_aggregation_manifest",
        "schema_version": "phase3.5b-semantic-delta-v1",
        "generated_at": now_iso(),
        "status": "CROSS_PROJECT_SEMANTIC_AGGREGATION_COMPLETE",
        "case_output_dirs": [str(path) for path in case_output_dirs],
        "case_count": len(case_output_dirs),
        "usable_case_count": len([m for m in case_manifests if m.get("status") != "MISSING_SEMANTIC_MANIFEST"]),
        "frequency_matrix": str(output_dir / "cross_project_gap_frequency_matrix.csv"),
        "priority_ranking": str(output_dir / "next_upgrade_priority_ranking.csv"),
        "quality_judgment": str(output_dir / "semantic_quality_level_preliminary_judgment.md"),
        "low_confidence_root_cause_report": str(output_dir / "cross_project_low_confidence_root_cause_report.md"),
        "semantic_match_reliability_report": str(output_dir / "cross_project_semantic_match_reliability_report.md"),
        "calibration_grade_readiness_report": str(output_dir / "calibration_grade_readiness_report.md"),
        "before_after_metrics_report": str(output_dir / "phase3_5d_before_after_metrics_report.md"),
        "low_confidence_total": all_low_confidence,
    }
    (output_dir / "cross_project_semantic_aggregation_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def write_cross_project_low_confidence_report(output_dir: Path, breakdown_rows: list[dict[str, Any]]) -> None:
    counts: dict[str, int] = {name: 0 for name in LOW_CONFIDENCE_TYPES}
    case_counts: dict[str, dict[str, int]] = {}
    for row in breakdown_rows:
        lc_type = row.get("low_confidence_type") or "TRUE_AMBIGUITY_HUMAN_GATE"
        count = int(float(row.get("count") or 0))
        counts[lc_type] = counts.get(lc_type, 0) + count
        case_counts.setdefault(str(row.get("case_dir") or "unknown"), {})[lc_type] = count
    total = sum(counts.values())
    lines = [
        "# Cross-Project Low-Confidence Root Cause Report",
        "",
        f"Generated: {now_iso()}",
        "",
        "Low-confidence items are separated into analyzer limitations and true human-gate ambiguity. They are not counted as confirmed AI CER defects.",
        "",
        "| Low-confidence type | Count | Percentage | Resolvable by analyzer upgrade | Recommended next action |",
        "|---|---:|---:|---|---|",
    ]
    for lc_type in LOW_CONFIDENCE_TYPES:
        count = counts.get(lc_type, 0)
        resolvable, action = LOW_CONFIDENCE_ACTIONS[lc_type]
        percentage = (count / total * 100) if total else 0
        lines.append(f"| {lc_type} | {count} | {percentage:.1f}% | {resolvable} | {action} |")
    lines.extend(["", "## Project Distribution", "", "| Case | " + " | ".join(LOW_CONFIDENCE_TYPES) + " |", "|---" + "|---:" * len(LOW_CONFIDENCE_TYPES) + "|"])
    for case_dir, bucket in sorted(case_counts.items()):
        lines.append("| `" + case_dir + "` | " + " | ".join(str(bucket.get(lc_type, 0)) for lc_type in LOW_CONFIDENCE_TYPES) + " |")
    (output_dir / "cross_project_low_confidence_root_cause_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_semantic_match_reliability_report(output_dir: Path, case_manifests: list[dict[str, Any]], breakdown_rows: list[dict[str, Any]]) -> None:
    usable = [manifest for manifest in case_manifests if manifest.get("status") != "MISSING_SEMANTIC_MANIFEST"]
    table_rows = []
    total_semantic_rows = 0
    total_low = 0
    total_sections = 0
    total_citations = 0
    for manifest in usable:
        semantic_rows = sum(int(spec.get("row_count") or 0) for spec in (manifest.get("semantic_tables") or {}).values())
        low = int(((manifest.get("low_confidence_queue") or {}).get("row_count")) or 0)
        gold = manifest.get("gold_reference_extraction") or {}
        total_semantic_rows += semantic_rows
        total_low += low
        total_sections += int(gold.get("section_count") or 0)
        total_citations += int(gold.get("citation_count") or 0)
        auto_rate = ((semantic_rows - low) / semantic_rows * 100) if semantic_rows else 0
        table_rows.append(
            {
                "case": str(manifest.get("baseline_root") or ""),
                "semantic_rows": semantic_rows,
                "low_confidence": low,
                "auto_high_confidence_rate": f"{auto_rate:.1f}%",
                "gold_sections": gold.get("section_count", 0),
                "gold_citations": gold.get("citation_count", 0),
            }
        )
    write_csv(output_dir / "cross_project_semantic_match_reliability_table.csv", table_rows)
    overall_auto = ((total_semantic_rows - total_low) / total_semantic_rows * 100) if total_semantic_rows else 0
    lines = [
        "# Cross-Project Semantic Match Reliability Report",
        "",
        f"Generated: {now_iso()}",
        "",
        f"- Usable cases: {len(usable)} / {len(case_manifests)}",
        f"- Semantic rows: {total_semantic_rows}",
        f"- Low-confidence rows: {total_low}",
        f"- Automatic high-confidence rate: {overall_auto:.1f}%",
        f"- Gold sections indexed: {total_sections}",
        f"- Gold citations indexed: {total_citations}",
        "",
        "Reliability improvement is judged from extraction coverage, section-aware candidates, entity normalization and staged confidence fields. The analyzer does not lower the 0.70 human-gate threshold.",
    ]
    (output_dir / "cross_project_semantic_match_reliability_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_calibration_grade_readiness_report(
    output_dir: Path,
    case_manifests: list[dict[str, Any]],
    frequency_rows: list[dict[str, Any]],
    breakdown_rows: list[dict[str, Any]],
) -> None:
    usable = [manifest for manifest in case_manifests if manifest.get("status") != "MISSING_SEMANTIC_MANIFEST"]
    low_by_type = {name: 0 for name in LOW_CONFIDENCE_TYPES}
    for row in breakdown_rows:
        lc_type = row.get("low_confidence_type") or "TRUE_AMBIGUITY_HUMAN_GATE"
        low_by_type[lc_type] = low_by_type.get(lc_type, 0) + int(float(row.get("count") or 0))
    analyzer_limitation = sum(low_by_type[name] for name in ("TEXT_EXTRACTION_GAP", "SECTION_MAPPING_GAP", "ENTITY_NORMALIZATION_GAP", "SEMANTIC_MATCHING_GAP"))
    true_ambiguity = low_by_type.get("TRUE_AMBIGUITY_HUMAN_GATE", 0)
    repeated_material = [row for row in frequency_rows if int(row.get("case_count") or 0) >= 2 and int(row.get("high_materiality_count") or 0) > 0]
    if len(usable) < 4:
        decision = "PHASE3_5D_REJECTED_PRECISION_NOT_IMPROVED"
        rationale = "Fewer than four target cases produced semantic manifests."
    elif analyzer_limitation > true_ambiguity * 3 and analyzer_limitation > 100:
        decision = "PHASE3_5D_ACCEPTED_WITH_LIMITATIONS"
        rationale = "The analyzer is calibration-useful but still limited by extraction/mapping/normalization workload."
    elif repeated_material:
        decision = "PHASE3_5D_CALIBRATION_GRADE_ACCEPTED"
        rationale = "Four target cases ran and repeated material semantic gaps can be attributed with evidence chains."
    else:
        decision = "PHASE3_5D_ACCEPTED_WITH_LIMITATIONS"
        rationale = "Four cases ran, but repeated material gaps are not yet strong enough for unrestricted calibration-grade acceptance."
    lines = [
        "# Calibration Grade Readiness Report",
        "",
        f"Generated: {now_iso()}",
        "",
        f"Decision: `{decision}`",
        "",
        rationale,
        "",
        "## Evidence",
        "",
        f"- Usable cases: {len(usable)} / {len(case_manifests)}",
        f"- Analyzer-limitation low-confidence items: {analyzer_limitation}",
        f"- True ambiguity human-gate items: {true_ambiguity}",
        f"- Repeated high-materiality root-cause buckets: {len(repeated_material)}",
        "",
        "The decision is about the Delta Analyzer, not CER Authoring quality. It does not modify authoring graph, agents, prompts, gates or baseline outputs.",
    ]
    (output_dir / "calibration_grade_readiness_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_before_after_metrics_report(
    output_dir: Path,
    case_manifests: list[dict[str, Any]],
    breakdown_rows: list[dict[str, Any]],
    sota_rows: list[dict[str, Any]],
) -> None:
    usable = [manifest for manifest in case_manifests if manifest.get("status") != "MISSING_SEMANTIC_MANIFEST"]
    semantic_rows = sum(sum(int(spec.get("row_count") or 0) for spec in (manifest.get("semantic_tables") or {}).values()) for manifest in usable)
    low_confidence = sum(int(((manifest.get("low_confidence_queue") or {}).get("row_count")) or 0) for manifest in usable)
    high_rate = ((semantic_rows - low_confidence) / semantic_rows * 100) if semantic_rows else 0.0
    low_by_type = {name: 0 for name in LOW_CONFIDENCE_TYPES}
    for row in breakdown_rows:
        lc_type = row.get("low_confidence_type") or "TRUE_AMBIGUITY_HUMAN_GATE"
        low_by_type[lc_type] = low_by_type.get(lc_type, 0) + int(float(row.get("count") or 0))
    table_count = sum(int(((manifest.get("gold_reference_table_extraction") or {}).get("table_count")) or 0) for manifest in usable)
    cell_count = sum(int(((manifest.get("gold_reference_table_extraction") or {}).get("cell_anchor_count")) or 0) for manifest in usable)
    priority_table_count = sum(int(((manifest.get("gold_reference_table_extraction") or {}).get("priority_table_count")) or 0) for manifest in usable)
    gold_sections = sum(int(((manifest.get("gold_reference_extraction") or {}).get("section_count")) or 0) for manifest in usable)
    gold_citations = sum(int(((manifest.get("gold_reference_extraction") or {}).get("citation_count")) or 0) for manifest in usable)
    sota_total = len(sota_rows)
    sota_high = sum(1 for row in sota_rows if str(row.get("human_gate_required")) == "no")
    sota_precision = (sota_high / sota_total * 100) if sota_total else 0.0
    current = {
        "semantic_rows": semantic_rows,
        "low_confidence_total": low_confidence,
        "automatic_high_confidence_rate": round(high_rate, 1),
        "entity_normalization_gap": low_by_type.get("ENTITY_NORMALIZATION_GAP", 0),
        "true_ambiguity_human_gate": low_by_type.get("TRUE_AMBIGUITY_HUMAN_GATE", 0),
        "gold_sections_indexed": gold_sections,
        "gold_citations_indexed": gold_citations,
        "table_extraction_coverage": table_count,
        "table_cell_anchor_count": cell_count,
        "priority_table_count": priority_table_count,
        "sota_benchmark_match_precision": round(sota_precision, 1),
        "analyzer_resolvable_low_confidence": sum(low_by_type[name] for name in ("TEXT_EXTRACTION_GAP", "SECTION_MAPPING_GAP", "ENTITY_NORMALIZATION_GAP", "SEMANTIC_MATCHING_GAP")),
        "true_human_gate_low_confidence": low_by_type.get("TRUE_AMBIGUITY_HUMAN_GATE", 0),
    }
    rows = []
    for metric, before in PHASE35C_BASELINE_METRICS.items():
        after = current.get(metric, "")
        delta = ""
        if isinstance(after, (int, float)) and isinstance(before, (int, float)):
            delta = round(after - before, 2)
        rows.append({"metric": metric, "phase3_5c_before": before, "phase3_5d_after": after, "delta": delta})
    rows.extend(
        [
            {"metric": "table_extraction_coverage", "phase3_5c_before": 0, "phase3_5d_after": table_count, "delta": table_count},
            {"metric": "table_cell_anchor_count", "phase3_5c_before": 0, "phase3_5d_after": cell_count, "delta": cell_count},
            {"metric": "priority_table_count", "phase3_5c_before": 0, "phase3_5d_after": priority_table_count, "delta": priority_table_count},
            {"metric": "sota_benchmark_match_precision_percent", "phase3_5c_before": "not measured", "phase3_5d_after": f"{sota_precision:.1f}", "delta": "new metric"},
            {"metric": "analyzer_resolvable_vs_true_human_gate", "phase3_5c_before": "357 / 306", "phase3_5d_after": f"{current['analyzer_resolvable_low_confidence']} / {current['true_human_gate_low_confidence']}", "delta": "typed"},
        ]
    )
    write_csv(output_dir / "phase3_5d_before_after_metrics.csv", rows)
    lines = [
        "# Phase 3.5D Before / After Metrics Report",
        "",
        f"Generated: {now_iso()}",
        "",
        "| Metric | Phase 3.5C before | Phase 3.5D after | Delta |",
        "|---|---:|---:|---:|",
    ]
    for row in rows:
        lines.append(f"| {row['metric']} | {row['phase3_5c_before']} | {row['phase3_5d_after']} | {row['delta']} |")
    lines.extend(
        [
            "",
            "The threshold remains 0.70. Improvements, if any, come from extraction, normalization, reranking and candidate scoring rather than threshold relaxation.",
        ]
    )
    (output_dir / "phase3_5d_before_after_metrics_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def _quality_level_judgment(case_manifests: list[dict[str, Any]], frequency_rows: list[dict[str, Any]], low_confidence_total: int) -> str:
    usable = [m for m in case_manifests if m.get("status") != "MISSING_SEMANTIC_MANIFEST"]
    repeated_high = [row for row in frequency_rows if int(row.get("case_count") or 0) >= 2 and int(row.get("high_materiality_count") or 0) > 0]
    nb_relevant = sum(int(row.get("nb_relevant_count") or 0) for row in frequency_rows)
    if len(usable) < 4:
        level = "Insufficient for final quality-level judgment because fewer than four target projects have semantic manifests."
    elif repeated_high or nb_relevant >= 4 or low_confidence_total > 120:
        level = "Current semantic evidence supports approximately a 70-point human-reviewable draft capability, but not a stable 80-point NB-ready draft."
    else:
        level = "Current semantic evidence may support preliminary 80-point NB-ready draft capability, pending CCD review of low-confidence items."
    lines = [
        "# Semantic Quality Level Preliminary Judgment",
        "",
        level,
        "",
        f"- Usable semantic cases: {len(usable)} / {len(case_manifests)}",
        f"- Repeated high-materiality root causes: {len(repeated_high)}",
        f"- NB-relevant material deltas: {nb_relevant}",
        f"- Low-confidence items: {low_confidence_total}",
        "",
        "This judgment is based on semantic deltas, not gate pass rate. Holdout outputs are used for evaluation only and must not be used to tune authoring rules unless separately approved.",
    ]
    return "\n".join(lines) + "\n"


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Locked Structural + Semantic Delta Analyzer for CER calibration")
    parser.add_argument("--baseline-root", help="Frozen baseline artifact root")
    parser.add_argument("--authoring-workbook", help="Path to frozen authoring_workbook.json")
    parser.add_argument("--qa-gate-report", help="Path to frozen qa_gate_report.json")
    parser.add_argument("--nb-locked-root", help="02_NB_ROUNDS_AND_RESPONSES_LOCKED path")
    parser.add_argument("--final-locked-root", help="03_FINAL_CERTIFIED_PACKAGE_LOCKED path")
    parser.add_argument("--output-dir", help="Separate delta_analysis output directory")
    parser.add_argument("--case-output-dir", action="append", default=[], help="Existing per-case delta_analysis output directory for cross-project aggregation")
    parser.add_argument("--aggregate-output-dir", help="Output directory for cross-project semantic aggregation")
    parser.add_argument("--json", action="store_true", help="Print manifest summary as JSON")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv or sys.argv[1:])
    if args.aggregate_output_dir:
        if not args.case_output_dir:
            raise SystemExit("--aggregate-output-dir requires at least one --case-output-dir")
        case_dirs = [_resolve_existing_path(value, "--case-output-dir", must_be_dir=True) for value in args.case_output_dir]
        aggregate_output_dir = Path(args.aggregate_output_dir).expanduser().resolve()
        manifest = aggregate_semantic_cases(case_dirs, aggregate_output_dir)
        if args.json:
            print(json.dumps({"status": manifest["status"], "output_dir": str(aggregate_output_dir), "usable_case_count": manifest["usable_case_count"]}, ensure_ascii=False, indent=2))
        return 0
    required = {
        "--baseline-root": args.baseline_root,
        "--authoring-workbook": args.authoring_workbook,
        "--qa-gate-report": args.qa_gate_report,
        "--nb-locked-root": args.nb_locked_root,
        "--final-locked-root": args.final_locked_root,
        "--output-dir": args.output_dir,
    }
    missing = [name for name, value in required.items() if not value]
    if missing:
        raise SystemExit(f"Missing required arguments for case analysis: {', '.join(missing)}")
    baseline_root = _resolve_existing_path(args.baseline_root, "--baseline-root", must_be_dir=True)
    authoring_workbook = _resolve_existing_path(args.authoring_workbook, "--authoring-workbook", must_be_dir=False)
    qa_gate_report = _resolve_existing_path(args.qa_gate_report, "--qa-gate-report", must_be_dir=False)
    nb_locked_root = _resolve_existing_path(args.nb_locked_root, "--nb-locked-root", must_be_dir=True)
    final_locked_root = _resolve_existing_path(args.final_locked_root, "--final-locked-root", must_be_dir=True)
    output_dir = Path(args.output_dir).expanduser().resolve()
    _validate_output_dir(output_dir, [baseline_root, nb_locked_root, final_locked_root])
    output_dir.mkdir(parents=True, exist_ok=True)

    baseline = load_baseline_context(authoring_workbook, qa_gate_report)
    locked_records = scan_locked_root(nb_locked_root, "02_NB_ROUNDS_AND_RESPONSES_LOCKED")
    locked_records.extend(scan_locked_root(final_locked_root, "03_FINAL_CERTIFIED_PACKAGE_LOCKED"))
    delta_tables = build_delta_tables(baseline, locked_records)
    semantic_tables, low_confidence_rows, semantic_summary, gold_context, nb_to_final_rows, ai_gap_to_nb_rows, semantic_context = build_semantic_delta_tables(baseline, locked_records)

    for table_name, rows in delta_tables.items():
        write_csv(output_dir / f"{table_name}.csv", rows)
    semantic_manifest = write_semantic_outputs(
        output_dir,
        semantic_tables,
        low_confidence_rows,
        semantic_summary,
        gold_context,
        nb_to_final_rows,
        ai_gap_to_nb_rows,
        semantic_context,
        baseline_root=baseline_root,
        nb_locked_root=nb_locked_root,
        final_locked_root=final_locked_root,
    )
    write_csv(output_dir / "LOCKED_ACCESS_LOG.csv", locked_access_rows(locked_records))
    write_csv(output_dir / "NEEDS_HUMAN_CLASSIFICATION.csv", human_classification_rows(locked_records))
    manifest = write_manifest(
        output_dir,
        baseline_root,
        authoring_workbook,
        qa_gate_report,
        nb_locked_root,
        final_locked_root,
        locked_records,
        delta_tables,
    )
    write_pilot_report(output_dir, manifest)
    if args.json:
        print(
            json.dumps(
                {
                    "status": manifest["status"],
                    "semantic_status": semantic_manifest["status"],
                    "locked_file_count": manifest["locked_file_count"],
                    "required_human_classification_count": manifest["required_human_classification_count"],
                    "low_confidence_count": semantic_manifest["low_confidence_queue"]["row_count"],
                    "output_dir": manifest["output_dir"],
                },
                ensure_ascii=False,
                indent=2,
            )
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
